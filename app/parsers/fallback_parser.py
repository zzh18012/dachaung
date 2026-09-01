"""降级解析器：pdfplumber（PDF）+ python-docx（DOCX）。

**为什么需要它**：Kreuzberg 4.10.2 实测对 DOCX 给不出 element-level 结构，
对手写最小 PDF 也给不出。本解析器直接调用底层库，提供：
- PDF：page + bbox + 文本 + 表格 + 图片（裁剪渲染）
- DOCX：段落、标题、表格、图片、题注（按文档顺序），section/paragraph_index/table_index
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from app.models import Document, Element, Relation, WarningRecord
from app.parsers.base import Parser, ParserError, detect_source_type, make_document_id
from app.parsers.table_linearize import linearize_table

# 这些库都是计划内依赖，且只在这个文件里 import（业务代码看不见）
try:
    import pdfplumber  # type: ignore[import-not-found]
    _PDFPLUMBER_VERSION = getattr(pdfplumber, "__version__", "unknown")
except ImportError as _e:  # pragma: no cover
    pdfplumber = None  # type: ignore[assignment]
    _PDFPLUMBER_VERSION = None
    _PDFPLUMBER_IMPORT_ERROR = str(_e)

try:
    import pypdfium2  # type: ignore[import-not-found]  # pdfplumber 的传递依赖
    _PDFIUM_VERSION = getattr(pypdfium2, "__version__", "unknown")
except ImportError as _e:  # pragma: no cover
    pypdfium2 = None  # type: ignore[assignment]
    _PDFIUM_VERSION = None
    _PDFIUM_IMPORT_ERROR = str(_e)

try:
    import docx  # type: ignore[import-not-found]  # python-docx 的 import 名是 docx
    from docx.oxml.ns import qn  # type: ignore[import-not-found]
    _DOCX_VERSION = getattr(docx, "__version__", "unknown")
except ImportError as _e:  # pragma: no cover
    docx = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    _DOCX_VERSION = None
    _DOCX_IMPORT_ERROR = str(_e)


# ---------- 通用 ----------

# 题注：以 "Table 1." / "Figure 1." / "表 1" / "图 1" 开头的段落
_CAPTION_RE = re.compile(
    r"^\s*(?:Table|Figure|Fig\.?|表|图)\s*[0-9０-９]+[\.、\s]",
    re.IGNORECASE,
)


def _is_caption(text: str) -> bool:
    return bool(_CAPTION_RE.match(text or ""))


# 批次 4 契约（docs/caption-relation-contract.md §1）：图题注前缀集，
# 数字限 ASCII；与元素分类用的 _CAPTION_RE（含全角）分工。
_FIGURE_CAPTION_RE = re.compile(
    r"^(?:Figure|Fig\.?|图)\s*[0-9]+[\.、\s]",
    re.IGNORECASE,
)

CAPTION_MAX_GAP_PT = 50.0


def _is_figure_caption(el: Element) -> bool:
    return el.type == "caption" and bool(
        _FIGURE_CAPTION_RE.match(el.content or "")
    )


def _sort_relations(rels: list[Relation]) -> list[Relation]:
    """契约 §2：relations 按 (type, from_id, to_id) 字典序稳定排序。"""
    return sorted(rels, key=lambda r: (r.type, r.from_id, r.to_id))


def match_caption_relations_docx(elements: list[Element]) -> list[Relation]:
    """契约 §3 docx 规则（纯函数）：image@P 仅当图题注@P+1 生成
    `image --has_caption--> caption`。"""
    captions = {
        e.source_locator.get("paragraph_index"): e
        for e in elements
        if _is_figure_caption(e)
    }
    rels: list[Relation] = []
    for e in elements:
        if e.type != "image":
            continue
        p = e.source_locator.get("paragraph_index")
        cap = captions.get(p + 1) if p is not None else None
        if cap is None:
            continue
        rels.append(
            Relation(
                type="has_caption",
                from_id=e.element_id,
                to_id=cap.element_id,
                metadata={"rule": "docx_adjacent_paragraph"},
            )
        )
    return _sort_relations(rels)


def match_caption_relations_pdf(elements: list[Element]) -> list[Relation]:
    """契约 §3 pdf 规则（纯函数）：同页 + 图题注在图下方
    0 < gap ≤ CAPTION_MAX_GAP_PT + x 区间相交 > 0；候选按
    (gap, image_id, caption_id) 升序贪心唯一配对。bbox 口径
    pdfplumber [x0, top, x1, bottom]。"""
    images = [e for e in elements if e.type == "image"]
    captions = [e for e in elements if _is_figure_caption(e)]
    candidates = []
    for img in images:
        iloc = img.source_locator
        ib = iloc.get("bbox")
        if not ib:
            continue
        for cap in captions:
            cloc = cap.source_locator
            cb = cloc.get("bbox")
            if not cb:
                continue
            if iloc.get("page") != cloc.get("page"):
                continue
            gap = cb[1] - ib[3]
            if not (0 < gap <= CAPTION_MAX_GAP_PT):
                continue
            if min(cb[2], ib[2]) - max(cb[0], ib[0]) <= 0:
                continue
            candidates.append((gap, img.element_id, cap.element_id))
    candidates.sort()
    used_img: set[str] = set()
    used_cap: set[str] = set()
    rels: list[Relation] = []
    for gap, img_id, cap_id in candidates:
        if img_id in used_img or cap_id in used_cap:
            continue
        used_img.add(img_id)
        used_cap.add(cap_id)
        rels.append(
            Relation(
                type="has_caption",
                from_id=img_id,
                to_id=cap_id,
                metadata={"rule": "pdf_geometry_below", "gap_pt": gap},
            )
        )
    return _sort_relations(rels)


# 批次 7 契约（docs/table-caption-relation-contract.md §2）：表题注前缀
# 集，与图题注前缀集互斥；数字限 ASCII。
_TABLE_CAPTION_RE = re.compile(
    r"^(?:Table|表格|表)\s*[0-9]+[\.、\s]",
    re.IGNORECASE,
)


def _is_table_caption(el: Element) -> bool:
    return el.type == "caption" and bool(
        _TABLE_CAPTION_RE.match(el.content or "")
    )


def match_table_caption_relations_docx(
    elements: list[Element],
) -> list[Relation]:
    """契约 §3 docx 规则（纯函数）：table@列表位置 i 仅当 i−1 元素是
    表题注 caption（表题注惯例在表上方；elements 列表顺序=body 顺序，
    table_index 与 paragraph_index 不同族不可数值比较）。"""
    rels: list[Relation] = []
    for i, e in enumerate(elements):
        if e.type != "table" or i == 0:
            continue
        prev = elements[i - 1]
        if _is_table_caption(prev):
            rels.append(
                Relation(
                    type="table_has_caption",
                    from_id=e.element_id,
                    to_id=prev.element_id,
                    metadata={"rule": "docx_adjacent_element_above"},
                )
            )
    return _sort_relations(rels)


def match_table_caption_relations_pdf(
    elements: list[Element],
) -> list[Relation]:
    """契约 §3 pdf 规则（纯函数）：同页 + 表题注在表上方
    0 < gap ≤ CAPTION_MAX_GAP_PT + x 区间相交 > 0；候选按
    (gap, table_id, caption_id) 升序贪心唯一配对（批次 4 下方规则的
    镜像；gap = table.top − caption.bottom）。"""
    tables = [e for e in elements if e.type == "table"]
    captions = [e for e in elements if _is_table_caption(e)]
    candidates = []
    for tbl in tables:
        tloc = tbl.source_locator
        tb = tloc.get("bbox")
        if not tb:
            continue
        for cap in captions:
            cloc = cap.source_locator
            cb = cloc.get("bbox")
            if not cb:
                continue
            if tloc.get("page") != cloc.get("page"):
                continue
            gap = tb[1] - cb[3]
            if not (0 < gap <= CAPTION_MAX_GAP_PT):
                continue
            if min(tb[2], cb[2]) - max(tb[0], cb[0]) <= 0:
                continue
            candidates.append((gap, tbl.element_id, cap.element_id))
    candidates.sort()
    used_tbl: set[str] = set()
    used_cap: set[str] = set()
    rels: list[Relation] = []
    for gap, tbl_id, cap_id in candidates:
        if tbl_id in used_tbl or cap_id in used_cap:
            continue
        used_tbl.add(tbl_id)
        used_cap.add(cap_id)
        rels.append(
            Relation(
                type="table_has_caption",
                from_id=tbl_id,
                to_id=cap_id,
                metadata={"rule": "pdf_geometry_above", "gap_pt": gap},
            )
        )
    return _sort_relations(rels)


def _rows_to_markdown(rows: list[list[Any]]) -> str:
    """把表格行渲染为 canonical markdown（批次 5 契约，共享实现）。"""
    return linearize_table(rows)


def _image_filename(document_id: str, prefix: str, index: int, ext: str = "png") -> str:
    """统一图片资源命名：image_<doc_id_short>_<prefix>_<idx>.<ext>"""
    safe_doc = document_id.replace("doc-", "")
    return f"image_{safe_doc}_{prefix}_{index:02d}.{ext}"


def _save_image(
    bytes_data: bytes,
    out_dir: Path,
    document_id: str,
    prefix: str,
    index: int,
    ext: str = "png",
) -> Path:
    """把图片字节保存到 out_dir，返回路径。目录不存在时自动创建。"""
    out_dir.mkdir(parents=True, exist_ok=True)
    name = _image_filename(document_id, prefix, index, ext)
    p = out_dir / name
    p.write_bytes(bytes_data)
    return p


# ---------- PDF ----------

def _group_words_to_paragraphs(words: list[dict]) -> list[dict[str, Any]]:
    """把 pdfplumber 的 words 列表按行/段落聚合。

    pdfplumber 的 word dict 用 `top`/`bottom` 表示 y 坐标（不是 y0/y1），
    `x0`/`x1` 表示 x 坐标。
    算法：
    1. 按 y 中点排序，y 差 < 行高一半视为同一行
    2. 行按 y 升序
    3. 行间距 > 1.5 * 中位行高 视为段落分隔
    """
    if not words:
        return []

    def top_of(w: dict) -> float:
        return float(w.get("top", 0.0))

    def bottom_of(w: dict) -> float:
        return float(w.get("bottom", 0.0))

    def y_center(w: dict) -> float:
        return (top_of(w) + bottom_of(w)) / 2

    sorted_words = sorted(words, key=lambda w: (y_center(w), w["x0"]))

    # 1. 行聚类
    lines: list[list[dict]] = []
    current_line: list[dict] = []
    current_y: float | None = None
    for w in sorted_words:
        yc = y_center(w)
        if current_y is None or abs(yc - current_y) <= 3.0:
            current_line.append(w)
            current_y = yc if current_y is None else (current_y + yc) / 2
        else:
            lines.append(current_line)
            current_line = [w]
            current_y = yc
    if current_line:
        lines.append(current_line)

    heights = [bottom_of(w) - top_of(w) for line in lines for w in line]
    median_h = sorted(heights)[len(heights) // 2] if heights else 12.0

    # 2. 段落聚类
    paragraphs: list[dict[str, Any]] = []
    para_lines: list[list[dict]] = []
    last_bottom: float | None = None
    for line in lines:
        line_top = min(top_of(w) for w in line)
        line_bottom = max(bottom_of(w) for w in line)
        if last_bottom is not None and (line_top - last_bottom) > 1.5 * median_h:
            paragraphs.append(_lines_to_para(para_lines))
            para_lines = []
        para_lines.append(line)
        last_bottom = line_bottom
    if para_lines:
        paragraphs.append(_lines_to_para(para_lines))
    return paragraphs


def _lines_to_para(lines: list[list[dict]]) -> dict[str, Any]:
    """把若干行 word 数组融合成 {text, bbox}。bbox 格式 [x0, top, x1, bottom]。"""
    all_words = [w for line in lines for w in line]
    if not all_words:
        return {"text": "", "bbox": None}
    text = " ".join(
        " ".join(w["text"] for w in sorted(line, key=lambda x: x["x0"]))
        for line in lines
    )
    bbox = [
        min(float(w["x0"]) for w in all_words),
        min(float(w.get("top", 0.0)) for w in all_words),
        max(float(w["x1"]) for w in all_words),
        max(float(w.get("bottom", 0.0)) for w in all_words),
    ]
    return {"text": text, "bbox": bbox}


def _classify_pdf_paragraph(text: str) -> tuple[str, dict[str, Any]]:
    """PDF 段落启发式分类：caption > heading > paragraph。"""
    t = text.strip()
    if not t:
        return "paragraph", {}
    if _is_caption(t):
        return "caption", {"heuristic": "caption_regex"}
    if len(t) <= 80 and not t.endswith(("。", ".", "!", "?", "！", "？")):
        return "heading", {"level": 0, "heuristic": "short_line"}
    return "paragraph", {}


def _render_pdf_image_region(
    pdf_path: Path, page_idx_0based: int, bbox: list[float], out_path: Path, dpi: int = 144
) -> bool:
    """向后兼容包装。"""
    return _render_pdf_image_region_verbose(pdf_path, page_idx_0based, bbox, out_path, dpi) is None


def _render_pdf_image_region_verbose(
    pdf_path: Path, page_idx_0based: int, bbox: list[float], out_path: Path, dpi: int = 144
) -> str | None:
    """用 pypdfium2 渲染整页，按 pdfplumber bbox（点，top 坐标系）裁剪。

    返回 None 表示写出成功；返回字符串表示错误描述。
    pdfplumber bbox 单位是 PDF 点（1/72 英寸），坐标系原点在左上（top-from-top）。
    pypdfium2 渲染后得到 PIL Image，左上原点一致，按 scale 换算即可。
    """
    if pypdfium2 is None:
        return f"pypdfium2 未安装：{_PDFIUM_IMPORT_ERROR}"
    try:
        doc = pypdfium2.PdfDocument(str(pdf_path))
    except Exception as e:
        return f"PdfDocument 打开失败: {type(e).__name__}: {e}"
    try:
        try:
            page = doc[page_idx_0based]
        except Exception as e:
            return f"取 page[{page_idx_0based}] 失败: {type(e).__name__}: {e}"
        try:
            scale = dpi / 72.0
            bitmap = page.render(scale=scale)
            pil = bitmap.to_pil()
        except Exception as e:
            return f"render/to_pil 失败: {type(e).__name__}: {e}"
        x0, top, x1, bottom = bbox
        crop = (
            max(0, int(x0 * scale)), max(0, int(top * scale)),
            min(pil.width, int(x1 * scale)), min(pil.height, int(bottom * scale)),
        )
        if crop[0] >= crop[2] or crop[1] >= crop[3]:
            return f"crop 退化 (0 size): bbox={bbox}, crop={crop}, page_size={pil.width}x{pil.height}"
        try:
            pil.crop(crop).save(out_path, format="PNG")
        except Exception as e:
            return f"PIL save 失败: {type(e).__name__}: {e}"
        return None
    finally:
        try:
            doc.close()
        except Exception:
            pass


_VECTOR_FIGURE_GAP_PT = 15.0
_VECTOR_FIGURE_MIN_W = 100.0
_VECTOR_FIGURE_MIN_H = 100.0


def _vector_figure_clusters(page) -> list[list[float]]:
    """识别页面上的矢量图形簇（批次 15，Option A 裁决）。

    pdfplumber page.images 只含栅格 XObject，矢量图形（Form XObject 或
    直接 path 绘制）不可见。此处把 rects/lines/curves 按 bbox 空间聚类
    （gap 15pt），簇含 ≥1 条 curve 且 bbox ≥100×100pt 计为 1 个矢量图：
    - 表格/公式框 = 纯 rects+lines 零 curves（全 devset 实证）
    - 公式根号 = 1 curve + 线但簇高 ≤51pt，被尺寸过滤排除

    已知边界（ADOPTION.md 批次 15）：
    - 同页重叠双图（如 003 p5/p7 图标+水印）计 1
    - 空间接近的多图（如 002 封面 logo+swoosh）计 1
    - 多段折线路径 pdfplumber 归类为 curve，其他语料的多段线表格可能误报
    """
    boxes: list[list[float]] = []
    for group, is_curve in ((page.curves, True), (page.rects, False), (page.lines, False)):
        try:
            objs = group or []
        except Exception:
            objs = []
        for o in objs:
            try:
                x0, x1 = float(o["x0"]), float(o["x1"])
                top, bottom = float(o["top"]), float(o["bottom"])
            except (KeyError, TypeError, ValueError):
                continue
            if x1 <= x0 or bottom <= top:
                continue
            boxes.append([x0, x1, top, bottom, 1.0 if is_curve else 0.0])
    gap = _VECTOR_FIGURE_GAP_PT
    changed = True
    while changed:
        changed = False
        merged: list[list[float]] = []
        for b in boxes:
            hit = None
            for t in merged:
                if not (
                    b[0] > t[1] + gap
                    or b[1] < t[0] - gap
                    or b[2] > t[3] + gap
                    or b[3] < t[2] - gap
                ):
                    hit = t
                    break
            if hit is not None:
                hit[0] = min(hit[0], b[0])
                hit[1] = max(hit[1], b[1])
                hit[2] = min(hit[2], b[2])
                hit[3] = max(hit[3], b[3])
                hit[4] += b[4]
                changed = True
            else:
                merged.append(b)
        boxes = merged
    return [
        [b[0], b[2], b[1], b[3]]
        for b in boxes
        if b[4] >= 1.0
        and (b[1] - b[0]) >= _VECTOR_FIGURE_MIN_W
        and (b[3] - b[2]) >= _VECTOR_FIGURE_MIN_H
    ]


def _parse_pdf(
    path: Path,
    source_hash: str,
    document_id: str,
    image_output_dir: Path | None,
) -> tuple[list[Element], list[WarningRecord]]:
    if pdfplumber is None:
        raise ParserError(
            code="pdfplumber_unavailable",
            message=f"pdfplumber 未安装：{_PDFPLUMBER_IMPORT_ERROR}",
        )
    elements: list[Element] = []
    warnings: list[WarningRecord] = []
    image_counter = 0
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                # 文本段落
                try:
                    words = page.extract_words(keep_blank_chars=False, use_text_flow=False)
                except Exception as e:
                    warnings.append(
                        WarningRecord(
                            code="pdfplumber_word_extract_failed",
                            reason=f"page {page_idx} extract_words 失败: {e}",
                            details={"page": page_idx},
                        )
                    )
                    words = []
                for para in _group_words_to_paragraphs(words):
                    text = para["text"].strip()
                    if not text:
                        continue
                    etype, meta = _classify_pdf_paragraph(text)
                    locator: dict[str, Any] = {"family": "page_geometry", "page": page_idx}
                    if para["bbox"] is not None:
                        locator["bbox"] = para["bbox"]
                    elements.append(
                        Element(
                            element_id=f"{document_id}::e{len(elements):04d}",
                            type=etype,
                            content=text,
                            parent_id=None,
                            source_locator=locator,
                            confidence=0.85,
                            metadata=meta,
                        )
                    )
                # 表格
                try:
                    tbls = page.find_tables() or []
                except Exception:
                    tbls = []
                for tbl in tbls:
                    try:
                        rows = tbl.extract() or []
                    except Exception:
                        rows = []
                    if not rows:
                        continue
                    md = _rows_to_markdown(rows)
                    tbl_bbox = getattr(tbl, "bbox", None)
                    tbl_locator: dict[str, Any] = {"family": "page_geometry", "page": page_idx}
                    if tbl_bbox:
                        tbl_locator["bbox"] = list(tbl_bbox)
                    elements.append(
                        Element(
                            element_id=f"{document_id}::e{len(elements):04d}",
                            type="table",
                            content=md,
                            parent_id=None,
                            source_locator=tbl_locator,
                            confidence=0.7,
                            metadata={
                                "row_count": len(rows),
                                "col_count": max((len(r) for r in rows), default=0),
                                "source": "pdfplumber",
                            },
                        )
                    )
                # 图片
                try:
                    pdf_images = page.images or []
                except Exception:
                    pdf_images = []
                for img in pdf_images:
                    x0 = float(img.get("x0", 0))
                    top = float(img.get("top", 0))
                    x1 = float(img.get("x1", 0))
                    bottom = float(img.get("bottom", 0))
                    if x1 <= x0 or bottom <= top:
                        continue
                    img_element_id = f"{document_id}::e{len(elements):04d}"
                    img_locator: dict[str, Any] = {
                        "family": "page_geometry",
                        "page": page_idx,
                        "bbox": [x0, top, x1, bottom],
                    }
                    resource_path: str | None = None
                    if image_output_dir is not None:
                        out_path = image_output_dir / _image_filename(
                            document_id, f"p{page_idx}", image_counter, "png"
                        )
                        try:
                            out_path.parent.mkdir(parents=True, exist_ok=True)
                        except OSError as e:
                            warnings.append(
                                WarningRecord(
                                    code="pdf_image_dir_failed",
                                    reason=f"无法创建图片输出目录: {e}",
                                )
                            )
                            out_path = None  # type: ignore[assignment]
                        if out_path is not None:
                            err = _render_pdf_image_region_verbose(
                                path, page_idx - 1, [x0, top, x1, bottom], out_path
                            )
                            if err is None:
                                resource_path = str(out_path)
                                image_counter += 1
                            else:
                                warnings.append(
                                    WarningRecord(
                                        code="pdf_image_render_failed",
                                        reason=f"page {page_idx} 图片渲染失败：{err}",
                                        details={
                                            "page": page_idx,
                                            "bbox": [x0, top, x1, bottom],
                                            "exception": err,
                                        },
                                    )
                                )
                    elements.append(
                        Element(
                            element_id=img_element_id,
                            type="image",
                            content=None,
                            resource_path=resource_path or "(unrendered)",
                            parent_id=None,
                            source_locator=img_locator,
                            confidence=0.6,
                            metadata={
                                "tag": img.get("tag"),
                                "srcsize": list(img.get("srcsize", [])),
                                "extracted_to_disk": resource_path is not None,
                            },
                        )
                    )
                # 矢量图形（批次 15）：曲线聚类检测，镜像栅格发射路径
                try:
                    vector_bboxes = _vector_figure_clusters(page)
                except Exception as e:
                    vector_bboxes = []
                    warnings.append(
                        WarningRecord(
                            code="pdf_vector_cluster_failed",
                            reason=f"page {page_idx} 矢量图聚类失败: {e}",
                            details={"page": page_idx},
                        )
                    )
                for vb in vector_bboxes:
                    vec_element_id = f"{document_id}::e{len(elements):04d}"
                    vec_locator: dict[str, Any] = {
                        "family": "page_geometry",
                        "page": page_idx,
                        "bbox": vb,
                    }
                    vec_resource: str | None = None
                    if image_output_dir is not None:
                        vec_path = image_output_dir / _image_filename(
                            document_id, f"p{page_idx}v", image_counter, "png"
                        )
                        try:
                            vec_path.parent.mkdir(parents=True, exist_ok=True)
                        except OSError as e:
                            warnings.append(
                                WarningRecord(
                                    code="pdf_image_dir_failed",
                                    reason=f"无法创建图片输出目录: {e}",
                                )
                            )
                            vec_path = None  # type: ignore[assignment]
                        if vec_path is not None:
                            err = _render_pdf_image_region_verbose(
                                path, page_idx - 1, vb, vec_path
                            )
                            if err is None:
                                vec_resource = str(vec_path)
                                image_counter += 1
                            else:
                                warnings.append(
                                    WarningRecord(
                                        code="pdf_image_render_failed",
                                        reason=f"page {page_idx} 矢量图渲染失败：{err}",
                                        details={"page": page_idx, "bbox": vb},
                                    )
                                )
                    elements.append(
                        Element(
                            element_id=vec_element_id,
                            type="image",
                            content=None,
                            resource_path=vec_resource or "(unrendered)",
                            parent_id=None,
                            source_locator=vec_locator,
                            confidence=0.5,
                            metadata={
                                "kind": "vector_cluster",
                                "extracted_to_disk": vec_resource is not None,
                            },
                        )
                    )
    except Exception as e:
        raise ParserError(
            code="pdfplumber_open_failed",
            message=f"pdfplumber 打开/解析 PDF 失败: {e}",
            details={"exception_type": type(e).__name__},
        ) from e
    if not elements:
        warnings.append(
            WarningRecord(
                code="pdf_no_text_extracted",
                reason="pdfplumber 未提取到任何文本/表格/图片（可能为扫描件，本阶段不支持 OCR）",
            )
        )
    return elements, warnings


# ---------- DOCX ----------

def _is_heading_style(style_name: str | None) -> tuple[bool, int]:
    """从 python-docx style 名判定是否是 heading 及级别。"""
    if not style_name:
        return False, 0
    s = style_name.strip().lower()
    if s == "title":
        return True, 1
    if s.startswith("heading"):
        try:
            level = int(s.replace("heading", "").strip())
            return True, max(1, level)
        except ValueError:
            return True, 1
    return False, 0


def _extract_inline_image_rids(paragraph_xml) -> list[str]:
    """从 paragraph 的 XML element 中找出所有内嵌图片的 rId。"""
    rids: list[str] = []
    if qn is None:
        return rids
    for drawing in paragraph_xml.iter(qn("w:drawing")):
        for blip in drawing.iter(qn("a:blip")):
            rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if rid:
                rids.append(rid)
    return rids


def _iter_flow_elements(el):
    """递归提取 body/sdtContent 内的段落和表格（保持文档顺序）。

    批次 14：w:body 顶层的 w:sdt（结构化文档标签容器）此前被主循环整块跳过，
    导致封面等内容（标题/图片/marker）静默丢失；此处深度优先下钻 w:sdtContent。
    已知边界：表格单元格（w:tc）内的 w:sdt 未处理（表格内容走 _rows_to_markdown 单独管线）。
    """
    for child in el.iterchildren():
        tag = child.tag
        if tag == qn("w:p") or tag == qn("w:tbl"):
            yield child
        elif tag == qn("w:sdt"):
            content = child.find(qn("w:sdtContent"))
            if content is not None:
                yield from _iter_flow_elements(content)  # 递归
        # 其余（w:sectPr 等）跳过


def _parse_docx(
    path: Path,
    source_hash: str,
    document_id: str,
    image_output_dir: Path | None,
) -> tuple[list[Element], list[WarningRecord]]:
    if docx is None:
        raise ParserError(
            code="python_docx_unavailable",
            message=f"python-docx 未安装：{_DOCX_IMPORT_ERROR}",
        )
    elements: list[Element] = []
    warnings: list[WarningRecord] = []
    image_counter = 0

    try:
        d = docx.Document(str(path))
    except Exception as e:
        raise ParserError(
            code="docx_open_failed",
            message=f"python-docx 打开 DOCX 失败: {e}",
            details={"exception_type": type(e).__name__},
        ) from e

    body = d.element.body
    para_counter = 0
    table_counter = 0
    section_idx = 0
    for child in _iter_flow_elements(body):
        tag = child.tag
        if tag == qn("w:p"):  # paragraph
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, d)
            text = (para.text or "").strip()
            style_name = para.style.name if para.style else None
            etype = "paragraph"
            level = 0
            if _is_caption(text):
                etype = "caption"
            elif style_name:
                is_h, lvl = _is_heading_style(style_name)
                if is_h:
                    etype = "heading"
                    level = lvl
            locator: dict[str, Any] = {
                "family": "structural_index",
                "paragraph_index": para_counter,
                "section": section_idx,
            }
            elements.append(
                Element(
                    element_id=f"{document_id}::e{len(elements):04d}",
                    type=etype,
                    content=text or "(空段落)",
                    parent_id=None,
                    source_locator=locator,
                    confidence=0.95,
                    metadata={
                        "level": level,
                        "style": style_name,
                        "empty": not text,
                    },
                )
            )
            # 处理段落内的内嵌图片
            for rid in _extract_inline_image_rids(child):
                rel = d.part.rels.get(rid)
                if rel is None or "image" not in rel.reltype:
                    continue
                target = rel.target_part
                ext = (target.partname.ext or "png").lower()
                resource_path: str | None = None
                if image_output_dir is not None:
                    try:
                        out_path = image_output_dir / _image_filename(
                            document_id, f"para{para_counter}", image_counter, ext
                        )
                        out_path.parent.mkdir(parents=True, exist_ok=True)
                        out_path.write_bytes(target.blob)
                        resource_path = str(out_path)
                        image_counter += 1
                    except OSError as e:
                        warnings.append(
                            WarningRecord(
                                code="docx_image_save_failed",
                                reason=f"图片保存失败 rId={rid}: {e}",
                                details={"rid": rid, "paragraph_index": para_counter},
                            )
                        )
                elements.append(
                    Element(
                        element_id=f"{document_id}::e{len(elements):04d}",
                        type="image",
                        content=None,
                        resource_path=resource_path or "(unsaved)",
                        parent_id=None,
                        source_locator={
                            "family": "structural_index",
                            "paragraph_index": para_counter,
                            "section": section_idx,
                            "relationship_id": rid,
                            "target_partname": str(target.partname),
                        },
                        confidence=0.95,
                        metadata={
                            "byte_size": len(target.blob),
                            "ext": ext,
                            "extracted_to_disk": resource_path is not None,
                        },
                    )
                )
            para_counter += 1
        elif tag == qn("w:tbl"):  # table
            from docx.table import Table
            tbl = Table(child, d)
            rows_data: list[list[str]] = []
            for row in tbl.rows:
                rows_data.append([(c.text or "").strip() for c in row.cells])
            md = _rows_to_markdown(rows_data)
            # 批次 5 契约 §2：0 行表不产出 element（静默跳过，对齐 pdf/html）
            if md:
                elements.append(
                    Element(
                        element_id=f"{document_id}::e{len(elements):04d}",
                        type="table",
                        content=md,
                        parent_id=None,
                        source_locator={
                            "family": "structural_index",
                            "table_index": table_counter,
                            "section": section_idx,
                        },
                        confidence=0.95,
                        metadata={
                            "row_count": len(rows_data),
                            "col_count": max((len(r) for r in rows_data), default=0),
                            "source": "python-docx",
                        },
                    )
                )
            table_counter += 1
    if not elements:
        warnings.append(
            WarningRecord(
                code="docx_no_content",
                reason="python-docx 未提取到段落或表格",
            )
        )
    return elements, warnings


# ---------- 顶层 Parser ----------

class FallbackParser(Parser):
    """PDF → pdfplumber + pypdfium2；DOCX → python-docx。

    优点：给出 element 级结构和精确 source_locator（PDF page/bbox，family=page_geometry；
    DOCX paragraph_index/table_index，family=structural_index；契约 docs/locator-kvfs-contract.md）。
    限制：不做 OCR、不识别图片内容文本、不重建跨页段落。
    """

    name = "fallback"
    version = f"pdfplumber={_PDFPLUMBER_VERSION},python-docx={_DOCX_VERSION},pypdfium2={_PDFIUM_VERSION}"
    supported_extensions = (".pdf", ".docx")
    priority = 10
    source_types = ("pdf", "docx")
    locator_family = None  # 两类型 family 不同（page_geometry/structural_index），按类型走内置绑定

    def __init__(self, image_output_dir: Path | str | None = None) -> None:
        """Args:
            image_output_dir: 若提供，把图片资源真实保存到此目录；element.resource_path 指向实存文件。
                              若为 None，图片 element 仅记录位置（resource_path 为占位符）。
        """
        self._image_output_dir = Path(image_output_dir) if image_output_dir else None

    def parse(self, path: str | Path, source_hash: str) -> Document:
        p = Path(path)
        if not p.is_file():
            raise ParserError(
                code="file_not_found",
                message=f"输入文件不存在: {p}",
                details={"path": str(p)},
            )
        source_type = detect_source_type(p)
        document_id = make_document_id(source_hash)
        if source_type == "pdf":
            elements, warnings = _parse_pdf(p, source_hash, document_id, self._image_output_dir)
            # 契约 §4：两类 relation 合并后按 (type, from_id, to_id) 排序
            relations = _sort_relations(
                match_caption_relations_pdf(elements)
                + match_table_caption_relations_pdf(elements)
            )
        else:
            elements, warnings = _parse_docx(p, source_hash, document_id, self._image_output_dir)
            relations = _sort_relations(
                match_caption_relations_docx(elements)
                + match_table_caption_relations_docx(elements)
            )
        return Document(
            document_id=document_id,
            source_path=str(p),
            source_type=source_type,
            source_hash=source_hash,
            parser_name=self.name,
            parser_version=self.version,
            elements=elements,
            chunks=[],
            relations=relations,
            warnings=warnings,
            errors=[],
            metadata={"fallback": True, "image_output_dir": str(self._image_output_dir) if self._image_output_dir else None},
        )


__all__ = ["FallbackParser"]
