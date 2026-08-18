"""evaluation/cli.py 第六百四十九轮 edges 测试（Round 1215）。

补强 cli edges152 未触及的角度（第五百八十七批，probe 实证）。

新角度（多节 DOCX 板 CLI 全链）：
- **docx run**——rc 0、成功 1/1、
  by_type {heading: 3, paragraph: 4}
- **stdout 汇总**——"pdf=0 docx=1"
  （docx 侧计数首锁）
- **inspect counts**——"elements=7
  chunks=3"
- **docx locator 行**——"docx_
  locator_valid_ratio             1.0000
  (ok)"（docx 文档的 inspect 指标行
  与 pdf 板互补）
- **pdf locator 行**——"pdf_locator_
  valid_ratio              null
  (not_pdf_document)"（非 pdf 文档
  的对照 null 首锁）
- **hbc 行**——"1.0000  (ok)"
- forbidden tokens 第五百七十五批（open 1）
"""

from __future__ import annotations

import inspect
import json

import evaluation.cli as cli_mod
from evaluation.cli import main


def _write_docx(path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Main Title", 0)
    doc.add_heading("1.1 First Sub Heading", level=2)
    doc.add_paragraph("Alpha paragraph under first sub.")
    doc.add_heading("1.2 Second Sub Heading", level=2)
    doc.add_paragraph("Beta paragraph under second sub.")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Gamma paragraph in section two.")
    doc.save(str(path))


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    _write_docx(tmp_path / "samples" / "sec.docx")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "sec", "path": "samples/sec.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _doc(tmp_path):
    from app.pipeline import process_single
    _board(tmp_path)
    doc, errors = process_single(
        tmp_path / "samples" / "sec.docx", tmp_path / "doc.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    return tmp_path / "doc.json"


def _run_cli(capsys, argv):
    rc = main(argv)
    out = capsys.readouterr().out
    return rc, out


# ---------- docx run ----------

def test_cli_run_docx_section_batch413(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    m = rep["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"heading": 3, "paragraph": 4}, "reason": None}
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {"success_count": 1, "total": 1,
                                "rate": 1.0}


def test_cli_run_docx_stdout_batch413(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    assert "[OK]" in out
    assert "documents=1（成功 1，失败 0）" in out
    assert "pdf=0 docx=1" in out


def test_cli_validate_report_batch413(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    rc2, out2 = _run_cli(capsys, [
        "validate-report", str(tmp_path / "r.json")])
    assert rc2 == 0
    assert "[OK]" in out2
    assert "通过" in out2


# ---------- inspect-doc ----------

def test_cli_inspect_counts_batch413(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert "counts:      elements=7 chunks=3" in out


def test_cli_inspect_by_type_batch413(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("element_count_by_type                "
            "heading=3, paragraph=4  (ok)") in out


def test_cli_inspect_docx_locator_batch413(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("docx_locator_valid_ratio             "
            "1.0000  (ok)") in out


def test_cli_inspect_pdf_locator_null_batch413(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("pdf_locator_valid_ratio              "
            "null  (not_pdf_document)") in out


def test_cli_inspect_hbc_batch413(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("heading_boundary_compliance          "
            "1.0000  (ok)") in out


def test_cli_inspect_type_docx_batch413(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert "type=docx" in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_identifier_counts_batch413():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


# ---------- forbidden tokens 第五百七十五批 ----------

def test_source_no_eval_batch413():
    assert "eval(" not in _src()


def test_source_no_exec_batch413():
    assert "exec(" not in _src()


def test_source_no_compile_batch413():
    assert "compile(" not in _src()


def test_source_no_globals_batch413():
    assert "globals(" not in _src()


def test_source_no_locals_batch413():
    assert "locals(" not in _src()


def test_source_no_os_system_batch413():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch413():
    assert "subprocess" not in _src()


def test_source_no_popen_batch413():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch413():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch413():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch413():
    assert "socket" not in _src()


def test_source_no_requests_batch413():
    assert "requests" not in _src()


def test_source_no_urllib_batch413():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch413():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch413():
    assert "yield" not in _src()


def test_source_no_async_await_batch413():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch413():
    assert _src().count("open(") == 1
