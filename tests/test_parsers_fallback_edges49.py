r"""app/parsers/fallback_parser.py 边角测试 - 第四十九轮（Round 1442）。

新角度（probe 实证）docx 域代码 + 制表符（历史 run 级考察
碰过 w:br/w:ins/hyperlink，域机制全空白）：
- w:fldSimple（简单域，内裹缓存结果 run '7'）：**文本整体
  丢失**——paragraph.text 只读直接 w:r，fldSimple 里的 run
  不可见 → '(空段落)' 占位（empty True）
- 复杂域链（begin → instrText ' NUMPAGES ' → separate →
  结果 run '42' → end）：**缓存结果保留** '42'——instrText
  不是 w:t 自动丢弃，域指令不留痕
- w:tab：字面 '\\t' 进 content（'before\\tafter'，对照
  w:br → '\\n'、分页符 → 无字符）
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _fldsimple_docx(tmp_path):
    d = Document()
    para = d.add_paragraph()
    fld = para._p.makeelement(
        qn("w:fldSimple"),
        {qn("w:instr"): "PAGE"})
    r = fld.makeelement(qn("w:r"), {})
    t = r.makeelement(qn("w:t"), {})
    t.text = "7"
    r.append(t)
    fld.append(r)
    para._p.append(fld)
    p = tmp_path / "fld.docx"
    d.save(str(p))
    return p


def _complex_field_docx(tmp_path):
    d = Document()
    para = d.add_paragraph()
    r1 = para.add_run()
    r1._r.append(r1._r.makeelement(
        qn("w:fldChar"),
        {qn("w:fldCharType"):
         "begin"}))
    r2 = para.add_run()
    it = r2._r.makeelement(
        qn("w:instrText"), {})
    it.text = " NUMPAGES "
    r2._r.append(it)
    r3 = para.add_run()
    r3._r.append(r3._r.makeelement(
        qn("w:fldChar"),
        {qn("w:fldCharType"):
         "separate"}))
    para.add_run("42")
    r5 = para.add_run()
    r5._r.append(r5._r.makeelement(
        qn("w:fldChar"),
        {qn("w:fldCharType"): "end"}))
    p = tmp_path / "cfld.docx"
    d.save(str(p))
    return p


def _tab_docx(tmp_path):
    d = Document()
    para = d.add_paragraph()
    para.add_run("before")
    r = para.add_run()
    r._r.append(r._r.makeelement(
        qn("w:tab"), {}))
    para.add_run("after")
    p = tmp_path / "tab.docx"
    d.save(str(p))
    return p


# ---------- fldSimple ----------

def test_fldsimple_text_lost(
        tmp_path):
    p = _fldsimple_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "(空段落)"]
    assert doc.elements[
        0].metadata["empty"] is True


def test_fldsimple_no_warning(
        tmp_path):
    p = _fldsimple_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_fldsimple_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _fldsimple_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


# ---------- 复杂域 ----------

def test_complex_field_cached(
        tmp_path):
    p = _complex_field_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "42"]
    assert doc.elements[
        0].type == "paragraph"


def test_complex_field_no_instr(
        tmp_path):
    p = _complex_field_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert "NUMPAGES" not in \
        doc.elements[0].content


def test_complex_field_chunk(
        tmp_path):
    p = _complex_field_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[
        0].text == "42"


def test_complex_field_schema(
        tmp_path):
    from app.schema import is_valid
    p = _complex_field_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


# ---------- w:tab ----------

def test_tab_literal(tmp_path):
    p = _tab_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].content == \
        "before\tafter"


def test_tab_chunk(tmp_path):
    p = _tab_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[
        0].text == "before\tafter"


def test_tab_no_warnings(
        tmp_path):
    p = _tab_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []
