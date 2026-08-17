"""evaluation/report.py 第五百二十五轮 edges 测试（Round 1081）。

补强 edges122-123 未触及的角度（第四百五十七批，probe 实证）。

新角度（真实标注板驱动的 boundary 三元组宏观分歧）：
- 四文档真实板（mc 40）：d1 长段+双锚 BBB/CCC →
  P/R/F1 全 1.0；d2 长 chunk 中段 marker → 全 0.0；
  d3 无标注 → 全 null no_annotation；d4 短段并成
  单 chunk 有标注 → **P/F1 null no_predicted_boundaries
  而 R 是 0.0**——recall 只需 gt 不需 predicted，同文档
  三元组两态并存
- **宏观三元组参与度分歧**：P {macro 0.5, 2 参 2 免}、
  R {macro 1/3, **3 参** 1 免}、F1 {macro 0.5, 2 参
  2 免}——d4 只进 R 的分母——真实值混账（1.0+0.0）
  均值 0.5，非全 1.0 均匀板
- 成功账本 {4, 4, 1.0}
- forbidden tokens 第五百五十二批（report 变体：15 项
  去 subprocess + open 0 + subprocess.run == 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    (tmp_path / "anns").mkdir()
    d1 = Document()
    for t in ("AAA first paragraph body one.",
              "BBB second paragraph body two.",
              "CCC third paragraph body three."):
        d1.add_paragraph(t)
    d1.save(str(tmp_path / "samples" / "good.docx"))
    d4 = Document()
    for t in ("AAA first", "BBB second", "CCC third"):
        d4.add_paragraph(t)
    d4.save(str(tmp_path / "samples" / "merged.docx"))
    d2 = Document()
    d2.add_paragraph("AAA " + "word " * 20)
    d2.add_paragraph("BBB tail end.")
    d2.save(str(tmp_path / "samples" / "miss.docx"))
    (tmp_path / "anns" / "a1.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "BBB", "position": "before"},
            {"marker": "CCC", "position": "before"}]}),
        encoding="utf-8")
    (tmp_path / "anns" / "a2.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d2",
        "chunk_boundary_anchors": [
            {"marker": "word", "position": "before"}]}),
        encoding="utf-8")
    (tmp_path / "anns" / "a4.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d4",
        "chunk_boundary_anchors": [
            {"marker": "BBB", "position": "before"}]}),
        encoding="utf-8")
    docs = [
        {"doc_id": "d1", "path": "samples/good.docx",
         "source_type": "docx",
         "annotation_file": "anns/a1.json"},
        {"doc_id": "d2", "path": "samples/miss.docx",
         "source_type": "docx",
         "annotation_file": "anns/a2.json"},
        {"doc_id": "d3", "path": "samples/good.docx",
         "source_type": "docx"},
        {"doc_id": "d4", "path": "samples/merged.docx",
         "source_type": "docx",
         "annotation_file": "anns/a4.json"}]
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": docs,
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(tmp_path / "m.json",
                                        tmp_path),
                          tmp_path / "o.json", max_chars=40)


_BT = ("chunk_boundary_precision",
       "chunk_boundary_recall", "chunk_boundary_f1")


# ---------- 逐文档边界值 ----------

def test_per_doc_boundary_values_batch280(tmp_path):
    rep = _run(tmp_path)
    m = {p["doc_id"]: p["metrics"] for p in rep["per_doc"]}
    for k in _BT:
        assert m["d1"][k] == {"value": 1.0, "reason": None}
        assert m["d2"][k] == {"value": 0.0, "reason": None}
        assert m["d3"][k] == {"value": None,
                              "reason": "no_annotation"}


# ---------- d4：R 0.0 与 P/F1 null 并存 ----------

def test_recall_without_predicted_batch280(tmp_path):
    rep = _run(tmp_path)
    m = {p["doc_id"]: p["metrics"] for p in rep["per_doc"]}
    assert m["d4"]["chunk_boundary_precision"] == {
        "value": None, "reason": "no_predicted_boundaries"}
    assert m["d4"]["chunk_boundary_recall"] == {
        "value": 0.0, "reason": None}
    assert m["d4"]["chunk_boundary_f1"] == {
        "value": None, "reason": "no_predicted_boundaries"}


# ---------- 宏观三元组参与度分歧 ----------

def test_boundary_trio_macro_divergence_batch280(tmp_path):
    ra = _run(tmp_path)["summary"]["ratio_macro_averages"]
    assert ra["chunk_boundary_precision"] == {
        "macro_average": 0.5, "participating_docs": 2,
        "not_evaluated": 2}
    assert ra["chunk_boundary_recall"] == {
        "macro_average": 0.3333333333333333,
        "participating_docs": 3, "not_evaluated": 1}
    assert ra["chunk_boundary_f1"] == {
        "macro_average": 0.5, "participating_docs": 2,
        "not_evaluated": 2}


# ---------- 成功账本 ----------

def test_success_all_four_batch280(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["success_rates"] == {
        "pipeline_success": {"success_count": 4,
                             "total": 4, "rate": 1.0}}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch280():
    src = _src()
    assert "macro = sum(values) / len(values)" in src
    assert ("not_eval = len(per_doc_results) "
            "- len(values)") in src


# ---------- forbidden tokens 第五百五十二批（report 变体） ----------

def test_source_no_eval_batch280():
    assert "eval(" not in _src()


def test_source_no_exec_batch280():
    assert "exec(" not in _src()


def test_source_no_compile_batch280():
    assert "compile(" not in _src()


def test_source_no_globals_batch280():
    assert "globals(" not in _src()


def test_source_no_locals_batch280():
    assert "locals(" not in _src()


def test_source_no_os_system_batch280():
    assert "os.system" not in _src()


def test_source_no_popen_batch280():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch280():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch280():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch280():
    assert "socket" not in _src()


def test_source_no_requests_batch280():
    assert "requests" not in _src()


def test_source_no_urllib_batch280():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch280():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch280():
    assert "yield" not in _src()


def test_source_no_async_await_batch280():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch280():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch280():
    assert _src().count("subprocess.run") == 2
