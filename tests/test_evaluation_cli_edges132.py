"""evaluation/cli.py 第五百零六轮 edges 测试（Round 1062）。

补强 edges129-131 未触及的角度（第四百三十八批，probe 实证）。

新角度（inspect-doc 对真实嵌图文档 JSON 的可读渲染）：
- inspect-doc 固定 image_base_dir=None，但 pipeline 落盘的
  resource_path 是**绝对路径**且真实存在 → 图片 ratio 照样
  1.0000 (ok)——inspect-doc 与 runner 对同一文档**口径一致**
  （绝对路径使 base dir 形同虚设）
- `_tolerance_chars` **伪指标泄漏**：chunk_boundary_prf 的
  私有记录在 runner 里被 pop、在 inspect-doc 里**不 pop**
  直接按 int 组渲染 "30 (ok)"——两条路径的真实不对称
- dict 分支真实首用：element_count_by_type 渲染
  "image=1, paragraph=3"（sorted + 逗号 join）
- null 组排尾 + reason=None 字面渲染 "(None)"（error_code）；
  parser 行 "fallback v" 直连逗号串版本号（无空格怪相）
- forbidden tokens 第五百三十三批（open 1）
"""

from __future__ import annotations

import inspect
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.cli as cli_mod
from app.pipeline import process_single
from evaluation.cli import main


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _inspect(tmp_path, capsys):
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("BBB third paragraph body.")
    src = tmp_path / "img.docx"
    d.save(str(src))
    doc, errors = process_single(
        src, tmp_path / "doc.json", parser_name="fallback",
        max_chars=200, write_json=True)
    assert errors == []
    rc = main(["inspect-doc", str(tmp_path / "doc.json")])
    out = capsys.readouterr().out
    return rc, out


# ---------- 基本面：rc 与元信息 ----------

def test_inspect_doc_image_doc_rc0_batch261(tmp_path, capsys):
    rc, out = _inspect(tmp_path, capsys)
    assert rc == 0
    assert "counts:      elements=4 chunks=1" in out
    assert "type=docx" in out


# ---------- 图片 ratio 口径一致 ----------

def test_inspect_doc_image_ratio_lit_batch261(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert ("  image_resource_exists_ratio"
            "          1.0000  (ok)") in out


# ---------- _tolerance_chars 伪指标泄漏 ----------

def test_inspect_doc_tolerance_leak_batch261(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert "  _tolerance_chars                     30  (ok)" in out


# ---------- dict 分支真实渲染 ----------

def test_inspect_doc_dict_metric_render_batch261(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert ("  element_count_by_type"
            "                image=1, paragraph=3  (ok)") in out


# ---------- null 组排尾 + "(None)" 字面 ----------

def test_inspect_doc_null_group_last_batch261(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    lines = out.splitlines()
    assert lines[-1] == ("  silent_drop_count"
                         "                    null"
                         "  (no_expectations)")
    assert "  error_code                           null  (None)" in out
    assert out.count("(no_annotation)") == 3
    assert (out.count(
        "(parser_does_not_emit_relations)") == 3)


# ---------- parser 行 v 直连怪相 ----------

def test_inspect_doc_parser_version_concat_batch261(
        tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert "parser:      fallback v" in out
    assert "v " not in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch261():
    src = _src()
    assert "image_base_dir=None" in src
    assert "_format_metric" in src


# ---------- forbidden tokens 第五百三十三批 ----------

def test_source_no_eval_batch261():
    assert "eval(" not in _src()


def test_source_no_exec_batch261():
    assert "exec(" not in _src()


def test_source_no_compile_batch261():
    assert "compile(" not in _src()


def test_source_no_globals_batch261():
    assert "globals(" not in _src()


def test_source_no_locals_batch261():
    assert "locals(" not in _src()


def test_source_no_os_system_batch261():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch261():
    assert "subprocess" not in _src()


def test_source_no_popen_batch261():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch261():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch261():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch261():
    assert "socket" not in _src()


def test_source_no_requests_batch261():
    assert "requests" not in _src()


def test_source_no_urllib_batch261():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch261():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch261():
    assert "yield" not in _src()


def test_source_no_async_await_batch261():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch261():
    assert _src().count("open(") == 1
