r"""集成金样板测试（Round 1400）。

一个富真 docx（H1/H2 + 长段 + 题注 + 空段 + inline 图片 +
2x2 表 + 尾段）穿全栈——parse → chunk(mc=120) → 落盘 JSON →
schema → metrics（带精确 expectations）→ 图片落盘目录——
单文件锁定整条纵向的行为契约（各环节单锁过，整链组合首次）：
- 10 元素类型序列（图片段产出第二个 '(空段落)' 占位）
- 6 chunk 精确文本（空占位两枚合并进同一 chunk）
- output 路径推导 image_output_dir → images-<sha16>/ 落盘
  真 PNG → irer 1.0
- sdc 0 / hbc 1.0 / docxloc 1.0 / crir 1.0
"""

from __future__ import annotations

import io
import json
import re
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.pipeline import process_single


P1 = ("Golden body paragraph one "
      "with more than enough "
      "characters to overflow a "
      "small budget alone here.")
P2 = ("Golden body paragraph two "
      "sits under the second "
      "heading and is long enough "
      "to split similarly.")

_EXPECT_ECT = {"heading": 2,
               "paragraph": 5,
               "caption": 1,
               "image": 1,
               "table": 1}


def _make_png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    return (sig
            + chunk(b"IHDR",
                    struct.pack(
                        ">IIBBBBB",
                        1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT",
                    zlib.compress(
                        b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b""))


def _build(tmp_path):
    d = Document()
    d.add_heading("Golden Root", 1)
    d.add_paragraph(P1)
    d.add_paragraph(
        "Figure 1: golden caption")
    d.add_paragraph("")
    d.add_picture(io.BytesIO(_make_png()),
                  width=914400)
    d.add_heading("Nested", 2)
    d.add_paragraph(P2)
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "h1"
    t.cell(0, 1).text = "h2"
    t.cell(1, 0).text = "v1"
    t.cell(1, 1).text = "v2"
    d.add_paragraph("tail paragraph")
    p = tmp_path / "gold.docx"
    d.save(str(p))
    return p


def _run(tmp_path):
    p = _build(tmp_path)
    out = tmp_path / "gold.json"
    doc, errors = process_single(
        p, out, parser_name="fallback",
        max_chars=120)
    return p, out, doc, errors


# ---------- 元素 ----------

def test_element_type_sequence(tmp_path):
    _, _, doc, errors = _run(tmp_path)
    assert errors == []
    assert [e.type for e in doc.elements
            ] == [
        "heading", "paragraph",
        "caption", "paragraph",
        "paragraph", "image",
        "heading", "paragraph",
        "table", "paragraph"]


def test_two_empty_placeholders(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    empties = [e for e in doc.elements
               if e.content ==
               "(空段落)"]
    assert len(empties) == 2
    assert all(e.metadata["empty"]
               is True
               for e in empties)


def test_element_ids_sequential(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    assert [e.element_id[-4:]
            for e in doc.elements] == [
        f"{i:04d}"
        for i in range(10)]


# ---------- 分块 ----------

def test_six_chunks(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    assert len(doc.chunks) == 6


def test_chunk_texts(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    assert [c.text for c in doc.chunks
            ] == [
        "Golden Root " + P1,
        "Figure 1: golden caption",
        "(空段落) (空段落)",
        "Nested " + P2,
        "| h1 | h2 |\n"
        "| --- | --- |\n"
        "| v1 | v2 |",
        "tail paragraph"]


def test_chunks_within_max(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    assert all(len(c.text) <= 120
               for c in doc.chunks)


def test_placeholder_chunk_sources(
        tmp_path):
    _, _, doc, _ = _run(tmp_path)
    c = doc.chunks[2]
    assert c.text == \
        "(空段落) (空段落)"
    assert len(c.source_element_ids) == 2


def test_table_chunk_isolated(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    assert len(
        doc.chunks[4]
        .source_element_ids) == 1


# ---------- 落盘 JSON ----------

def test_json_written(tmp_path):
    _, out, _, _ = _run(tmp_path)
    assert out.is_file()


def test_json_shape(tmp_path):
    _, out, _, _ = _run(tmp_path)
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert data["schema_version"] == \
        "0.1.0"
    assert len(data["elements"]) == 10
    assert len(data["chunks"]) == 6
    assert data["errors"] == []
    assert data["source_type"] == \
        "docx"


def test_json_schema_valid(tmp_path):
    from app.schema import is_valid
    _, out, _, _ = _run(tmp_path)
    assert is_valid(
        json.loads(
            out.read_text(
                encoding="utf-8")))


def test_document_id_pattern(tmp_path):
    _, _, doc, _ = _run(tmp_path)
    assert re.fullmatch(
        r"doc-[0-9a-f]{16}",
        doc.document_id)


def test_source_hash_matches(tmp_path):
    p, _, doc, _ = _run(tmp_path)
    assert doc.source_hash == \
        compute_file_hash(p)


# ---------- 图片落盘（output 推导 dir） ----------

def test_images_dir_created(tmp_path):
    _, out, doc, _ = _run(tmp_path)
    sha = doc.source_hash[:16]
    img_dir = out.parent / (
        f"images-{sha}")
    assert img_dir.is_dir()
    pngs = list(img_dir.glob("*.png"))
    assert len(pngs) == 1
    assert pngs[0].read_bytes()[:8] == \
        b"\x89PNG\r\n\x1a\n"


def test_image_resource_matches_dir(
        tmp_path):
    _, out, doc, _ = _run(tmp_path)
    img = [e for e in doc.elements
           if e.type == "image"][0]
    assert Path(
        img.resource_path).is_file()


# ---------- 指标 ----------

def test_metrics_all_green(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    _, _, doc, _ = _run(tmp_path)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        {"element_count_by_type":
         _EXPECT_ECT})
    assert m["silent_drop_count"] == {
        "value": 0, "reason": None}
    assert m[
        "heading_boundary_compliance"] \
        == {"value": 1.0,
            "reason": None}
    assert m[
        "docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}
    assert m[
        "chunk_reference_intact_ratio"] \
        == {"value": 1.0,
            "reason": None}


def test_metrics_irer_one(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    _, _, doc, _ = _run(tmp_path)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "image_resource_exists_ratio"] \
        == {"value": 1.0,
            "reason": None}


def test_metrics_ect(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    _, _, doc, _ = _run(tmp_path)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "element_count_by_type"][
        "value"] == _EXPECT_ECT
