r"""app/parsers/fallback_parser.py 边角测试 - 第九轮（Round 1382）。

补强未触达面（probe 实证，历史 fallback docx 板只用普通段落，
表格与列表样式零覆盖）：
- docx 表格 → table 元素（markdown 渲染）+ table_index locator
- paragraph_index 与 table_index 各自独立计数（p0 t0 p1 t1 p2）
- List Bullet / List Number 样式 → 普通 paragraph（不产
  list_item）
- 合并单元格：文本在合并跨度两端重复（'| merged | merged |'）
- 表格与段落在文档流中保序
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser


def _parse(tmp_path, builder, name="d.docx"):
    p = tmp_path / name
    builder(p)
    return FallbackParser().parse(
        p, compute_file_hash(p))


def _basic(p):
    d = Document()
    d.add_heading("Title", 1)
    d.add_paragraph("plain para")
    d.add_paragraph("bullet item",
                    style="List Bullet")
    d.add_paragraph("numbered item",
                    style="List Number")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "a"
    t.cell(0, 1).text = "b"
    t.cell(1, 0).text = "c"
    t.cell(1, 1).text = "d"
    d.add_paragraph("after table")
    d.save(str(p))


def _interleaved(p):
    d = Document()
    d.add_paragraph("p0")
    t1 = d.add_table(rows=1, cols=2)
    t1.cell(0, 0).text = "x1"
    t1.cell(0, 1).text = "y1"
    d.add_paragraph("p1")
    t2 = d.add_table(rows=1, cols=1)
    t2.cell(0, 0).text = "solo"
    d.add_paragraph("p2")
    d.save(str(p))


def _merged(p):
    d = Document()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).merge(t.cell(0, 1))
    t.cell(0, 0).text = "merged"
    t.cell(1, 0).text = "c"
    t.cell(1, 1).text = "d"
    d.save(str(p))


# ---------- 列表样式 → 普通 paragraph ----------

def test_list_bullet_is_paragraph(tmp_path):
    doc = _parse(tmp_path, _basic)
    styles = [e for e in doc.elements
              if e.content == "bullet item"]
    assert len(styles) == 1
    assert styles[0].type == "paragraph"


def test_list_number_is_paragraph(tmp_path):
    doc = _parse(tmp_path, _basic)
    styles = [e for e in doc.elements
              if e.content == "numbered item"]
    assert styles[0].type == "paragraph"


def test_no_list_item_type_emitted(tmp_path):
    doc = _parse(tmp_path, _basic)
    assert all(e.type != "list_item"
               for e in doc.elements)


def test_basic_six_elements(tmp_path):
    doc = _parse(tmp_path, _basic)
    assert [e.type for e in doc.elements] == [
        "heading", "paragraph", "paragraph",
        "paragraph", "table", "paragraph"]


# ---------- 表格渲染 ----------

def test_table_md_render(tmp_path):
    doc = _parse(tmp_path, _basic)
    table = [e for e in doc.elements
             if e.type == "table"][0]
    assert table.content == (
        "| a | b |\n| --- | --- |\n| c | d |")


def test_table_locator_table_index(tmp_path):
    doc = _parse(tmp_path, _basic)
    table = [e for e in doc.elements
             if e.type == "table"][0]
    assert table.source_locator == {
        "table_index": 0, "section": 0}


# ---------- 双表交错计数 ----------

def test_interleaved_order(tmp_path):
    doc = _parse(tmp_path, _interleaved)
    assert [e.type for e in doc.elements] == [
        "paragraph", "table", "paragraph",
        "table", "paragraph"]


def test_paragraph_index_independent(tmp_path):
    doc = _parse(tmp_path, _interleaved)
    paras = [e for e in doc.elements
             if e.type == "paragraph"]
    assert [e.source_locator[
        "paragraph_index"]
        for e in paras] == [0, 1, 2]


def test_table_index_independent(tmp_path):
    doc = _parse(tmp_path, _interleaved)
    tables = [e for e in doc.elements
              if e.type == "table"]
    assert [e.source_locator["table_index"]
            for e in tables] == [0, 1]


def test_second_table_render(tmp_path):
    doc = _parse(tmp_path, _interleaved)
    tables = [e for e in doc.elements
              if e.type == "table"]
    assert tables[1].content == (
        "| solo |\n| --- |")


# ---------- 合并单元格 ----------

def test_merged_cell_text_duplicated(tmp_path):
    doc = _parse(tmp_path, _merged)
    assert doc.elements[0].content == (
        "| merged | merged |\n"
        "| --- | --- |\n| c | d |")


def test_merged_table_single_element(tmp_path):
    doc = _parse(tmp_path, _merged)
    assert len(doc.elements) == 1
    assert doc.elements[0].type == "table"


# ---------- 无告警 + schema ----------

def test_basic_no_warnings(tmp_path):
    doc = _parse(tmp_path, _basic)
    assert doc.warnings == []


def test_interleaved_passes_schema(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path, _interleaved)
    assert is_valid(doc.to_dict())


def test_merged_passes_schema(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path, _merged)
    assert is_valid(doc.to_dict())
