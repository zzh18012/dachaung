"""evaluation/report.py 第五百八十一轮 edges 测试（Round 1342）。

补强 edges149 未触及的角度（第七百一十四批，probe 实证）。

新角度（多级标题 docx / complete 状态透传）：
- **三级标题板**——
  level 1/2/3 各一
  + 三段 → ecbt
  {heading:3,
  paragraph:3}、
  ect 6
- **hbc 满分**——
  多级标题不扰
  heading boundary
  合规 {1.0, None}
- **complete 透传**
  ——manifest
  devset_status
  'complete' →
  devset.status
  'complete'（首次
  非 incomplete 锁）
- **单 docx 计数**
  ——file 1 /
  docx 1 / pdf 0 /
  groups 1
- forbidden tokens 第七百八十四批（open 0）
"""

from __future__ import annotations

import inspect
import json
import tempfile
from pathlib import Path

import evaluation.report as report_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _run(tmp_path):
    d = Document()
    d.add_heading("H1 top", level=1)
    d.add_paragraph("para one.")
    d.add_heading("H2 mid", level=2)
    d.add_paragraph("para two.")
    d.add_heading("H3 deep", level=3)
    d.add_paragraph("para three.")
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "complete",
        "documents": [
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx"}]}),
        encoding="utf-8")
    mf = load_manifest(tmp_path / "m.json",
                       project_root=tmp_path)
    return run_evaluation(mf, tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=800)


# ---------- 三级标题板 ----------

def test_ecbt_three_headings_batch540(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "element_count_by_type"] == {
        "value": {"heading": 3,
                  "paragraph": 3},
        "reason": None}


def test_ect_six_batch540(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "element_count_total"] == {
        "value": 6, "reason": None}


def test_counts_sum_six_batch540(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"][
        "element_count_total"] == {
        "sum": 6, "participating_docs": 1}


# ---------- hbc 满分 ----------

def test_hbc_one_batch540(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


def test_hbc_macro_one_batch540(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "heading_boundary_compliance"] == {
        "macro_average": 1.0,
        "participating_docs": 1,
        "not_evaluated": 0}


# ---------- complete 透传 ----------

def test_complete_status_passthrough_batch540(
        tmp_path):
    assert _run(tmp_path)["devset"][
        "status"] == "complete"


def test_devset_section_batch540(tmp_path):
    assert _run(tmp_path)["devset"] == {
        "status": "complete", "file_count": 1,
        "pdf_count": 0, "docx_count": 1,
        "content_group_count": 1,
        "categories_covered": []}


# ---------- 单 docx 计数 ----------

def test_docx_only_devset_batch540(tmp_path):
    d = _run(tmp_path)["devset"]
    assert d["docx_count"] == 1
    assert d["pdf_count"] == 0


# ---------- 全绿复核 ----------

def test_all_green_batch540(tmp_path):
    r = _run(tmp_path)
    ra = r["summary"]["ratio_macro_averages"]
    for k in ("schema_valid",
              "text_preservation_equal",
              "docx_locator_valid_ratio"):
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 1,
                         "not_evaluated": 0}, k


def test_success_one_batch540(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 1, "total": 1,
        "rate": 1.0}


def test_report_schema_batch540(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch540():
    src = _src()
    assert "def aggregate_summary(" in src
    assert "_RATIO_METRICS" in src
    assert "def build_devset_section(" in src


# ---------- forbidden tokens 第七百八十四批 ----------

def test_source_no_eval_batch540():
    assert "eval(" not in _src()


def test_source_no_exec_batch540():
    assert "exec(" not in _src()


def test_source_no_compile_batch540():
    assert "compile(" not in _src()


def test_source_no_globals_batch540():
    assert "globals(" not in _src()


def test_source_no_locals_batch540():
    assert "locals(" not in _src()


def test_source_no_os_system_batch540():
    assert "os.system" not in _src()


def test_source_subprocess_run_two_batch540():
    assert _src().count("subprocess.run") == 2


def test_source_no_popen_batch540():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch540():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch540():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch540():
    assert "socket" not in _src()


def test_source_no_requests_batch540():
    assert "requests" not in _src()


def test_source_no_urllib_batch540():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch540():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch540():
    assert "yield" not in _src()


def test_source_no_async_await_batch540():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch540():
    assert ".call(" not in _src()


def test_source_open_count_is_0_batch540():
    assert _src().count("open(") == 0
