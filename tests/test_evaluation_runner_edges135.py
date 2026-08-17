"""evaluation/runner.py 第五百二十六轮 edges 测试（Round 1082）。

补强 edges132-134 未触及的角度（第四百五十八批，probe 实证）。

新角度（tolerance_chars 真值翻转 + 公开指标面 20 键名册）：
- **同板同锚、容差翻转真值**：marker "w5" 距最近预测边界
  落在 (7, 30] 区间——tol 7 → P/R/F1 全 0.0（无一命中）；
  tol 30 → P 0.5 / R 1.0 / F1 0.6666666666666666——
  run_evaluation 的 tolerance_chars 不只是透传，同一
  真实板两容差下指标值整体翻转（edges114 只验内部行
  伪指标，未验值翻转）
- 宏观 R 随之翻转：tol 7 {macro 0.0, 1 参} vs tol 30
  {macro 1.0, 1 参}
- **公开 metrics 面 20 键全名册**（sorted == 精确比对）：
  私有键 _tolerance_chars/_missing_markers 已剥离、
  公开面恰 20 项
- forbidden tokens 第五百五十三批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation

_METRICS_ROSTER = [
    "chunk_boundary_f1", "chunk_boundary_precision",
    "chunk_boundary_recall", "chunk_reference_intact_ratio",
    "docx_locator_valid_ratio", "element_count_by_type",
    "element_count_total", "error_code", "figure_caption_f1",
    "figure_caption_precision", "figure_caption_recall",
    "heading_boundary_compliance",
    "image_resource_exists_ratio",
    "pdf_locator_valid_ratio", "pipeline_success",
    "schema_valid", "silent_drop_count",
    "text_char_multiset_precision",
    "text_char_multiset_recall", "text_preservation_equal"]


def _run(tmp_path, tol):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph(
        "AAA " + " ".join(f"w{i}" for i in range(1, 21)))
    d.add_paragraph("BBB tail end.")
    d.save(str(tmp_path / "samples" / "knee.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "w5", "position": "before"}]}),
        encoding="utf-8")
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/knee.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json"}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(tmp_path / "m.json",
                                        tmp_path),
                          tmp_path / "o.json", max_chars=40,
                          tolerance_chars=tol)


_BT = ("chunk_boundary_precision",
       "chunk_boundary_recall", "chunk_boundary_f1")


# ---------- 容差翻转真值 ----------

def test_tolerance_flips_values_batch281(tmp_path):
    m7 = _run(tmp_path, 7)["per_doc"][0]["metrics"]
    for k in _BT:
        assert m7[k] == {"value": 0.0, "reason": None}
    m30 = _run(tmp_path, 30)["per_doc"][0]["metrics"]
    assert m30["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m30["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m30["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 宏观 R 随容差翻转 ----------

def test_macro_recall_flips_batch281(tmp_path):
    ra7 = _run(tmp_path, 7)["summary"][
        "ratio_macro_averages"]
    assert ra7["chunk_boundary_recall"] == {
        "macro_average": 0.0, "participating_docs": 1,
        "not_evaluated": 0}
    ra30 = _run(tmp_path, 30)["summary"][
        "ratio_macro_averages"]
    assert ra30["chunk_boundary_recall"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 0}


# ---------- 公开指标面 20 键名册 ----------

def test_public_metrics_roster_batch281(tmp_path):
    rep = _run(tmp_path, 30)
    m = rep["per_doc"][0]["metrics"]
    assert sorted(m) == _METRICS_ROSTER
    assert len(m) == 20
    blob = json.dumps(rep["per_doc"])
    assert "_tolerance_chars" not in blob
    assert "_missing_markers" not in blob


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch281():
    src = _src()
    assert "public_per_doc.append(" in src
    assert '"per_doc": public_per_doc,' in src


# ---------- forbidden tokens 第五百五十三批 ----------

def test_source_no_eval_batch281():
    assert "eval(" not in _src()


def test_source_no_exec_batch281():
    assert "exec(" not in _src()


def test_source_no_compile_batch281():
    assert "compile(" not in _src()


def test_source_no_globals_batch281():
    assert "globals(" not in _src()


def test_source_no_locals_batch281():
    assert "locals(" not in _src()


def test_source_no_os_system_batch281():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch281():
    assert "subprocess" not in _src()


def test_source_no_popen_batch281():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch281():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch281():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch281():
    assert "socket" not in _src()


def test_source_no_requests_batch281():
    assert "requests" not in _src()


def test_source_no_urllib_batch281():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch281():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch281():
    assert "yield" not in _src()


def test_source_no_async_await_batch281():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch281():
    assert _src().count("open(") == 2
