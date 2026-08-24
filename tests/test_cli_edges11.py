r"""app/cli.py 边角测试 - 第十一轮（Round 1402）。

新角度（probe 实证）：R1400/R1401 金样 docx/PDF 双胞胎在
CLI inspect 渲染层的精确算术（历史只锁过摘要行与 '(none)'
缺席，从未锁过具体 span 值/统计公式/9 宽类型填充）：
- element 行格式 '- [heading  ] doc-<sha16>::e0000  | ...'
  （类型左对齐 9 宽；element_id 用 :: 双冒号连接）
- chunk text/refs 统计行精确算术（docx 11/109/49/296 +
  1/2/1.5；pdf 7/108/47/283 + 1/2/1.3——avg 截断到 0/1 位）
- 每条 span 精确 [start:end]（docx 9 条、pdf 8 条；
  end == len(content)；image 元素永远无 span）
"""

from __future__ import annotations

import io
import os
import re
import struct
import subprocess
import sys
import zlib
from pathlib import Path

from docx import Document

VENV_PYTHON = str(
    Path(__file__).resolve().parent.parent
    / ".venv" / "Scripts" / "python.exe")
PROJECT_ROOT = \
    Path(__file__).resolve().parent.parent
_PYTHON = (VENV_PYTHON
           if Path(VENV_PYTHON).is_file()
           else sys.executable)

P1D = ("Golden body paragraph one "
       "with more than enough "
       "characters to overflow a "
       "small budget alone here.")
P2D = ("Golden body paragraph two "
       "sits under the second "
       "heading and is long enough "
       "to split similarly.")
P1P = ("Golden pdf body paragraph "
       "one with more than enough "
       "characters to overflow a "
       "small chunk budget.")
P2P = ("Golden pdf body paragraph "
       "two sits under the second "
       "heading and is long enough "
       "to split as well.")


def _make_png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    return (sig
            + chunk(b"IHDR",
                    struct.pack(
                        ">IIBBBBB",
                        1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT",
                    zlib.compress(
                        b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b""))


def _build_docx(tmp_path):
    d = Document()
    d.add_heading("Golden Root", 1)
    d.add_paragraph(P1D)
    d.add_paragraph(
        "Figure 1: golden caption")
    d.add_paragraph("")
    d.add_picture(io.BytesIO(_make_png()),
                  width=914400)
    d.add_heading("Nested", 2)
    d.add_paragraph(P2D)
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "h1"
    t.cell(0, 1).text = "h2"
    t.cell(1, 0).text = "v1"
    t.cell(1, 1).text = "v2"
    d.add_paragraph("tail paragraph")
    p = tmp_path / "gold.docx"
    d.save(str(p))
    return p


def _build_pdf(tmp_path):
    def text(x, y, t):
        return (f"BT /F1 12 Tf {x} {y} "
                f"Td ({t}) Tj ET")

    grid = [
        "1 w 0 0 0 RG",
        "100 400 m 340 400 l S",
        "100 350 m 340 350 l S",
        "100 400 m 100 350 l S",
        "220 400 m 220 350 l S",
        "340 400 m 340 350 l S",
    ]
    c1 = "\n".join(grid + [
        text(72, 700, "Golden Root"),
        text(72, 640, P1P),
        text(72, 580,
             "Figure 1: golden pdf "
             "caption"),
        text(110, 365, "ph1"),
        text(230, 365, "ph2"),
        "q 80 0 0 80 450 600 cm "
        "/Im1 Do Q",
        text(72, 300, "Nested"),
        text(72, 240, P2P),
        text(72, 180,
             "tail short")]).encode()
    img = b"\x00\xff\x00"
    objs = {
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] /Count 1 >>"),
        3: (b"<< /Type /Page /Parent 2 0 R "
            b"/MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 "
            b"6 0 R >> /XObject << /Im1 "
            b"5 0 R >> >> /Contents "
            b"4 0 R >>"),
        5: (b"<< /Type /XObject /Subtype "
            b"/Image /Width 1 /Height 1 "
            b"/ColorSpace /DeviceRGB "
            b"/BitsPerComponent 8 "
            b"/Length 3 >>\nstream\n"
            + img + b"\nendstream"),
        6: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        4: (b"<< /Length "
            + str(len(c1)).encode()
            + b" >>\nstream\n" + c1
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n".encode()
                + objs[oid] + b"\nendobj\n")
    xref_pos = len(out)
    out += (b"xref\n0 7\n"
            b"0000000000 65535 f \n")
    for oid in range(1, 7):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 7 "
            b"/Root 1 0 R >>\nstartxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    p = tmp_path / "gold.pdf"
    p.write_bytes(bytes(out))
    return p


def _run_cli(args):
    env = {
        **os.environ,
        "PYTHONIOENCODING": "utf-8",
        "PYTHONPATH":
            str(PROJECT_ROOT)
            + os.pathsep
            + os.environ.get("PYTHONPATH", ""),
    }
    proc = subprocess.run(
        [_PYTHON, "-X", "utf8",
         "-m", "app.cli", *args],
        capture_output=True, text=True,
        encoding="utf-8",
        errors="replace", env=env)
    return (proc.returncode,
            proc.stdout, proc.stderr)


def _inspect(tmp_path, kind):
    p = (_build_docx(tmp_path)
         if kind == "docx"
         else _build_pdf(tmp_path))
    out = tmp_path / (kind + ".json")
    rc, _, err = _run_cli(
        ["parse", str(p), "-o", str(out),
         "--max-chars", "120"])
    assert rc == 0, err
    rc, so, err = _run_cli(
        ["inspect", str(out),
         "--elements", "--chunks",
         "--spans", "--limit", "0"])
    assert rc == 0, err
    return so


_SPAN_RE = re.compile(
    r"span: doc-[0-9a-f]{16}::"
    r"(e\d{4})\[(\d+):(\d+)\]")


def _spans(so):
    return _SPAN_RE.findall(so)


# ---------- docx：element 行 ----------

def test_docx_element_line_format(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert re.search(
        r"^  - \[heading  \] "
        r"doc-[0-9a-f]{16}::e0000  "
        r"\| Golden Root$",
        so, re.M)


def test_docx_type_padding_9wide(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    for padded in ("[heading  ]",
                   "[paragraph]",
                   "[caption  ]",
                   "[image    ]",
                   "[table    ]"):
        assert padded in so


def test_docx_element_count_line(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert "elements (10):" in so
    assert len(re.findall(
        r"^  - \[", so, re.M)) == 10


def test_docx_image_empty_preview(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert re.search(
        r"\[image    \] "
        r"doc-[0-9a-f]{16}::e0005  \| $",
        so, re.M)


def test_docx_table_preview_flattened(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert ("| h1 | h2 | "
            "| --- | --- | "
            "| v1 | v2 |") in so


# ---------- docx：统计行 ----------

def test_docx_chunk_text_stats(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert ("chunk text:  min=11 "
            "max=109 avg=49 total=296"
            ) in so


def test_docx_chunk_refs_stats(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert ("chunk refs:  min=1 "
            "max=2 avg=1.5") in so


# ---------- docx：span 算术 ----------

def test_docx_span_values(tmp_path):
    so = _inspect(tmp_path, "docx")
    assert _spans(so) == [
        ("e0000", "0", "11"),
        ("e0001", "0", "97"),
        ("e0002", "0", "24"),
        ("e0003", "0", "5"),
        ("e0004", "0", "5"),
        ("e0006", "0", "6"),
        ("e0007", "0", "94"),
        ("e0008", "0", "37"),
        ("e0009", "0", "14")]


def test_docx_span_count(tmp_path):
    so = _inspect(tmp_path, "docx")
    assert len(_spans(so)) == 9
    assert "spans: (none)" not in so


def test_docx_image_has_no_span(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert "e0005[" not in so


def test_limit0_no_more_hint(
        tmp_path):
    so = _inspect(tmp_path, "docx")
    assert ("use --limit 0 to see all"
            not in so)


def test_no_parent_shown(tmp_path):
    so = _inspect(tmp_path, "docx")
    assert " parent=" not in so


# ---------- pdf：element 行 ----------

def test_pdf_element_count_line(
        tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert "elements (9):" in so
    assert "chunks (6):" in so
    assert len(re.findall(
        r"^  - \[", so, re.M)) == 9


def test_pdf_word_group_element(
        tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert re.search(
        r"\[heading  \] "
        r"doc-[0-9a-f]{16}::e0003  "
        r"\| ph1 ph2$",
        so, re.M)


def test_pdf_image_last_empty_preview(
        tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert re.search(
        r"\[image    \] "
        r"doc-[0-9a-f]{16}::e0008  \| $",
        so, re.M)


# ---------- pdf：统计行 + span ----------

def test_pdf_chunk_text_stats(
        tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert ("chunk text:  min=7 "
            "max=108 avg=47 total=283"
            ) in so


def test_pdf_chunk_refs_stats(
        tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert ("chunk refs:  min=1 "
            "max=2 avg=1.3") in so


def test_pdf_span_values(tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert _spans(so) == [
        ("e0000", "0", "11"),
        ("e0001", "0", "96"),
        ("e0002", "0", "28"),
        ("e0003", "0", "7"),
        ("e0004", "0", "6"),
        ("e0005", "0", "96"),
        ("e0006", "0", "10"),
        ("e0007", "0", "27")]


def test_pdf_image_has_no_span(
        tmp_path):
    so = _inspect(tmp_path, "pdf")
    assert "e0008[" not in so
    assert "spans: (none)" not in so
