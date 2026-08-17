"""evaluation/report.py 第五百零四轮 edges 测试（Round 1060）。

补强 edges119/120 未触及的角度（第四百三十六批，probe 实证）。

新角度（好+坏同列 documents 的真实混合板，参与语义三分）：
- edges119 的坏文档走 ef（被排除，rate 1.0）；本批让
  corrupt docx 以**普通 document** 身份入场，与好文档
  同列——summary 上三种参与语义同屏：
  success_rates total 数**所有**文档（2）→ rate 0.5；
  counts 只数非 None（participating 1，sum 2）；
  ratio macro 只数非 null（视指标而定）
- 12 个 ratio 指标在同一 run 里按参与度精确分成
  6/6 两半：schema_valid/docx_locator/chunk_ref/
  text_* 六项 {1.0, 1, 1}；pdf_locator（无 pdf 文档）
  /image（无图片）/heading（无标题）/chunk_boundary_*
  （无标注）六项 {None, 0, 2}——三种 null 成因
  （pipeline_failed / 源类型门控 / 结构缺失）一次显形
- corrupt 文档的 error_code 是 **metrics 成员**而非
  顶层键（{value: "docx_open_failed"}），per_doc 顺序
  保 manifest documents 序
- forbidden tokens 第五百三十一批（report 变体：15 项
  去 subprocess + open 0 + subprocess.run == 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("BBB second paragraph body.")
    d.save(str(tmp_path / "samples" / "good.docx"))
    (tmp_path / "samples" / "bad.docx").write_bytes(
        b"not a docx")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/good.docx",
             "source_type": "docx"},
            {"doc_id": "d2", "path": "samples/bad.docx",
             "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json", max_chars=200)


# ---------- success：total 数所有文档 ----------

def test_mixed_rate_half_batch259(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {"success_count": 1,
                                "total": 2, "rate": 0.5}


# ---------- counts：只数非 None ----------

def test_mixed_counts_participating_one_batch259(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["counts"] == {
        "element_count_total": {"sum": 2,
                                "participating_docs": 1}}


# ---------- ratio 参与度 6/6 精确分半 ----------

def test_participation_partition_batch259(tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    full = {"schema_valid", "docx_locator_valid_ratio",
            "chunk_reference_intact_ratio",
            "text_preservation_equal",
            "text_char_multiset_precision",
            "text_char_multiset_recall"}
    empty = {"pdf_locator_valid_ratio",
             "image_resource_exists_ratio",
             "heading_boundary_compliance",
             "chunk_boundary_precision",
             "chunk_boundary_recall",
             "chunk_boundary_f1"}
    assert set(ra) == full | empty
    for k in full:
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 1,
                         "not_evaluated": 1}, k
    for k in empty:
        assert ra[k] == {"macro_average": None,
                         "participating_docs": 0,
                         "not_evaluated": 2}, k


# ---------- 三种 null 成因同屏 ----------

def test_null_causes_triple_batch259(tmp_path):
    rep = _run(tmp_path)
    d2 = rep["per_doc"][1]["metrics"]
    assert d2["element_count_total"] == {
        "value": None, "reason": "pipeline_failed"}
    ra = rep["summary"]["ratio_macro_averages"]
    assert ra["pdf_locator_valid_ratio"] == {
        "macro_average": None, "participating_docs": 0,
        "not_evaluated": 2}
    assert ra["chunk_boundary_f1"] == {
        "macro_average": None, "participating_docs": 0,
        "not_evaluated": 2}


# ---------- 顺序与 error_code 位置 ----------

def test_per_doc_order_and_error_metric_batch259(tmp_path):
    rep = _run(tmp_path)
    assert [(p["doc_id"],
             p["metrics"]["pipeline_success"]["value"])
            for p in rep["per_doc"]] == [("d1", True),
                                         ("d2", False)]
    assert "error_code" not in rep["per_doc"][1]
    assert rep["per_doc"][1]["metrics"]["error_code"] == {
        "value": "docx_open_failed", "reason": None}


# ---------- 无 expectations → silent 全空 ----------

def test_silent_none_no_expectations_batch259(tmp_path):
    assert _run(tmp_path)["summary"][
        "silent_drop_total"] is None


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch259():
    src = _src()
    assert "if total else None" in src
    assert '"participating_docs": len(values)' in src


# ---------- forbidden tokens 第五百三十一批（report 变体） ----------

def test_source_no_eval_batch259():
    assert "eval(" not in _src()


def test_source_no_exec_batch259():
    assert "exec(" not in _src()


def test_source_no_compile_batch259():
    assert "compile(" not in _src()


def test_source_no_globals_batch259():
    assert "globals(" not in _src()


def test_source_no_locals_batch259():
    assert "locals(" not in _src()


def test_source_no_os_system_batch259():
    assert "os.system" not in _src()


def test_source_no_popen_batch259():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch259():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch259():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch259():
    assert "socket" not in _src()


def test_source_no_requests_batch259():
    assert "requests" not in _src()


def test_source_no_urllib_batch259():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch259():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch259():
    assert "yield" not in _src()


def test_source_no_async_await_batch259():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch259():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch259():
    assert _src().count("subprocess.run") == 2
