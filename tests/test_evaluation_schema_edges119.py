"""evaluation/schema.py 第四百八十七轮 edges 测试（Round 1043）。

补强 edges118 未触及的角度（第四百一十九批，probe 实证）。

新角度（真实报告单点突变矩阵）：
- 此前 schema 测试的手工 payload 最小化（缺键/错键
  常并发出多错）；edges75 的 run_evaluation 报告是
  patch process_single 的空 metrics 板。本批用真实
  docx 穿真实管线产出的报告（per_doc metrics 恰 20
  键、provenance/devset 全真值）做单键突变——其余
  全真，每次恰 1 错、路径精确
- A devset.status→"bogus"：enum 错 @ ['devset',
  'status']；B per_doc[0].source_type→"txt"：enum 错
  @ ['per_doc', 0, 'source_type']（含数组下标）；
  C provenance 删 parser_name：required 错 @
  ['provenance']；D summary.silent_drop_total→"x"：
  类型错（'integer', 'null' 联合）
- 四突变共享同一真实基底：真实管线的全绿产物对
  schema 是紧贴合，任何单键扰动即单点报错
- forbidden tokens 第五百一十四批（open 2）
"""

from __future__ import annotations

import copy
import inspect
import json

import pytest
from docx import Document

import evaluation.schema as schema_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import EvalSchemaError, validate


def _real_report(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_paragraph("Hello world paragraph one.")
    d.add_paragraph("Second paragraph here.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/a.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json", max_chars=50)


def _mut(tmp_path, fn):
    rep = _real_report(tmp_path)
    fn(rep)
    try:
        validate(rep, "evaluation-report.schema.json")
        return None
    except EvalSchemaError as e:
        return e


# ---------- 真实基底全绿 ----------

def test_real_base_passes_batch241(tmp_path):
    rep = _real_report(tmp_path)
    assert len(rep["per_doc"][0]["metrics"]) == 20
    validate(rep, "evaluation-report.schema.json")


# ---------- 单点突变矩阵 ----------

def test_mutate_devset_status_batch241(tmp_path):
    e = _mut(tmp_path, lambda r: r["devset"].__setitem__(
        "status", "bogus"))
    assert len(e.errors) == 1
    assert e.errors[0]["path"] == ["devset", "status"]
    assert e.errors[0]["message"] == (
        "'bogus' is not one of ['complete', 'incomplete']")


def test_mutate_per_doc_source_type_batch241(tmp_path):
    e = _mut(tmp_path, lambda r: r["per_doc"][0].__setitem__(
        "source_type", "txt"))
    assert len(e.errors) == 1
    assert e.errors[0]["path"] == ["per_doc", 0, "source_type"]
    assert e.errors[0]["message"] == \
        "'txt' is not one of ['pdf', 'docx']"


def test_mutate_drop_parser_name_batch241(tmp_path):
    e = _mut(tmp_path, lambda r: r["provenance"].pop(
        "parser_name"))
    assert len(e.errors) == 1
    assert e.errors[0]["path"] == ["provenance"]
    assert e.errors[0]["message"] == \
        "'parser_name' is a required property"


def test_mutate_silent_drop_type_batch241(tmp_path):
    e = _mut(tmp_path, lambda r: r["summary"].__setitem__(
        "silent_drop_total", "x"))
    assert len(e.errors) == 1
    assert e.errors[0]["path"] == ["summary",
                                   "silent_drop_total"]
    assert e.errors[0]["message"] == (
        "'x' is not of type 'integer', 'null'")


def test_mutations_isolated_batch241(tmp_path):
    base = _real_report(tmp_path)
    for fn in (
        lambda r: r["devset"].__setitem__("status", "bogus"),
        lambda r: r["provenance"].pop("parser_name"),
    ):
        r = copy.deepcopy(base)
        fn(r)
        with pytest.raises(EvalSchemaError):
            validate(r, "evaluation-report.schema.json")
    validate(base, "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch241():
    src = _src()
    assert "sorted(validator.iter_errors(instance)" in src
    assert "key=lambda e: list(e.absolute_path)" in src


# ---------- forbidden tokens 第五百一十四批 ----------

def test_source_no_eval_batch241():
    assert "eval(" not in _src()


def test_source_no_exec_batch241():
    assert "exec(" not in _src()


def test_source_no_compile_batch241():
    assert "compile(" not in _src()


def test_source_no_globals_batch241():
    assert "globals(" not in _src()


def test_source_no_locals_batch241():
    assert "locals(" not in _src()


def test_source_no_os_system_batch241():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch241():
    assert "subprocess" not in _src()


def test_source_no_popen_batch241():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch241():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch241():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch241():
    assert "socket" not in _src()


def test_source_no_requests_batch241():
    assert "requests" not in _src()


def test_source_no_urllib_batch241():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch241():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch241():
    assert "yield" not in _src()


def test_source_no_async_await_batch241():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch241():
    assert _src().count("open(") == 2
