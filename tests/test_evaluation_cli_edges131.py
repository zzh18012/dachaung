"""evaluation/cli.py 第四百九十九轮 edges 测试（Round 1055）。

补强 edges130 未触及的角度（第四百三十一批，probe 实证）。

新角度（--max-chars 地板膝关节 31/32）：
- StructuralChunker 有绝对地板 max_chars >= 32
  （< 32 抛 ValueError "max_chars 过小"）——此前
  R1041 用 50/30 两点对比，本批锁死精确膝关节：
  31 → chunker_failed（成功 0，失败 1）、32 → 成功
  （成功 1，失败 0），rc 均 0
- 地板与文档内容无关：10+11 字符的小文档（内容
  远低于 31）在 mc 31 照样失败——失败读的是旗标
  不是内容
- forbidden tokens 第五百二十六批（open 1）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document

import evaluation.cli as cli_mod
from evaluation.cli import main


def _setup(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_paragraph("Short one.")
    d.add_paragraph("Short two.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/a.docx",
                       "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")


def _run(tmp_path, mc):
    out, err = io.StringIO(), io.StringIO()
    rp = tmp_path / f"o{mc}.json"
    with contextlib.redirect_stdout(out), \
            contextlib.redirect_stderr(err):
        rc = main(["run", "--manifest",
                   str(tmp_path / "m.json"),
                   "--output", str(rp),
                   "--max-chars", str(mc)])
    rep = json.loads(rp.read_text(encoding="utf-8"))
    return rc, out.getvalue(), err.getvalue(), rep


# ---------- 膝关节下沿 31 ----------

def test_mc31_floor_fails_batch253(tmp_path):
    _setup(tmp_path)
    rc, out, err, rep = _run(tmp_path, 31)
    assert rc == 0
    assert err == ""
    assert "documents=1（成功 0，失败 1）" in out
    m = rep["per_doc"][0]["metrics"]
    assert m["error_code"] == {"value": "chunker_failed",
                               "reason": None}
    assert m["pipeline_success"] == {"value": False,
                                     "reason": None}


# ---------- 膝关节上沿 32 ----------

def test_mc32_floor_passes_batch253(tmp_path):
    _setup(tmp_path)
    rc, out, err, rep = _run(tmp_path, 32)
    assert rc == 0
    assert "documents=1（成功 1，失败 0）" in out
    m = rep["per_doc"][0]["metrics"]
    assert m["error_code"] == {"value": None, "reason": None}
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}


# ---------- 地板与内容无关 ----------

def test_floor_content_independent_batch253(tmp_path):
    _setup(tmp_path)
    _, _, _, rep31 = _run(tmp_path, 31)
    _, _, _, rep32 = _run(tmp_path, 32)
    assert rep31["per_doc"][0]["metrics"][
        "element_count_total"] == {
        "value": None, "reason": "pipeline_failed"}
    assert rep32["per_doc"][0]["metrics"][
        "element_count_total"]["value"] == 2
    rates = [rep31["summary"]["success_rates"]
             ["pipeline_success"]["rate"],
             rep32["summary"]["success_rates"]
             ["pipeline_success"]["rate"]]
    assert rates == [0.0, 1.0]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch253():
    src = _src()
    assert '"--max-chars"' in src
    assert "default=800" in src


# ---------- forbidden tokens 第五百二十六批 ----------

def test_source_no_eval_batch253():
    assert "eval(" not in _src()


def test_source_no_exec_batch253():
    assert "exec(" not in _src()


def test_source_no_compile_batch253():
    assert "compile(" not in _src()


def test_source_no_globals_batch253():
    assert "globals(" not in _src()


def test_source_no_locals_batch253():
    assert "locals(" not in _src()


def test_source_no_os_system_batch253():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch253():
    assert "subprocess" not in _src()


def test_source_no_popen_batch253():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch253():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch253():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch253():
    assert "socket" not in _src()


def test_source_no_requests_batch253():
    assert "requests" not in _src()


def test_source_no_urllib_batch253():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch253():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch253():
    assert "yield" not in _src()


def test_source_no_async_await_batch253():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch253():
    assert _src().count("open(") == 1
