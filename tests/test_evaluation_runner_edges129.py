"""evaluation/runner.py 第四百八十四轮 edges 测试（Round 1040）。

补强 edges128 未触及的角度（第四百一十六批，probe 实证）。

新角度（真实 docx 穿透真实 process_single 全链）：
- 此前 runner 测试 107/125 patch 掉 process_single、
  其余喂 b"x" 伪字节；真实 python-docx 文档穿过
  未打补丁管线（fallback parser → 结构分块 → 校验
  → 指标 → 报告 RS）从未在 runner 层锁过
- max_chars 50 全绿板：pipeline/schema_valid True、
  element_count_total 2、docx_locator 1.0（两段落
  paragraph_index 均 >= 1）、text_preservation_equal
  True、报告过 evaluation-report RS
- 同板加 annotation（"Second" before）→ 两段合并单
  chunk 无预测边界 → precision/F1 null
  no_predicted_boundaries + recall 0.0（真实文本的
  分母塌缩，非手工 metrics dict）
- max_chars 30/25：真实结构分块器在小上限下失败 →
  error_code chunker_failed、pipeline_success False、
  element_count_total/schema_valid 双 null
  pipeline_failed、汇总 success rate 0.0
- forbidden tokens 第五百一十一批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate_file


def _root(tmp_path, max_chars, anchors=None):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_paragraph("Hello world paragraph one.")
    d.add_paragraph("Second paragraph here.")
    d.save(str(tmp_path / "samples" / "real.docx"))
    entry = {"doc_id": "d1", "path": "samples/real.docx",
             "source_type": "docx"}
    if anchors is not None:
        (tmp_path / "anns").mkdir()
        (tmp_path / "anns" / "ann.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": "d1",
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        entry["annotation_file"] = "anns/ann.json"
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [entry],
        "expected_failures": []}), encoding="utf-8")
    out = tmp_path / "o.json"
    rep = run_evaluation(load_manifest(mf, tmp_path), out,
                         max_chars=max_chars)
    return rep, out


# ---------- 真实 docx 全绿板 ----------

def test_real_docx_happy_board_batch238(tmp_path):
    rep, _ = _root(tmp_path, 50)
    m = rep["per_doc"][0]["metrics"]
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}
    assert m["schema_valid"] == {"value": True, "reason": None}
    assert m["element_count_total"] == {"value": 2,
                                        "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}


def test_real_docx_report_rs_valid_batch238(tmp_path):
    _, out = _root(tmp_path, 50)
    validate_file(out, "evaluation-report.schema.json")


# ---------- 单 chunk annotation 分母塌缩 ----------

def test_annotation_single_chunk_nulls_batch238(tmp_path):
    rep, _ = _root(tmp_path, 50, anchors=[
        {"marker": "Second", "position": "before"}])
    m = rep["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": None, "reason": "no_predicted_boundaries"}
    assert m["chunk_boundary_recall"] == {"value": 0.0,
                                          "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": None, "reason": "no_predicted_boundaries"}


# ---------- 小 max_chars 真实分块失败 ----------

def test_mc30_chunker_failed_board_batch238(tmp_path):
    rep, _ = _root(tmp_path, 30)
    m = rep["per_doc"][0]["metrics"]
    assert m["pipeline_success"] == {"value": False,
                                     "reason": None}
    assert m["error_code"] == {"value": "chunker_failed",
                               "reason": None}
    assert m["element_count_total"] == {
        "value": None, "reason": "pipeline_failed"}
    assert m["schema_valid"] == {
        "value": None, "reason": "pipeline_failed"}


def test_mc25_success_rate_zero_batch238(tmp_path):
    rep, _ = _root(tmp_path, 25)
    assert rep["per_doc"][0]["metrics"]["error_code"] == {
        "value": "chunker_failed", "reason": None}
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {"success_count": 0, "total": 1,
                                "rate": 0.0}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch238():
    src = _src()
    assert "def run_evaluation(" in src
    assert "document, errors = process_single(" in src
    assert "_load_annotation" in src


# ---------- forbidden tokens 第五百一十一批 ----------

def test_source_no_eval_batch238():
    assert "eval(" not in _src()


def test_source_no_exec_batch238():
    assert "exec(" not in _src()


def test_source_no_compile_batch238():
    assert "compile(" not in _src()


def test_source_no_globals_batch238():
    assert "globals(" not in _src()


def test_source_no_locals_batch238():
    assert "locals(" not in _src()


def test_source_no_os_system_batch238():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch238():
    assert "subprocess" not in _src()


def test_source_no_popen_batch238():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch238():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch238():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch238():
    assert "socket" not in _src()


def test_source_no_requests_batch238():
    assert "requests" not in _src()


def test_source_no_urllib_batch238():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch238():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch238():
    assert "yield" not in _src()


def test_source_no_async_await_batch238():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch238():
    assert _src().count("open(") == 2
