"""evaluation/runner.py 第五百七十轮 edges 测试（Round 1126）。

补强 edges144 未触及的角度（第五百零二批，probe 实证）。

新角度（配对 DOCX+PDF 真跑）：
- **devset 精确六键**——双向 paired_with 的 docx+pdf 真跑 →
  devset {status incomplete, file_count 2,
  content_group_count 1, pdf_count 1, docx_count 1,
  categories_covered [cat-a, cat-b]}——配对计数经真实
  load_manifest + run 到报告（旧锁 edges20 直构 dataclass
  只断言 > 0，真跑精确值首锁）
- **混合成败行**——docx 行全活；空白最小 PDF 行
  pipeline_success False + error_code
  no_extracted_elements + ect null pipeline_failed——
  同板两命运
- **summary 半成功**——success {1, 2, 0.5} +
  pdf_locator macro null / participating 0 /
  not_evaluated 2（唯一 pdf 失败 + docx 本就不参评）
- forbidden tokens 第五百九十八批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate

_PDF = (b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 100 100]>>endobj\n"
        b"xref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n"
        b"0000000052 00000 n \n0000000101 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n164\n%%EOF")


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Docx side.")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    (tmp_path / "samples" / "p.pdf").write_bytes(_PDF)
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "dx", "path": "samples/g.docx",
             "source_type": "docx", "categories": ["cat-a"],
             "paired_with": "pp"},
            {"doc_id": "pp", "path": "samples/p.pdf",
             "source_type": "pdf", "categories": ["cat-b"],
             "paired_with": "dx"}]}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _run(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    return r


# ---------- devset 精确六键 ----------

def test_paired_mixed_devset_exact_batch325(tmp_path):
    r = _run(tmp_path)
    assert r["devset"] == {
        "status": "incomplete",
        "file_count": 2,
        "content_group_count": 1,
        "pdf_count": 1,
        "docx_count": 1,
        "categories_covered": ["cat-a", "cat-b"]}


# ---------- 混合成败行 ----------

def test_paired_mixed_rows_batch325(tmp_path):
    r = _run(tmp_path)
    assert [p["doc_id"] for p in r["per_doc"]] == ["dx", "pp"]
    assert [p["source_type"] for p in r["per_doc"]] == [
        "docx", "pdf"]
    ok = [p["metrics"]["pipeline_success"]["value"]
          for p in r["per_doc"]]
    assert ok == [True, False]
    pp = r["per_doc"][1]["metrics"]
    assert pp["error_code"] == {"value": "no_extracted_elements",
                                "reason": None}
    assert pp["element_count_total"] == {
        "value": None, "reason": "pipeline_failed"}


# ---------- summary 半成功 ----------

def test_paired_mixed_summary_batch325(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 1, "total": 2, "rate": 0.5}
    assert r["summary"]["ratio_macro_averages"][
        "pdf_locator_valid_ratio"] == {
        "macro_average": None,
        "participating_docs": 0,
        "not_evaluated": 2}


# ---------- 报告照过 Schema ----------

def test_paired_mixed_schema_ok_batch325(tmp_path):
    r = _run(tmp_path)
    validate(r, "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch325():
    src = _src()
    assert "计时只记 total（用 time.perf_counter 包住 process_single）" in src
    assert "失败文档（errors 非空）也写入 per_doc" in src


# ---------- forbidden tokens 第五百九十八批 ----------

def test_source_no_eval_batch325():
    assert "eval(" not in _src()


def test_source_no_exec_batch325():
    assert "exec(" not in _src()


def test_source_no_compile_batch325():
    assert "compile(" not in _src()


def test_source_no_globals_batch325():
    assert "globals(" not in _src()


def test_source_no_locals_batch325():
    assert "locals(" not in _src()


def test_source_no_os_system_batch325():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch325():
    assert "subprocess" not in _src()


def test_source_no_popen_batch325():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch325():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch325():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch325():
    assert "socket" not in _src()


def test_source_no_requests_batch325():
    assert "requests" not in _src()


def test_source_no_urllib_batch325():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch325():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch325():
    assert "yield" not in _src()


def test_source_no_async_await_batch325():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch325():
    assert _src().count("open(") == 2
