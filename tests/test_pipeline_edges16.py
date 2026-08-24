r"""app/pipeline.py 边角测试 - 第十六轮（Round 1393）。

新角度（probe 实证）：真实文件上的 heading 锚定分块几何
（历史 heading 边界只用 md/html/ipynb 或 monkeypatch，真
PDF/真 DOCX 字节从未验证）——双标题 + 五长段、mc=150：
- 5 chunk：heading 并入后随长段（2 srcs），每个长段独立
  （单句不可再分），PDF 与 DOCX 双板形状完全一致
- 每 chunk ≤ max_chars
- hbc 1.0（两个 heading 都锚住 chunk 开头）、tpe True
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.pipeline import process_single


P1 = ("First body paragraph with "
      "plenty of characters to fill "
      "a small chunk budget well "
      "beyond one hundred chars "
      "total.")
P2 = ("Second body paragraph also "
      "long enough to exceed the "
      "tiny max chars budget on its "
      "own without any help needed.")
P3 = ("Third body paragraph closes "
      "section one with more text "
      "that again exceeds the small "
      "chunk budget alone.")
P4 = ("Fourth paragraph lives under "
      "the second heading and is "
      "long enough to split the "
      "budget yet again here.")
P5 = ("Fifth paragraph closes the "
      "document under heading two "
      "with enough characters for "
      "another overflow split.")


def _build_pdf(pages_lines):
    n_pages = len(pages_lines)
    page_ids = [3 + i * 2
                for i in range(n_pages)]
    content_ids = [4 + i * 2
                   for i in range(n_pages)]
    font_id = 3 + n_pages * 2
    objs = {font_id: b"<< /Type /Font "
                    b"/Subtype /Type1 "
                    b"/BaseFont "
                    b"/Helvetica >>"}
    objs[1] = (b"<< /Type /Catalog "
               b"/Pages 2 0 R >>")
    kids = " ".join(f"{pid} 0 R"
                    for pid in page_ids)
    objs[2] = (f"<< /Type /Pages /Kids "
               f"[{kids}] /Count "
               f"{n_pages} >>").encode()
    for i, lines in enumerate(pages_lines):
        pid = page_ids[i]
        cid = content_ids[i]
        objs[pid] = (
            f"<< /Type /Page /Parent 2 0 R "
            f"/MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 "
            f"{font_id} 0 R >> >> /Contents "
            f"{cid} 0 R >>").encode()
        blocks = []
        for (y, line) in lines:
            esc = line.replace(
                "\\", r"\\").replace(
                "(", r"\(").replace(
                ")", r"\)")
            blocks.append(
                f"BT /F1 12 Tf 72 {y} Td "
                f"({esc}) Tj ET")
        stream = " ".join(blocks).encode()
        objs[cid] = (
            b"<< /Length "
            + str(len(stream)).encode()
            + b" >>\nstream\n" + stream
            + b"\nendstream")
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n".encode()
                + objs[oid] + b"\nendobj\n")
    xref_pos = len(out)
    maxid = max(objs)
    out += (f"xref\n0 {maxid + 1}\n"
            .encode()
            + b"0000000000 65535 f \n")
    for oid in range(1, maxid + 1):
        out += (
            f"{offsets[oid]:010d} 00000 n \n"
            .encode() if oid in offsets
            else b"0000000000 65535 f \n")
    out += (f"trailer\n<< /Size "
            f"{maxid + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF"
            ).encode()
    return bytes(out)


def _pdf_doc(tmp_path):
    p = tmp_path / "geo.pdf"
    p.write_bytes(_build_pdf([[
        (700, "Section One Title"),
        (640, P1), (580, P2), (520, P3),
        (460, "Section Two Title"),
        (400, P4), (340, P5)]]))
    return process_single(
        p, None, parser_name="fallback",
        max_chars=150)


def _docx_doc(tmp_path):
    d = Document()
    d.add_heading("Section One Title", 1)
    d.add_paragraph(P1)
    d.add_paragraph(P2)
    d.add_paragraph(P3)
    d.add_heading("Section Two Title", 1)
    d.add_paragraph(P4)
    d.add_paragraph(P5)
    p = tmp_path / "geo.docx"
    d.save(str(p))
    return process_single(
        p, None, parser_name="fallback",
        max_chars=150)


_SHAPE = [
    ("Section One Title First body",
     2),
    (P2, 1),
    (P3, 1),
    ("Section Two Title Fourth "
     "paragraph", 2),
    (P5, 1)]


# ---------- PDF：形状 ----------

def test_pdf_five_chunks(tmp_path):
    doc, errors = _pdf_doc(tmp_path)
    assert errors == []
    assert len(doc.chunks) == 5


def test_pdf_chunk_texts(tmp_path):
    doc, _ = _pdf_doc(tmp_path)
    assert [c.text for c in doc.chunks
            ] == [
        "Section One Title " + P1,
        P2, P3,
        "Section Two Title " + P4,
        P5]


def test_pdf_chunk_source_counts(tmp_path):
    doc, _ = _pdf_doc(tmp_path)
    assert [len(c.source_element_ids)
            for c in doc.chunks] == [
        2, 1, 1, 2, 1]


def test_pdf_chunks_within_max(tmp_path):
    doc, _ = _pdf_doc(tmp_path)
    assert all(len(c.text) <= 150
               for c in doc.chunks)


def test_pdf_chunk_ids_sequential(tmp_path):
    doc, _ = _pdf_doc(tmp_path)
    assert [c.chunk_id[-4:]
            for c in doc.chunks] == [
        "0000", "0001", "0002",
        "0003", "0004"]


# ---------- DOCX：同形状 ----------

def test_docx_five_chunks(tmp_path):
    doc, errors = _docx_doc(tmp_path)
    assert errors == []
    assert len(doc.chunks) == 5


def test_docx_chunk_texts(tmp_path):
    doc, _ = _docx_doc(tmp_path)
    assert [c.text for c in doc.chunks
            ] == [
        "Section One Title " + P1,
        P2, P3,
        "Section Two Title " + P4,
        P5]


def test_docx_chunk_source_counts(tmp_path):
    doc, _ = _docx_doc(tmp_path)
    assert [len(c.source_element_ids)
            for c in doc.chunks] == [
        2, 1, 1, 2, 1]


# ---------- 跨板一致 ----------

def test_pdf_docx_same_shape(tmp_path):
    pdf_doc, _ = _pdf_doc(tmp_path)
    docx_doc, _ = _docx_doc(tmp_path)
    assert ([c.text for c in pdf_doc.chunks]
            == [c.text for c
                in docx_doc.chunks])


# ---------- 指标 ----------

def test_pdf_hbc_one(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    doc, _ = _pdf_doc(tmp_path)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "pdf",
        None)
    assert m[
        "heading_boundary_compliance"] \
        == {"value": 1.0, "reason": None}
    assert m["text_preservation_equal"] \
        == {"value": True,
            "reason": None}


def test_docx_hbc_one(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    doc, _ = _docx_doc(tmp_path)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "heading_boundary_compliance"] \
        == {"value": 1.0, "reason": None}
    assert m["text_preservation_equal"] \
        == {"value": True,
            "reason": None}
