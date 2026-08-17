"""evaluation/cli.py 第五百二十七轮 edges 测试（Round 1083）。

补强 edges132-134 未触及的角度（第四百五十九批，probe 实证）。

新角度（run 子命令 tolerance 旗标真值翻转 + 报告文件取证）：
- **CLI run --tolerance-chars 7** 在 straddle 板（marker
  距最近预测边界落在 (7, 30]）→ 报告文件 P/R/F1 全
  0.0；缺省（30）同一板 → 0.5 / 1.0 / 0.6666666666666666
  ——run 的容差旗标不只是被接受（基础测试只验跑通），
  同板两容差报告文件值整体翻转
- run stdout：'documents=1（成功 1，失败 0）' 与
  'devset_status=incomplete file_count=1 groups=1
  pdf=0 docx=1'
- validate-report 对全 0.0 报告照过（0.0 是合法 value）
- forbidden tokens 第五百五十四批（open 1）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document

import evaluation.cli as cli_mod
from evaluation.cli import main


def _board(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    (tmp_path / "anns").mkdir()
    d = Document()
    d.add_paragraph(
        "AAA " + " ".join(f"w{i}" for i in range(1, 21)))
    d.add_paragraph("BBB tail end.")
    d.save(str(tmp_path / "samples" / "knee.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "w5", "position": "before"}]}),
        encoding="utf-8")
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/knee.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json"}],
        "expected_failures": []}), encoding="utf-8")


def _run_cli(tmp_path, extra, out_name):
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["run", "--manifest", str(tmp_path / "m.json"),
                   "--output", str(tmp_path / out_name),
                   "--max-chars", "40"] + extra)
    return rc, buf.getvalue()


_BT = ("chunk_boundary_precision",
       "chunk_boundary_recall", "chunk_boundary_f1")


# ---------- 容差旗标翻转报告文件值 ----------

def test_run_tolerance_flips_report_batch282(tmp_path):
    _board(tmp_path)
    rc, _ = _run_cli(tmp_path, ["--tolerance-chars", "7"],
                     "r7.json")
    assert rc == 0
    m = json.loads(
        (tmp_path / "r7.json").read_text(encoding="utf-8")
    )["per_doc"][0]["metrics"]
    for k in _BT:
        assert m[k] == {"value": 0.0, "reason": None}
    rc, _ = _run_cli(tmp_path, [], "r30.json")
    assert rc == 0
    m = json.loads(
        (tmp_path / "r30.json").read_text(encoding="utf-8")
    )["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- run stdout 行 ----------

def test_run_stdout_lines_batch282(tmp_path):
    _board(tmp_path)
    _, out = _run_cli(tmp_path, ["--tolerance-chars", "7"],
                      "r7.json")
    assert "documents=1（成功 1，失败 0）" in out
    assert ("devset_status=incomplete file_count=1"
            " groups=1 pdf=0 docx=1") in out


# ---------- 全 0.0 报告照过 validate-report ----------

def test_validate_report_all_zero_batch282(tmp_path):
    _board(tmp_path)
    _run_cli(tmp_path, ["--tolerance-chars", "7"], "r7.json")
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["validate-report",
                   str(tmp_path / "r7.json")])
    assert rc == 0
    assert "[OK]" in buf.getvalue()


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch282():
    src = _src()
    assert src.count(
        "tolerance_chars=args.tolerance_chars") == 2
    assert '"--tolerance-chars",' in src


# ---------- forbidden tokens 第五百五十四批 ----------

def test_source_no_eval_batch282():
    assert "eval(" not in _src()


def test_source_no_exec_batch282():
    assert "exec(" not in _src()


def test_source_no_compile_batch282():
    assert "compile(" not in _src()


def test_source_no_globals_batch282():
    assert "globals(" not in _src()


def test_source_no_locals_batch282():
    assert "locals(" not in _src()


def test_source_no_os_system_batch282():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch282():
    assert "subprocess" not in _src()


def test_source_no_popen_batch282():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch282():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch282():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch282():
    assert "socket" not in _src()


def test_source_no_requests_batch282():
    assert "requests" not in _src()


def test_source_no_urllib_batch282():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch282():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch282():
    assert "yield" not in _src()


def test_source_no_async_await_batch282():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch282():
    assert _src().count("open(") == 1
