"""evaluation/manifest.py 第五百零三轮 edges 测试（Round 1059）。

补强 edges130 未触及的角度（第四百三十五批，probe 实证）。

新角度（清单字段 → 报告字段全链路可追溯）：
- R1052 只锁 loader 层；本批让每个 manifest 可选
  字段在**一次真实 run** 里的下游效果全部显形：
  expectations {paragraph 5} → per-doc silent 3 →
  汇总 silent_drop_total 3；annotation（marker
  before）→ boundary P/R/F1 全 1.0；categories
  ["zeta","alpha"] → devset.categories_covered 排序
  ["alpha","zeta"]——加载、运行、聚合三段贯穿
- devset 六键全屏（status/file_count/groups/
  pdf/docx/categories_covered）真实值锁定
- forbidden tokens 第五百三十批（open 1）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    (tmp_path / "anns").mkdir()
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("CCC third paragraph body.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    (tmp_path / "anns" / "ann.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "CCC third",
             "position": "before"}]}), encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/a.docx",
            "source_type": "docx",
            "categories": ["zeta", "alpha"],
            "expectations": {"element_count_by_type":
                             {"paragraph": 5}},
            "annotation_file": "anns/ann.json"}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json", max_chars=40)


# ---------- expectations → silent 链 ----------

def test_expectations_to_silent_chain_batch258(tmp_path):
    rep = _run(tmp_path)
    m = rep["per_doc"][0]["metrics"]
    assert m["silent_drop_count"] == {"value": 3,
                                      "reason": None}
    assert rep["summary"]["silent_drop_total"] == 3


# ---------- annotation → boundary 链 ----------

def test_annotation_to_boundary_chain_batch258(tmp_path):
    rep = _run(tmp_path)
    m = rep["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {"value": 1.0,
                                             "reason": None}
    assert m["chunk_boundary_recall"] == {"value": 1.0,
                                          "reason": None}
    assert m["chunk_boundary_f1"] == {"value": 1.0,
                                      "reason": None}


# ---------- categories → devset 链 ----------

def test_categories_to_devset_chain_batch258(tmp_path):
    rep = _run(tmp_path)
    assert rep["devset"] == {
        "status": "incomplete", "file_count": 1,
        "content_group_count": 1, "pdf_count": 0,
        "docx_count": 1,
        "categories_covered": ["alpha", "zeta"]}


# ---------- 三链同屏 ----------

def test_all_chains_one_run_batch258(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["silent_drop_total"] == 3
    assert rep["per_doc"][0]["metrics"][
        "chunk_boundary_f1"]["value"] == 1.0
    assert rep["devset"]["categories_covered"] == [
        "alpha", "zeta"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch258():
    src = _src()
    assert "expectations=d.get(\"expectations\")" in src
    assert "categories_covered" in src


# ---------- forbidden tokens 第五百三十批 ----------

def test_source_no_eval_batch258():
    assert "eval(" not in _src()


def test_source_no_exec_batch258():
    assert "exec(" not in _src()


def test_source_no_compile_batch258():
    assert "compile(" not in _src()


def test_source_no_globals_batch258():
    assert "globals(" not in _src()


def test_source_no_locals_batch258():
    assert "locals(" not in _src()


def test_source_no_os_system_batch258():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch258():
    assert "subprocess" not in _src()


def test_source_no_popen_batch258():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch258():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch258():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch258():
    assert "socket" not in _src()


def test_source_no_requests_batch258():
    assert "requests" not in _src()


def test_source_no_urllib_batch258():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch258():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch258():
    assert "yield" not in _src()


def test_source_no_async_await_batch258():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch258():
    assert _src().count("open(") == 1
