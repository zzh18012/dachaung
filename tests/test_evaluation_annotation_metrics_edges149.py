"""evaluation/annotation_metrics.py 第五百八十轮 edges 测试（Round 1253）。

补强 edges148 未触及的角度（第六百二十五批，probe 实证）。

新角度（同文档异 chunking 的锚几何 / 三界分母）：
- **mc45 板结构**——4 块 [标题独块,
  段+占位, 表格, 尾+节] srcs [1,2,
  1,3] → 预测界升为 3
- **三界分母新值**——单锚 tol 0 →
  P 1/3 / R 1.0 / F1 0.5（分母 3
  首锁，前史全 2 界）
- **标题锚 d 0**——mc45 标题独块
  → "Chapter One Title" after 恰
  界 1（mc120 合并板此锚漏）
- **双锚 tol 21 翻转**——tol 20 →
  P 1/3 / F1 0.4；tol 21 → P 2/3 /
  R 1.0 / F1 0.8（第二占位距表尾
  21，同距不同分母）
- **mc60 板**——3 块 srcs [3,1,3]
  → 双锚 tol 0 全 0.5
- forbidden tokens 第七百一十六批（open 0）
"""

from __future__ import annotations

import inspect

import evaluation.annotation_metrics as am_mod
from evaluation.annotation_metrics import chunk_boundary_prf


def _ks_docx(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    p = tmp_path / "ks.docx"
    doc.save(str(p))
    return p


def _doc(tmp_path, mc):
    from app.pipeline import process_single
    d, errors = process_single(_ks_docx(tmp_path), tmp_path / "o.json",
                               parser_name="fallback", max_chars=mc)
    assert errors == []
    return d.to_dict()


def _ann(*pairs):
    return {"chunk_boundary_anchors": [
        {"marker": m, "position": pos} for m, pos in pairs]}


def _prf(dd, pairs, tol):
    r = chunk_boundary_prf(dd, _ann(*pairs), tol)
    return {k: r[k]["value"] for k in (
        "chunk_boundary_precision", "chunk_boundary_recall",
        "chunk_boundary_f1")}


# ---------- mc45 板结构 ----------

def test_mc45_four_chunks_batch451(tmp_path):
    dd = _doc(tmp_path, 45)
    assert len(dd["chunks"]) == 4
    assert [c["text"] for c in dd["chunks"][:2]] == [
        "Chapter One Title",
        "First para under chapter one. (空段落)"]
    assert dd["chunks"][2]["text"] == "| L | R |\n| --- | --- |"


def test_mc45_source_counts_batch451(tmp_path):
    dd = _doc(tmp_path, 45)
    assert [len(c["source_element_ids"]) for c in dd["chunks"]] == [
        1, 2, 1, 3]


# ---------- 三界分母新值 ----------

def test_mc45_single_anchor_third_batch451(tmp_path):
    assert _prf(_doc(tmp_path, 45),
                [("(空段落)", "after")], 0) == {
        "chunk_boundary_precision": 0.3333333333333333,
        "chunk_boundary_recall": 1.0,
        "chunk_boundary_f1": 0.5}


def test_mc45_heading_anchor_d0_batch451(tmp_path):
    assert _prf(_doc(tmp_path, 45),
                [("Chapter One Title", "after")], 0) == {
        "chunk_boundary_precision": 0.3333333333333333,
        "chunk_boundary_recall": 1.0,
        "chunk_boundary_f1": 0.5}


# ---------- 双锚 tol 21 翻转 ----------

def test_mc45_dual_tol20_batch451(tmp_path):
    assert _prf(_doc(tmp_path, 45),
                [("(空段落)", "after"), ("(空段落)", "after")], 20) == {
        "chunk_boundary_precision": 0.3333333333333333,
        "chunk_boundary_recall": 0.5,
        "chunk_boundary_f1": 0.4}


def test_mc45_dual_tol21_batch451(tmp_path):
    assert _prf(_doc(tmp_path, 45),
                [("(空段落)", "after"), ("(空段落)", "after")], 21) == {
        "chunk_boundary_precision": 0.6666666666666666,
        "chunk_boundary_recall": 1.0,
        "chunk_boundary_f1": 0.8}


def test_mc45_second_ph_d21_batch451(tmp_path):
    dd = _doc(tmp_path, 45)
    joined = " ".join(c["text"] for c in dd["chunks"])
    assert joined.startswith("(空段落)", 93)
    table_end = joined.index("| --- | --- |") + len("| --- | --- |")
    assert 93 + len("(空段落)") - table_end == 21


# ---------- 尾块内锚漏 ----------

def test_mc45_runsplit_miss_batch451(tmp_path):
    assert _prf(_doc(tmp_path, 45),
                [("run split", "after")], 0) == {
        "chunk_boundary_precision": 0.0,
        "chunk_boundary_recall": 0.0,
        "chunk_boundary_f1": 0.0}


# ---------- mc60 板 ----------

def test_mc60_three_chunks_batch451(tmp_path):
    dd = _doc(tmp_path, 60)
    assert [len(c["source_element_ids"]) for c in dd["chunks"]] == [
        3, 1, 3]
    assert dd["chunks"][0]["text"] == (
        "Chapter One Title First para under chapter one. (空段落)")


def test_mc60_dual_tol0_half_batch451(tmp_path):
    assert _prf(_doc(tmp_path, 60),
                [("(空段落)", "after"), ("(空段落)", "after")], 0) == {
        "chunk_boundary_precision": 0.5,
        "chunk_boundary_recall": 0.5,
        "chunk_boundary_f1": 0.5}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch451():
    src = _src()
    assert "matched = 0" in src
    assert "precision = matched / num_predicted" in src
    assert "recall = matched / num_anchors" in src


# ---------- forbidden tokens 第七百一十六批 ----------

def test_source_no_eval_batch451():
    assert "eval(" not in _src()


def test_source_no_exec_batch451():
    assert "exec(" not in _src()


def test_source_no_compile_batch451():
    assert "compile(" not in _src()


def test_source_no_globals_batch451():
    assert "globals(" not in _src()


def test_source_no_locals_batch451():
    assert "locals(" not in _src()


def test_source_no_os_system_batch451():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch451():
    assert "subprocess" not in _src()


def test_source_no_popen_batch451():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch451():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch451():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch451():
    assert "socket" not in _src()


def test_source_no_requests_batch451():
    assert "requests" not in _src()


def test_source_no_urllib_batch451():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch451():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch451():
    assert "yield" not in _src()


def test_source_no_async_await_batch451():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch451():
    assert "open(" not in _src()
