"""evaluation/cli.py 第五百九十九轮 edges 测试（Round 1155）。

补强 cli edges142 未触及的角度（第五百二十七批，probe 实证）。

新角度（表格 PDF 的 CLI 链 / 容差经 CLI 传导）：
- **表格板 run stdout 汇总行**——2×2 网格板经 CLI
  run → rc 0 + stdout "[OK] 评测完成" +
  "documents=1（成功 1，失败 0）" + "pdf=1 docx=0"
  （成功汇总行首锁，旧断言只读报告 JSON）
- **inspect-doc counts 行**——表格 doc 检视输出含
  "counts:      elements=2 chunks=2"——检视通道的
  真实元素/块计数行首锁
- **--tolerance-chars 传导刀锋**——DOCX Section
  Title 板：CLI tol 30 → 报告 F1 0.0、tol 31 →
  1.0——容差参数经 CLI 全链生效（首锁）
- forbidden tokens 第六百二十七批（open 1）
"""

from __future__ import annotations

import inspect
import json

import evaluation.cli as cli_mod
from docx import Document
from evaluation.cli import main


def _build_pdf(objects, n_obj) -> bytes:
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objects[num] + b"endobj\n")
    xref_pos = len(out)
    out += b"xref\n0 " + str(n_obj).encode() + b"\n"
    out += b"0000000000 65535 f \n"
    for num in range(1, n_obj):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size " + str(n_obj).encode()
            + b"/Root 1 0 R>>\nstartxref\n"
            + str(xref_pos).encode() + b"\n%%EOF\n")
    return bytes(out)


def _grid_pdf() -> bytes:
    s = (b"1 w 0 G\n"
         b"10 40 100 40 re S\n60 40 0 40 re S\n"
         b"10 60 100 0 re S\n"
         b"BT /F1 10 Tf 15 65 Td (Aa Bb) Tj ET\n"
         b"BT /F1 10 Tf 65 65 Td (Cc Dd) Tj ET\n"
         b"BT /F1 10 Tf 15 45 Td (Ee Ff) Tj ET\n"
         b"BT /F1 10 Tf 65 45 Td (Gg Hh) Tj ET")
    return _build_pdf({
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 100]"
            b"/Resources<</Font<</F1 5 0 R>>>>"
            b"/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }, 6)


def _pdf_board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "samples" / "t.pdf").write_bytes(_grid_pdf())
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "tb", "path": "samples/t.pdf",
                       "source_type": "pdf"}]}),
        encoding="utf-8")
    return mf


def _docx_board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("Section Title", level=1)
    d.add_paragraph("Body follows the heading here.")
    d.add_heading("Sub Section", level=2)
    d.add_paragraph("More body text after sub.")
    d.save(str(tmp_path / "samples" / "h.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "hd",
        "chunk_boundary_anchors": [
            {"marker": "Section Title", "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m2.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "hd", "path": "samples/h.docx",
             "source_type": "docx",
             "annotation_file": "anns/a.json"}]}),
        encoding="utf-8")
    return mf


def _run_cli(capsys, argv):
    rc = main(argv)
    out = capsys.readouterr().out
    return rc, out


# ---------- 表格板 run stdout 汇总行 ----------

def test_cli_run_table_stdout_batch353(tmp_path, capsys):
    mf = _pdf_board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "32"])
    assert rc == 0
    assert "[OK] 评测完成" in out
    assert "documents=1（成功 1，失败 0）" in out
    assert "pdf=1 docx=0" in out


def test_cli_run_table_report_batch353(tmp_path, capsys):
    mf = _pdf_board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "32"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"][
        "element_count_by_type"] == {
        "value": {"heading": 1, "table": 1}, "reason": None}


def test_cli_validate_table_report_batch353(tmp_path, capsys):
    mf = _pdf_board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "32"])
    assert rc == 0
    rc2, out2 = _run_cli(capsys, [
        "validate-report", str(tmp_path / "r.json")])
    assert rc2 == 0
    assert "[OK]" in out2


# ---------- inspect-doc counts 行 ----------

def test_cli_inspect_doc_table_counts_batch353(tmp_path, capsys):
    from app.pipeline import process_single
    _pdf_board(tmp_path)
    doc, errors = process_single(
        tmp_path / "samples" / "t.pdf", tmp_path / "doc.json",
        parser_name="fallback", max_chars=32)
    assert errors == []
    rc, out = _run_cli(capsys, [
        "inspect-doc", str(tmp_path / "doc.json")])
    assert rc == 0
    assert "counts:      elements=2 chunks=2" in out


def test_cli_inspect_doc_tolerance_flag_batch353(tmp_path, capsys):
    from app.pipeline import process_single
    _pdf_board(tmp_path)
    doc, errors = process_single(
        tmp_path / "samples" / "t.pdf", tmp_path / "doc.json",
        parser_name="fallback", max_chars=32)
    assert errors == []
    rc, out = _run_cli(capsys, [
        "inspect-doc", str(tmp_path / "doc.json"),
        "--tolerance-chars", "31"])
    assert rc == 0
    assert "counts:" in out


# ---------- --tolerance-chars 传导刀锋 ----------

def test_cli_tolerance_30_miss_batch353(tmp_path, capsys):
    mf = _docx_board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r30.json"),
        "--max-chars", "200", "--tolerance-chars", "30"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r30.json").read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"][
        "chunk_boundary_f1"] == {"value": 0.0, "reason": None}


def test_cli_tolerance_31_hit_batch353(tmp_path, capsys):
    mf = _docx_board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r31.json"),
        "--max-chars", "200", "--tolerance-chars", "31"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r31.json").read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"][
        "chunk_boundary_f1"] == {"value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_identifier_counts_batch353():
    src = _src()
    assert src.count("tolerance_chars") == 4
    assert src.count("inspect") == 9
    assert src.count("validate") == 7


# ---------- forbidden tokens 第六百二十七批 ----------

def test_source_no_eval_batch353():
    assert "eval(" not in _src()


def test_source_no_exec_batch353():
    assert "exec(" not in _src()


def test_source_no_compile_batch353():
    assert "compile(" not in _src()


def test_source_no_globals_batch353():
    assert "globals(" not in _src()


def test_source_no_locals_batch353():
    assert "locals(" not in _src()


def test_source_no_os_system_batch353():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch353():
    assert "subprocess" not in _src()


def test_source_no_popen_batch353():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch353():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch353():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch353():
    assert "socket" not in _src()


def test_source_no_requests_batch353():
    assert "requests" not in _src()


def test_source_no_urllib_batch353():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch353():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch353():
    assert "yield" not in _src()


def test_source_no_async_await_batch353():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch353():
    assert _src().count("open(") == 1
