"""evaluation/runner.py 第五百四十七轮 edges 测试（Round 1103）。

补强 edges135-137 未触及的角度（第四百七十九批，probe 实证）。

新角度（双通道同文档 / 复用确定性 / 裸文件名输出）：
- **双通道同文档**：annotation_file + expectations
  同挂一个 document——silent_drop 3（期望通道：
  期望 5 段实际 2 段）与 boundary null
  no_predicted_boundaries（标注通道：mc 200 合并
  单 chunk 无预测边界）并存——两条通道互不干扰
  各自入账（单通道形态均已锁，同屏首锁）
- **manifest 复用确定性**：同一 load_manifest 对象
  连跑两次 → summary 逐字相等（确定性；与 edges18
  的"两次调用不同 dict（不缓存）"互补：对象不同、
  内容确定）；换 max_chars 40 → 小板 summary 仍
  相等（微文档对 max_chars 不敏感）但 provenance
  如实记录 40——参数入档与聚合脱钩
- **裸文件名输出**：output_path=Path("bare.json") →
  output_root 为 Path(".") 照常 mkdir → 报告写
  CWD（monkeypatch.chdir 后落 tmp）
- forbidden tokens 第五百七十四批（open 2）
"""

from __future__ import annotations

import inspect
import json
import pathlib

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first body.")
    d.add_paragraph("BBB second body.")
    d.save(str(tmp_path / "samples" / "g.docx"))
    (tmp_path / "anns" / "a.json").write_text(
        json.dumps({
            "annotation_version": "1.0", "doc_id": "d1",
            "chunk_boundary_anchors": [
                {"marker": "BBB",
                 "position": "before"}]}),
        encoding="utf-8")
    m = {
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/g.docx",
            "source_type": "docx",
            "annotation_file": "anns/a.json",
            "expectations": {
                "element_count_by_type": {
                    "paragraph": 5}}}],
    }
    mf = tmp_path / "manifest.json"
    mf.write_text(json.dumps(m), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 双通道同文档 ----------

def test_dual_channel_one_doc_batch302(tmp_path):
    rep = run_evaluation(
        _board(tmp_path), tmp_path / "r.json",
        parser_name="fallback", max_chars=200)
    mts = rep["per_doc"][0]["metrics"]
    assert mts["silent_drop_count"] == {
        "value": 3, "reason": None}
    assert mts["chunk_boundary_f1"] == {
        "value": None,
        "reason": "no_predicted_boundaries"}


# ---------- manifest 复用确定性 ----------

def test_manifest_reuse_deterministic_batch302(tmp_path):
    man = _board(tmp_path)
    rep1 = run_evaluation(
        man, tmp_path / "r1.json",
        parser_name="fallback", max_chars=200)
    rep2 = run_evaluation(
        man, tmp_path / "r2.json",
        parser_name="fallback", max_chars=200)
    assert rep1["summary"] == rep2["summary"]
    rep3 = run_evaluation(
        man, tmp_path / "r3.json",
        parser_name="fallback", max_chars=40)
    assert rep3["summary"] == rep1["summary"]
    assert rep3["provenance"]["max_chars"] == 40
    assert rep1["provenance"]["max_chars"] == 200


# ---------- 裸文件名输出 ----------

def test_bare_filename_output_batch302(tmp_path,
                                       monkeypatch):
    man = _board(tmp_path)
    monkeypatch.chdir(tmp_path)
    run_evaluation(
        man, pathlib.Path("bare.json"),
        parser_name="fallback", max_chars=200)
    assert (tmp_path / "bare.json").exists()
    blob = json.loads(
        (tmp_path / "bare.json").read_text(
            encoding="utf-8"))
    assert blob["report_version"] == "1.1"


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch302():
    src = _src()
    assert "per_doc_results.append(" in src
    assert ("expected_failure_results: "
            "list[dict[str, Any]] = []") in src


# ---------- forbidden tokens 第五百七十四批 ----------

def test_source_no_eval_batch302():
    assert "eval(" not in _src()


def test_source_no_exec_batch302():
    assert "exec(" not in _src()


def test_source_no_compile_batch302():
    assert "compile(" not in _src()


def test_source_no_globals_batch302():
    assert "globals(" not in _src()


def test_source_no_locals_batch302():
    assert "locals(" not in _src()


def test_source_no_os_system_batch302():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch302():
    assert "subprocess" not in _src()


def test_source_no_popen_batch302():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch302():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch302():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch302():
    assert "socket" not in _src()


def test_source_no_requests_batch302():
    assert "requests" not in _src()


def test_source_no_urllib_batch302():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch302():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch302():
    assert "yield" not in _src()


def test_source_no_async_await_batch302():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch302():
    assert _src().count("open(") == 2
