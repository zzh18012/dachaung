"""evaluation/annotation_metrics.py 第四百九十三轮 edges 测试（Round 1049）。

补强 edges130 未触及的角度（第四百二十五批，probe 实证）。

新角度（真实 docx × 真实分块 × 真实 anchor 金路径）：
- annotation 测试 130 轮全部手工构造 chunks；本批
  process_single 真实解析 3 段 docx，用真实结构分块器
  在两个 max_chars 下产出真实切分（mc 55/60 → 2
  chunk [p1+p2][p3]；mc 40/50 → 3 chunk 逐段），
  真实文本 anchor "CCC third" before 直喂
  chunk_boundary_prf
- 金路径：2 chunk 板 P/R/F1 全 1.0（真实切分点与
  真实标注在容差内重合——此前从未用非手造文本锁过）
- 欠预测板：3 chunk → 2 预测边界对 1 标注 →
  P 0.5 / R 1.0 / F1 0.6666666666666666（真实
  过切分的精确分数）
- 同一真实文档 figure_caption_* 三键全
  parser_does_not_emit_relations（真实 fallback
  产物无 relations）
- forbidden tokens 第五百二十批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

import evaluation.annotation_metrics as am_mod
from evaluation.annotation_metrics import (
    chunk_boundary_prf, figure_caption_prf)


def _real_doc(tmp_path, mc):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("BBB second paragraph body.")
    d.add_paragraph("CCC third paragraph body.")
    d.save(str(p))
    from app.pipeline import process_single
    doc, errors = process_single(p, tmp_path / "s.json",
                                 parser_name="fallback",
                                 max_chars=mc,
                                 write_json=False)
    assert errors == []
    return doc.to_dict()


_ANN = {"annotation_version": "1.0", "doc_id": "x",
        "chunk_boundary_anchors": [
            {"marker": "CCC third", "position": "before"}]}


# ---------- 真实切分形态 ----------

def test_real_split_shapes_batch247(tmp_path):
    two = _real_doc(tmp_path, 55)
    assert len(two["chunks"]) == 2
    assert two["chunks"][0]["text"].startswith("AAA first")
    assert two["chunks"][1]["text"].startswith("CCC third")
    three = _real_doc(tmp_path, 40)
    assert [c["text"][:3] for c in three["chunks"]] == \
        ["AAA", "BBB", "CCC"]


# ---------- 金路径：2 chunk 全 1.0 ----------

def test_real_gold_path_batch247(tmp_path):
    dd = _real_doc(tmp_path, 55)
    out = chunk_boundary_prf(dd, _ANN)
    assert out["chunk_boundary_precision"] == {"value": 1.0,
                                               "reason": None}
    assert out["chunk_boundary_recall"] == {"value": 1.0,
                                            "reason": None}
    assert out["chunk_boundary_f1"] == {"value": 1.0,
                                        "reason": None}
    assert "_missing_markers" not in out
    assert out["_tolerance_chars"] == {"value": 30,
                                       "reason": None}


# ---------- 欠预测板：3 chunk P 0.5 ----------

def test_real_oversplit_board_batch247(tmp_path):
    dd = _real_doc(tmp_path, 40)
    out = chunk_boundary_prf(dd, _ANN)
    assert out["chunk_boundary_precision"] == {"value": 0.5,
                                               "reason": None}
    assert out["chunk_boundary_recall"] == {"value": 1.0,
                                            "reason": None}
    assert out["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 真实文档 figure_caption ----------

def test_real_figure_caption_nulls_batch247(tmp_path):
    dd = _real_doc(tmp_path, 55)
    fc = figure_caption_prf(dd, _ANN)
    assert set(fc) == {"figure_caption_precision",
                       "figure_caption_recall",
                       "figure_caption_f1"}
    for v in fc.values():
        assert v["value"] is None
        assert v["reason"] == \
            "parser_does_not_emit_relations"


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch247():
    src = _src()
    assert "parser_does_not_emit_relations" in src
    assert "def chunk_boundary_prf(" in src
    assert "def figure_caption_prf(" in src


# ---------- forbidden tokens 第五百二十批 ----------

def test_source_no_eval_batch247():
    assert "eval(" not in _src()


def test_source_no_exec_batch247():
    assert "exec(" not in _src()


def test_source_no_compile_batch247():
    assert "compile(" not in _src()


def test_source_no_globals_batch247():
    assert "globals(" not in _src()


def test_source_no_locals_batch247():
    assert "locals(" not in _src()


def test_source_no_os_system_batch247():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch247():
    assert "subprocess" not in _src()


def test_source_no_popen_batch247():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch247():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch247():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch247():
    assert "socket" not in _src()


def test_source_no_requests_batch247():
    assert "requests" not in _src()


def test_source_no_urllib_batch247():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch247():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch247():
    assert "yield" not in _src()


def test_source_no_async_await_batch247():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch247():
    assert "open(" not in _src()
