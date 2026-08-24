r"""app/parsers/fallback_parser.py 边角测试 - 第十六轮（Round 1392）。

新角度（probe 实证）：真 docx 页眉/页脚与表内图片的可见性：
- sections[0].header/.footer 文本对 parser 完全不可见
  （只有 body 元素）
- 表格单元格内 inline 图片被静默丢弃——无 image 元素、
  不落盘、无告警，cell 渲染为空 '|  |'
- paragraph_index 只数 body 段（表格不占号）
- 管线：表格 isolated chunk；irer null no_image_elements
"""

from __future__ import annotations

import io
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser


def _make_png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    return (sig
            + chunk(b"IHDR",
                    struct.pack(
                        ">IIBBBBB",
                        1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT",
                    zlib.compress(
                        b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b""))


def _build(tmp_path, name="hf.docx"):
    d = Document()
    d.add_heading("HF Doc", 1)
    d.sections[0].header.paragraphs[
        0].text = "HEADER TEXT VISIBLE?"
    d.sections[0].footer.paragraphs[
        0].text = "footer text here"
    d.add_paragraph("body before table")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "cell text"
    cp = t.cell(0, 1).paragraphs[0]
    cp.add_run().add_picture(
        io.BytesIO(_make_png()))
    d.add_paragraph("body after table")
    p = tmp_path / name
    d.save(str(p))
    return p


def _parse(tmp_path):
    p = _build(tmp_path)
    img_dir = tmp_path / "imgs"
    img_dir.mkdir()
    return FallbackParser(
        image_output_dir=str(img_dir)
    ).parse(p, compute_file_hash(p)), \
        img_dir


# ---------- 页眉/页脚不可见 ----------

def test_header_text_absent(tmp_path):
    doc, _ = _parse(tmp_path)
    assert all(
        "HEADER TEXT" not in
        (e.content or "")
        for e in doc.elements)


def test_footer_text_absent(tmp_path):
    doc, _ = _parse(tmp_path)
    assert all(
        "footer text" not in
        (e.content or "")
        for e in doc.elements)


def test_element_types(tmp_path):
    doc, _ = _parse(tmp_path)
    assert [e.type for e in doc.elements
            ] == ["heading", "paragraph",
                  "table", "paragraph"]


# ---------- 表内图片 ----------

def test_cell_image_not_extracted(tmp_path):
    doc, _ = _parse(tmp_path)
    assert all(e.type != "image"
               for e in doc.elements)


def test_cell_image_no_file(tmp_path):
    _, img_dir = _parse(tmp_path)
    assert list(img_dir.iterdir()) == []


def test_cell_image_no_warning(tmp_path):
    doc, _ = _parse(tmp_path)
    assert doc.warnings == []


def test_image_cell_renders_empty(tmp_path):
    doc, _ = _parse(tmp_path)
    table = [e for e in doc.elements
             if e.type == "table"][0]
    assert table.content == (
        "| cell text |  |\n"
        "| --- | --- |")


# ---------- locator ----------

def test_paragraph_indexes_body_only(
        tmp_path):
    doc, _ = _parse(tmp_path)
    paras = [e for e in doc.elements
             if e.type == "paragraph"]
    assert [e.source_locator[
        "paragraph_index"]
        for e in paras] == [1, 2]


def test_table_index_zero(tmp_path):
    doc, _ = _parse(tmp_path)
    table = [e for e in doc.elements
             if e.type == "table"][0]
    assert table.source_locator == {
        "table_index": 0, "section": 0}


# ---------- schema ----------

def test_passes_schema(tmp_path):
    from app.schema import is_valid
    doc, _ = _parse(tmp_path)
    assert is_valid(doc.to_dict())


# ---------- 管线 + 指标 ----------

def test_pipeline_three_chunks(tmp_path):
    from app.pipeline import \
        process_single
    p = _build(tmp_path)
    doc, errors = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text for c in doc.chunks
            ] == [
        "HF Doc body before table",
        "| cell text |  |\n"
        "| --- | --- |",
        "body after table"]


def test_pipeline_table_isolated(tmp_path):
    from app.pipeline import \
        process_single
    p = _build(tmp_path)
    doc, _ = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    assert len(
        doc.chunks[1].source_element_ids
    ) == 1


def test_metrics_irer_no_images(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    from app.pipeline import \
        process_single
    p = _build(tmp_path)
    doc, _ = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "image_resource_exists_ratio"] \
        == {"value": None,
            "reason": "no_image_elements"}
    assert m[
        "element_count_by_type"][
        "value"] == {
        "heading": 1, "paragraph": 2,
        "table": 1}
    assert m[
        "chunk_reference_intact_ratio"] \
        == {"value": 1.0, "reason": None}
