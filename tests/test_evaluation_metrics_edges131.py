"""evaluation/metrics.py 第五百一十六轮 edges 测试（Round 1072）。

补强 edges127-130 未触及的角度（第四百四十八批，probe 实证）。

新角度（真实 caption 分型 + 双图全局计数器，corpus 首次
产出 caption 类型元素）：
- parser 的 _CAPTION_RE（^(Table|Figure|Fig\.?|表|图)
  \s*[0-9０-９]+[\.、:\s]）在真实 docx 上生效：
  "Figure 1: ..." → type **caption**；中文分支 "图 1
  ..." / "表 2 ..." 同样 caption；无编号前缀的普通段
  落保持 paragraph——caption 类型首次以真实产物现身
- **双图同 doc**：image_counter 是**文档级全局**计数
  （para0_00.png 与 para2_01.png——前缀随段落、序号
  跨段累加），双图 ratio 1.0
- ecbt 三分 {paragraph 2, image 2, caption 2}、ect 6；
  题注文本与图片占位符同权计入保持性账本（全绿）
- forbidden tokens 第五百四十三批（open 0）
"""

from __future__ import annotations

import inspect
import struct
import zlib
from io import BytesIO
from pathlib import Path

from docx import Document

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from evaluation.metrics import compute_automatic_metrics


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _dual_doc(tmp_path):
    d = Document()
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("Figure 1: a real caption line.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("Figure 2: another caption line.")
    src = tmp_path / "cap.docx"
    d.save(str(src))
    doc, errors = process_single(
        src, tmp_path / "s.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    return doc.to_dict()


def _m(dd):
    return compute_automatic_metrics(dd, None, "docx",
                                     None,
                                     image_base_dir=None)


# ---------- 真实 caption 分型 ----------

def test_real_caption_typing_batch271(tmp_path):
    dd = _dual_doc(tmp_path)
    assert [e["type"] for e in dd["elements"]] == [
        "paragraph", "image", "caption",
        "paragraph", "image", "caption"]
    assert dd["elements"][2]["content"] == \
        "Figure 1: a real caption line."
    assert dd["elements"][5]["content"] == \
        "Figure 2: another caption line."


# ---------- 双图全局计数器 ----------

def test_dual_image_global_counter_batch271(tmp_path):
    dd = _dual_doc(tmp_path)
    names = [Path(e["resource_path"]).name
             for e in dd["elements"]
             if e["type"] == "image"]
    assert names[0].endswith("_para0_00.png")
    assert names[1].endswith("_para2_01.png")
    for n in names:
        assert Path(next(
            p for p in [dd["elements"][1]["resource_path"],
                        dd["elements"][4][
                            "resource_path"]]
            if Path(p).name == n)).is_file()


# ---------- ecbt 三分 ----------

def test_caption_ecbt_triple_batch271(tmp_path):
    m = _m(_dual_doc(tmp_path))
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 2, "image": 2,
                  "caption": 2},
        "reason": None}
    assert m["element_count_total"] == {"value": 6,
                                        "reason": None}
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 中文题注分支 ----------

def test_chinese_caption_batch271(tmp_path):
    d = Document()
    d.add_paragraph("图 1 中文题注内容")
    d.add_paragraph("表 2 表格题注")
    d.add_paragraph("普通段落没有编号前缀")
    src = tmp_path / "c.docx"
    d.save(str(src))
    doc, errors = process_single(
        src, tmp_path / "s2.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    dd = doc.to_dict()
    assert [e["type"] for e in dd["elements"]] == [
        "caption", "caption", "paragraph"]


# ---------- 题注计入保持性 ----------

def test_caption_preservation_batch271(tmp_path):
    m = _m(_dual_doc(tmp_path))
    assert m["text_preservation_equal"] == {
        "value": True, "reason": None}
    assert m["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert m["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch271():
    src = _src()
    assert ('metrics["element_count_by_type"] = '
            '{"value": by_type, "reason": None}') in src
    assert "from collections import Counter" in src


# ---------- forbidden tokens 第五百四十三批 ----------

def test_source_no_eval_batch271():
    assert "eval(" not in _src()


def test_source_no_exec_batch271():
    assert "exec(" not in _src()


def test_source_no_compile_batch271():
    assert "compile(" not in _src()


def test_source_no_globals_batch271():
    assert "globals(" not in _src()


def test_source_no_locals_batch271():
    assert "locals(" not in _src()


def test_source_no_os_system_batch271():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch271():
    assert "subprocess" not in _src()


def test_source_no_popen_batch271():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch271():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch271():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch271():
    assert "socket" not in _src()


def test_source_no_requests_batch271():
    assert "requests" not in _src()


def test_source_no_urllib_batch271():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch271():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch271():
    assert "yield" not in _src()


def test_source_no_async_await_batch271():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch271():
    assert "open(" not in _src()
