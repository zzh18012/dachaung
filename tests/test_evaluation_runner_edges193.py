"""evaluation/runner.py 第六百二十四轮 edges 测试（Round 1180）。

补强 edges192 未触及的角度（第五百五十二批，probe 实证）。

新角度（页眉页脚不可见 / 缺锚原因码）：
- **页眉页脚不可见**——header/footer 段落文
  本不入 elements（body-only 提取域首锁）；
  paragraph_index 仍从 0 连续（页眉不占位）
- **缺锚原因码**——marker 不在流中 → P 0.0 /
  R null+no_ground_truth_anchors_in_stream /
  F1 null+precision_or_recall_not_evaluated
  （两个原因码首锁）
- **部分缺锚容忍**——present+missing 双锚 →
  缺锚静默丢弃、剩余锚照常匹配 → 全 1.0
- forbidden tokens 第六百五十二批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None, split=False):
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "a").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Body paragraph one.")
    d.sections[0].header.paragraphs[0].text = \
        "Running Header Text"
    d.sections[0].footer.paragraphs[0].text = \
        "Page Footer Text"
    if split:
        d.add_heading("Split Heading", level=1)
    d.add_paragraph("Body paragraph two.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    docs = [{"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}]
    if anchors is not None:
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        docs[0]["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 页眉页脚不可见 ----------

def test_header_footer_invisible_batch378(tmp_path):
    _board(tmp_path, "hf")
    doc, errors = process_single(
        tmp_path / "s" / "hf.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    dd = doc.to_dict()
    assert [e["content"] for e in dd["elements"]] == [
        "Body paragraph one.", "Body paragraph two."]
    assert [e["source_locator"]["paragraph_index"]
            for e in dd["elements"]] == [0, 1]
    assert [c["text"] for c in dd["chunks"]] == [
        "Body paragraph one. Body paragraph two."]


def test_header_split_chunks_batch378(tmp_path):
    _board(tmp_path, "hf2", split=True)
    doc, errors = process_single(
        tmp_path / "s" / "hf2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    assert [c["text"] for c in doc.to_dict()["chunks"]] == [
        "Body paragraph one.",
        "Split Heading Body paragraph two."]


# ---------- 缺锚原因码 ----------

def test_present_anchor_batch378(tmp_path):
    r = run_evaluation(_board(tmp_path, "hf3", split=True,
        anchors=[{"marker": "one.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_missing_anchor_reasons_batch378(tmp_path):
    r = run_evaluation(_board(tmp_path, "hf4", split=True,
        anchors=[{"marker": "Header Text",
                  "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": None,
        "reason": "no_ground_truth_anchors_in_stream"}
    assert m["chunk_boundary_f1"] == {
        "value": None,
        "reason": "precision_or_recall_not_evaluated"}


# ---------- 部分缺锚容忍 ----------

def test_one_of_two_anchors_batch378(tmp_path):
    r = run_evaluation(_board(tmp_path, "hf5", split=True,
        anchors=[{"marker": "one.", "position": "after"},
                 {"marker": "Header Text",
                  "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


# ---------- 指标 ----------

def test_header_by_type_batch378(tmp_path):
    r = run_evaluation(_board(tmp_path, "hf6", split=True),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 2, "heading": 1},
        "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch378():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("annotation") == 10
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百五十二批 ----------

def test_source_no_eval_batch378():
    assert "eval(" not in _src()


def test_source_no_exec_batch378():
    assert "exec(" not in _src()


def test_source_no_compile_batch378():
    assert "compile(" not in _src()


def test_source_no_globals_batch378():
    assert "globals(" not in _src()


def test_source_no_locals_batch378():
    assert "locals(" not in _src()


def test_source_no_os_system_batch378():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch378():
    assert "subprocess" not in _src()


def test_source_no_popen_batch378():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch378():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch378():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch378():
    assert "socket" not in _src()


def test_source_no_requests_batch378():
    assert "requests" not in _src()


def test_source_no_urllib_batch378():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch378():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch378():
    assert "yield" not in _src()


def test_source_no_async_await_batch378():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch378():
    assert _src().count("open(") == 2
