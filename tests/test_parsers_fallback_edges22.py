r"""app/parsers/fallback_parser.py 边角测试 - 第二十二轮（Round 1407）。

新角度（probe 实证）两个未锁的 docx 结构事实：
- 嵌套表（cell 内 add_table）：内表文本被完全静默丢弃
  （'ic0'/'ic1' 不出现在任何元素/渲染里、无告警），外表
  照常 '| oc00 | oc01 |' 渲染
- add_section 分节：分节符产生 '(空段落)' 占位元素
  （paragraph_index 顺延），但 locator 的 section 字段
  不随节走（第二节的段仍 section 0）；三段并进单 chunk
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _nested(tmp_path):
    d = Document()
    d.add_paragraph("before outer")
    outer = d.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "oc00"
    outer.cell(0, 1).text = "oc01"
    outer.cell(1, 0).text = "oc10"
    outer.cell(1, 1).text = "oc11"
    inner = outer.cell(0, 0).add_table(
        rows=1, cols=2)
    inner.cell(0, 0).text = "ic0"
    inner.cell(0, 1).text = "ic1"
    d.add_paragraph("after tables")
    p = tmp_path / "nest.docx"
    d.save(str(p))
    return p


def _sections(tmp_path):
    d = Document()
    d.add_paragraph("section one para")
    d.add_section(WD_SECTION.NEW_PAGE)
    d.add_paragraph("section two para")
    p = tmp_path / "sec.docx"
    d.save(str(p))
    return p


def _parse(p):
    return FallbackParser().parse(
        p, compute_file_hash(p))


# ---------- 嵌套表 ----------

def test_nested_three_elements(
        tmp_path):
    doc = _parse(_nested(tmp_path))
    assert [e.type for e in
            doc.elements] == [
        "paragraph", "table",
        "paragraph"]


def test_nested_inner_text_dropped(
        tmp_path):
    doc = _parse(_nested(tmp_path))
    for e in doc.elements:
        assert "ic0" not in e.content
        assert "ic1" not in e.content
    assert doc.elements[
        1].content == (
        "| oc00 | oc01 |\n"
        "| --- | --- |\n"
        "| oc10 | oc11 |")


def test_nested_silent_no_warning(
        tmp_path):
    doc = _parse(_nested(tmp_path))
    assert doc.warnings == []


def test_nested_paragraph_indexes(
        tmp_path):
    doc = _parse(_nested(tmp_path))
    assert doc.elements[
        0].source_locator == {
        "paragraph_index": 0,
        "section": 0}
    assert doc.elements[
        2].source_locator == {
        "paragraph_index": 1,
        "section": 0}


def test_nested_table_locator(
        tmp_path):
    doc = _parse(_nested(tmp_path))
    assert doc.elements[
        1].source_locator == {
        "table_index": 0,
        "section": 0}


def test_nested_schema_valid(tmp_path):
    from app.schema import is_valid
    doc = _parse(_nested(tmp_path))
    assert is_valid(doc.to_dict())


def test_nested_pipeline_chunks(
        tmp_path):
    p = _nested(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text for c in doc.chunks
            ] == [
        "before outer",
        "| oc00 | oc01 |\n"
        "| --- | --- |\n"
        "| oc10 | oc11 |",
        "after tables"]
    assert len(doc.chunks[1]
               .source_element_ids) == 1


# ---------- 分节 ----------

def test_section_placeholder_element(
        tmp_path):
    doc = _parse(_sections(tmp_path))
    assert [e.content
            for e in doc.elements] == [
        "section one para",
        "(空段落)",
        "section two para"]
    assert doc.elements[
        1].metadata["empty"] is True


def test_section_indexes(tmp_path):
    doc = _parse(_sections(tmp_path))
    assert [e.source_locator[
        "paragraph_index"]
        for e in doc.elements
        ] == [0, 1, 2]


def test_section_field_stays_zero(
        tmp_path):
    """add_section 后第二节段落
    locator 的 section 仍为 0。"""
    doc = _parse(_sections(tmp_path))
    assert [e.source_locator["section"]
            for e in doc.elements
            ] == [0, 0, 0]


def test_section_schema_valid(tmp_path):
    from app.schema import is_valid
    doc = _parse(_sections(tmp_path))
    assert is_valid(doc.to_dict())


def test_section_pipeline_one_chunk(
        tmp_path):
    p = _sections(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert len(doc.chunks) == 1
    assert doc.chunks[0].text == (
        "section one para "
        "(空段落) section two para")
    assert len(doc.chunks[0]
               .source_element_ids) == 3


def test_section_docxloc_green(
        tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    p = _sections(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}
