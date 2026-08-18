"""evaluation/annotation_metrics.py 第五百七十九轮 edges 测试（Round 1247）。

补强 edges147 未触及的角度（第六百一十九批，probe 实证）。

新角度（真板锚点方位/距离全像）：
- **尾块收尾锚漏**——"Second section
  body text." after → 流尾非预测界
  → tol 30 仍全 0.0（最终块末不计
  界首锁——与 edges146 占位 after
  d 0 恒中成对照）
- **段尾 d 6 翻转**——"First para
  under chapter one." after → 距界 1
  = len(" (空段落)") = 6 → tol 5 漏 /
  tol 6 中
- **表格内 "R" d 8 翻转**——gt 61 距
  界 1（53）恰 8 → tol 7 漏 / tol 8 中
- **before d 5 翻转**——"(空段落)"
  before → 首 48 距界 1（53）恰 5 →
  tol 4 漏 / tol 5 中（真板 before
  阈值首锁）
- **流首 before 漏**——"Chapter One
  Title" before → pos 0 无界 → 0.0
- forbidden tokens 第七百一十一批（open 0）
"""

from __future__ import annotations

import inspect

import evaluation.annotation_metrics as am_mod
from evaluation.annotation_metrics import chunk_boundary_prf


def _base_doc(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from app.pipeline import process_single
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    p = tmp_path / "ks.docx"
    doc.save(str(p))
    d, errors = process_single(p, tmp_path / "o.json",
                               parser_name="fallback",
                               max_chars=120)
    assert errors == []
    return d.to_dict()


def _ann(*pairs):
    return {"chunk_boundary_anchors": [
        {"marker": m, "position": pos} for m, pos in pairs]}


def _prf(dd, pairs, tol):
    r = chunk_boundary_prf(dd, _ann(*pairs), tol)
    return {k: r[k]["value"] for k in (
        "chunk_boundary_precision", "chunk_boundary_recall",
        "chunk_boundary_f1")}


_ALL_ZERO = {
    "chunk_boundary_precision": 0.0,
    "chunk_boundary_recall": 0.0,
    "chunk_boundary_f1": 0.0,
}

_TWO_THIRDS = {
    "chunk_boundary_precision": 0.5,
    "chunk_boundary_recall": 1.0,
    "chunk_boundary_f1": 0.6666666666666666,
}


# ---------- 尾块收尾锚漏 ----------

def test_stream_end_after_miss_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path),
                [("Second section body text.", "after")], 30) == \
        _ALL_ZERO


# ---------- 段尾 d 6 翻转 ----------

def test_firstpara_after_tol5_miss_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path),
                [("First para under chapter one.", "after")], 5) == \
        _ALL_ZERO


def test_firstpara_after_tol6_hit_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path),
                [("First para under chapter one.", "after")], 6) == \
        _TWO_THIRDS


def test_firstpara_d6_arith_batch445(tmp_path):
    dd = _base_doc(tmp_path)
    joined = " ".join(c["text"] for c in dd["chunks"])
    end = joined.index("First para under chapter one.") + \
        len("First para under chapter one.")
    assert joined[end:end + 6] == " (空段落)"
    assert len(" (空段落)") == 6


# ---------- 表格内 "R" d 8 翻转 ----------

def test_r_after_tol7_miss_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path), [("R", "after")], 7) == \
        _ALL_ZERO


def test_r_after_tol8_hit_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path), [("R", "after")], 8) == \
        _TWO_THIRDS


def test_table_chunk2_text_batch445(tmp_path):
    dd = _base_doc(tmp_path)
    assert dd["chunks"][1]["text"] == "| L | R |\n| --- | --- |"


def test_r_gt_sixty_one_batch445(tmp_path):
    dd = _base_doc(tmp_path)
    joined = " ".join(c["text"] for c in dd["chunks"])
    assert joined.index("R") + 1 == 61


# ---------- before d 5 翻转 ----------

def test_ph_before_tol4_miss_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path), [("(空段落)", "before")], 4) == \
        _ALL_ZERO


def test_ph_before_tol5_hit_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path), [("(空段落)", "before")], 5) == \
        _TWO_THIRDS


def test_ph_starts_forty_eight_batch445(tmp_path):
    dd = _base_doc(tmp_path)
    joined = " ".join(c["text"] for c in dd["chunks"])
    assert joined.startswith("(空段落)", 48)
    assert joined.startswith("(空段落)", 93)


# ---------- 流首 before 漏 ----------

def test_chapter_before_miss_batch445(tmp_path):
    assert _prf(_base_doc(tmp_path),
                [("Chapter One Title", "before")], 30) == _ALL_ZERO


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch445():
    src = _src()
    assert "d = abs(pv - gv)" in src
    assert "if d <= tolerance_chars:" in src
    assert ("find_pos = stream.find(marker, search_from) "
            "if marker else -1" in src)


# ---------- forbidden tokens 第七百一十一批 ----------

def test_source_no_eval_batch445():
    assert "eval(" not in _src()


def test_source_no_exec_batch445():
    assert "exec(" not in _src()


def test_source_no_compile_batch445():
    assert "compile(" not in _src()


def test_source_no_globals_batch445():
    assert "globals(" not in _src()


def test_source_no_locals_batch445():
    assert "locals(" not in _src()


def test_source_no_os_system_batch445():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch445():
    assert "subprocess" not in _src()


def test_source_no_popen_batch445():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch445():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch445():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch445():
    assert "socket" not in _src()


def test_source_no_requests_batch445():
    assert "requests" not in _src()


def test_source_no_urllib_batch445():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch445():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch445():
    assert "yield" not in _src()


def test_source_no_async_await_batch445():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch445():
    assert "open(" not in _src()
