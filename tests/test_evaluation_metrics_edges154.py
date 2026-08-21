"""evaluation/metrics.py 第五百七十五轮 edges 测试（Round 1314）。

补强 edges153 未触及的角度（第六百八十六批，probe 实证）。

新角度（metrics 级真 DOCX 文档面）：
- **locator 分型**——
  dlvr {1.0, None} +
  plvr {None,
  not_pdf_document}
  （compute_
  automatic_metrics
  直调 docx 首锁）
- **hbc 跨型一致**——
  docx heading 亦 1.0
  （双 mc 皆同）
- **mc 晶格**——mc32
  11 块 [sequential,
  split×10]；mc10000
  1 块 text 322
  [sequential]（heading
  + 段落全合体）
- **抽取面**——ecbt
  {heading:1,
  paragraph:1} + ect
  2；tpe True；crir
  1.0；tcmp 1.0/1.0
- **sdc docx 面**——
  {heading:2}→1、
  {paragraph:5}→4
  （docx 期望同口径）
- **figure_caption**
  ——annotation 侧
  恒 {None, parser_
  does_not_emit_
  relations}（双 mc）
- forbidden tokens 第七百六十一批（open 0）
"""

from __future__ import annotations

import inspect

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from docx import Document
from evaluation.annotation_metrics import \
    figure_caption_prf
from evaluation.metrics import compute_automatic_metrics


def _doc(tmp_path, mc):
    d = Document()
    d.add_heading("HeadingTitle", level=1)
    d.add_paragraph(" ".join("Sent%d." % i
                             for i in range(40)))
    d.save(str(tmp_path / "c.docx"))
    doc, errors = process_single(tmp_path / "c.docx",
                                 tmp_path / "o.json",
                                 parser_name="fallback",
                                 max_chars=mc)
    assert errors == []
    return doc.to_dict()


def _m(dd, exp=None):
    return compute_automatic_metrics(dd, None, "docx",
                                     exp)


# ---------- locator 分型 ----------

def test_dlvr_one_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


def test_plvr_not_pdf_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["pdf_locator_valid_ratio"] == {
        "value": None, "reason": "not_pdf_document"}


# ---------- hbc 跨型一致 ----------

def test_hbc_32_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


def test_hbc_10000_batch512(tmp_path):
    m = _m(_doc(tmp_path, 10000))
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


# ---------- mc 晶格 ----------

def test_chunk_count_32_batch512(tmp_path):
    assert len(_doc(tmp_path, 32)["chunks"]) == 11


def test_strategies_32_batch512(tmp_path):
    dd = _doc(tmp_path, 32)
    assert [c["metadata"]["strategy"]
            for c in dd["chunks"][:3]] == [
        "sequential",
        "long_paragraph_sentence_split",
        "long_paragraph_sentence_split"]


def test_chunk_count_10000_batch512(tmp_path):
    assert len(_doc(tmp_path, 10000)["chunks"]) == 1


def test_single_chunk_text_322_batch512(tmp_path):
    dd = _doc(tmp_path, 10000)
    assert len(dd["chunks"][0]["text"]) == 322
    assert dd["chunks"][0]["metadata"][
        "strategy"] == "sequential"


def test_single_chunk_two_ids_batch512(tmp_path):
    dd = _doc(tmp_path, 10000)
    assert len(dd["chunks"][0][
        "source_element_ids"]) == 2


# ---------- 抽取面 ----------

def test_ecbt_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["element_count_by_type"]["value"] == {
        "heading": 1, "paragraph": 1}


def test_ect_two_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["element_count_total"]["value"] == 2


def test_tpe_true_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["text_preservation_equal"]["value"] \
        is True


def test_crir_one_batch512(tmp_path):
    m = _m(_doc(tmp_path, 32))
    assert m["chunk_reference_intact_ratio"] == {
        "value": 1.0, "reason": None}


def test_tcmp_both_one_batch512(tmp_path):
    m = _m(_doc(tmp_path, 10000))
    assert m["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert m["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}


# ---------- sdc docx 面 ----------

def test_sdc_heading_two_batch512(tmp_path):
    m = _m(_doc(tmp_path, 10000),
           {"element_count_by_type": {"heading": 2}})
    assert m["silent_drop_count"] == {"value": 1,
                                      "reason": None}


def test_sdc_paragraph_five_batch512(tmp_path):
    m = _m(_doc(tmp_path, 10000),
           {"element_count_by_type": {
               "paragraph": 5}})
    assert m["silent_drop_count"] == {"value": 4,
                                      "reason": None}


# ---------- figure_caption ----------

def test_fcp_null_32_batch512(tmp_path):
    dd = _doc(tmp_path, 32)
    assert figure_caption_prf(dd, None)[
        "figure_caption_precision"] == {
        "value": None,
        "reason": "parser_does_not_emit_relations"}


def test_fcp_null_10000_batch512(tmp_path):
    dd = _doc(tmp_path, 10000)
    assert figure_caption_prf(dd, None)[
        "figure_caption_f1"] == {
        "value": None,
        "reason": "parser_does_not_emit_relations"}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch512():
    src = _src()
    assert '_PDF_BBOX_REQUIRED_TYPES = ("heading", ' \
           '"paragraph", "caption", "list_item")' in src
    assert "drops += (exp - actual)" in src


# ---------- forbidden tokens 第七百六十一批 ----------

def test_source_no_eval_batch512():
    assert "eval(" not in _src()


def test_source_no_exec_batch512():
    assert "exec(" not in _src()


def test_source_no_compile_batch512():
    assert "compile(" not in _src()


def test_source_no_globals_batch512():
    assert "globals(" not in _src()


def test_source_no_locals_batch512():
    assert "locals(" not in _src()


def test_source_no_os_system_batch512():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch512():
    assert "subprocess" not in _src()


def test_source_no_popen_batch512():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch512():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch512():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch512():
    assert "socket" not in _src()


def test_source_no_requests_batch512():
    assert "requests" not in _src()


def test_source_no_urllib_batch512():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch512():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch512():
    assert "yield" not in _src()


def test_source_no_async_await_batch512():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_0_batch512():
    assert _src().count("open(") == 0
