"""evaluation/report.py 第五百七十八轮 edges 测试（Round 1324）。

补强 edges146 未触及的角度（第六百九十六批，probe 实证）。

新角度（3-doc 异构 devset 聚合全景 / 多类型 expectations）：
- **counts sum 8**——combo
  PDF(2) + 1P 无标题
  PDF(1) + 图片 DOCX(5)
  跨三型求和首锁
  （participating 3）
- **sdt 4 多类型**——
  docx expectations
  {heading:2, paragraph:5,
  image:2} vs 实际
  {1,3,1} → 1+2+1=4
  （report 级多类型
  求和首锁）
- **无标注 cb 面**——
  h1 无 annotation_file
  → cbp {None,
  no_annotation}；
  macro trio 恰
  participating 2 +
  not_evaluated 1
  （从全参与到带豁免
  首锁）
- **12 ratio 全景**——
  schema_valid /
  crir / tpe / tcmp /
  tcmr 全 {1.0, 3, 0}；
  dlvr/irer {1.0,1,2}；
  plvr/hbc {1.0,2,1}
- **cbp 宏 (1/15+1)/2**
  ——g1 per-doc 1/15、
  d1 1.0；cbr 1.0、
  cbf 0.5625
- **无 expectations**——
  pdf 两文档 sdc
  {None, no_expectations}
- **devset 段**——
  file 3 / pdf 2 / docx 1
  / content_group 3 /
  categories [] /
  incomplete
- forbidden tokens 第七百六十九批（open 0）
"""

from __future__ import annotations

import inspect
import json
import struct
import tempfile
import zlib
from pathlib import Path

import evaluation.report as report_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


def _chunk(t: bytes, data: bytes) -> bytes:
    c = struct.pack(">I", len(data)) + t + data
    return c + struct.pack(
        ">I", zlib.crc32(t + data) & 0xffffffff)


PNG = (b"\x89PNG\r\n\x1a\n"
       + _chunk(b"IHDR", struct.pack(">IIBBBBB",
                                     1, 1, 8, 2, 0, 0, 0))
       + _chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
       + _chunk(b"IEND", b""))

LONG = " ".join("Word%d." % i for i in range(60))
COMBO = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
         % ("A" * 80)
         + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
         % LONG).encode()
ONEP = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
        % LONG).encode()


def _board(tmp_path):
    (tmp_path / "g.pdf").write_bytes(_wrap(COMBO))
    (tmp_path / "h.pdf").write_bytes(_wrap(ONEP))
    (tmp_path / "img.png").write_bytes(PNG)
    d = Document()
    d.add_heading("PicDoc", level=1)
    d.add_paragraph("Text before picture.")
    d.add_picture(str(tmp_path / "img.png"))
    d.add_paragraph("Text after picture.")
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "ann").mkdir(exist_ok=True)
    (tmp_path / "ann" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "g1",
        "chunk_boundary_anchors": [
            {"marker": "Word3.", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "ann" / "b.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "before", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "g1", "path": "g.pdf",
             "source_type": "pdf",
             "annotation_file": "ann/a.json"},
            {"doc_id": "h1", "path": "h.pdf",
             "source_type": "pdf"},
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx",
             "annotation_file": "ann/b.json",
             "expectations": {
                 "element_count_by_type": {
                     "heading": 2, "paragraph": 5,
                     "image": 2}}}]}),
        encoding="utf-8")
    return load_manifest((tmp_path / "m.json"),
                         project_root=tmp_path)


def _run(tmp_path):
    return run_evaluation(_board(tmp_path),
                          tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=32)


# ---------- counts sum 8 ----------

def test_counts_sum_eight_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"][
        "element_count_total"] == {
        "sum": 8, "participating_docs": 3}


# ---------- sdt 多类型 ----------

def test_sdt_multi_type_four_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["silent_drop_total"] == 4


def test_d1_sdc_four_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][2]["metrics"][
        "silent_drop_count"] == {"value": 4,
                                 "reason": None}


def test_pdf_sdc_no_expectations_batch522(tmp_path):
    r = _run(tmp_path)
    for i in (0, 1):
        assert r["per_doc"][i]["metrics"][
            "silent_drop_count"] == {
            "value": None,
            "reason": "no_expectations"}


# ---------- 无标注 cb 面 ----------

def test_h1_cbp_no_annotation_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "chunk_boundary_precision"] == {
        "value": None, "reason": "no_annotation"}


def test_cbp_macro_exemption_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_precision"] == {
        "macro_average": (1 / 15 + 1.0) / 2,
        "participating_docs": 2, "not_evaluated": 1}


def test_cbr_macro_exemption_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_recall"] == {
        "macro_average": 1.0,
        "participating_docs": 2, "not_evaluated": 1}


def test_cbf_macro_exemption_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"] == {
        "macro_average": 0.5625,
        "participating_docs": 2, "not_evaluated": 1}


def test_g1_per_doc_cb_values_batch522(tmp_path):
    r = _run(tmp_path)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1 / 15, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.125, "reason": None}


# ---------- 12 ratio 全景 ----------

def test_all_doc_ratios_three_zero_batch522(tmp_path):
    r = _run(tmp_path)
    ra = r["summary"]["ratio_macro_averages"]
    for k in ("schema_valid",
              "chunk_reference_intact_ratio",
              "text_preservation_equal",
              "text_char_multiset_precision",
              "text_char_multiset_recall"):
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 3,
                         "not_evaluated": 0}


def test_docx_only_ratios_batch522(tmp_path):
    r = _run(tmp_path)
    ra = r["summary"]["ratio_macro_averages"]
    for k in ("docx_locator_valid_ratio",
              "image_resource_exists_ratio"):
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 1,
                         "not_evaluated": 2}


def test_two_participating_ratios_batch522(tmp_path):
    r = _run(tmp_path)
    ra = r["summary"]["ratio_macro_averages"]
    for k in ("pdf_locator_valid_ratio",
              "heading_boundary_compliance"):
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 2,
                         "not_evaluated": 1}


def test_h1_hbc_no_heading_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "heading_boundary_compliance"] == {
        "value": None,
        "reason": "no_heading_elements"}


def test_d1_plvr_not_pdf_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][2]["metrics"][
        "pdf_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_pdf_document"}


def test_ratio_key_count_twelve_batch522(tmp_path):
    r = _run(tmp_path)
    assert len(r["summary"][
        "ratio_macro_averages"]) == 12


# ---------- success / 顺序 ----------

def test_success_three_thirds_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 3, "total": 3, "rate": 1.0}


def test_per_doc_manifest_order_batch522(tmp_path):
    r = _run(tmp_path)
    assert [p["doc_id"] for p in r["per_doc"]] == [
        "g1", "h1", "d1"]
    assert [p["source_type"] for p
            in r["per_doc"]] == ["pdf", "pdf", "docx"]


def test_wall_time_total_float_batch522(tmp_path):
    r = _run(tmp_path)
    for p in r["per_doc"]:
        assert isinstance(
            p["wall_time_seconds"]["total"], float)


# ---------- devset 段 ----------

def test_devset_section_batch522(tmp_path):
    r = _run(tmp_path)
    assert r["devset"] == {
        "status": "incomplete", "file_count": 3,
        "pdf_count": 2, "docx_count": 1,
        "content_group_count": 3,
        "categories_covered": []}


# ---------- 报告合法性 ----------

def test_report_schema_batch522(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch522():
    src = _src()
    assert "def aggregate_summary(" in src
    assert "_RATIO_METRICS" in src
    assert '"text_char_multiset_precision"' in src
    assert '"chunk_boundary_f1"' in src


# ---------- forbidden tokens 第七百六十九批 ----------

def test_source_no_eval_batch522():
    assert "eval(" not in _src()


def test_source_no_exec_batch522():
    assert "exec(" not in _src()


def test_source_no_compile_batch522():
    assert "compile(" not in _src()


def test_source_no_globals_batch522():
    assert "globals(" not in _src()


def test_source_no_locals_batch522():
    assert "locals(" not in _src()


def test_source_no_os_system_batch522():
    assert "os.system" not in _src()


def test_source_subprocess_run_two_batch522():
    assert _src().count("subprocess.run") == 2


def test_source_no_popen_batch522():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch522():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch522():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch522():
    assert "socket" not in _src()


def test_source_no_requests_batch522():
    assert "requests" not in _src()


def test_source_no_urllib_batch522():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch522():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch522():
    assert "yield" not in _src()


def test_source_no_async_await_batch522():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch522():
    assert ".call(" not in _src()


def test_source_open_count_is_0_batch522():
    assert _src().count("open(") == 0
