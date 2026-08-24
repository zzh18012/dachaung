r"""evaluation/cli.py 边角第一百八十四轮（Round 1414）。

新角度（probe 实证）金样双胞胎配对穿评测（R1400/R1401
单测过管线，paired_with 组合 + 双侧图片落盘首次）：
- devset {file_count 2, content_group_count 1, pdf/docx
  各 1, categories ['golden']}
- success 2/2；sdt 0；ect_sum 19 参与 2（10+9）
- **irer macro 1.0 参与 2**——runner 双侧真渲染：docx
  images-<sha>/image_<sha>_para4_00.png（段基命名）、
  pdf images-<sha>/image_<sha>_p1_00.png（页基命名），
  目录按各自 source hash 分开
"""

from __future__ import annotations

import contextlib
import io
import json
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document

from evaluation.cli import main


def _make_png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    return (sig
            + chunk(b"IHDR",
                    struct.pack(
                        ">IIBBBBB",
                        1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT",
                    zlib.compress(
                        b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b""))


P1 = ("Golden body paragraph one "
      "with more than enough "
      "characters to overflow a "
      "small budget alone here.")
P2 = ("Golden body paragraph two "
      "sits under the second "
      "heading and is long enough "
      "to split similarly.")
P1P = ("Golden pdf body paragraph "
       "one with more than enough "
       "characters to overflow a "
       "small chunk budget.")
P2P = ("Golden pdf body paragraph "
       "two sits under the second "
       "heading and is long enough "
       "to split as well.")


def _stage(tmp_path):
    s = tmp_path / "samples"
    s.mkdir()
    d = Document()
    d.add_heading("Golden Root", 1)
    d.add_paragraph(P1)
    d.add_paragraph(
        "Figure 1: golden caption")
    d.add_paragraph("")
    d.add_picture(
        io.BytesIO(_make_png()),
        width=914400)
    d.add_heading("Nested", 2)
    d.add_paragraph(P2)
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "h1"
    t.cell(0, 1).text = "h2"
    t.cell(1, 0).text = "v1"
    t.cell(1, 1).text = "v2"
    d.add_paragraph("tail paragraph")
    d.save(str(s / "gold.docx"))

    def text(x, y, t_):
        return (f"BT /F1 12 Tf {x} {y} "
                f"Td ({t_}) Tj ET")

    grid = [
        "1 w 0 0 0 RG",
        "100 400 m 340 400 l S",
        "100 350 m 340 350 l S",
        "100 400 m 100 350 l S",
        "220 400 m 220 350 l S",
        "340 400 m 340 350 l S",
    ]
    c1 = "\n".join(grid + [
        text(72, 700, "Golden Root"),
        text(72, 640, P1P),
        text(72, 580,
             "Figure 1: golden pdf "
             "caption"),
        text(110, 365, "ph1"),
        text(230, 365, "ph2"),
        "q 80 0 0 80 450 600 cm "
        "/Im1 Do Q",
        text(72, 300, "Nested"),
        text(72, 240, P2P),
        text(72, 180,
             "tail short")]).encode()
    img = b"\x00\xff\x00"
    objs = {
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] "
            b"/Count 1 >>"),
        3: (b"<< /Type /Page /Parent "
            b"2 0 R /MediaBox "
            b"[0 0 612 792] "
            b"/Resources << /Font "
            b"<< /F1 6 0 R >> "
            b"/XObject << /Im1 "
            b"5 0 R >> >> /Contents "
            b"4 0 R >>"),
        5: (b"<< /Type /XObject "
            b"/Subtype /Image "
            b"/Width 1 /Height 1 "
            b"/ColorSpace /DeviceRGB "
            b"/BitsPerComponent 8 "
            b"/Length 3 >>\nstream\n"
            + img + b"\nendstream"),
        6: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        4: (b"<< /Length "
            + str(len(c1)).encode()
            + b" >>\nstream\n" + c1
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    out += (b"xref\n0 7\n"
            b"0000000000 65535 f \n")
    for oid in range(1, 7):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 7 "
            b"/Root 1 0 R >>\n"
            b"startxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    (s / "gold.pdf").write_bytes(
        bytes(out))

    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "gdocx",
             "path": "samples/gold.docx",
             "source_type": "docx",
             "paired_with": "gpdf",
             "categories": ["golden"],
             "expectations": {
                 "element_count_by_type":
                     {"heading": 2,
                      "paragraph": 5,
                      "caption": 1,
                      "image": 1,
                      "table": 1}}},
            {"doc_id": "gpdf",
             "path": "samples/gold.pdf",
             "source_type": "pdf",
             "paired_with": "gdocx",
             "categories": ["golden"],
             "expectations": {
                 "element_count_by_type":
                     {"heading": 4,
                      "paragraph": 2,
                      "caption": 1,
                      "table": 1,
                      "image": 1}}},
        ]}, ensure_ascii=False),
        encoding="utf-8")
    return mf


def _run(tmp_path):
    mf = _stage(tmp_path)
    rep = tmp_path / "r.json"
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["run", "--manifest",
                   str(mf),
                   "--output", str(rep),
                   "--parser", "fallback",
                   "--max-chars", "120"])
    data = json.loads(
        rep.read_text(encoding="utf-8"))
    return rc, data, rep


def test_run_rc0(tmp_path):
    rc, _, _ = _run(tmp_path)
    assert rc == 0


def test_devset_group_one(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["devset"] == {
        "status": "incomplete",
        "file_count": 2,
        "content_group_count": 1,
        "pdf_count": 1,
        "docx_count": 1,
        "categories_covered":
            ["golden"]}


def test_success_both(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "success_rates"][
        "pipeline_success"] == {
        "success_count": 2,
        "total": 2, "rate": 1.0}


def test_silent_drop_zero(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "silent_drop_total"] == 0


def test_ect_sum_nineteen(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "counts"][
        "element_count_total"] == {
        "sum": 19,
        "participating_docs": 2}


def test_irer_macro_one(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "image_resource_"
        "exists_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 0}


def test_per_doc_irer(tmp_path):
    _, data, _ = _run(tmp_path)
    for d in data["per_doc"]:
        assert d["metrics"][
            "image_resource_"
            "exists_ratio"] == {
            "value": 1.0,
            "reason": None}


def test_two_image_dirs(tmp_path):
    _, _, rep = _run(tmp_path)
    dirs = sorted(
        (rep.parent / "_per_doc")
        .glob("images-*"))
    assert len(dirs) == 2
    names = [d.name for d in dirs]
    assert all(
        n.startswith("images-")
        for n in names)
    assert names[0] != names[1]


def test_docx_image_para_naming(
        tmp_path):
    _, _, rep = _run(tmp_path)
    pngs = list(
        (rep.parent / "_per_doc")
        .glob("images-*/"
              "*_para4_00.png"))
    assert len(pngs) == 1


def test_pdf_image_page_naming(
        tmp_path):
    _, _, rep = _run(tmp_path)
    pngs = list(
        (rep.parent / "_per_doc")
        .glob("images-*/"
              "*_p1_00.png"))
    assert len(pngs) == 1


def test_png_signatures(tmp_path):
    _, _, rep = _run(tmp_path)
    pngs = list(
        (rep.parent / "_per_doc")
        .glob("images-*/*.png"))
    assert len(pngs) == 2
    for p in pngs:
        assert p.read_bytes()[:8] == \
            b"\x89PNG\r\n\x1a\n"


def test_dir_sha_matches_file(
        tmp_path):
    """目录 sha16 == 目录内文件名
    sha16。"""
    _, _, rep = _run(tmp_path)
    for d in (rep.parent / "_per_doc"
              ).glob("images-*"):
        sha = d.name[
            len("images-"):]
        for f in d.glob("*.png"):
            assert sha in f.name


def test_validate_report_rc0(
        tmp_path):
    _, _, rep = _run(tmp_path)
    buf = io.StringIO()
    with contextlib.redirect_stdout(
            buf):
        rc = main(["validate-report",
                   str(rep)])
    assert rc == 0
