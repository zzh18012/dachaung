"""evaluation/runner.py 第六百五十六轮 edges 测试（Round 1235）。

补强 edges220 未触及的角度（第六百零七批，probe 实证）。

新角度（DOCX 厨房水槽板：交互组合网）：
- **七元素四类型混排**——heading +
  paragraph + 空白占位 + table +
  run 拼接段 + 分节占位 + 节 2 段
  ——五类原子（标题/正文/占位/
  表格/run 拼接）单一板首合
- **双源占位同文**——"   "（空白）
  与 add_section 插入占位同为
  "(空段落)"（idx 2 与 idx 4）
- **paragraph_index 跳过表**——
  段族 0-5 连续，表走 table_index 0
- **section 恒 0**——add_section 后
  段落 section 仍 0（R1214 单锁在
  组合板复证）
- **三块交互**——标题前向合并组 /
  表隔离块 / 尾段跨节合并组（e0004-
  e0006 横跨分节符首锁）
- forbidden tokens 第七百零二批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _docx(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    (tmp_path / "s").mkdir(exist_ok=True)
    p = tmp_path / "s" / "ks.docx"
    doc.save(str(p))
    return p


def _board(tmp_path):
    _docx(tmp_path)
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "ks", "path": "s/ks.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _doc(tmp_path):
    from app.pipeline import process_single
    doc, errors = process_single(
        _docx(tmp_path), tmp_path / "o.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    return doc.to_dict()


# ---------- 元素序列 ----------

def test_element_sequence_batch433(tmp_path):
    dd = _doc(tmp_path)
    assert [e["type"] for e in dd["elements"]] == [
        "heading", "paragraph", "paragraph", "table",
        "paragraph", "paragraph", "paragraph"]


def test_two_placeholder_origins_batch433(tmp_path):
    dd = _doc(tmp_path)
    assert dd["elements"][2]["content"] == "(空段落)"
    assert dd["elements"][5]["content"] == "(空段落)"


def test_heading_and_tail_contents_batch433(tmp_path):
    dd = _doc(tmp_path)
    assert dd["elements"][0]["content"] == "Chapter One Title"
    assert dd["elements"][4]["content"] == "Tail run split"
    assert dd["elements"][6]["content"] == \
        "Second section body text."


def test_table_markdown_and_locator_batch433(tmp_path):
    dd = _doc(tmp_path)
    tbl = dd["elements"][3]
    assert tbl["content"] == "| L | R |\n| --- | --- |"
    assert tbl["source_locator"] == {"table_index": 0,
                                     "section": 0}


# ---------- 定位键家族 ----------

def test_paragraph_index_skips_table_batch433(tmp_path):
    dd = _doc(tmp_path)
    paras = [e for e in dd["elements"]
             if "paragraph_index" in e["source_locator"]]
    assert [e["source_locator"]["paragraph_index"]
            for e in paras] == [0, 1, 2, 3, 4, 5]


def test_section_stays_zero_batch433(tmp_path):
    dd = _doc(tmp_path)
    assert all(e["source_locator"]["section"] == 0
               for e in dd["elements"])


# ---------- 三块交互 ----------

def test_chunk1_heading_merge_batch433(tmp_path):
    dd = _doc(tmp_path)
    c = dd["chunks"][0]
    assert c["text"] == \
        "Chapter One Title First para under chapter one. (空段落)"
    assert [s.split("::")[-1]
            for s in c["source_element_ids"]] == \
        ["e0000", "e0001", "e0002"]


def test_chunk2_table_isolated_batch433(tmp_path):
    dd = _doc(tmp_path)
    c = dd["chunks"][1]
    assert c["text"] == "| L | R |\n| --- | --- |"
    assert [s.split("::")[-1]
            for s in c["source_element_ids"]] == ["e0003"]


def test_chunk3_section_crossing_batch433(tmp_path):
    dd = _doc(tmp_path)
    assert len(dd["chunks"]) == 3
    c = dd["chunks"][2]
    assert c["text"] == \
        "Tail run split (空段落) Second section body text."
    assert [s.split("::")[-1]
            for s in c["source_element_ids"]] == \
        ["e0004", "e0005", "e0006"]


# ---------- runner 级指标 ----------

def test_runner_metrics_batch433(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_total"] == {"value": 7,
                                        "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"heading": 1, "paragraph": 5,
                  "table": 1}, "reason": None}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


def test_runner_tpe_and_success_batch433(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["schema_valid"] == {"value": True, "reason": None}
    assert r["summary"]["success_rates"]["pipeline_success"] == {
        "success_count": 1, "total": 1, "rate": 1.0}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch433():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第七百零二批 ----------

def test_source_no_eval_batch433():
    assert "eval(" not in _src()


def test_source_no_exec_batch433():
    assert "exec(" not in _src()


def test_source_no_compile_batch433():
    assert "compile(" not in _src()


def test_source_no_globals_batch433():
    assert "globals(" not in _src()


def test_source_no_locals_batch433():
    assert "locals(" not in _src()


def test_source_no_os_system_batch433():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch433():
    assert "subprocess" not in _src()


def test_source_no_popen_batch433():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch433():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch433():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch433():
    assert "socket" not in _src()


def test_source_no_requests_batch433():
    assert "requests" not in _src()


def test_source_no_urllib_batch433():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch433():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch433():
    assert "yield" not in _src()


def test_source_no_async_await_batch433():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch433():
    assert _src().count("open(") == 2
