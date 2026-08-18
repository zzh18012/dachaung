"""evaluation/annotation_metrics.py 第五百七十八轮 edges 测试（Round 1243）。

补强 edges146 未触及的角度（第六百一十五批，probe 实证）。

新角度（第二占位锚精确距离 21 的阈值翻转）：
- **d 恰 21**——chunk3 前缀
  "Tail run split (空段落)" 20
  字符 + join 空格 1 = 第二
  "(空段落)" after 的 gt 恰距
  界 2 21 字符（probe 阈值扫描
  实证）
- **tol 20 漏**——F1 0.5（第一
  锚 d 0 恒中、第二锚差 1 字符
  漏）
- **tol 21 中**——F1 1.0（翻转
  点首锁，真板最紧距离锁）
- **单锚 tol 0 不受累**——仅第
  一锚 → P 0.5 / R 1.0 / F1 2/3
  （d 0 免容差）
- forbidden tokens 第七百零八批（open 0）
"""

from __future__ import annotations

import inspect
import tempfile
from pathlib import Path

import evaluation.annotation_metrics as am_mod
from evaluation.annotation_metrics import chunk_boundary_prf


def _base_doc(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from app.pipeline import process_single
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    p = tmp_path / "ks.docx"
    doc.save(str(p))
    d, errors = process_single(p, tmp_path / "o.json",
                               parser_name="fallback",
                               max_chars=120)
    assert errors == []
    return d.to_dict()


def _ann(*pairs):
    return {"chunk_boundary_anchors": [
        {"marker": m, "position": pos} for m, pos in pairs]}


_TWO = ("(空段落)", "after"), ("(空段落)", "after")


# ---------- tol 20 漏 ----------

def test_tol20_half_batch441(tmp_path):
    r = chunk_boundary_prf(_base_doc(tmp_path), _ann(*_TWO), 20)
    assert r["chunk_boundary_f1"] == {
        "value": 0.5, "reason": None}


def test_tol20_pr_batch441(tmp_path):
    r = chunk_boundary_prf(_base_doc(tmp_path), _ann(*_TWO), 20)
    assert r["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert r["chunk_boundary_recall"] == {
        "value": 0.5, "reason": None}


# ---------- tol 21 中 ----------

def test_tol21_all_batch441(tmp_path):
    r = chunk_boundary_prf(_base_doc(tmp_path), _ann(*_TWO), 21)
    assert r["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_tol21_pr_batch441(tmp_path):
    r = chunk_boundary_prf(_base_doc(tmp_path), _ann(*_TWO), 21)
    assert r["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert r["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}


# ---------- 距离结构 ----------

def test_chunk3_prefix_len_batch441(tmp_path):
    dd = _base_doc(tmp_path)
    assert dd["chunks"][2]["text"].startswith(
        "Tail run split (空段落)")
    assert len("Tail run split (空段落)") == 20


# ---------- 单锚 tol 0 不受累 ----------

def test_single_anchor_tol0_batch441(tmp_path):
    r = chunk_boundary_prf(
        _base_doc(tmp_path), _ann(("(空段落)", "after")), 0)
    assert r["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert r["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert r["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch441():
    src = _src()
    assert 'joined_raw = " ".join(norm_chunks)' in src
    assert "predicted.append(end)" in src
    assert "gt_positions.append(find_pos)" in src
    assert ("gt_positions.append(find_pos + len(marker))"
            in src)
    assert "if find_pos < 0:" in src


# ---------- forbidden tokens 第七百零八批 ----------

def test_source_no_eval_batch441():
    assert "eval(" not in _src()


def test_source_no_exec_batch441():
    assert "exec(" not in _src()


def test_source_no_compile_batch441():
    assert "compile(" not in _src()


def test_source_no_globals_batch441():
    assert "globals(" not in _src()


def test_source_no_locals_batch441():
    assert "locals(" not in _src()


def test_source_no_os_system_batch441():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch441():
    assert "subprocess" not in _src()


def test_source_no_popen_batch441():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch441():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch441():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch441():
    assert "socket" not in _src()


def test_source_no_requests_batch441():
    assert "requests" not in _src()


def test_source_no_urllib_batch441():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch441():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch441():
    assert "yield" not in _src()


def test_source_no_async_await_batch441():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch441():
    assert "open(" not in _src()
