"""evaluation/metrics.py 第五百三十轮 edges 测试（Round 1086）。

补强 edges130-132 未触及的角度（第四百六十二批，probe 实证）。

新角度（图片双候选解析 + 多类型 silent 混账算术）：
- **image_resource_exists_ratio 双候选解析**：
  候选一 Path(rp) 原样（相对即 CWD 相对——chdir 后
  无 base_dir 也命中）；候选二 image_base_dir /
  **basename**（丢弃目录）——resource_path 带子目录
  "images-abc/img_00.png" 时：文件落在 base 直下
  img_00.png → 1.0（候选二救场）；文件只在
  base/images-abc/ 下 → 0.0（basename 拼接把目录丢
  了）——三条解析路径同屏
- **多类型 silent 混账**：{para 2, heading 5} 对实际
  {para 3, heading 1} → silent **4**（under-demand 每
  类型原谅、over 逐类型相加，非笼统差）；双过索
  {9, 9} → 14（6+8）
- forbidden tokens 第五百五十七批（open 0）
"""

from __future__ import annotations

import copy
import inspect

from docx import Document

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from evaluation.metrics import compute_automatic_metrics


def _heading_doc(tmp_path):
    p = tmp_path / "h.docx"
    d = Document()
    d.add_paragraph("AAA intro paragraph before heading.")
    d.add_heading("Late Title", level=1)
    d.add_paragraph("BBB body after heading one.")
    d.add_paragraph("CCC body after heading two.")
    d.save(str(p))
    doc, errors = process_single(
        p, tmp_path / "s.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    return doc.to_dict()


def _img_doc(tmp_path):
    d2 = copy.deepcopy(_heading_doc(tmp_path))
    d2["elements"].append({
        "element_id": "e-img", "type": "image",
        "parent_id": None,
        "source_locator": {"paragraph_index": 9,
                           "relationship_id": "r1",
                           "target_partname": "/x.png"},
        "content": None,
        "resource_path": "images-abc/img_00.png",
        "confidence": 1.0, "metadata": {}})
    return d2


def _silent(dd, exp):
    return compute_automatic_metrics(
        dd, None, "docx",
        {"element_count_by_type": exp})["silent_drop_count"]


# ---------- 多类型 silent 混账 ----------

def test_multi_type_mixed_silent_batch285(tmp_path):
    dd = _heading_doc(tmp_path)
    assert _silent(dd, {"paragraph": 2, "heading": 5}) == {
        "value": 4, "reason": None}
    assert _silent(dd, {"paragraph": 9, "heading": 9}) == {
        "value": 14, "reason": None}


# ---------- 候选二：base / basename 救场 ----------

def test_basename_join_hit_batch285(tmp_path):
    d2 = _img_doc(tmp_path)
    (tmp_path / "img_00.png").write_bytes(b"\x89PNG fake")
    m = compute_automatic_metrics(d2, None, "docx", None,
                                  image_base_dir=tmp_path)
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- basename 拼接丢目录 ----------

def test_subdir_not_resolved_batch285(tmp_path):
    d2 = _img_doc(tmp_path)
    sub = tmp_path / "images-abc"
    sub.mkdir()
    (sub / "img_00.png").write_bytes(b"\x89PNG fake")
    m = compute_automatic_metrics(d2, None, "docx", None,
                                  image_base_dir=tmp_path)
    assert m["image_resource_exists_ratio"] == {
        "value": 0.0, "reason": None}


# ---------- 候选一：Path(rp) 原样 CWD 相对 ----------

def test_cwd_verbatim_candidate_batch285(tmp_path, monkeypatch):
    d2 = _img_doc(tmp_path)
    sub = tmp_path / "images-abc"
    sub.mkdir()
    (sub / "img_00.png").write_bytes(b"\x89PNG fake")
    monkeypatch.chdir(tmp_path)
    m = compute_automatic_metrics(d2, None, "docx", None)
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch285():
    src = _src()
    assert ("candidates.append(image_base_dir / "
            "Path(rp).name)") in src
    assert "for p in candidates:" in src


# ---------- forbidden tokens 第五百五十七批 ----------

def test_source_no_eval_batch285():
    assert "eval(" not in _src()


def test_source_no_exec_batch285():
    assert "exec(" not in _src()


def test_source_no_compile_batch285():
    assert "compile(" not in _src()


def test_source_no_globals_batch285():
    assert "globals(" not in _src()


def test_source_no_locals_batch285():
    assert "locals(" not in _src()


def test_source_no_os_system_batch285():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch285():
    assert "subprocess" not in _src()


def test_source_no_popen_batch285():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch285():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch285():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch285():
    assert "socket" not in _src()


def test_source_no_requests_batch285():
    assert "requests" not in _src()


def test_source_no_urllib_batch285():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch285():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch285():
    assert "yield" not in _src()


def test_source_no_async_await_batch285():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch285():
    assert "open(" not in _src()
