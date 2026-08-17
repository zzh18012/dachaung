"""evaluation/annotation_metrics.py 第五百轮 edges 测试（Round 1056）。

补强 edges131 未触及的角度（第四百三十二批，probe 实证）。

新角度（markdown 表格锚 + 真实文本精确膝关节）：
- 真实三类型 docx 的 table 被 parser 渲染成
  markdown（"| cell one |\\n| --- |"）成独立 chunk；
  标注 marker "| cell" before 直接命中 markdown
  文本 → P/R/F1 全 1.0——锚对**渲染后形态**（非
  原始 cell 文本）工作，markdown 锚首次锁定
- 真实文本精确膝关节：2 chunk 板 marker "AAA
  first" after → gt 9、pred 25、距离恰 16——
  tolerance 15 → 全 0.0、16 → 全 1.0；真实文本
  的 d=|len(chunk)-len(marker)| 代数此前从未验算
- 完整 marker（== chunk 文本）d=0 任意容差全 1.0
  （对照：距离为零时膝关节消失）
- forbidden tokens 第五百二十七批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

import evaluation.annotation_metrics as am_mod
from evaluation.annotation_metrics import chunk_boundary_prf


def _doc(tmp_path, builder, mc):
    p = tmp_path / "a.docx"
    builder().save(str(p))
    from app.pipeline import process_single
    doc, errors = process_single(p, tmp_path / "s.json",
                                 parser_name="fallback",
                                 max_chars=mc,
                                 write_json=False)
    assert errors == []
    return doc.to_dict()


def _rich():
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA first paragraph body.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "cell one"
    return d


def _two():
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("CCC third paragraph body.")
    return d


# ---------- markdown 表格锚 ----------

def test_markdown_table_anchor_batch254(tmp_path):
    dd = _doc(tmp_path, _rich, 200)
    assert dd["chunks"][1]["text"].startswith("| cell one |")
    out = chunk_boundary_prf(dd, {"chunk_boundary_anchors": [
        {"marker": "| cell", "position": "before"}]})
    assert out["chunk_boundary_precision"] == {"value": 1.0,
                                               "reason": None}
    assert out["chunk_boundary_recall"] == {"value": 1.0,
                                            "reason": None}
    assert out["chunk_boundary_f1"] == {"value": 1.0,
                                        "reason": None}


# ---------- 真实文本精确膝关节 ----------

def test_real_knee_at_16_batch254(tmp_path):
    dd = _doc(tmp_path, _two, 40)
    below = chunk_boundary_prf(dd, {
        "chunk_boundary_anchors": [
            {"marker": "AAA first",
             "position": "after"}]}, tolerance_chars=15)
    at = chunk_boundary_prf(dd, {
        "chunk_boundary_anchors": [
            {"marker": "AAA first",
             "position": "after"}]}, tolerance_chars=16)
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall", "chunk_boundary_f1"):
        assert below[k] == {"value": 0.0, "reason": None}
        assert at[k] == {"value": 1.0, "reason": None}


def test_knee_algebra_batch254(tmp_path):
    dd = _doc(tmp_path, _two, 40)
    chunk_len = len(dd["chunks"][0]["text"])
    marker_len = len("AAA first")
    assert chunk_len - marker_len == 16


# ---------- 完整 marker 距离零 ----------

def test_full_marker_zero_distance_batch254(tmp_path):
    dd = _doc(tmp_path, _two, 40)
    for tol in (0, 1):
        out = chunk_boundary_prf(dd, {
            "chunk_boundary_anchors": [
                {"marker": "AAA first paragraph body.",
                 "position": "after"}]}, tolerance_chars=tol)
        assert out["chunk_boundary_f1"] == {"value": 1.0,
                                            "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch254():
    src = _src()
    assert "if d <= tolerance_chars:" in src
    assert "gt_positions.append(find_pos + len(marker))" in src


# ---------- forbidden tokens 第五百二十七批 ----------

def test_source_no_eval_batch254():
    assert "eval(" not in _src()


def test_source_no_exec_batch254():
    assert "exec(" not in _src()


def test_source_no_compile_batch254():
    assert "compile(" not in _src()


def test_source_no_globals_batch254():
    assert "globals(" not in _src()


def test_source_no_locals_batch254():
    assert "locals(" not in _src()


def test_source_no_os_system_batch254():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch254():
    assert "subprocess" not in _src()


def test_source_no_popen_batch254():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch254():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch254():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch254():
    assert "socket" not in _src()


def test_source_no_requests_batch254():
    assert "requests" not in _src()


def test_source_no_urllib_batch254():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch254():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch254():
    assert "yield" not in _src()


def test_source_no_async_await_batch254():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch254():
    assert "open(" not in _src()
