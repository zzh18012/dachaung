"""evaluation/runner.py 第六百一十四轮 edges 测试（Round 1170）。

补强 edges183 未触及的角度（第五百四十二批，probe 实证）。

新角度（跨源五型 devset 聚合）：
- **PDF+DOCX 五型同 devset**——两文档各五型 →
  element_count_total {sum: 10, participating_docs:
  2}；pipeline_success {2, 2, 1.0}（跨源聚合首锁）
- **定位率按源分流**——pdf_locator_valid_ratio
  macro {1.0, 1 参与, 1 未评}；docx_locator_
  valid_ratio 镜像对称——PDF 文档 docx_loc 为
  null、DOCX 文档 pdf_loc 为 null（分母只算
  同源文档）
- **devset 元数据**——file_count 2、groups 2、
  pdf 1、docx 1、categories_covered []（无
  categories 字段时空列表）
- **chunk_boundary 全未评**——无标注双文档 →
  {null, 0, 2}；silent_drop_total null
- forbidden tokens 第六百四十二批（open 2）
"""

from __future__ import annotations

import inspect
import io
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _build_pdf(objects, n_obj) -> bytes:
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objects[num] + b"endobj\n")
    xref_pos = len(out)
    out += b"xref\n0 " + str(n_obj).encode() + b"\n"
    out += b"0000000000 65535 f \n"
    for num in range(1, n_obj):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size " + str(n_obj).encode()
            + b"/Root 1 0 R>>\nstartxref\n"
            + str(xref_pos).encode() + b"\n%%EOF\n")
    return bytes(out)


def _five_pdf() -> bytes:
    s = (b"1 w 0 G\n"
         b"10 180 100 50 re S\n60 180 0 50 re S\n"
         b"10 230 100 0 re S\n"
         b"q 40 0 0 40 200 300 cm /Im0 Do Q\n"
         b"BT /F1 10 Tf 15 205 Td (Ga) Tj ET\n"
         b"BT /F1 10 Tf 65 205 Td (Gb) Tj ET\n"
         b"BT /F1 12 Tf 10 390 Td "
         b"(Figure 3: pdf caption text.) Tj ET\n"
         b"BT /F1 12 Tf 10 330 Td "
         b"(Regular paragraph with a period.) Tj ET")
    return _build_pdf({
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 400]"
            b"/Resources<</XObject<</Im0 6 0 R>>"
            b"/Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        6: (b"<</Type/XObject/Subtype/Image/Width 1/Height 1"
            b"/ColorSpace/DeviceRGB/BitsPerComponent 8"
            b"/Length 3>>stream\n" + b"\xff\x00\x00"
            + b"\nendstream "),
    }, 7)


PNG = (b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01'
       b'\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89'
       b'\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01'
       b'\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82')


def _board(tmp_path):
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "s" / "a.pdf").write_bytes(_five_pdf())
    d = Document()
    d.add_heading("Doc Heading L1", level=1)
    p = d.add_paragraph("Body text before the image. ")
    p.add_run().add_picture(io.BytesIO(PNG))
    d.add_paragraph("Figure 2: docx caption text below.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "X1"
    t.cell(0, 1).text = "Y1"
    d.save(str(tmp_path / "s" / "b.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "pf", "path": "s/a.pdf",
             "source_type": "pdf"},
            {"doc_id": "df", "path": "s/b.docx",
             "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- devset 元数据 ----------

def test_cross_devset_meta_batch368(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    assert r["devset"] == {
        "status": "incomplete", "file_count": 2,
        "content_group_count": 2, "pdf_count": 1,
        "docx_count": 1, "categories_covered": []}


# ---------- 聚合计数 ----------

def test_cross_devset_counts_batch368(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    s = r["summary"]
    assert s["counts"]["element_count_total"] == {
        "sum": 10, "participating_docs": 2}
    assert s["success_rates"]["pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}


# ---------- 定位率按源分流 ----------

def test_cross_devset_source_split_batch368(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    ra = r["summary"]["ratio_macro_averages"]
    assert ra["pdf_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}
    assert ra["docx_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}
    assert ra["image_resource_exists_ratio"] == {
        "macro_average": 1.0, "participating_docs": 2,
        "not_evaluated": 0}


# ---------- 未评通道 ----------

def test_cross_devset_not_evaluated_batch368(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    ra = r["summary"]["ratio_macro_averages"]
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall", "chunk_boundary_f1"):
        assert ra[k] == {"macro_average": None,
                         "participating_docs": 0,
                         "not_evaluated": 2}
    assert r["summary"]["silent_drop_total"] is None


# ---------- per_doc 定位分流 ----------

def test_cross_devset_per_doc_split_batch368(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    by_id = {pd["doc_id"]: pd["metrics"]
             for pd in r["per_doc"]}
    assert by_id["pf"]["pdf_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}
    assert by_id["pf"]["docx_locator_valid_ratio"] == {
        "value": None, "reason": "not_docx_document"}
    assert by_id["df"]["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}
    assert by_id["df"]["pdf_locator_valid_ratio"] == {
        "value": None, "reason": "not_pdf_document"}
    assert all(m["pipeline_success"] == {
        "value": True, "reason": None}
        for m in by_id.values())


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch368():
    src = _src()
    assert src.count("manifest") == 5
    assert src.count("per_doc") == 12
    assert src.count("error_code") == 4


# ---------- forbidden tokens 第六百四十二批 ----------

def test_source_no_eval_batch368():
    assert "eval(" not in _src()


def test_source_no_exec_batch368():
    assert "exec(" not in _src()


def test_source_no_compile_batch368():
    assert "compile(" not in _src()


def test_source_no_globals_batch368():
    assert "globals(" not in _src()


def test_source_no_locals_batch368():
    assert "locals(" not in _src()


def test_source_no_os_system_batch368():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch368():
    assert "subprocess" not in _src()


def test_source_no_popen_batch368():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch368():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch368():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch368():
    assert "socket" not in _src()


def test_source_no_requests_batch368():
    assert "requests" not in _src()


def test_source_no_urllib_batch368():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch368():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch368():
    assert "yield" not in _src()


def test_source_no_async_await_batch368():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch368():
    assert _src().count("open(") == 2
