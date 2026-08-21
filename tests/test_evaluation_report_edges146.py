"""evaluation/report.py 第五百七十七轮 edges 测试（Round 1318）。

补强 edges145 未触及的角度（第六百九十批，probe 实证）。

新角度（跨型异构宏 / locator 互补劈叉）：
- **跨型 cbp 宏**——
  pdf 1/15 + docx
  1/10 → (1/15+1/10)/2
  = 1/12（跨文档类型
  宏平均首锁）
- **cbf 跨型**——
  (0.125 + 2/11)/2 =
  0.1534090909…（pdf
  f1 2/16 + docx f1
  2/11 跨型首锁）
- **cbr 双满**——
  {1.0, 2 参与, 0}
- **locator 互补劈叉**
  ——dlvr {1.0, 1 参与,
  1 未评}（pdf 侧
  not_docx）同时 plvr
  {1.0, 1 参与, 1 未评}
  （docx 侧 not_pdf）
  ——双键各半参与对
  称面首锁
- **全绿双参与**——
  tpe/hbc {1.0, 2, 0}
  （docx heading 亦
  参与 hbc）
- **counts 跨型同和**
  ——{sum 4, 2}
- forbidden tokens 第七百六十四批（open 0）
"""

from __future__ import annotations

import inspect
import json

import evaluation.report as report_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
STREAM = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
          % ("A" * 80)
          + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
          % LONG).encode()


def _run(tmp_path):
    (tmp_path / "g.pdf").write_bytes(_wrap(STREAM))
    d = Document()
    d.add_heading("HeadingTitle", level=1)
    d.add_paragraph(" ".join("Sent%d." % i
                             for i in range(40)))
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "ann").mkdir(exist_ok=True)
    (tmp_path / "ann" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "g1",
        "chunk_boundary_anchors": [
            {"marker": "Word3.", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "ann" / "b.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "Sent3.", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "g1", "path": "g.pdf",
             "source_type": "pdf",
             "annotation_file": "ann/a.json"},
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx",
             "annotation_file": "ann/b.json"}]}),
        encoding="utf-8")
    mf = load_manifest((tmp_path / "m.json"),
                       project_root=tmp_path)
    return run_evaluation(mf, tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=32)


# ---------- 跨型 cbp 宏 ----------

def test_cbp_macro_cross_type_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_precision"] == {
        "macro_average": (1 / 15 + 1 / 10) / 2,
        "participating_docs": 2,
        "not_evaluated": 0}


def test_cbp_per_doc_values_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "chunk_boundary_precision"]["value"] == 1 / 15
    assert r["per_doc"][1]["metrics"][
        "chunk_boundary_precision"]["value"] == 0.1


# ---------- cbf / cbr ----------

def test_cbf_macro_cross_type_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"] == {
        "macro_average": (0.125 + 2 / 11) / 2,
        "participating_docs": 2,
        "not_evaluated": 0}


def test_cbr_macro_full_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_recall"] == {
        "macro_average": 1.0, "participating_docs": 2,
        "not_evaluated": 0}


# ---------- locator 互补劈叉 ----------

def test_dlvr_half_participation_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "docx_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


def test_plvr_half_participation_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "pdf_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


def test_dlvr_pdf_not_evaluated_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "docx_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_docx_document"}


def test_plvr_docx_not_evaluated_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "pdf_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_pdf_document"}


# ---------- 全绿双参与 ----------

def test_tpe_two_participating_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "text_preservation_equal"] == {
        "macro_average": 1.0, "participating_docs": 2,
        "not_evaluated": 0}


def test_hbc_two_participating_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "heading_boundary_compliance"] == {
        "macro_average": 1.0, "participating_docs": 2,
        "not_evaluated": 0}


# ---------- counts / success ----------

def test_counts_cross_type_sum_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"][
        "element_count_total"] == {
        "sum": 4, "participating_docs": 2}


def test_success_full_batch516(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}


def test_ratio_keys_twelve_batch516(tmp_path):
    r = _run(tmp_path)
    assert len(r["summary"][
        "ratio_macro_averages"]) == 12


# ---------- 报告合法性 ----------

def test_report_schema_batch516(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch516():
    src = _src()
    assert "def aggregate_summary(" in src
    assert "_RATIO_METRICS" in src


# ---------- forbidden tokens 第七百六十四批 ----------

def test_source_no_eval_batch516():
    assert "eval(" not in _src()


def test_source_no_exec_batch516():
    assert "exec(" not in _src()


def test_source_no_compile_batch516():
    assert "compile(" not in _src()


def test_source_no_globals_batch516():
    assert "globals(" not in _src()


def test_source_no_locals_batch516():
    assert "locals(" not in _src()


def test_source_no_os_system_batch516():
    assert "os.system" not in _src()


def test_source_no_subprocess_call_batch516():
    assert ".call(" not in _src()


def test_source_no_popen_batch516():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch516():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch516():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch516():
    assert "socket" not in _src()


def test_source_no_requests_batch516():
    assert "requests" not in _src()


def test_source_no_urllib_batch516():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch516():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch516():
    assert "yield" not in _src()


def test_source_no_async_await_batch516():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch516():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch516():
    assert _src().count("subprocess.run") == 2
