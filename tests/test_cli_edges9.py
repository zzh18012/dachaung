r"""app/cli.py 边角测试 - 第九轮（Round 1386）。

全新角度：真实文件穿 CLI 子进程全纵向（probe 实证，历史
CLI 测试只用 dummy 字节或 markdown，从未有合法 xref 的
手工 PDF / python-docx 真文件完成 parse→validate→inspect）：
- parse 真实 PDF：rc 0、[OK] 行、INFO 自动选 fallback
- parse 真实 DOCX：rc 0、heading+paragraph 落盘
- JSON 顶层键全集（14 键）、source_type 按扩展名
- validate 产出的 JSON：rc 0、'通过 Schema 校验'
- inspect：schema 0.1.0、document_id doc-<16hex>
- chunk metadata 三键、source_element_ids 非空
"""

from __future__ import annotations

import json
import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document


VENV_PYTHON = str(
    Path(__file__).resolve().parent.parent
    / ".venv" / "Scripts" / "python.exe")
PROJECT_ROOT = \
    Path(__file__).resolve().parent.parent
_PYTHON = (VENV_PYTHON
           if Path(VENV_PYTHON).is_file()
           else sys.executable)


def _run_cli(args, cwd=None):
    env = {
        **os.environ,
        "PYTHONIOENCODING": "utf-8",
        "PYTHONPATH":
            str(PROJECT_ROOT)
            + os.pathsep
            + os.environ.get("PYTHONPATH", ""),
    }
    proc = subprocess.run(
        [_PYTHON, "-X", "utf8", "-m", "app.cli",
         *args],
        capture_output=True, text=True,
        encoding="utf-8", errors="replace",
        cwd=str(cwd) if cwd else None,
        env=env)
    return (proc.returncode, proc.stdout,
            proc.stderr)


def _build_pdf(pages_lines):
    n_pages = len(pages_lines)
    page_ids = [3 + i * 2
                for i in range(n_pages)]
    content_ids = [4 + i * 2
                   for i in range(n_pages)]
    font_id = 3 + n_pages * 2
    objs = {font_id: b"<< /Type /Font "
                    b"/Subtype /Type1 "
                    b"/BaseFont /Helvetica >>"}
    objs[1] = (b"<< /Type /Catalog "
               b"/Pages 2 0 R >>")
    kids = " ".join(f"{pid} 0 R"
                    for pid in page_ids)
    objs[2] = (f"<< /Type /Pages /Kids "
               f"[{kids}] /Count "
               f"{n_pages} >>").encode()
    for i, lines in enumerate(pages_lines):
        pid = page_ids[i]
        cid = content_ids[i]
        objs[pid] = (
            f"<< /Type /Page /Parent 2 0 R "
            f"/MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 "
            f"{font_id} 0 R >> >> /Contents "
            f"{cid} 0 R >>").encode()
        blocks = []
        for (y, line) in lines:
            esc = line.replace(
                "\\", r"\\").replace(
                "(", r"\(").replace(
                ")", r"\)")
            blocks.append(
                f"BT /F1 12 Tf 72 {y} Td "
                f"({esc}) Tj ET")
        stream = " ".join(blocks).encode()
        objs[cid] = (
            b"<< /Length "
            + str(len(stream)).encode()
            + b" >>\nstream\n" + stream
            + b"\nendstream")
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n".encode()
                + objs[oid] + b"\nendobj\n")
    xref_pos = len(out)
    maxid = max(objs)
    out += (f"xref\n0 {maxid + 1}\n"
            .encode()
            + b"0000000000 65535 f \n")
    for oid in range(1, maxid + 1):
        out += (
            f"{offsets[oid]:010d} 00000 n \n"
            .encode() if oid in offsets
            else b"0000000000 65535 f \n")
    out += (f"trailer\n<< /Size "
            f"{maxid + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF"
            ).encode()
    return bytes(out)


def _make_pdf(tmp_path):
    p = tmp_path / "real.pdf"
    p.write_bytes(_build_pdf([[
        (700, "CLI Real Heading"),
        (640, "Body text parsed by "
              "real pdfplumber through "
              "the CLI subprocess.")]]))
    return p


def _make_docx(tmp_path):
    p = tmp_path / "real.docx"
    d = Document()
    d.add_heading("Docx CLI Title", 1)
    d.add_paragraph(
        "A simple docx paragraph.")
    d.save(str(p))
    return p


# ---------- parse 真实 PDF ----------

def test_parse_real_pdf_rc0(tmp_path):
    out = tmp_path / "o.json"
    rc, so, se = _run_cli(
        ["parse", str(_make_pdf(tmp_path)),
         "-o", str(out)])
    assert rc == 0
    assert so.startswith("[OK]")


def test_parse_real_pdf_autoselect_info(
        tmp_path):
    out = tmp_path / "o.json"
    _, _, se = _run_cli(
        ["parse", str(_make_pdf(tmp_path)),
         "-o", str(out)])
    assert "INFO" in se
    assert "fallback" in se


def test_parse_real_pdf_json_created(
        tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    assert out.exists()
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert data["errors"] == []


def test_pdf_json_top_keys(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert sorted(data.keys()) == [
        "chunks", "document_id",
        "elements", "errors",
        "metadata", "parser_name",
        "parser_version", "relations",
        "schema_version", "source_hash",
        "source_path", "source_type",
        "warnings"]


def test_pdf_json_elements(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert [(e["type"], e["content"])
            for e in data["elements"]] == [
        ("heading", "CLI Real Heading"),
        ("paragraph",
         "Body text parsed by real "
         "pdfplumber through the CLI "
         "subprocess.")]


def test_pdf_json_source_type(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert data["source_type"] == "pdf"
    assert data["parser_name"] == \
        "fallback"


def test_pdf_json_document_id_pattern(
        tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert re.fullmatch(
        r"doc-[0-9a-f]{16}",
        data["document_id"])


def test_pdf_json_chunk_shape(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert len(data["chunks"]) == 1
    ch = data["chunks"][0]
    assert sorted(
        ch["metadata"].keys()) == [
        "char_count", "max_chars",
        "strategy"]
    assert ch["source_element_ids"]


# ---------- parse 真实 DOCX ----------

def test_parse_real_docx_rc0(tmp_path):
    out = tmp_path / "o.json"
    rc, so, _ = _run_cli(
        ["parse", str(_make_docx(tmp_path)),
         "-o", str(out)])
    assert rc == 0
    assert so.startswith("[OK]")


def test_docx_json_elements(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_docx(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert [(e["type"], e["content"])
            for e in data["elements"]] == [
        ("heading", "Docx CLI Title"),
        ("paragraph",
         "A simple docx paragraph.")]


def test_docx_json_source_type(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_docx(tmp_path)),
              "-o", str(out)])
    data = json.loads(
        out.read_text(encoding="utf-8"))
    assert data["source_type"] == "docx"


# ---------- validate 产出的 JSON ----------

def test_validate_pdf_output(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    rc, so, _ = _run_cli(
        ["validate", str(out)])
    assert rc == 0
    assert "通过 Schema 校验" in so


def test_validate_docx_output(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_docx(tmp_path)),
              "-o", str(out)])
    rc, so, _ = _run_cli(
        ["validate", str(out)])
    assert rc == 0
    assert "通过 Schema 校验" in so


# ---------- inspect 产出的 JSON ----------

def test_inspect_pdf_output(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    rc, so, _ = _run_cli(
        ["inspect", str(out)])
    assert rc == 0
    assert "schema:      0.1.0" in so
    assert "document_id: doc-" in so


def test_inspect_pdf_chunks(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    rc, so, _ = _run_cli(
        ["inspect", str(out),
         "--chunks", "--limit", "1"])
    assert rc == 0
    assert "::c0000" in so
    assert "chars=80 refs=2" in so


def test_inspect_pdf_elements(tmp_path):
    out = tmp_path / "o.json"
    _run_cli(["parse",
              str(_make_pdf(tmp_path)),
              "-o", str(out)])
    rc, so, _ = _run_cli(
        ["inspect", str(out),
         "--elements"])
    assert rc == 0
    assert "CLI Real Heading" in so
