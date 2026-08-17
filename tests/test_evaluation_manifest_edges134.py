"""evaluation/manifest.py 第五百二十四轮 edges 测试（Round 1080）。

补强 edges131-133 未触及的角度（第四百五十六批，probe 实证）。

新角度（heading/未产出类型期望 + 清单退化形态）：
- expectations 带 **heading 类型**首次走真实 run：
  {heading 1, paragraph 3} 精确相抵 → silent {0}、
  total 0；{heading 2} 过索 → silent 1
- **未产出类型的期望全额跌落**：{list_item 2} →
  silent 2、{table 1, header 1} → silent 2——枚举里
  合法但 fallback parser 从不产出的类型，期望几份
  就几份全 silent
- **element_count_by_type 内层键自由**：塞
  bogus_type: 1 照过 load_manifest → silent 1——与
  edges84 的 expectations 层闭仓（外层未知键即拒）
  构成两层纪律对照：外层闭、内层开
- **documents: [] 退化 run**：per_doc []、success
  {success_count 0, total 0, rate null}（分母 0 出
  null 不出 1.0）、counts {sum null, participating 0}、
  silent_drop_total None——全空清单照常出报告
- **doc_id 不查重**：两条同 doc_id 同 path 条目照收，
  per_doc 出两条同名记录
- forbidden tokens 第五百五十一批（open 1）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path, docs, name="m.json"):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA intro paragraph before the heading.")
    d.add_heading("Late Title", level=1)
    d.add_paragraph("BBB body after the heading one.")
    d.add_paragraph("CCC body after the heading two.")
    d.save(str(tmp_path / "samples" / "h.docx"))
    (tmp_path / name).write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": docs,
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(tmp_path / name,
                                        tmp_path),
                          tmp_path / "o.json",
                          max_chars=200)


def _exp_doc(doc_id, ecbt):
    e = {"doc_id": doc_id, "path": "samples/h.docx",
         "source_type": "docx"}
    if ecbt is not None:
        e["expectations"] = {"element_count_by_type": ecbt}
    return e


# ---------- heading 期望：精确相抵 ----------

def test_heading_expectations_exact_batch279(tmp_path):
    rep = _run(tmp_path, [_exp_doc("d1", {"heading": 1,
                                          "paragraph": 3})])
    assert [p["metrics"]["silent_drop_count"]
            for p in rep["per_doc"]] == [
        {"value": 0, "reason": None}]
    assert rep["summary"]["silent_drop_total"] == 0


# ---------- heading 期望：过索 1 ----------

def test_heading_expectations_over_batch279(tmp_path):
    rep = _run(tmp_path, [_exp_doc("d1", {"heading": 2})])
    assert [p["metrics"]["silent_drop_count"]
            for p in rep["per_doc"]] == [
        {"value": 1, "reason": None}]


# ---------- 未产出类型全额跌落 ----------

def test_never_produced_types_drop_batch279(tmp_path):
    rep = _run(tmp_path, [_exp_doc("d1", {"list_item": 2})])
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 2, "reason": None}
    rep2 = _run(tmp_path, [_exp_doc("d1", {"table": 1,
                                           "header": 1})])
    assert rep2["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 2, "reason": None}


# ---------- 内层键自由（对照 edges84 外层闭仓） ----------

def test_interior_keys_freeform_batch279(tmp_path):
    rep = _run(tmp_path, [_exp_doc("d1",
                                   {"bogus_type": 1})])
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 1, "reason": None}


# ---------- documents: [] 退化 run ----------

def test_empty_documents_degenerate_batch279(tmp_path):
    rep = _run(tmp_path, [])
    assert rep["per_doc"] == []
    assert rep["summary"]["success_rates"] == {
        "pipeline_success": {"success_count": 0,
                             "total": 0, "rate": None}}
    assert rep["summary"]["counts"] == {
        "element_count_total": {"sum": None,
                                "participating_docs": 0}}
    assert rep["summary"]["silent_drop_total"] is None


# ---------- doc_id 不查重 ----------

def test_duplicate_doc_id_accepted_batch279(tmp_path):
    rep = _run(tmp_path, [
        {"doc_id": "dup", "path": "samples/h.docx",
         "source_type": "docx"},
        {"doc_id": "dup", "path": "samples/h.docx",
         "source_type": "docx"}])
    assert len(rep["per_doc"]) == 2
    assert [p["doc_id"] for p in rep["per_doc"]] == [
        "dup", "dup"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch279():
    src = _src()
    assert "def _resolve_relative_path(" in src
    assert 'raise ManifestError(f"{field_name} 为空")' \
        in src


# ---------- forbidden tokens 第五百五十一批 ----------

def test_source_no_eval_batch279():
    assert "eval(" not in _src()


def test_source_no_exec_batch279():
    assert "exec(" not in _src()


def test_source_no_compile_batch279():
    assert "compile(" not in _src()


def test_source_no_globals_batch279():
    assert "globals(" not in _src()


def test_source_no_locals_batch279():
    assert "locals(" not in _src()


def test_source_no_os_system_batch279():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch279():
    assert "subprocess" not in _src()


def test_source_no_popen_batch279():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch279():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch279():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch279():
    assert "socket" not in _src()


def test_source_no_requests_batch279():
    assert "requests" not in _src()


def test_source_no_urllib_batch279():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch279():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch279():
    assert "yield" not in _src()


def test_source_no_async_await_batch279():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch279():
    assert _src().count("open(") == 1
