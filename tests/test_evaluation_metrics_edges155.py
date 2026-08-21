"""evaluation/metrics.py 第五百七十六轮 edges 测试（Round 1320）。

补强 edges154 未触及的角度（第六百九十二批，probe 实证）。

新角度（真 DOCX 内嵌图片抽取面）：
- **图片元素抽取**——
  add_picture 1×1 PNG
  → elements [heading,
  paragraph,
  paragraph, image,
  paragraph]，image
  带 resource_path
  指向抽取出的 images-
  <hash>/ PNG（真
  DOCX 图片全链首锁）
- **ecbt 含 image**——
  {heading:1,
  paragraph:3,
  image:1}；ect 5
- **irer docx 1.0**——
  抽取文件实存 →
  {1.0, None}（docx
  图片资源面首锁，
  区别 no_image_
  elements null）
- **irer 文件删除后**
  ——删抽取 PNG →
  {0.0, None}（存在
  性硬核验首锁）
- **sdc image 型**——
  {image:2} → 1
- **chunks=2**——mc32
  合并几何
- **tpe/dlvr 全绿**——
  图片不扰保真面
- forbidden tokens 第七百六十六批（open 0）
"""

from __future__ import annotations

import inspect
import struct
import zlib

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from docx import Document
from evaluation.metrics import \
    compute_automatic_metrics


def _chunk(t: bytes, data: bytes) -> bytes:
    c = struct.pack(">I", len(data)) + t + data
    return c + struct.pack(
        ">I", zlib.crc32(t + data) & 0xffffffff)


def _png() -> bytes:
    return (b"\x89PNG\r\n\x1a\n"
            + _chunk(b"IHDR",
                     struct.pack(">IIBBBBB",
                                 1, 1, 8, 2, 0,
                                 0, 0))
            + _chunk(b"IDAT",
                     zlib.compress(b"\x00\xff\x00\x00"))
            + _chunk(b"IEND", b""))


def _doc(tmp_path):
    (tmp_path / "img.png").write_bytes(_png())
    d = Document()
    d.add_heading("PicDoc", level=1)
    d.add_paragraph("Text before picture.")
    d.add_picture(str(tmp_path / "img.png"))
    d.add_paragraph("Text after picture.")
    d.save(str(tmp_path / "c.docx"))
    doc, errors = process_single(tmp_path / "c.docx",
                                 tmp_path / "o.json",
                                 parser_name="fallback",
                                 max_chars=32)
    assert errors == []
    return doc.to_dict()


def _m(dd, exp=None):
    return compute_automatic_metrics(dd, None, "docx",
                                     exp)


# ---------- 图片元素抽取 ----------

def test_element_types_with_image_batch518(tmp_path):
    dd = _doc(tmp_path)
    assert [e["type"] for e in dd["elements"]] == [
        "heading", "paragraph", "paragraph",
        "image", "paragraph"]


def test_image_resource_path_batch518(tmp_path):
    dd = _doc(tmp_path)
    rp = dd["elements"][3]["resource_path"]
    assert rp is not None
    assert rp.endswith(".png")
    assert "images-" in rp


def test_image_file_exists_batch518(tmp_path):
    dd = _doc(tmp_path)
    from pathlib import Path
    assert Path(dd["elements"][3][
        "resource_path"]).is_file()


# ---------- ecbt / ect ----------

def test_ecbt_with_image_batch518(tmp_path):
    m = _m(_doc(tmp_path))
    assert m["element_count_by_type"]["value"] == {
        "heading": 1, "paragraph": 3, "image": 1}


def test_ect_five_batch518(tmp_path):
    m = _m(_doc(tmp_path))
    assert m["element_count_total"]["value"] == 5


# ---------- irer ----------

def test_irer_one_batch518(tmp_path):
    m = _m(_doc(tmp_path))
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


def test_irer_zero_after_delete_batch518(tmp_path):
    dd = _doc(tmp_path)
    from pathlib import Path
    Path(dd["elements"][3]["resource_path"]).unlink()
    m = _m(dd)
    assert m["image_resource_exists_ratio"] == {
        "value": 0.0, "reason": None}


# ---------- sdc image 型 ----------

def test_sdc_image_two_batch518(tmp_path):
    m = _m(_doc(tmp_path),
           {"element_count_by_type": {"image": 2}})
    assert m["silent_drop_count"] == {"value": 1,
                                      "reason": None}


def test_sdc_image_one_zero_batch518(tmp_path):
    m = _m(_doc(tmp_path),
           {"element_count_by_type": {"image": 1}})
    assert m["silent_drop_count"] == {"value": 0,
                                      "reason": None}


# ---------- chunks / tpe / dlvr ----------

def test_chunks_two_batch518(tmp_path):
    assert len(_doc(tmp_path)["chunks"]) == 2


def test_tpe_true_with_image_batch518(tmp_path):
    m = _m(_doc(tmp_path))
    assert m["text_preservation_equal"]["value"] \
        is True


def test_dlvr_one_with_image_batch518(tmp_path):
    m = _m(_doc(tmp_path))
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


def test_hbc_one_batch518(tmp_path):
    m = _m(_doc(tmp_path))
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch518():
    src = _src()
    assert "st_size" in src
    assert "no_image_elements" in src


# ---------- forbidden tokens 第七百六十六批 ----------

def test_source_no_eval_batch518():
    assert "eval(" not in _src()


def test_source_no_exec_batch518():
    assert "exec(" not in _src()


def test_source_no_compile_batch518():
    assert "compile(" not in _src()


def test_source_no_globals_batch518():
    assert "globals(" not in _src()


def test_source_no_locals_batch518():
    assert "locals(" not in _src()


def test_source_no_os_system_batch518():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch518():
    assert "subprocess" not in _src()


def test_source_no_popen_batch518():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch518():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch518():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch518():
    assert "socket" not in _src()


def test_source_no_requests_batch518():
    assert "requests" not in _src()


def test_source_no_urllib_batch518():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch518():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch518():
    assert "yield" not in _src()


def test_source_no_async_await_batch518():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_0_batch518():
    assert _src().count("open(") == 0
