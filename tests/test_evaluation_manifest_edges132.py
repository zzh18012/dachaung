"""evaluation/manifest.py 第五百一十轮 edges 测试（Round 1066）。

补强 edges129-131 未触及的角度（第四百四十二批，probe 实证）。

新角度（图片通路的清单可追溯性，真实嵌图 run）：
- expectations 带 **image 类型**首次走真实 run：
  {image 1, paragraph 3} 与真实板精确相抵 → per-doc
  silent {0, None}、汇总 silent_drop_total 0；
  过索 image 2 → silent 1（图片也能造 silent drop）
- 真实标注锚定 "(空段落)" 占位符经 run_evaluation：
  boundary P/R/F1 全 1.0（R1063 直调层结论在整装
  run 复现——占位符 gold 在装载→运行全链路成立）
- 图片板本身：ecbt {paragraph 3, image 1} +
  image ratio 1.0
- 无 categories 时 devset.categories_covered 是
  **空列表**（[] 而非 null）——六键全屏真实值
- forbidden tokens 第五百三十七批（open 1）
"""

from __future__ import annotations

import inspect
import json
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _setup(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body with "
                    "enough text to split nicely here.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("BBB third paragraph body with "
                    "enough text to split nicely here.")
    d.save(str(tmp_path / "samples" / "img.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "(空段落)",
             "position": "before"}]}), encoding="utf-8")


def _run(tmp_path, expectations):
    mf = tmp_path / "m.json"
    doc = {"doc_id": "d1", "path": "samples/img.docx",
           "source_type": "docx"}
    if expectations is not None:
        doc["expectations"] = {
            "element_count_by_type": expectations}
        doc["annotation_file"] = "anns/a.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [doc],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json", max_chars=80)


# ---------- image expectations 精确相抵 ----------

def test_image_expectations_exact_batch265(tmp_path):
    _setup(tmp_path)
    rep = _run(tmp_path, {"image": 1, "paragraph": 3})
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 0,
                                 "reason": None}
    assert rep["summary"]["silent_drop_total"] == 0


# ---------- 过索图片 → silent 1 ----------

def test_over_demand_image_silent_batch265(tmp_path):
    _setup(tmp_path)
    rep = _run(tmp_path, {"image": 2})
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 1,
                                 "reason": None}


# ---------- 占位符标注锚经整装 run ----------

def test_placeholder_annotation_gold_run_batch265(
        tmp_path):
    _setup(tmp_path)
    m = _run(tmp_path, {"image": 1})["per_doc"][0][
        "metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


# ---------- 图片板本身 ----------

def test_image_board_run_batch265(tmp_path):
    _setup(tmp_path)
    m = _run(tmp_path, {"image": 1})["per_doc"][0][
        "metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 3, "image": 1},
        "reason": None}
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 无 categories → 空列表 ----------

def test_devset_empty_categories_batch265(tmp_path):
    _setup(tmp_path)
    assert _run(tmp_path, {"image": 1})["devset"] == {
        "status": "incomplete", "file_count": 1,
        "content_group_count": 1, "pdf_count": 0,
        "docx_count": 1, "categories_covered": []}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch265():
    src = _src()
    assert "def categories_covered(self)" in src
    assert "def file_count(self)" in src


# ---------- forbidden tokens 第五百三十七批 ----------

def test_source_no_eval_batch265():
    assert "eval(" not in _src()


def test_source_no_exec_batch265():
    assert "exec(" not in _src()


def test_source_no_compile_batch265():
    assert "compile(" not in _src()


def test_source_no_globals_batch265():
    assert "globals(" not in _src()


def test_source_no_locals_batch265():
    assert "locals(" not in _src()


def test_source_no_os_system_batch265():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch265():
    assert "subprocess" not in _src()


def test_source_no_popen_batch265():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch265():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch265():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch265():
    assert "socket" not in _src()


def test_source_no_requests_batch265():
    assert "requests" not in _src()


def test_source_no_urllib_batch265():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch265():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch265():
    assert "yield" not in _src()


def test_source_no_async_await_batch265():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch265():
    assert _src().count("open(") == 1
