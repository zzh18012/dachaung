"""表格 → Markdown 线性化契约测试（Stage 6 批次 5）。

契约：docs/table-linearization-contract.md（2026-08-30 冻结）。逐条映射：
- §2 canonical 管线：None→"" / CR 规整 / \\n→<br> / |→\\| / strip，
  顺序固定；首行=表头；短行补齐；0 行→""；全空表仍产结构字符串
- §3 接入：三 parser 走共享 linearize_table；md 仅反转义 \\|；docx
  0 行表静默跳过；合并单元格重复语义
- §4 确定性 + md roundtrip 幂等（linearize(split(rendered))==rendered）
- §5 ipynb markdown cell 表格路径断言（裁决⑦）；text 永不产 table；
  ipynb code cell 不产 table
"""

from __future__ import annotations

import json
from pathlib import Path

import docx as pydocx

from app.parsers.html_parser import HtmlParser
from app.parsers.ipynb_parser import IpynbParser
from app.parsers.markdown_parser import MarkdownParser
from app.parsers.table_linearize import linearize_table
from app.parsers.text_parser import TextParser

# ---------- §2 共享纯函数 ----------


def test_pipe_in_cell_escaped():
    assert linearize_table([["a|b", "c"]]) == "| a\\|b | c |\n| --- | --- |"


def test_newline_kinds_become_br():
    assert linearize_table([["a\nb"]]) == "| a<br>b |\n| --- |"
    assert linearize_table([["a\r\nb"]]) == "| a<br>b |\n| --- |"
    assert linearize_table([["a\rb"]]) == "| a<br>b |\n| --- |"


def test_none_cell_becomes_empty():
    assert linearize_table([["a", None]]) == "| a |  |\n| --- | --- |"


def test_short_row_padded():
    assert linearize_table([["a", "b", "c"], ["x"]]) == (
        "| a | b | c |\n| --- | --- | --- |\n| x |  |  |"
    )


def test_single_row_table():
    assert linearize_table([["only"]]) == "| only |\n| --- |"


def test_all_empty_cells_still_renders_structure():
    md = linearize_table([["", ""], ["", ""]])
    assert md == "|  |  |\n| --- | --- |\n|  |  |"


def test_zero_rows_returns_empty_string():
    assert linearize_table([]) == ""


def test_cell_stripped():
    assert linearize_table([["  x  ", " y "]]) == "| x | y |\n| --- | --- |"


def test_no_unicode_normalization():
    # 全角字符原样保留（不做 NFC/NFKC）
    s = "ＡＢ"
    assert linearize_table([[s]]) == f"| {s} |\n| --- |"


def test_deterministic_same_input_same_output():
    rows = [["a|b", None], ["x\ny"]]
    assert linearize_table(rows) == linearize_table(list(rows))


# ---------- §3 md 反转义（仅 \\|，其余反斜杠保留） ----------


def test_md_escaped_pipe_unescape_and_reescape(tmp_path: Path):
    """roundtrip 幂等：content 重新作为 md 源解析再渲染，结果不变。"""
    src = "| a \\| b | c |\n| --- | --- |\n| 1 | 2 |\n"
    doc = MarkdownParser().parse(
        _md_file(tmp_path, src), source_hash="a" * 64)
    e = doc.elements[0]
    assert e.type == "table"
    assert e.content == "| a \\| b | c |\n| --- | --- |\n| 1 | 2 |"
    # 二次 roundtrip：content 本身就是合法 md 表格源
    doc2 = MarkdownParser().parse(
        _md_file(tmp_path, e.content + "\n"), source_hash="a" * 64)
    assert doc2.elements[0].content == e.content


def test_md_other_backslash_sequences_preserved(tmp_path: Path):
    """裁决③：仅反转义 \\|，其余反斜杠序列原样保留。"""
    src = "| a\\nb | c |\n| --- | --- |\n"
    doc = MarkdownParser().parse(
        _md_file(tmp_path, src), source_hash="a" * 64)
    e = doc.elements[0]
    assert e.content == "| a\\nb | c |\n| --- | --- |"


# ---------- §3 三 parser 一致性 ----------


def _md_file(tmp_path: Path, text: str) -> Path:
    p = tmp_path / "t.md"
    p.write_text(text, encoding="utf-8")
    return p


def _html_file(tmp_path: Path, text: str) -> Path:
    p = tmp_path / "t.html"
    p.write_text(text, encoding="utf-8")
    return p


def test_three_parsers_same_rows_same_content(tmp_path: Path):
    md_doc = MarkdownParser().parse(
        _md_file(tmp_path, "| h1 | h2 |\n| --- | --- |\n| a\\|b | c |\n"),
        source_hash="a" * 64)
    html_doc = HtmlParser().parse(
        _html_file(
            tmp_path,
            "<table><tr><td>h1</td><td>h2</td></tr>"
            "<tr><td>a|b</td><td>c</td></tr></table>"),
        source_hash="a" * 64)
    md_tbl = [e for e in md_doc.elements if e.type == "table"]
    html_tbl = [e for e in html_doc.elements if e.type == "table"]
    assert len(md_tbl) == 1 and len(html_tbl) == 1
    assert md_tbl[0].content == html_tbl[0].content == (
        "| h1 | h2 |\n| --- | --- |\n| a\\|b | c |")


def test_fallback_rows_to_markdown_delegates_shared():
    from app.parsers.fallback_parser import _rows_to_markdown

    assert _rows_to_markdown([["a|b"]]) == linearize_table([["a|b"]])


# ---------- §3 docx：多段落单元格 / 合并单元格 / 0 行表 ----------


def _docx_file(tmp_path: Path, build) -> Path:
    d = pydocx.Document()
    build(d)
    p = tmp_path / "t.docx"
    d.save(str(p))
    return p


def test_docx_multiparagraph_cell_becomes_br(tmp_path: Path):
    def build(d):
        t = d.add_table(rows=2, cols=2)
        c = t.cell(0, 0)
        c.text = "p1"
        c.add_paragraph("p2")
        t.cell(0, 1).text = "h2"
        t.cell(1, 0).text = "x"
        t.cell(1, 1).text = "y"

    from app.parsers.fallback_parser import FallbackParser

    doc = FallbackParser().parse(
        _docx_file(tmp_path, build), source_hash="a" * 64)
    tbl = [e for e in doc.elements if e.type == "table"]
    assert len(tbl) == 1
    assert tbl[0].content == (
        "| p1<br>p2 | h2 |\n| --- | --- |\n| x | y |")


def test_docx_merged_cell_repeats_content(tmp_path: Path):
    def build(d):
        t = d.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "m"
        t.cell(0, 0).merge(t.cell(0, 1))
        t.cell(1, 0).text = "a"
        t.cell(1, 1).text = "b"

    from app.parsers.fallback_parser import FallbackParser

    doc = FallbackParser().parse(
        _docx_file(tmp_path, build), source_hash="a" * 64)
    tbl = [e for e in doc.elements if e.type == "table"]
    assert len(tbl) == 1
    # 契约 §3：合并单元格保持库给定重复语义，不去重不标记
    assert tbl[0].content == (
        "| m | m |\n| --- | --- |\n| a | b |")


def test_docx_zero_row_table_skipped_silently(tmp_path: Path):
    def build(d):
        d.add_paragraph("before")
        d.add_table(rows=0, cols=2)
        d.add_paragraph("after")

    from app.parsers.fallback_parser import FallbackParser

    doc = FallbackParser().parse(
        _docx_file(tmp_path, build), source_hash="a" * 64)
    tbl = [e for e in doc.elements if e.type == "table"]
    assert tbl == []
    assert doc.warnings == []  # 静默跳过，不产 warning（裁决④）
    paras = [e.content for e in doc.elements if e.type == "paragraph"]
    assert "before" in paras and "after" in paras


# ---------- §5 裁决⑦：ipynb 表格路径 + text 零表格 ----------


def _ipynb_file(tmp_path: Path, cells: list[dict]) -> Path:
    p = tmp_path / "t.ipynb"
    p.write_text(json.dumps({
        "nbformat": 4, "nbformat_minor": 5, "metadata": {},
        "cells": cells,
    }), encoding="utf-8")
    return p


def test_ipynb_markdown_cell_table_path(tmp_path: Path):
    """ipynb markdown cell 的 pipe 表格经共享线性化产出 canonical content。"""
    nb = _ipynb_file(tmp_path, [
        {"cell_type": "markdown",
         "source": ["| a \\| b | c |\n", "| --- | --- |\n", "| 1 | 2 |\n"]},
        {"cell_type": "code", "source": ["print(1)\n"]},
    ])
    doc = IpynbParser().parse(nb, source_hash="a" * 64)
    tbls = [e for e in doc.elements if e.type == "table"]
    assert len(tbls) == 1
    assert tbls[0].content == "| a \\| b | c |\n| --- | --- |\n| 1 | 2 |"
    assert tbls[0].source_locator["cell_type"] == "markdown"


def test_ipynb_code_cell_never_table(tmp_path: Path):
    nb = _ipynb_file(tmp_path, [
        {"cell_type": "code",
         "source": ["| a | b |", "| --- | --- |"]},
    ])
    doc = IpynbParser().parse(nb, source_hash="a" * 64)
    assert [e for e in doc.elements if e.type == "table"] == []


def test_text_parser_never_table(tmp_path: Path):
    p = tmp_path / "t.txt"
    p.write_text("| a | b |\n| --- | --- |\n", encoding="utf-8")
    doc = TextParser().parse(p, source_hash="a" * 64)
    assert [e for e in doc.elements if e.type == "table"] == []
    assert all(e.type == "paragraph" for e in doc.elements)
