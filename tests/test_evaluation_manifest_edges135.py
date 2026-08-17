"""evaluation/manifest.py 第五百三十一轮 edges 测试（Round 1087）。

补强 edges131-134 未触及的角度（第四百六十三批，probe 实证）。

新角度（ef 条目闭仓名册 + devset_status 枚举双值）：
- **ef 条目没有 annotation_file 通道**：expected_failure
  def 闭仓——塞 annotation_file 即拒 "Additional
  properties are not allowed"；塞 expectations 同拒
  ——标注与期望都是 document 账本专属，ef 账本只收
  [doc_id, path, expected_error_code]
- **expected_error_code minLength**：空串 "" →
  "'' should be non-empty"
- **devset_status 是枚举 ['complete', 'incomplete']**：
  "whatever" 拒（enum 报文）；**"complete" 照收**——
  语料库全程 incomplete，complete 是合法但从未用过
  的另一半；complete 经真实 run 流入报告
  devset.status == "complete"
- forbidden tokens 第五百五十八批（open 1）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import (
    ManifestError, load_manifest)
from evaluation.runner import run_evaluation


def _base(tmp_path, ef, status="incomplete"):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": status,
        "documents": [], "expected_failures": ef}),
        encoding="utf-8")


def _expect_reject(tmp_path, frag):
    try:
        load_manifest(tmp_path / "m.json", tmp_path)
        raised = False
    except (ManifestError, Exception) as e:
        raised = True
        assert frag in str(e)
    assert raised


# ---------- ef 无 annotation 通道 ----------

def test_ef_closed_no_annotation_batch286(tmp_path):
    _base(tmp_path, [{
        "doc_id": "f1", "path": "samples/a.docx",
        "expected_error_code": "docx_open_failed",
        "annotation_file": "anns/a.json"}])
    _expect_reject(tmp_path,
                   "Additional properties are not allowed")


# ---------- ef 无 expectations 通道 ----------

def test_ef_closed_no_expectations_batch286(tmp_path):
    _base(tmp_path, [{
        "doc_id": "f1", "path": "samples/a.docx",
        "expected_error_code": "E",
        "expectations": {"element_count_by_type": {}}}])
    _expect_reject(tmp_path,
                   "Additional properties are not allowed")


# ---------- expected_error_code 空串 ----------

def test_ef_empty_code_rejected_batch286(tmp_path):
    _base(tmp_path, [{
        "doc_id": "f1", "path": "samples/a.docx",
        "expected_error_code": ""}])
    _expect_reject(tmp_path, "'' should be non-empty")


# ---------- devset_status 枚举 ----------

def test_devset_status_enum_batch286(tmp_path):
    _base(tmp_path, [], status="whatever")
    try:
        load_manifest(tmp_path / "m.json", tmp_path)
        raised = False
    except (ManifestError, Exception) as e:
        raised = True
        assert "'whatever' is not one of ['complete', " \
               "'incomplete']" in str(e)
    assert raised


# ---------- complete 照收并流入报告 ----------

def test_complete_status_flows_batch286(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "complete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/a.docx",
                       "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    rep = run_evaluation(load_manifest(tmp_path / "m.json",
                                       tmp_path),
                         tmp_path / "o.json", max_chars=200)
    assert rep["devset"] == {
        "status": "complete", "file_count": 1,
        "content_group_count": 1, "pdf_count": 0,
        "docx_count": 1, "categories_covered": []}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch286():
    src = _src()
    assert 'annotation_file_str=d.get("annotation_file")' \
        in src
    assert "def _resolve_relative_path(" in src


# ---------- forbidden tokens 第五百五十八批 ----------

def test_source_no_eval_batch286():
    assert "eval(" not in _src()


def test_source_no_exec_batch286():
    assert "exec(" not in _src()


def test_source_no_compile_batch286():
    assert "compile(" not in _src()


def test_source_no_globals_batch286():
    assert "globals(" not in _src()


def test_source_no_locals_batch286():
    assert "locals(" not in _src()


def test_source_no_os_system_batch286():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch286():
    assert "subprocess" not in _src()


def test_source_no_popen_batch286():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch286():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch286():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch286():
    assert "socket" not in _src()


def test_source_no_requests_batch286():
    assert "requests" not in _src()


def test_source_no_urllib_batch286():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch286():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch286():
    assert "yield" not in _src()


def test_source_no_async_await_batch286():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch286():
    assert _src().count("open(") == 1
