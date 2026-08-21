"""evaluation/runner.py 第六百七十轮 edges 测试（Round 1315）。

补强 edges234 未触及的角度（第六百八十七批，probe 实证）。

新角度（DOCX 锚定 / 缺文件文档 / 同文档双 ef）：
- **DOCX 锚定 cbp**——
  docx 文档带
  annotation_file →
  cbp 0.1 = 1/10（11
  块 10 边界；marker
  匹配型无关首锁）
  → 宏 {0.1, 1 参与,
  2 未评}
- **缺文件文档**——
  path 指向不存在
  文件 → error_code
  file_not_found +
  success False +
  ect null/pipeline_
  failed（清单内缺文
  件面首锁）
- **同文档双 ef**——
  bad1 两条期望（一
  中一错码）→ 同
  actual，matches
  [True, False]（ef
  按 entry 独立评首锁）
- **docx 期望面**——
  heading:2 → sdc 1；
  sdt 恰 1
- **success 1/3**——
  好坏缺三分账
- forbidden tokens 第七百六十二批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


def _board(tmp_path):
    (tmp_path / "bad.pdf").write_bytes(b"not a pdf")
    d = Document()
    d.add_heading("HeadingTitle", level=1)
    d.add_paragraph(" ".join("Sent%d." % i
                             for i in range(40)))
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "ann").mkdir(exist_ok=True)
    (tmp_path / "ann" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "Sent3.", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx",
             "annotation_file": "ann/a.json",
             "expectations": {
                 "element_count_by_type": {
                     "heading": 2}}},
            {"doc_id": "miss", "path": "gone.pdf",
             "source_type": "pdf"},
            {"doc_id": "bad1", "path": "bad.pdf",
             "source_type": "pdf"}],
        "expected_failures": [
            {"doc_id": "bad1", "path": "bad.pdf",
             "expected_error_code":
             "pdfplumber_open_failed"},
            {"doc_id": "bad1", "path": "bad.pdf",
             "expected_error_code": "other_code"}]}),
        encoding="utf-8")
    return load_manifest((tmp_path / "m.json"),
                         project_root=tmp_path)


def _run(tmp_path):
    return run_evaluation(_board(tmp_path),
                          tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=32)


# ---------- DOCX 锚定 cbp ----------

def test_docx_anchored_cbp_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "chunk_boundary_precision"] == {
        "value": 0.1, "reason": None}


def test_docx_cbp_macro_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_precision"] == {
        "macro_average": 0.1, "participating_docs": 1,
        "not_evaluated": 2}


# ---------- 缺文件文档 ----------

def test_missing_doc_error_code_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "error_code"] == {
        "value": "file_not_found", "reason": None}


def test_missing_doc_success_false_batch513(
        tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "pipeline_success"] == {"value": False,
                                "reason": None}


def test_missing_doc_ect_null_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "element_count_total"] == {
        "value": None, "reason": "pipeline_failed"}


def test_missing_doc_tpe_null_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "text_preservation_equal"] == {
        "value": None, "reason": "pipeline_failed"}


# ---------- 同文档双 ef ----------

def test_double_ef_matches_batch513(tmp_path):
    r = _run(tmp_path)
    assert [e["matches"]
            for e in r["expected_failures"]] == [
        True, False]


def test_double_ef_same_actual_batch513(tmp_path):
    r = _run(tmp_path)
    assert [e["actual_error_code"]
            for e in r["expected_failures"]] == [
        "pdfplumber_open_failed",
        "pdfplumber_open_failed"]


def test_double_ef_codes_batch513(tmp_path):
    r = _run(tmp_path)
    assert [e["expected_error_code"]
            for e in r["expected_failures"]] == [
        "pdfplumber_open_failed", "other_code"]


# ---------- docx 期望面 ----------

def test_docx_sdc_one_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 1,
                                 "reason": None}


def test_sdt_one_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["silent_drop_total"] == 1


# ---------- 成功率 ----------

def test_success_one_third_batch513(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 1, "total": 3,
        "rate": 1 / 3}


# ---------- 报告合法性 ----------

def test_report_schema_batch513(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_counts_batch513():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("per_doc") == 12
    assert src.count("open(") == 2


# ---------- forbidden tokens 第七百六十二批 ----------

def test_source_no_eval_batch513():
    assert "eval(" not in _src()


def test_source_no_exec_batch513():
    assert "exec(" not in _src()


def test_source_no_compile_batch513():
    assert "compile(" not in _src()


def test_source_no_globals_batch513():
    assert "globals(" not in _src()


def test_source_no_locals_batch513():
    assert "locals(" not in _src()


def test_source_no_os_system_batch513():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch513():
    assert "subprocess" not in _src()


def test_source_no_popen_batch513():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch513():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch513():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch513():
    assert "socket" not in _src()


def test_source_no_requests_batch513():
    assert "requests" not in _src()


def test_source_no_urllib_batch513():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch513():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch513():
    assert "yield" not in _src()


def test_source_no_async_await_batch513():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch513():
    assert ".call(" not in _src()
