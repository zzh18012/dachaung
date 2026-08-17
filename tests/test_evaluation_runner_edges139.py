"""evaluation/runner.py 第五百五十四轮 edges 测试（Round 1110）。

补强 edges138 未触及的角度（第四百八十六批，probe 实证）。

新角度（容差沉默 / 双通道 stub 清理 / 容差达聚合层）：
- **容差沉默**：tolerance_chars 0/30 两跑，公开报告
  json.dumps 全文不含 "tolerance" 字样——私有 _tolerance_chars
  被剥离且 provenance 不记录：容差塑造数字但姓名不入档
  （全 blob 首锁；旧锁只查 per_doc 单键）
- **双通道 stub 清理**：documents d1 + expected_failures ef1
  全跑后 _per_doc 目录留存、d1.json 与 ef1.json 两个 stub
  双双 unlink——ef 通道 inline 写 stub 也清理（_process_one
  级清理已锁，全跑双通道首锁）
- **容差达聚合层**：marker "head" before 距预测边界 10 字符
  ——tol 0 全 miss（P/R/F1 全 0.0）、tol 30 命中（P 0.5 /
  R 1.0 / F1 0.6667）；summary macro 同幅翻转而 counts
  逐字相等——容差改 ratio 不改 counts（edges135 锁了
  per-doc 翻转，本批锁 summary 层）
- forbidden tokens 第五百八十二批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("head TAIL")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    (tmp_path / "anns" / "a.json").write_text(
        json.dumps({
            "annotation_version": "1.0", "doc_id": "d1",
            "chunk_boundary_anchors": [
                {"marker": "head", "position": "before"}]}),
        encoding="utf-8")
    mf = tmp_path / "manifest.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/g.docx",
            "source_type": "docx",
            "annotation_file": "anns/a.json"}],
        "expected_failures": [{
            "doc_id": "ef1", "path": "samples/ghost.docx",
            "expected_error_code": "file_not_found"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _run(tmp_path, tol):
    return run_evaluation(
        _board(tmp_path), tmp_path / f"r{tol}.json",
        parser_name="fallback", max_chars=200,
        tolerance_chars=tol)


# ---------- 容差沉默 ----------

def test_tolerance_silent_in_public_report_batch309(tmp_path):
    r0 = _run(tmp_path, 0)
    r30 = _run(tmp_path, 30)
    assert "tolerance" not in json.dumps(r0)
    assert "tolerance" not in json.dumps(r30)


# ---------- 双通道 stub 清理 ----------

def test_full_run_stub_cleanup_both_channels_batch309(tmp_path):
    _run(tmp_path, 30)
    pd = tmp_path / "_per_doc"
    assert pd.is_dir()
    assert not (pd / "d1.json").is_file()
    assert not (pd / "ef1.json").is_file()


# ---------- 容差达聚合层 ----------

def test_tolerance_flip_reaches_aggregate_batch309(tmp_path):
    r0 = _run(tmp_path, 0)
    r30 = _run(tmp_path, 30)
    m0 = r0["per_doc"][0]["metrics"]
    m30 = r30["per_doc"][0]["metrics"]
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall", "chunk_boundary_f1"):
        assert m0[k] == {"value": 0.0, "reason": None}
    assert m30["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m30["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m30["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}
    f1_0 = r0["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"]
    f1_30 = r30["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"]
    assert f1_0 == {"macro_average": 0.0,
                    "participating_docs": 1,
                    "not_evaluated": 0}
    assert f1_30["macro_average"] == 0.6666666666666666
    assert (r0["summary"]["counts"] ==
            r30["summary"]["counts"])


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch309():
    src = _src()
    assert "避免不同 doc" in src
    assert "空目录留作" in src


# ---------- forbidden tokens 第五百八十二批 ----------

def test_source_no_eval_batch309():
    assert "eval(" not in _src()


def test_source_no_exec_batch309():
    assert "exec(" not in _src()


def test_source_no_compile_batch309():
    assert "compile(" not in _src()


def test_source_no_globals_batch309():
    assert "globals(" not in _src()


def test_source_no_locals_batch309():
    assert "locals(" not in _src()


def test_source_no_os_system_batch309():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch309():
    assert "subprocess" not in _src()


def test_source_no_popen_batch309():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch309():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch309():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch309():
    assert "socket" not in _src()


def test_source_no_requests_batch309():
    assert "requests" not in _src()


def test_source_no_urllib_batch309():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch309():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch309():
    assert "yield" not in _src()


def test_source_no_async_await_batch309():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch309():
    assert _src().count("open(") == 2
