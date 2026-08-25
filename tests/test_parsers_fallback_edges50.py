r"""app/parsers/fallback_parser.py 边角测试 - 第五十轮（Round 1443）。

新角度（probe 实证）文档级外围：空文档 / 批注 / 默认值 /
元数据（历史全在"有内容"场景，空体与外围 part 未考察）：
- 空 docx（无任何段落表格）：0 元素 + docx_no_content 告警；
  管线层 no_extracted_elements 结构化错误（details.warnings
  携带原因链、doc=None）
- zip 手术注入 word/comments.xml + commentReference：批注
  文本**完全不可见**（对照 R1418 脚注同款），无告警
- 元素默认值：confidence 0.95、parent_id None；doc.relations
  恒 []（本期不发射关系）
- 文档级 metadata：PDF /Info 与 docx core_properties（title/
  author）**都不进模型**——恒 {'fallback': True,
  'image_output_dir': None}
"""

from __future__ import annotations

import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _empty_docx(tmp_path):
    p = tmp_path / "empty.docx"
    Document().save(str(p))
    return p


def _comment_docx(tmp_path):
    d = Document()
    d.add_paragraph(
        "Body with comment anchor.")
    src = tmp_path / "src.docx"
    d.save(str(src))
    dst = tmp_path / "cmt.docx"
    with zipfile.ZipFile(src) as zin:
        names = zin.namelist()
        doc_xml = zin.read(
            "word/document.xml").decode(
            "utf-8")
        ct = zin.read(
            "[Content_Types].xml").decode(
            "utf-8")
        rels = zin.read(
            "word/_rels/"
            "document.xml.rels").decode(
            "utf-8")
        m = re.search(
            r"(<w:p\b[^>]*>.*?"
            r"<w:t[^>]*>Body with "
            r"comment anchor\."
            r"</w:t>)(.*?)"
            r"(</w:p>)",
            doc_xml, re.S)
        assert m
        doc_xml = (
            doc_xml[:m.end(2)]
            + '<w:r><w:commentReference '
              'w:id="0"/></w:r>'
            + doc_xml[m.end(2):])
        ct = ct.replace(
            "</Types>",
            '<Override PartName='
            '"/word/comments.xml" '
            'ContentType='
            '"application/vnd.'
            'openxmlformats-'
            'officedocument.'
            'wordprocessingml.'
            'comments+xml"/>'
            "</Types>")
        rels = rels.replace(
            "</Relationships>",
            '<Relationship Id='
            '"rIdCmt1" Type='
            '"http://schemas.'
            'openxmlformats.org/'
            'officeDocument/2006/'
            'relationships/'
            'comments" Target='
            '"comments.xml"/>'
            "</Relationships>")
        comments = (
            '<?xml version="1.0" '
            'encoding="UTF-8" '
            'standalone="yes"?>'
            '<w:comments xmlns:w='
            '"http://schemas.'
            'openxmlformats.org/'
            'wordprocessingml/2006/'
            'main">'
            '<w:comment w:id="0" '
            'w:author="A" '
            'w:date="2024-01-01T00:'
            '00:00Z" w:initials="A">'
            '<w:p><w:r><w:t>This is '
            'the comment text.'
            '</w:t></w:r></w:p>'
            '</w:comment>'
            '</w:comments>')
        with zipfile.ZipFile(
                dst, "w",
                zipfile.ZIP_DEFLATED) \
                as zout:
            for n in names:
                if n == "word/document.xml":
                    zout.writestr(
                        n, doc_xml)
                elif n == ("[Content_"
                           "Types].xml"):
                    zout.writestr(n, ct)
                elif n == ("word/_rels/"
                           "document.xml"
                           ".rels"):
                    zout.writestr(
                        n, rels)
                else:
                    zout.writestr(
                        n, zin.read(n))
            zout.writestr(
                "word/comments.xml",
                comments)
    return dst


# ---------- 空 docx ----------

def test_empty_no_elements(
        tmp_path):
    p = _empty_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements == []
    assert [w.code
            for w in doc.warnings] == [
        "docx_no_content"]


def test_empty_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _empty_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_empty_pipeline_error(
        tmp_path):
    p = _empty_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert doc is None
    assert [e.code
            for e in errors] == [
        "no_extracted_elements"]
    det = errors[0].to_dict()[
        "details"]
    assert det["warnings"][0][
        "code"] == "docx_no_content"
    assert det["source_type"] == \
        "docx"


# ---------- 批注 ----------

def test_comment_invisible(
        tmp_path):
    p = _comment_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "Body with comment anchor."]
    assert doc.warnings == []


def test_comment_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _comment_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_comment_chunk(
        tmp_path):
    p = _comment_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[
        0].text == \
        "Body with comment anchor."


# ---------- 默认值 ----------

def test_element_defaults(
        tmp_path):
    d = Document()
    d.add_paragraph("x")
    p = tmp_path / "conf.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    e = doc.elements[0]
    assert e.confidence == 0.95
    assert e.parent_id is None


def test_relations_empty(
        tmp_path):
    d = Document()
    d.add_paragraph("x")
    p = tmp_path / "rel.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.relations == []
    assert doc.to_dict()[
        "relations"] == []


# ---------- 文档 metadata ----------

def test_metadata_constant_docx(
        tmp_path):
    d = Document()
    d.add_paragraph("meta para")
    cp = d.core_properties
    cp.title = "My Title"
    cp.author = "Author Name"
    p = tmp_path / "m.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.metadata == {
        "fallback": True,
        "image_output_dir": None}
    assert doc.to_dict()[
        "metadata"] == {
        "fallback": True,
        "image_output_dir": None}


def test_metadata_constant_pdf(
        tmp_path):
    from app.hash import \
        compute_file_hash as h
    content = (b"BT /F1 12 Tf 72 700"
               b" Td (Meta text) Tj"
               b" ET")
    objs = {
        5: (b"<< /Type /Font /"
            b"Subtype /Type1 /"
            b"BaseFont /"
            b"Helvetica >>"),
        1: (b"<< /Type /Catalog /"
            b"Pages 2 0 R >>"),
        2: (b"<< /Type /Pages /"
            b"Kids [3 0 R] /"
            b"Count 1 >>"),
        3: (b"<< /Type /Page /"
            b"Parent 2 0 R /"
            b"MediaBox [0 0 612 "
            b"792] /Resources <<"
            b" /Font << /F1 5 0 R"
            b" >> >> /Contents "
            b"4 0 R >>"),
        4: (b"<< /Length "
            + str(len(content))
            .encode()
            + b" >>\nstream\n"
            + content
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    out += (b"xref\n0 6\n"
            b"0000000000 65535 f \n")
    for oid in range(1, 6):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 6 "
            b"/Root 1 0 R >>\n"
            b"startxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    p = tmp_path / "m.pdf"
    p.write_bytes(bytes(out))
    doc = FallbackParser().parse(
        p, h(p))
    assert doc.metadata == {
        "fallback": True,
        "image_output_dir": None}
