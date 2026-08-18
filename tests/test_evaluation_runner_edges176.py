"""evaluation/runner.py 第六百零四轮 edges 测试（Round 1160）。

补强 edges175 未触及的角度（第五百三十二批，probe 实证）。

新角度（DOCX 真表格通道）：
- **python-docx add_table 真跑**——DOCX 内表格 →
  table 元素 markdown "| R1C1 | R1C2 |…| R2C1 |
  R2C2 |" + {row_count 2, col_count 2,
  source python-docx}（DOCX 表格通道 runner 级
  首锁，历史 DOCX 板全是段落）
- **文档序不重排**——DOCX 元素按文档序
  [paragraph, table, paragraph]——与 PDF 的
  "文本前置表格殿后"（edges174）成对照
- **DOCX 表格切文本流**——3 chunks [sequential,
  isolated_table, sequential]：表前后段落各成块
  ——修正 edges166 的 PDF 板解读：表格确会
  flush 顺序缓冲，PDF 板的"跨表合块"实为表格
  元素殿后的排序产物
- forbidden tokens 第六百三十二批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Intro paragraph before the table.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "R1C1"
    t.cell(0, 1).text = "R1C2"
    t.cell(1, 0).text = "R2C1"
    t.cell(1, 1).text = "R2C2"
    d.add_paragraph("Outro paragraph after the table.")
    d.save(str(tmp_path / "samples" / f"{doc_id}.docx"))
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"samples/{doc_id}.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- python-docx add_table 真跑 ----------

def test_docx_table_element_batch358(tmp_path):
    _board(tmp_path, "dt")
    doc, errors = process_single(
        tmp_path / "samples" / "dt.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "paragraph", "table", "paragraph"]
    assert els[1]["content"] == (
        "| R1C1 | R1C2 |\n| --- | --- |\n| R2C1 | R2C2 |")
    assert els[1]["metadata"] == {
        "row_count": 2, "col_count": 2,
        "source": "python-docx"}


def test_docx_table_locators_batch358(tmp_path):
    _board(tmp_path, "dt2")
    doc, errors = process_single(
        tmp_path / "samples" / "dt2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    els = doc.to_dict()["elements"]
    assert els[0]["source_locator"] == {
        "paragraph_index": 0, "section": 0}
    assert els[1]["source_locator"] == {
        "table_index": 0, "section": 0}
    assert els[2]["source_locator"] == {
        "paragraph_index": 1, "section": 0}


# ---------- DOCX 表格切文本流 ----------

def test_docx_table_cuts_text_flow_batch358(tmp_path):
    _board(tmp_path, "dt3")
    doc, errors = process_single(
        tmp_path / "samples" / "dt3.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert len(chunks) == 3
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "isolated_table", "sequential"]
    assert chunks[0]["text"] == \
        "Intro paragraph before the table."
    assert chunks[2]["text"] == \
        "Outro paragraph after the table."
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks)


# ---------- 指标 ----------

def test_docx_table_metrics_batch358(tmp_path):
    r = run_evaluation(_board(tmp_path, "dt4"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 2, "table": 1}, "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["chunk_reference_intact_ratio"] == {
        "value": 1.0, "reason": None}
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch358():
    src = _src()
    assert src.count("process_single") == 6
    assert src.count("metrics") == 13
    assert src.count("chunk") == 9


# ---------- forbidden tokens 第六百三十二批 ----------

def test_source_no_eval_batch358():
    assert "eval(" not in _src()


def test_source_no_exec_batch358():
    assert "exec(" not in _src()


def test_source_no_compile_batch358():
    assert "compile(" not in _src()


def test_source_no_globals_batch358():
    assert "globals(" not in _src()


def test_source_no_locals_batch358():
    assert "locals(" not in _src()


def test_source_no_os_system_batch358():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch358():
    assert "subprocess" not in _src()


def test_source_no_popen_batch358():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch358():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch358():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch358():
    assert "socket" not in _src()


def test_source_no_requests_batch358():
    assert "requests" not in _src()


def test_source_no_urllib_batch358():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch358():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch358():
    assert "yield" not in _src()


def test_source_no_async_await_batch358():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch358():
    assert _src().count("open(") == 2
