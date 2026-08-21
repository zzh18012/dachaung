"""evaluation/runner.py 第六百六十九轮 edges 测试（Round 1309）。

补强 edges233 未触及的角度（第六百八十一批，probe 实证）。

新角度（真 DOCX 混合 devset / complete 状态 / 未知 ef 文档）：
- **complete 透传**——
  devset_status complete →
  devset.status "complete"
  （枚举另一支首锁）
- **混合计数**——
  pdf+docx+坏 pdf →
  {file 3, group 2,
  pdf 2, docx 1}（docx_
  count 非零首锁）
- **locator 分型空值**
  ——pdf 文档 dlvr=
  not_docx_document /
  docx 文档 plvr=
  not_pdf_document /
  坏文档双键 pipeline_
  failed（跨型三态首锁）
- **docx 抽取面**——
  python-docx heading+
  paragraph → ecbt
  {heading:1,
  paragraph:1}；tpe
  True；dlvr 1.0
- **未知 ef 文档**——
  doc_id 不在 documents
  → actual file_not_
  found + matches False
  （非 None 首锁）
- **ef 混合结果**——
  bad1 命中 True + zzz
  落空 False 同列表
- forbidden tokens 第七百五十七批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
STREAM = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
          % ("A" * 80)
          + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
          % LONG).encode()


def _board(tmp_path):
    (tmp_path / "g.pdf").write_bytes(_wrap(STREAM))
    (tmp_path / "bad.pdf").write_bytes(b"not a pdf")
    d = Document()
    d.add_heading("DocTitle", level=1)
    d.add_paragraph(" ".join("Sent%d." % i
                             for i in range(40)))
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "complete",
        "documents": [
            {"doc_id": "g1", "path": "g.pdf",
             "source_type": "pdf"},
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx",
             "paired_with": "g1"},
            {"doc_id": "bad1", "path": "bad.pdf",
             "source_type": "pdf"}],
        "expected_failures": [
            {"doc_id": "bad1", "path": "bad.pdf",
             "expected_error_code":
             "pdfplumber_open_failed"},
            {"doc_id": "zzz", "path": "nowhere.pdf",
             "expected_error_code":
             "pdfplumber_open_failed"}]}),
        encoding="utf-8")
    return load_manifest((tmp_path / "m.json"),
                         project_root=tmp_path)


def _run(tmp_path):
    return run_evaluation(_board(tmp_path),
                          tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=32)


# ---------- complete 透传 ----------

def test_devset_status_complete_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["devset"]["status"] == "complete"


# ---------- 混合计数 ----------

def test_mixed_counts_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["devset"]["file_count"] == 3
    assert r["devset"]["pdf_count"] == 2
    assert r["devset"]["docx_count"] == 1


def test_group_count_with_cross_pair_batch507(
        tmp_path):
    r = _run(tmp_path)
    assert r["devset"]["content_group_count"] == 2


# ---------- locator 分型空值 ----------

def test_pdf_doc_dlvr_not_docx_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "docx_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_docx_document"}


def test_docx_doc_dlvr_one_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


def test_docx_doc_plvr_not_pdf_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "pdf_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_pdf_document"}


def test_bad_doc_double_locator_failed_batch507(
        tmp_path):
    r = _run(tmp_path)
    m = r["per_doc"][2]["metrics"]
    assert m["docx_locator_valid_ratio"] == {
        "value": None, "reason": "pipeline_failed"}
    assert m["pdf_locator_valid_ratio"] == {
        "value": None, "reason": "pipeline_failed"}


def test_dlvr_macro_one_third_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "docx_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 2}


def test_plvr_macro_one_third_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "pdf_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 2}


# ---------- docx 抽取面 ----------

def test_docx_ecbt_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "element_count_by_type"]["value"] == {
        "heading": 1, "paragraph": 1}


def test_docx_tpe_true_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "text_preservation_equal"]["value"] is True


# ---------- 未知 ef 文档 ----------

def test_ef_unknown_doc_file_not_found_batch507(
        tmp_path):
    r = _run(tmp_path)
    assert r["expected_failures"][1] == {
        "doc_id": "zzz",
        "expected_error_code":
        "pdfplumber_open_failed",
        "actual_error_code": "file_not_found",
        "matches": False}


def test_ef_mixed_results_batch507(tmp_path):
    r = _run(tmp_path)
    assert [e["matches"]
            for e in r["expected_failures"]] == [
        True, False]


def test_ef_order_preserved_batch507(tmp_path):
    r = _run(tmp_path)
    assert [e["doc_id"]
            for e in r["expected_failures"]] == [
        "bad1", "zzz"]


# ---------- 成功率 ----------

def test_success_two_thirds_batch507(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 2, "total": 3,
        "rate": 2 / 3}


# ---------- 报告合法性 ----------

def test_report_schema_batch507(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_counts_batch507():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("per_doc") == 12
    assert src.count("open(") == 2


# ---------- forbidden tokens 第七百五十七批 ----------

def test_source_no_eval_batch507():
    assert "eval(" not in _src()


def test_source_no_exec_batch507():
    assert "exec(" not in _src()


def test_source_no_compile_batch507():
    assert "compile(" not in _src()


def test_source_no_globals_batch507():
    assert "globals(" not in _src()


def test_source_no_locals_batch507():
    assert "locals(" not in _src()


def test_source_no_os_system_batch507():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch507():
    assert "subprocess" not in _src()


def test_source_no_popen_batch507():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch507():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch507():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch507():
    assert "socket" not in _src()


def test_source_no_requests_batch507():
    assert "requests" not in _src()


def test_source_no_urllib_batch507():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch507():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch507():
    assert "yield" not in _src()


def test_source_no_async_await_batch507():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch507():
    assert ".call(" not in _src()
