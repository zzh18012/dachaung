"""evaluation/schema.py 第五百八十二轮 edges 测试（Round 1311）。

补强 edges146 未触及的角度（第六百八十三批，probe 实证）。

新角度（document.schema.json 真文档变异面）：
- **sei 非空链**——[]
  → '[] should be
  non-empty' @ [chunks,
  0, source_element_ids]
  （每 chunk 非空 sei
  不变量 schema 级首锁）
- **sei 空串项**——['']
  → '' should be
  non-empty @ [.., 0]
- **chunk 缺 text**——
  required @ [chunks, 0]
- **element 面**——缺
  source_locator /
  空 element_id 均
  required/non-empty
- **locator 类型**——
  paragraph_index
  '3' → not integer；
  None → not object
- **无 if/then 区分**——
  docx 文档带 pdf 形
  locator（page+bbox）
  VALID（locator 分型
  靠 parser 约定而非
  schema 强制——宽松
  面首锁）
- **srctype 宽枚举**——
  document 级 [pdf,
  docx, markdown,
  html, text, ipynb]
  （manifest 严域
  [pdf, docx] 不对称
  首锁）
- **element type 枚举**
  ——8 值严域
- forbidden tokens 第七百五十八批（open 2）
"""

from __future__ import annotations

import copy
import inspect

import evaluation.schema as schema_mod
from docx import Document
from app.pipeline import process_single
from evaluation.schema import EvalSchemaError, validate


def _base():
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        tp = __import__("pathlib").Path(td)
        d = Document()
        d.add_heading("T", level=1)
        d.add_paragraph("Body text here.")
        d.save(str(tp / "c.docx"))
        doc, errors = process_single(tp / "c.docx",
                                     tp / "o.json",
                                     parser_name="fallback",
                                     max_chars=32)
        assert errors == []
        return doc.to_dict()


def _rej(base, mutate, message, path):
    d = copy.deepcopy(base)
    mutate(d)
    try:
        validate(d, "document.schema.json")
    except EvalSchemaError as e:
        assert e.errors[0]["message"] == message
        assert list(e.errors[0]["path"]) == path
    else:
        raise AssertionError("expected rejection")


def _acc(base, mutate):
    d = copy.deepcopy(base)
    mutate(d)
    validate(d, "document.schema.json")


# ---------- sei 非空链 ----------

def test_sei_empty_rejected_batch509():
    b = _base()
    _rej(b, lambda d: d["chunks"][0].__setitem__(
             "source_element_ids", []),
         "[] should be non-empty",
         ["chunks", 0, "source_element_ids"])


def test_sei_empty_string_item_batch509():
    b = _base()
    _rej(b, lambda d: d["chunks"][0].__setitem__(
             "source_element_ids", [""]),
         "'' should be non-empty",
         ["chunks", 0, "source_element_ids", 0])


def test_sei_two_ids_baseline_batch509():
    b = _base()
    assert len(b["chunks"][0]["source_element_ids"]) \
        == 2


# ---------- chunk 面 ----------

def test_chunk_missing_text_batch509():
    b = _base()
    _rej(b, lambda d: d["chunks"][0].pop("text"),
         "'text' is a required property",
         ["chunks", 0])


def test_chunk_missing_metadata_batch509():
    b = _base()
    _rej(b, lambda d: d["chunks"][0].pop("metadata"),
         "'metadata' is a required property",
         ["chunks", 0])


# ---------- element 面 ----------

def test_element_missing_locator_batch509():
    b = _base()
    _rej(b, lambda d: d["elements"][0].pop(
             "source_locator"),
         "'source_locator' is a required property",
         ["elements", 0])


def test_element_empty_id_batch509():
    b = _base()
    _rej(b, lambda d: d["elements"][0].__setitem__(
             "element_id", ""),
         "'' should be non-empty",
         ["elements", 0, "element_id"])


def test_element_missing_type_batch509():
    b = _base()
    _rej(b, lambda d: d["elements"][0].pop("type"),
         "'type' is a required property",
         ["elements", 0])


# ---------- locator 类型 ----------

def test_pi_string_rejected_batch509():
    b = _base()
    _rej(b, lambda d: d["elements"][0][
             "source_locator"].__setitem__(
        "paragraph_index", "3"),
         "'3' is not of type 'integer'",
         ["elements", 0, "source_locator",
          "paragraph_index"])


def test_locator_null_rejected_batch509():
    b = _base()
    _rej(b, lambda d: d["elements"][0].__setitem__(
             "source_locator", None),
         "None is not of type 'object'",
         ["elements", 0, "source_locator"])


# ---------- 无 if/then 区分 ----------

def test_pdf_locator_on_docx_valid_batch509():
    b = _base()
    _acc(b, lambda d: d["elements"][0].__setitem__(
        "source_locator",
        {"page": 1, "bbox": [0, 0, 10, 10]}))


def test_docx_locator_baseline_batch509():
    b = _base()
    assert b["elements"][0]["source_locator"] == {
        "section": 0, "paragraph_index": 0}


# ---------- srctype 宽枚举 ----------

def test_srctype_text_requires_line_batch509():
    b = _base()
    _rej(b, lambda d: d.__setitem__("source_type",
                                    "text"),
         "'line' is a required property",
         ["elements", 0, "source_locator"])


def test_srctype_pdf_requires_page_batch509():
    b = _base()
    _rej(b, lambda d: d.__setitem__("source_type",
                                    "pdf"),
         "'page' is a required property",
         ["elements", 0, "source_locator"])


def test_srctype_ipynb_requires_cell_batch509():
    b = _base()
    _rej(b, lambda d: d.__setitem__("source_type",
                                    "ipynb"),
         "'cell_index' is a required property",
         ["elements", 0, "source_locator"])


def test_srctype_txt_rejected_batch509():
    b = _base()
    _rej(b, lambda d: d.__setitem__("source_type",
                                    "txt"),
         "'txt' is not one of ['pdf', 'docx', "
         "'markdown', 'html', 'text', 'ipynb']",
         ["source_type"])


# ---------- element type 枚举 ----------

def test_element_type_enum_batch509():
    b = _base()
    _rej(b, lambda d: d["elements"][0].__setitem__(
             "type", "bogus"),
         "'bogus' is not one of ['heading', "
         "'paragraph', 'list_item', 'table', "
         "'image', 'caption', 'header', 'footer']",
         ["elements", 0, "type"])


def test_element_type_header_valid_batch509():
    b = _base()
    _acc(b, lambda d: d["elements"][0].__setitem__(
        "type", "header"))


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch509():
    src = _src()
    assert "class EvalSchemaError(Exception):" in src
    assert "def validate(instance" in src


# ---------- forbidden tokens 第七百五十八批 ----------

def test_source_no_eval_batch509():
    assert "eval(" not in _src()


def test_source_no_exec_batch509():
    assert "exec(" not in _src()


def test_source_no_compile_batch509():
    assert "compile(" not in _src()


def test_source_no_globals_batch509():
    assert "globals(" not in _src()


def test_source_no_locals_batch509():
    assert "locals(" not in _src()


def test_source_no_os_system_batch509():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch509():
    assert "subprocess" not in _src()


def test_source_no_popen_batch509():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch509():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch509():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch509():
    assert "socket" not in _src()


def test_source_no_requests_batch509():
    assert "requests" not in _src()


def test_source_no_urllib_batch509():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch509():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch509():
    assert "yield" not in _src()


def test_source_no_async_await_batch509():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch509():
    assert _src().count("open(") == 2
