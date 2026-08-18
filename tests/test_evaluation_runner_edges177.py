"""evaluation/runner.py 第六百零五轮 edges 测试（Round 1161）。

补强 edges176 未触及的角度（第五百三十三批，probe 实证）。

新角度（DOCX 五型混排全通道）：
- **五型文档序**——heading + paragraph + 内嵌图
  + caption + table 单 DOCX → elements 恰按文档序
  五型各一（五型同文档首锁）
- **内嵌图真元素**——add_picture → image 元素
  content None + locator {paragraph_index,
  relationship_id rId9, target_partname
  /word/media/image1.png} + {byte_size 67,
  ext png, extracted_to_disk True}
- **Figure 前缀 caption**——Normal 样式段落
  "Figure 2: …" → caption（DOCX 侧 caption 通道
  首锁）→ isolated_caption 块
- **图片不产块**——3 chunks 皆不引用 image 元素
  id；image_resource_exists_ratio 真跑 1.0
- forbidden tokens 第六百三十三批（open 2）
"""

from __future__ import annotations

import inspect
import json
from io import BytesIO

import evaluation.runner as runner_mod
from docx import Document
from docx.shared import Inches
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single

_PNG = (b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01'
        b'\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89'
        b'\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01'
        b'\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82')


def _board(tmp_path, doc_id):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("Doc Heading", level=1)
    p = d.add_paragraph("Body text with inline picture next.")
    p.add_run().add_picture(BytesIO(_PNG), width=Inches(1))
    d.add_paragraph("Figure 2: docx caption text below.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "X1"
    t.cell(0, 1).text = "Y1"
    d.save(str(tmp_path / "samples" / f"{doc_id}.docx"))
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"samples/{doc_id}.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 五型文档序 ----------

def test_docx_five_types_order_batch359(tmp_path):
    _board(tmp_path, "mi")
    doc, errors = process_single(
        tmp_path / "samples" / "mi.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "heading", "paragraph", "image", "caption", "table"]
    assert els[0]["content"] == "Doc Heading"
    assert els[1]["content"] == \
        "Body text with inline picture next."


# ---------- 内嵌图真元素 ----------

def test_docx_inline_image_element_batch359(tmp_path):
    _board(tmp_path, "mi2")
    doc, errors = process_single(
        tmp_path / "samples" / "mi2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    img = [e for e in doc.to_dict()["elements"]
           if e["type"] == "image"][0]
    assert img["content"] is None
    assert img["source_locator"]["paragraph_index"] == 1
    assert img["source_locator"]["section"] == 0
    assert img["source_locator"]["relationship_id"] == "rId9"
    assert img["source_locator"]["target_partname"] == \
        "/word/media/image1.png"
    assert img["metadata"]["byte_size"] == 67
    assert img["metadata"]["ext"] == "png"
    assert img["metadata"]["extracted_to_disk"] is True


# ---------- Figure 前缀 caption ----------

def test_docx_caption_regex_batch359(tmp_path):
    _board(tmp_path, "mi3")
    doc, errors = process_single(
        tmp_path / "samples" / "mi3.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    els = doc.to_dict()["elements"]
    cap = els[3]
    assert cap["type"] == "caption"
    assert cap["content"] == "Figure 2: docx caption text below."
    assert cap["metadata"]["style"] == "Normal"
    chunks = doc.to_dict()["chunks"]
    assert chunks[1]["metadata"]["strategy"] == \
        "isolated_caption"
    assert chunks[1]["text"] == \
        "Figure 2: docx caption text below."


# ---------- 图片不产块 ----------

def test_docx_image_no_chunk_batch359(tmp_path):
    _board(tmp_path, "mi4")
    doc, errors = process_single(
        tmp_path / "samples" / "mi4.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    dd = doc.to_dict()
    img_id = [e["element_id"] for e in dd["elements"]
              if e["type"] == "image"][0]
    assert len(dd["chunks"]) == 3
    for c in dd["chunks"]:
        assert img_id not in c["source_element_ids"]


# ---------- 指标 ----------

def test_docx_five_types_metrics_batch359(tmp_path):
    r = run_evaluation(_board(tmp_path, "mi5"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"heading": 1, "paragraph": 1, "image": 1,
                  "caption": 1, "table": 1},
        "reason": None}
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch359():
    src = _src()
    assert src.count("annotation") == 10
    assert src.count("per_doc") == 12
    assert src.count("run_evaluation") == 2


# ---------- forbidden tokens 第六百三十三批 ----------

def test_source_no_eval_batch359():
    assert "eval(" not in _src()


def test_source_no_exec_batch359():
    assert "exec(" not in _src()


def test_source_no_compile_batch359():
    assert "compile(" not in _src()


def test_source_no_globals_batch359():
    assert "globals(" not in _src()


def test_source_no_locals_batch359():
    assert "locals(" not in _src()


def test_source_no_os_system_batch359():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch359():
    assert "subprocess" not in _src()


def test_source_no_popen_batch359():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch359():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch359():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch359():
    assert "socket" not in _src()


def test_source_no_requests_batch359():
    assert "requests" not in _src()


def test_source_no_urllib_batch359():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch359():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch359():
    assert "yield" not in _src()


def test_source_no_async_await_batch359():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch359():
    assert _src().count("open(") == 2
