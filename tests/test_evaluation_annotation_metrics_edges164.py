"""evaluation/annotation_metrics.py 第五百九十五轮 edges 测试（Round 1343）。

补强 edges163 未触及的角度（第七百一十五批，probe 实证）。

新角度（抽取流边界 / 单 chunk 无预测 / unicode）：
- **匹配流=抽取文本**
  ——marker 'BT'
  （PDF 内容流前缀
  但抽取文本无）→
  missing trio（流
  非原始 content
  stream 首锁）
- **regex 字符字面**
  ——marker
  '(Word3.' 含括号
  → 照样字面 find
  → missing
- **跨空格 marker**
  ——'Word3. Word4'
  双词整段 → HIT
  {1/14, 1.0, 2/15}
- **单 chunk 无预测**
  ——docx mc800
  单 chunk →
  cbp/f1 {None,
  no_predicted_
  boundaries} +
  cbr {0.0, None}
  （新 reason 首锁）
- **unicode 命中/未命中**
  ——中文 marker
  同样走无预测面
  （单 chunk 板
  marker 命中与否
  无差别首锁）
- forbidden tokens 第七百八十五批（open 0）
"""

from __future__ import annotations

import inspect
import tempfile
from pathlib import Path

import evaluation.annotation_metrics as am_mod
from app.pipeline import process_single
from docx import Document
from evaluation.annotation_metrics import \
    chunk_boundary_prf


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
ONEP = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
        % LONG).encode()

MISSING = ({"value": 0.0, "reason": None},
           {"value": None,
            "reason":
                "no_ground_truth_anchors_in_stream"},
           {"value": None,
            "reason":
                "precision_or_recall_not_evaluated"})
HIT14 = ({"value": 1 / 14, "reason": None},
         {"value": 1.0, "reason": None},
         {"value": 2 / 15, "reason": None})


def _pdf():
    with tempfile.TemporaryDirectory() as td:
        tp = Path(td)
        (tp / "c.pdf").write_bytes(_wrap(ONEP))
        doc, errors = process_single(
            tp / "c.pdf", tp / "o.json",
            parser_name="fallback", max_chars=32)
        assert errors == []
        return doc.to_dict()


def _docx():
    with tempfile.TemporaryDirectory() as td:
        tp = Path(td)
        d = Document()
        d.add_heading("标题一", level=1)
        d.add_paragraph("这是第一段内容，包含中文文本。")
        d.add_paragraph("Second paragraph.")
        d.save(str(tp / "c.docx"))
        doc, errors = process_single(
            tp / "c.docx", tp / "o.json",
            parser_name="fallback", max_chars=800)
        assert errors == []
        return doc.to_dict()


def _trio(doc, anchors, tol=30):
    ann = {"annotation_version": "1.0",
           "doc_id": "x",
           "chunk_boundary_anchors": anchors}
    r = chunk_boundary_prf(doc, ann, tol)
    return (r["chunk_boundary_precision"],
            r["chunk_boundary_recall"],
            r["chunk_boundary_f1"])


# ---------- 匹配流=抽取文本 ----------

def test_bt_marker_missing_batch541():
    assert _trio(_pdf(), [
        {"marker": "BT",
         "position": "before"}]) == MISSING


def test_bt_after_missing_batch541():
    assert _trio(_pdf(), [
        {"marker": "BT",
         "position": "after"}]) == MISSING


# ---------- regex 字符字面 ----------

def test_paren_marker_missing_batch541():
    assert _trio(_pdf(), [
        {"marker": "(Word3.",
         "position": "after"}]) == MISSING


def test_long_text_found_batch541():
    assert _trio(_pdf(), [
        {"marker": "Word3.",
         "position": "after"}]) == HIT14


# ---------- 跨空格 marker ----------

def test_span_space_marker_hit_batch541():
    assert _trio(_pdf(), [
        {"marker": "Word3. Word4",
         "position": "after"}]) == HIT14


def test_span_space_after_only_batch541():
    t = _trio(_pdf(), [
        {"marker": "Word3. Word4",
         "position": "after"}])
    assert t[0]["reason"] is None


# ---------- 单 chunk 无预测 ----------

def test_single_chunk_no_pred_batch541():
    t = _trio(_docx(), [
        {"marker": "中文文本",
         "position": "after"}])
    assert t[0] == {"value": None,
                    "reason":
                        "no_predicted_boundaries"}
    assert t[2] == {"value": None,
                    "reason":
                        "no_predicted_boundaries"}


def test_single_chunk_cbr_zero_batch541():
    t = _trio(_docx(), [
        {"marker": "中文文本",
         "position": "after"}])
    assert t[1] == {"value": 0.0,
                    "reason": None}


def test_docx_one_chunk_batch541():
    assert len(_docx()["chunks"]) == 1


# ---------- unicode 命中/未命中 ----------

def test_unicode_miss_same_face_batch541():
    hit = _trio(_docx(), [
        {"marker": "中文文本",
         "position": "after"}])
    miss = _trio(_docx(), [
        {"marker": "不存在",
         "position": "after"}])
    assert hit == miss


def test_unicode_miss_cbr_zero_batch541():
    t = _trio(_docx(), [
        {"marker": "不存在",
         "position": "after"}])
    assert t[1] == {"value": 0.0,
                    "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_no_pred_reason_batch541():
    assert '"no_predicted_boundaries"' in _src()


def test_source_marker_or_batch541():
    assert "if marker else -1" in _src()


# ---------- forbidden tokens 第七百八十五批 ----------

def test_source_no_eval_batch541():
    assert "eval(" not in _src()


def test_source_no_exec_batch541():
    assert "exec(" not in _src()


def test_source_no_compile_batch541():
    assert "compile(" not in _src()


def test_source_no_globals_batch541():
    assert "globals(" not in _src()


def test_source_no_locals_batch541():
    assert "locals(" not in _src()


def test_source_no_os_system_batch541():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch541():
    assert "subprocess" not in _src()


def test_source_no_popen_batch541():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch541():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch541():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch541():
    assert "socket" not in _src()


def test_source_no_requests_batch541():
    assert "requests" not in _src()


def test_source_no_urllib_batch541():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch541():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch541():
    assert "yield" not in _src()


def test_source_no_async_await_batch541():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_0_batch541():
    assert _src().count("open(") == 0
