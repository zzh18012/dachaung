"""evaluation/runner.py 第六百四十轮 edges 测试（Round 1196）。

补强 edges206 未触及的角度（第五百六十八批，probe 实证）。

新角度（超链接文本透传 / 容差三明治）：
- **hyperlink 文本透传**——w:hyperlink
  内 w:t 经 python-docx 1.2.0 的
  paragraph.text 原样入流（"Hidden
  Link Text" 内联不丢首锁；run 与
  hyperlink 内容直接拼接）
- **容差三明治**——同段两中位锚：
  "Hidden" after（距界 24 ≤ 30）→
  全 1.0 vs "Before" after（距界
  40 > 30）→ 全 0.0（容差界两侧
  对照首锁）
- **一命中一偏离**——[Before, link.]
  → P 1.0 / R 0.5 / F1 2/3
- forbidden tokens 第六百六十八批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "a").mkdir(exist_ok=True)
    d = Document()
    p = d.add_paragraph()
    p.add_run("Before link ")
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), 'rId99')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "Hidden Link Text"
    r.append(t)
    hyperlink.append(r)
    p._p.append(hyperlink)
    p.add_run(" after link.")
    d.add_paragraph("Second plain paragraph follows here.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    entry = {"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}
    if anchors is not None:
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        entry["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": [entry]}),
                  encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


_P1 = "Before link Hidden Link Text after link."
_P2 = "Second plain paragraph follows here."


# ---------- 超链接文本透传 ----------

def test_hyperlink_text_flows_batch394(tmp_path):
    _board(tmp_path, "hl")
    doc, errors = process_single(
        tmp_path / "s" / "hl.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=60)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == ["paragraph"] * 2
    assert els[0]["content"] == _P1
    assert "Hidden Link Text" in els[0]["content"]


def test_hyperlink_chunks_batch394(tmp_path):
    _board(tmp_path, "hl2")
    doc, errors = process_single(
        tmp_path / "s" / "hl2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=60)
    chunks = doc.to_dict()["chunks"]
    assert [c["text"] for c in chunks] == [_P1, _P2]
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks)


# ---------- 容差三明治 ----------

def test_link_anchor_batch394(tmp_path):
    r = run_evaluation(_board(tmp_path, "a1", [
        {"marker": "link.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=60)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_hidden_anchor_batch394(tmp_path):
    r = run_evaluation(_board(tmp_path, "a2", [
        {"marker": "Hidden", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=60)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_before_anchor_batch394(tmp_path):
    r = run_evaluation(_board(tmp_path, "a3", [
        {"marker": "Before", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=60)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.0, "reason": None}


def test_both_anchors_batch394(tmp_path):
    r = run_evaluation(_board(tmp_path, "a4", [
        {"marker": "Before", "position": "after"},
        {"marker": "link.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=60)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 指标 ----------

def test_hyperlink_metrics_batch394(tmp_path):
    r = run_evaluation(_board(tmp_path, "mx"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=60)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 2}, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch394():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百六十八批 ----------

def test_source_no_eval_batch394():
    assert "eval(" not in _src()


def test_source_no_exec_batch394():
    assert "exec(" not in _src()


def test_source_no_compile_batch394():
    assert "compile(" not in _src()


def test_source_no_globals_batch394():
    assert "globals(" not in _src()


def test_source_no_locals_batch394():
    assert "locals(" not in _src()


def test_source_no_os_system_batch394():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch394():
    assert "subprocess" not in _src()


def test_source_no_popen_batch394():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch394():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch394():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch394():
    assert "socket" not in _src()


def test_source_no_requests_batch394():
    assert "requests" not in _src()


def test_source_no_urllib_batch394():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch394():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch394():
    assert "yield" not in _src()


def test_source_no_async_await_batch394():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch394():
    assert _src().count("open(") == 2
