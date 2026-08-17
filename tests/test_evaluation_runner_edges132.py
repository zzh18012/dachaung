"""evaluation/runner.py 第五百零五轮 edges 测试（Round 1061）。

补强 edges129-131 未触及的角度（第四百三十七批，probe 实证）。

新角度（**真实嵌入图片**走通全链路，corpus 首次 add_picture）：
- 手工构造 1x1 PNG（struct+zlib，零新依赖）嵌入 docx，
  经 fallback parser → pipeline 落盘 → run_evaluation 度量：
  image_resource_exists_ratio {1.0, None}，summary
  {macro 1.0, participating 1, not_evaluated 0}——
  图片指标第一次用真实磁盘文件点亮
- 落盘位置真相：`_per_doc/images-<sha>/image_<sha8>_
  para1_00.png`（images 目录**直接**在 _per_doc 下，
  非 <doc_id> 子目录）；文件字节与构造 PNG **逐字节
  相等**（blob 原样透传）
- 承载图片的段落同时产生两个元素："(空段落)" 文本
  元素 + image 元素（content None、resource_path 非
  null），二者 locator 的 paragraph_index 同为 1；
  image 计入 element_count_total（4）
- forbidden tokens 第五百三十二批（open 2）
"""

from __future__ import annotations

import inspect
import json
import struct
import zlib
from io import BytesIO
from pathlib import Path

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        crc = zlib.crc32(c) & 0xFFFFFFFF
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", crc))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("BBB third paragraph body.")
    d.save(str(tmp_path / "samples" / "img.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/img.docx",
                       "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    rep = run_evaluation(load_manifest(mf, tmp_path),
                         tmp_path / "o.json", max_chars=200)
    return rep


# ---------- 图片指标端到端点亮 ----------

def test_image_ratio_real_end_to_end_batch260(tmp_path):
    rep = _run(tmp_path)
    assert rep["per_doc"][0]["metrics"][
        "image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}
    assert rep["summary"]["ratio_macro_averages"][
        "image_resource_exists_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 0}


# ---------- 元素形态：空段落 + image 双元素 ----------

def test_image_element_shape_batch260(tmp_path):
    from app.pipeline import process_single
    _run(tmp_path)
    doc, errors = process_single(
        tmp_path / "samples" / "img.docx",
        tmp_path / "_per_doc" / "s.json",
        parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    dd = doc.to_dict()
    assert [e["type"] for e in dd["elements"]] == [
        "paragraph", "paragraph", "image", "paragraph"]
    host = dd["elements"][1]
    img = dd["elements"][2]
    assert host["content"] == "(空段落)"
    assert img["content"] is None
    assert img["resource_path"] is not None
    assert (host["source_locator"]["paragraph_index"]
            == img["source_locator"]["paragraph_index"] == 1)


# ---------- 落盘位置与字节保真 ----------

def test_image_file_on_disk_batch260(tmp_path):
    _run(tmp_path)
    files = [q for q in (tmp_path / "_per_doc").rglob("*")
             if q.is_file()]
    assert len(files) == 1
    p = files[0]
    assert p.name.endswith("_para1_00.png")
    assert p.parent.name.startswith("images-")
    assert p.parent.parent.name == "_per_doc"
    assert p.read_bytes() == _png_bytes()
    assert p.read_bytes()[:8] == b"\x89PNG\r\n\x1a\n"


# ---------- image 计入元素总数 ----------

def test_ect_counts_image_batch260(tmp_path):
    rep = _run(tmp_path)
    assert rep["per_doc"][0]["metrics"][
        "element_count_total"] == {"value": 4,
                                   "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch260():
    src = _src()
    assert "image_dir.is_dir()" in src
    assert "_per_doc" in src


# ---------- forbidden tokens 第五百三十二批 ----------

def test_source_no_eval_batch260():
    assert "eval(" not in _src()


def test_source_no_exec_batch260():
    assert "exec(" not in _src()


def test_source_no_compile_batch260():
    assert "compile(" not in _src()


def test_source_no_globals_batch260():
    assert "globals(" not in _src()


def test_source_no_locals_batch260():
    assert "locals(" not in _src()


def test_source_no_os_system_batch260():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch260():
    assert "subprocess" not in _src()


def test_source_no_popen_batch260():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch260():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch260():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch260():
    assert "socket" not in _src()


def test_source_no_requests_batch260():
    assert "requests" not in _src()


def test_source_no_urllib_batch260():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch260():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch260():
    assert "yield" not in _src()


def test_source_no_async_await_batch260():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch260():
    assert _src().count("open(") == 2
