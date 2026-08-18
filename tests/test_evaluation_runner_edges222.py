"""evaluation/runner.py 第六百五十七轮 edges 测试（Round 1241）。

补强 edges221 未触及的角度（第六百一十三批，probe 实证）。

新角度（水槽板 + 真标注文件全链）：
- **双占位标注全中经 manifest**
  ——annotation_file 挂 "(空段落)"
  after ×2 → run_evaluation 报
  P/R/F1 全 1.0（标注文件加载 +
  annotation_metrics + runner 三层
  一链首锁；R1238 是直调
  chunk_boundary_prf）
- **summary 三 macro**——
  precision/recall/f1 各
  {1.0, 1, 0}
- **异构指标共存**——boundary
  全中之外 ect 7 / hbc 1.0 照出
  （标注在场不扰其余指标）
- forbidden tokens 第七百零七批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _docx(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    (tmp_path / "s").mkdir(exist_ok=True)
    p = tmp_path / "s" / "ks.docx"
    doc.save(str(p))
    return p


def _board(tmp_path):
    _docx(tmp_path)
    (tmp_path / "a").mkdir(exist_ok=True)
    (tmp_path / "a" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "ks",
        "chunk_boundary_anchors": [
            {"marker": "(空段落)", "position": "after"},
            {"marker": "(空段落)", "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "ks", "path": "s/ks.docx",
                       "source_type": "docx",
                       "annotation_file": "a/a.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 双占位标注全中经 manifest ----------

def test_anchor_per_doc_all_hit_batch439(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


# ---------- summary 三 macro ----------

def test_summary_f1_macro_batch439(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 0}


def test_summary_precision_macro_batch439(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_precision"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 0}


# ---------- 标注装载 ----------

def test_annotation_resolved_exists_batch439(tmp_path):
    m = _board(tmp_path)
    entry = m.documents[0]
    assert entry.annotation_resolved is not None
    assert entry.annotation_resolved.is_file()


# ---------- 异构指标共存 ----------

def test_other_metrics_still_full_batch439(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_total"] == {"value": 7,
                                        "reason": None}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


def test_schema_still_valid_batch439(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["schema_valid"] == {"value": True, "reason": None}
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch439():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第七百零七批 ----------

def test_source_no_eval_batch439():
    assert "eval(" not in _src()


def test_source_no_exec_batch439():
    assert "exec(" not in _src()


def test_source_no_compile_batch439():
    assert "compile(" not in _src()


def test_source_no_globals_batch439():
    assert "globals(" not in _src()


def test_source_no_locals_batch439():
    assert "locals(" not in _src()


def test_source_no_os_system_batch439():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch439():
    assert "subprocess" not in _src()


def test_source_no_popen_batch439():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch439():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch439():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch439():
    assert "socket" not in _src()


def test_source_no_requests_batch439():
    assert "requests" not in _src()


def test_source_no_urllib_batch439():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch439():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch439():
    assert "yield" not in _src()


def test_source_no_async_await_batch439():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch439():
    assert _src().count("open(") == 2
