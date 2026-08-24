r"""app/chunkers/structural.py 边角测试 - 第十一轮（Round 1441）。

新角度（probe 实证）表格元素与 max_chars 的交互（历史分块
考察全是纯文本元素，表格元素从没过量）：
- 表格元素**拒绝切分**：8×3 表格 markdown 491 字符，max_chars
  100/200/490/491/492 全部原样单 chunk 491（markdown 结构
  不可断，max_chars 对单表格元素是软约束）
- 长文本元素**空格优先切**：690 字符段落 mc=100 → 7×95+31
  （95 < 100，词边界回退），mc=200 → 4×199
- 无空格长词 300×a mc=100 → **精确 100×3 字符级硬切**
  （无边界可回退时不留整）
- 表格是**chunk 孤岛**：para(10)+table(15)+para(9) mc=100
  → 3 个独立 chunk（对照三短段落会并成单 chunk），表格
  不与相邻文本合并
- PDF 长文本（>80 归 paragraph）同样切分：318 字符 mc=100
  → 3×95+31
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.pipeline import process_single


def _big_table_docx(tmp_path):
    d = Document()
    d.add_paragraph("intro")
    t = d.add_table(rows=8, cols=3)
    for r in range(8):
        for c in range(3):
            t.cell(r, c).text = \
                "cell%02d%02d content" % (r, c)
    p = tmp_path / "big.docx"
    d.save(str(p))
    return p


def _long_para_docx(tmp_path):
    d = Document()
    d.add_paragraph(" ".join(
        "word%03d" % i
        for i in range(100)))
    p = tmp_path / "lp.docx"
    d.save(str(p))
    return p


# ---------- 表格拒切 ----------

def test_table_unsplit_mc100(
        tmp_path):
    p = _big_table_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=100)
    assert errors == []
    assert [len(c.text)
            for c in doc.chunks] == [
        5, 491]


def test_table_unsplit_mc200(
        tmp_path):
    p = _big_table_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=200)
    assert errors == []
    assert len(doc.chunks) == 2
    assert len(doc.chunks[1].text) \
        == 491


def test_table_boundary_exact(
        tmp_path):
    p = _big_table_docx(tmp_path)
    doc0, _ = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    tlen = len(doc0.chunks[1].text)
    assert tlen == 491
    for mc in (tlen - 1, tlen,
               tlen + 1):
        doc, errors = process_single(
            p, None,
            parser_name="fallback",
            max_chars=mc)
        assert errors == []
        assert len(doc.chunks) == 2
        assert len(
            doc.chunks[1].text) == 491


# ---------- 文本空格优先 ----------

def test_para_split_mc100(
        tmp_path):
    p = _long_para_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=100)
    assert errors == []
    lens = [len(c.text)
            for c in doc.chunks]
    assert lens == [95] * 8 + [31]
    assert len(doc.elements[
        0].content) == 799
    assert doc.chunks[
        0].text.startswith(
        "word000 word001")


def test_para_split_mc200(
        tmp_path):
    p = _long_para_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=200)
    assert errors == []
    assert [len(c.text)
            for c in doc.chunks] == [
        199, 199, 199, 199]
    assert doc.chunks[
        3].text.endswith(
        "word098 word099")


# ---------- 无空格硬切 ----------

def test_unbroken_hard_split(
        tmp_path):
    d = Document()
    d.add_paragraph("a" * 300)
    p = tmp_path / "lw.docx"
    d.save(str(p))
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=100)
    assert errors == []
    assert [len(c.text)
            for c in doc.chunks] == [
        100, 100, 100]


# ---------- 表格孤岛 ----------

def _mix_docx(tmp_path):
    d = Document()
    d.add_paragraph("alpha beta")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "gam"
    d.add_paragraph("delta eps")
    p = tmp_path / "mix.docx"
    d.save(str(p))
    return p


def test_table_island(tmp_path):
    p = _mix_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=100)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "alpha beta",
        "| gam |\n| --- |",
        "delta eps"]
    assert all(
        len(c.source_element_ids) == 1
        for c in doc.chunks)


def test_paras_would_merge(
        tmp_path):
    d = Document()
    d.add_paragraph("alpha beta")
    d.add_paragraph("delta eps")
    p = tmp_path / "twop.docx"
    d.save(str(p))
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=100)
    assert errors == []
    assert len(doc.chunks) == 1
    assert doc.chunks[
        0].text == \
        "alpha beta delta eps"


# ---------- 通用 ----------

def test_big_table_no_drop(
        tmp_path):
    p = _big_table_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=100)
    assert errors == []
    joined = "".join(
        c.text for c in doc.chunks)
    assert "cell0700 content" in joined
    assert "cell0702 content" in joined
