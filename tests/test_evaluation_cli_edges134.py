"""evaluation/cli.py 第五百二十轮 edges 测试（Round 1076）。

补强 edges129-133 未触及的角度（第四百五十二批，probe 实证）。

新角度（inspect-doc 渲染真实双图双题注文档）：
- 双图双题注（一英 "Figure 1: ..." 一中 "图 2 ..."）
  文档 JSON：counts 行 "elements=6 chunks=4"（6 元素
  切成 4 chunk 的真实形态）
- dict 分支按码点排序渲染——**caption 排首**：
  "caption=2, image=2, paragraph=2"
- **figure_caption 三项在"最接近能用"的文档上仍是
  null parser_does_not_emit_relations**——文档里图片
  与题注文本俱在，指标依然全黑（关系不由 parser 发
  出，不启发式凑合）——该不变量的最强真实对照
- heading_boundary null no_heading_elements（题注不
  是 heading）
- forbidden tokens 第五百四十七批（open 1）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.cli as cli_mod
from app.pipeline import process_single
from evaluation.cli import main


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _inspect(tmp_path, capsys):
    d = Document()
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("Figure 1: a real caption line.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("图 2 中文题注内容")
    d.save(str(tmp_path / "cap.docx"))
    doc, errors = process_single(
        tmp_path / "cap.docx", tmp_path / "doc.json",
        parser_name="fallback", max_chars=200,
        write_json=True)
    assert errors == []
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["inspect-doc",
                   str(tmp_path / "doc.json")])
    return rc, buf.getvalue()


# ---------- counts 行真实形态 ----------

def test_caption_doc_counts_line_batch275(tmp_path,
                                           capsys):
    rc, out = _inspect(tmp_path, capsys)
    assert rc == 0
    assert "counts:      elements=6 chunks=4" in out


# ---------- dict 分支 caption 排首 ----------

def test_ecbt_caption_first_batch275(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert ("  element_count_by_type"
            "                caption=2, image=2,"
            " paragraph=2  (ok)") in out


# ---------- figure_caption 在最接近能用的文档上仍黑 ----------

def test_figure_caption_dark_with_captions_batch275(
        tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert ("  figure_caption_f1"
            "                    null  "
            "(parser_does_not_emit_relations)") in out


# ---------- 题注不是 heading ----------

def test_heading_null_no_heading_batch275(tmp_path,
                                           capsys):
    _, out = _inspect(tmp_path, capsys)
    assert ("  heading_boundary_compliance"
            "          null  "
            "(no_heading_elements)") in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch275():
    src = _src()
    assert ("def _format_metric(name: str, "
            "metric: dict)") in src
    assert "sorted(value.items())" in src


# ---------- forbidden tokens 第五百四十七批 ----------

def test_source_no_eval_batch275():
    assert "eval(" not in _src()


def test_source_no_exec_batch275():
    assert "exec(" not in _src()


def test_source_no_compile_batch275():
    assert "compile(" not in _src()


def test_source_no_globals_batch275():
    assert "globals(" not in _src()


def test_source_no_locals_batch275():
    assert "locals(" not in _src()


def test_source_no_os_system_batch275():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch275():
    assert "subprocess" not in _src()


def test_source_no_popen_batch275():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch275():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch275():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch275():
    assert "socket" not in _src()


def test_source_no_requests_batch275():
    assert "requests" not in _src()


def test_source_no_urllib_batch275():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch275():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch275():
    assert "yield" not in _src()


def test_source_no_async_await_batch275():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch275():
    assert _src().count("open(") == 1
