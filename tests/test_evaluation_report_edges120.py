"""evaluation/report.py 第四百九十七轮 edges 测试（Round 1053）。

补强 edges119 未触及的角度（第四百二十九批，probe 实证）。

新角度（聚合键闭包：20 进 14 出，余 6 键永不出席）：
- 真实三类型 docx（heading+2 段+table）全 run 的
  per_doc metrics 恰 20 键全名册首次锁定（含
  text_char_multiset_precision/recall）
- 汇总触达 14 键 = ra 12 + counts 1 +
  success_rates 1；**余 6 键**（element_count_by_type、
  error_code、figure_caption_*、silent_drop_count）
  在 summary 任何 section 都不出现——聚合键宇宙
  闭包一次锁死
- silent_drop_count 不进任何 dict 但直喂顶层
  silent_drop_total（真实 expectations 欠供 3 →
  total 3）
- 真实 heading_boundary_compliance 1.0 流入 ra
  {macro 1.0, participating 1, not_evaluated 0}
- forbidden tokens 第五百二十四批（open 0）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.report as rpt
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("CCC third paragraph body.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "cell one"
    d.save(str(tmp_path / "samples" / "a.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/a.docx",
            "source_type": "docx",
            "expectations": {"element_count_by_type":
                             {"paragraph": 5}}}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json", max_chars=200)


_RA = {"chunk_boundary_f1", "chunk_boundary_precision",
       "chunk_boundary_recall", "chunk_reference_intact_ratio",
       "docx_locator_valid_ratio",
       "heading_boundary_compliance",
       "image_resource_exists_ratio",
       "pdf_locator_valid_ratio", "schema_valid",
       "text_char_multiset_precision",
       "text_char_multiset_recall",
       "text_preservation_equal"}


# ---------- 20 键全名册 ----------

def test_metrics_twenty_keys_batch251(tmp_path):
    rep = _run(tmp_path)
    assert sorted(rep["per_doc"][0]["metrics"]) == [
        "chunk_boundary_f1", "chunk_boundary_precision",
        "chunk_boundary_recall",
        "chunk_reference_intact_ratio",
        "docx_locator_valid_ratio",
        "element_count_by_type", "element_count_total",
        "error_code", "figure_caption_f1",
        "figure_caption_precision",
        "figure_caption_recall",
        "heading_boundary_compliance",
        "image_resource_exists_ratio",
        "pdf_locator_valid_ratio", "pipeline_success",
        "schema_valid", "silent_drop_count",
        "text_char_multiset_precision",
        "text_char_multiset_recall",
        "text_preservation_equal"]


# ---------- ra 恰 12 键 + 真实 heading 参与 ----------

def test_ratio_twelve_real_heading_batch251(tmp_path):
    rep = _run(tmp_path)
    ra = rep["summary"]["ratio_macro_averages"]
    assert set(ra) == _RA
    assert ra["heading_boundary_compliance"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 0}


# ---------- 余 6 键闭包 ----------

def test_residual_six_never_aggregated_batch251(tmp_path):
    rep = _run(tmp_path)
    s = rep["summary"]
    touched = (set(s["counts"]) | set(s["success_rates"])
               | set(s["ratio_macro_averages"]))
    residual = {"element_count_by_type", "error_code",
                "figure_caption_f1",
                "figure_caption_precision",
                "figure_caption_recall",
                "silent_drop_count"}
    assert residual & touched == set()
    assert (set(rep["per_doc"][0]["metrics"])
            - residual) - {"pipeline_success",
                           "element_count_total"} == _RA


# ---------- silent_drop 直喂 total ----------

def test_silent_drop_feeds_total_batch251(tmp_path):
    rep = _run(tmp_path)
    m = rep["per_doc"][0]["metrics"]
    assert m["silent_drop_count"] == {"value": 3,
                                      "reason": None}
    assert rep["summary"]["silent_drop_total"] == 3
    assert "silent_drop_count" not in \
        rep["summary"]["ratio_macro_averages"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(rpt)


def test_source_key_lines_batch251():
    src = _src()
    assert "_RATIO_METRICS" in src
    assert "silent_drop_total" in src
    assert "def aggregate_summary(" in src


# ---------- forbidden tokens 第五百二十四批 ----------

def test_source_no_eval_batch251():
    assert "eval(" not in _src()


def test_source_no_exec_batch251():
    assert "exec(" not in _src()


def test_source_no_compile_batch251():
    assert "compile(" not in _src()


def test_source_no_globals_batch251():
    assert "globals(" not in _src()


def test_source_no_locals_batch251():
    assert "locals(" not in _src()


def test_source_no_os_system_batch251():
    assert "os.system" not in _src()


def test_source_no_popen_batch251():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch251():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch251():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch251():
    assert "socket" not in _src()


def test_source_no_requests_batch251():
    assert "requests" not in _src()


def test_source_no_urllib_batch251():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch251():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch251():
    assert "yield" not in _src()


def test_source_no_async_await_batch251():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch251():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch251():
    assert _src().count("subprocess.run") == 2
