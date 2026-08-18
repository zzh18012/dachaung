"""evaluation/runner.py 第六百五十轮 edges 测试（Round 1214）。

补强 edges214 未触及的角度（第五百八十六批，probe 实证）。

新角度（多节 DOCX 编号标题 / 节断空段）：
- **节断空段**——add_section(NEW_
  PAGE) 在元素流中落一个 content=
  "(空段落)" 的 paragraph 元素（占位
  内容非空串首锁）
- **section 恒 0**——节后段落
  paragraph_index 连续（5、6）但
  section 仍是 0（单文档流不分节计
  数首锁）
- **标题三块**——mc120：[Title 单
  独块 10 字符 1 源, 1.1+Alpha 54
  字符 2 源, 1.2+Beta+空段+Gamma
  93 字符 4 源]（标题向后成垒向前
  合并，空段照合入）
- **hbc 1.0**——三块首源全是
  heading
- **双锚半查**——"sub." ×2 →
  P/R/F1 全 0.5（一锚近界命中一锚
  远界落空）
- forbidden tokens 第六百八十五批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _write_docx(path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Main Title", 0)
    doc.add_heading("1.1 First Sub Heading", level=2)
    doc.add_paragraph("Alpha paragraph under first sub.")
    doc.add_heading("1.2 Second Sub Heading", level=2)
    doc.add_paragraph("Beta paragraph under second sub.")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Gamma paragraph in section two.")
    doc.save(str(path))


def _docx_path(tmp_path, doc_id):
    (tmp_path / "s").mkdir(exist_ok=True)
    p = tmp_path / "s" / f"{doc_id}.docx"
    _write_docx(p)
    return p


def _board(tmp_path, doc_id, anchors=None):
    _docx_path(tmp_path, doc_id)
    docs = [{"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}]
    if anchors is not None:
        (tmp_path / "a").mkdir(exist_ok=True)
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        docs[0]["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 节断空段 ----------

def test_section_break_empty_para_batch412(tmp_path):
    doc, errors = process_single(
        _docx_path(tmp_path, "sb"), tmp_path / "o.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "heading", "heading", "paragraph",
        "heading", "paragraph", "paragraph", "paragraph"]
    assert els[5]["content"] == "(空段落)"
    assert els[6]["content"] == "Gamma paragraph in section two."


def test_section_field_stays_zero_batch412(tmp_path):
    doc, errors = process_single(
        _docx_path(tmp_path, "sz"), tmp_path / "o.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["source_locator"]["section"]
            for e in els] == [0] * 7
    assert [e["source_locator"]["paragraph_index"]
            for e in els] == list(range(7))


# ---------- 标题三块 ----------

def test_chunks_three_batch412(tmp_path):
    doc, errors = process_single(
        _docx_path(tmp_path, "c3"), tmp_path / "o.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    chunks = doc.to_dict()["chunks"]
    assert [len(c["text"]) for c in chunks] == [10, 54, 93]
    assert [len(c["source_element_ids"]) for c in chunks] == \
        [1, 2, 4]
    assert chunks[0]["text"] == "Main Title"
    assert chunks[2]["text"].startswith("1.2 Second Sub Heading")


def test_heading_first_sources_batch412(tmp_path):
    doc, errors = process_single(
        _docx_path(tmp_path, "hf"), tmp_path / "o.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    d = doc.to_dict()
    by_id = {e["element_id"]: e for e in d["elements"]}
    for c in d["chunks"]:
        first = by_id[c["source_element_ids"][0]]
        assert first["type"] == "heading"


# ---------- 锚 ----------

def test_anchor_sub_twice_batch412(tmp_path):
    r = run_evaluation(_board(tmp_path, "a1", [
        {"marker": "sub.", "position": "after"},
        {"marker": "sub.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.5, "reason": None}


# ---------- 指标 ----------

def test_metrics_docx_batch412(tmp_path):
    r = run_evaluation(_board(tmp_path, "mx"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=120)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"heading": 3, "paragraph": 4},
        "reason": None}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch412():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百八十五批 ----------

def test_source_no_eval_batch412():
    assert "eval(" not in _src()


def test_source_no_exec_batch412():
    assert "exec(" not in _src()


def test_source_no_compile_batch412():
    assert "compile(" not in _src()


def test_source_no_globals_batch412():
    assert "globals(" not in _src()


def test_source_no_locals_batch412():
    assert "locals(" not in _src()


def test_source_no_os_system_batch412():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch412():
    assert "subprocess" not in _src()


def test_source_no_popen_batch412():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch412():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch412():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch412():
    assert "socket" not in _src()


def test_source_no_requests_batch412():
    assert "requests" not in _src()


def test_source_no_urllib_batch412():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch412():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch412():
    assert "yield" not in _src()


def test_source_no_async_await_batch412():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch412():
    assert _src().count("open(") == 2
