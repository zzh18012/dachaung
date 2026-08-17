"""evaluation/manifest.py 第五百三十八轮 edges 测试（Round 1094）。

补强 edges133-135 未触及的角度（第四百七十批，probe 实证）。

新角度（鬼影路径全链路：加载照收 / 运行才爆）：
- **load 不查存在性**：samples/ghost.docx 不存在，
  load_manifest 照常返回——存在性执法在 run 侧
  （edges106 的 ghost 是 resolution 单元；这里走完
  documents 台账全链）
- **鬼影文档三件套**：pipeline_success False /
  error_code 'file_not_found' / element_count_total
  null+pipeline_failed——失败也有结构化指标
- **鬼影 ef 期望命中**：expected file_not_found →
  matches True（4-key 精确 entry）——缺失是被
  预期的失败
- **期望错码即失配**：expected docx_open_failed →
  matches False——missing ≠ corrupt，错误码语义
  有区分度
- **混合板半成功率**：good + ghost 双文档 →
  success {1, 2, 0.5}，per-doc codes
  [None, 'file_not_found']——真实台账的失败摊薄
- forbidden tokens 第五百六十五批（open 1）
"""

from __future__ import annotations

import inspect
import json
import pathlib

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _good_docx(p: pathlib.Path) -> None:
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("BBB second paragraph body.")
    d.save(str(p))


def _board(tmp_path, docs, ef=None):
    (tmp_path / "samples").mkdir(exist_ok=True)
    for doc in docs:
        if doc.get("real"):
            _good_docx(tmp_path / doc["path"])
    m = {
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": doc["doc_id"], "path": doc["path"],
             "source_type": "docx"} for doc in docs],
    }
    if ef is not None:
        m["expected_failures"] = ef
    mp = tmp_path / "manifest.json"
    mp.write_text(json.dumps(m), encoding="utf-8")
    loaded = load_manifest(mp, project_root=tmp_path)
    return loaded, mp


def _run(tmp_path, docs, ef=None):
    loaded, _ = _board(tmp_path, docs, ef)
    return run_evaluation(
        loaded, tmp_path / "report.json",
        parser_name="fallback", max_chars=200)


# ---------- 加载不查存在性 ----------

def test_ghost_doc_load_accepted_batch293(tmp_path):
    loaded, _ = _board(tmp_path, [
        {"doc_id": "d1", "path": "samples/ghost.docx",
         "real": False}])
    assert [d.doc_id for d in loaded.documents] == ["d1"]


# ---------- 鬼影文档运行三件套 ----------

def test_ghost_doc_run_trio_batch293(tmp_path):
    rep = _run(tmp_path, [
        {"doc_id": "d1", "path": "samples/ghost.docx",
         "real": False}])
    m = rep["per_doc"][0]["metrics"]
    assert m["pipeline_success"] == {
        "value": False, "reason": None}
    assert m["error_code"] == {
        "value": "file_not_found", "reason": None}
    assert m["element_count_total"] == {
        "value": None, "reason": "pipeline_failed"}
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 0, "total": 1, "rate": 0.0}


# ---------- 鬼影 ef 期望命中 ----------

def test_ghost_ef_matches_batch293(tmp_path):
    rep = _run(tmp_path, [], [
        {"doc_id": "f1", "path": "samples/ghost.docx",
         "source_type": "docx",
         "expected_error_code": "file_not_found"}])
    assert rep["expected_failures"] == [{
        "doc_id": "f1",
        "expected_error_code": "file_not_found",
        "actual_error_code": "file_not_found",
        "matches": True}]


# ---------- 期望错码即失配 ----------

def test_ghost_ef_wrong_code_batch293(tmp_path):
    rep = _run(tmp_path, [], [
        {"doc_id": "f1", "path": "samples/ghost.docx",
         "source_type": "docx",
         "expected_error_code": "docx_open_failed"}])
    assert rep["expected_failures"] == [{
        "doc_id": "f1",
        "expected_error_code": "docx_open_failed",
        "actual_error_code": "file_not_found",
        "matches": False}]


# ---------- 混合板半成功率 ----------

def test_mixed_ghost_board_batch293(tmp_path):
    rep = _run(tmp_path, [
        {"doc_id": "good1", "path": "samples/good.docx",
         "real": True},
        {"doc_id": "ghost1", "path": "samples/ghost.docx",
         "real": False}])
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 1, "total": 2, "rate": 0.5}
    codes = [d["metrics"]["error_code"]["value"]
             for d in rep["per_doc"]]
    assert codes == [None, "file_not_found"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch293():
    src = _src()
    assert "def _has_backslash(" in src
    assert "def _is_absolute_like(" in src


# ---------- forbidden tokens 第五百六十五批 ----------

def test_source_no_eval_batch293():
    assert "eval(" not in _src()


def test_source_no_exec_batch293():
    assert "exec(" not in _src()


def test_source_no_compile_batch293():
    assert "compile(" not in _src()


def test_source_no_globals_batch293():
    assert "globals(" not in _src()


def test_source_no_locals_batch293():
    assert "locals(" not in _src()


def test_source_no_os_system_batch293():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch293():
    assert "subprocess" not in _src()


def test_source_no_popen_batch293():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch293():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch293():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch293():
    assert "socket" not in _src()


def test_source_no_requests_batch293():
    assert "requests" not in _src()


def test_source_no_urllib_batch293():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch293():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch293():
    assert "yield" not in _src()


def test_source_no_async_await_batch293():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch293():
    assert _src().count("open(") == 1
