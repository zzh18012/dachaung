"""evaluation/runner.py 第四百九十八轮 edges 测试（Round 1054）。

补强 edges130 未触及的角度（第四百三十批，probe 实证）。

新角度（双真实 docx 并跑：成功而无残骸）：
- 两份真实 python-docx 文档（3 类型富板 + 2 段素板）
  同一 run 穿真实管线双双成功——最真实的双文档
  评测形态首次在 runner 层锁定
- 成功路径 _per_doc **空**：doc stub 双双 unlink、
  真实 docx 无图片 → 无 images 目录 → 成功不留
  任何残骸（与 R1034 patch 板留 images-<hash16>
  正交、与 R1047 失败路径空目录互为成因镜像：
  一个"写了又删"、一个"从未写"）
- 汇总屏：counts {sum 5, participating 2}（3+2 跨
  文档求和）、rate {2, 2, 1.0}、docx_locator
  {1.0, 2, 0} 双参与；报告 RS 有效；双 wall_time
  total 均 float
- forbidden tokens 第五百二十五批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate_file


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    a = Document()
    a.add_heading("Real Title", level=1)
    a.add_paragraph("AAA first paragraph body.")
    t = a.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "cell"
    a.save(str(tmp_path / "samples" / "rich.docx"))
    b = Document()
    b.add_paragraph("Plain one.")
    b.add_paragraph("Plain two.")
    b.save(str(tmp_path / "samples" / "plain.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/rich.docx",
             "source_type": "docx"},
            {"doc_id": "d2", "path": "samples/plain.docx",
             "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    out = tmp_path / "o.json"
    rep = run_evaluation(load_manifest(mf, tmp_path), out,
                         max_chars=200)
    return rep, out, tmp_path / "_per_doc"


# ---------- 双真实文档双双成功 ----------

def test_two_real_docs_both_succeed_batch252(tmp_path):
    rep, _, _ = _run(tmp_path)
    board = [(r["doc_id"],
              r["metrics"]["element_count_total"]["value"],
              r["metrics"]["pipeline_success"]["value"])
             for r in rep["per_doc"]]
    assert board == [("d1", 3, True), ("d2", 2, True)]


# ---------- 成功而无残骸 ----------

def test_success_leaves_no_artifacts_batch252(tmp_path):
    _, _, per = _run(tmp_path)
    assert per.is_dir()
    assert sorted(p.name for p in per.iterdir()) == []


# ---------- 汇总屏 ----------

def test_two_doc_summary_batch252(tmp_path):
    rep, _, _ = _run(tmp_path)
    s = rep["summary"]
    assert s["counts"]["element_count_total"] == {
        "sum": 5, "participating_docs": 2}
    assert s["success_rates"]["pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}
    assert s["ratio_macro_averages"][
        "docx_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 2,
        "not_evaluated": 0}


def test_report_rs_and_wall_batch252(tmp_path):
    rep, out, _ = _run(tmp_path)
    validate_file(out, "evaluation-report.schema.json")
    for r in rep["per_doc"]:
        assert isinstance(r["wall_time_seconds"]["total"],
                          float)


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch252():
    src = _src()
    assert "out_stub.unlink()" in src
    assert "except OSError:" in src
    assert "write_json=False" in src


# ---------- forbidden tokens 第五百二十五批 ----------

def test_source_no_eval_batch252():
    assert "eval(" not in _src()


def test_source_no_exec_batch252():
    assert "exec(" not in _src()


def test_source_no_compile_batch252():
    assert "compile(" not in _src()


def test_source_no_globals_batch252():
    assert "globals(" not in _src()


def test_source_no_locals_batch252():
    assert "locals(" not in _src()


def test_source_no_os_system_batch252():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch252():
    assert "subprocess" not in _src()


def test_source_no_popen_batch252():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch252():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch252():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch252():
    assert "socket" not in _src()


def test_source_no_requests_batch252():
    assert "requests" not in _src()


def test_source_no_urllib_batch252():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch252():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch252():
    assert "yield" not in _src()


def test_source_no_async_await_batch252():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch252():
    assert _src().count("open(") == 2
