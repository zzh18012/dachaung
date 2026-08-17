"""evaluation/report.py 第五百五十七轮 edges 测试（Round 1113）。

补强 edges127 未触及的角度（第四百八十九批，probe 实证）。

新角度（双 parser 指标层全等 / summary 宽严分歧）：
- **双 parser 指标层全等**：同一 manifest 真跑 fallback 与
  kreuzberg → summary 与 per_doc 两个 dict 逐字相等——
  parser 差异止步于 provenance，聚合层完全可互换（旧锁
  在 CLI 层只比大四指标，本批锁全 dict）
- **summary 宽严分歧**：summary.counts 塞 {"bogus":
  "garbage-string"} → 报告照过——counts/success_rates/
  ratio_macro_averages 在 schema 里只是 type object，内层
  结构不设防；同 summary 的 silent_drop_total 塞串 → 拒
  "'not-an-int' is not of type 'integer', 'null'"——
  四键中唯一强类型（宽严同屏首锁）
- forbidden tokens 第五百八十五批（open 0，报告变体
  15 条 + subprocess.run 计数 2）
"""

from __future__ import annotations

import copy
import inspect
import json

from docx import Document

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import EvalSchemaError, validate


def _manifest(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA kreuzberg probe body.")
    d.add_paragraph("BBB second line.")
    d.save(str(tmp_path / "samples" / "g.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/g.docx",
            "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 双 parser 指标层全等 ----------

def test_dual_parser_metric_layer_equality_batch312(tmp_path):
    man = _manifest(tmp_path)
    fb = run_evaluation(man, tmp_path / "r-fb.json",
                        parser_name="fallback", max_chars=200)
    kz = run_evaluation(man, tmp_path / "r-kz.json",
                        parser_name="kreuzberg", max_chars=200)
    assert fb["summary"] == kz["summary"]
    assert ([p["metrics"] for p in fb["per_doc"]] ==
            [p["metrics"] for p in kz["per_doc"]])
    assert fb["devset"] == kz["devset"]
    assert (fb["provenance"]["parser_name"] ==
            "fallback")
    assert (kz["provenance"]["parser_name"] ==
            "kreuzberg")
    assert (kz["provenance"]["parser_version"] ==
            "4.10.2")


# ---------- summary 宽严分歧 ----------

def test_summary_loose_counts_typed_silent_batch312(tmp_path):
    man = _manifest(tmp_path)
    r = run_evaluation(man, tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    loose = copy.deepcopy(r)
    loose["summary"]["counts"] = {
        "bogus": "garbage-string"}
    validate(loose, "evaluation-report.schema.json")
    typed = copy.deepcopy(r)
    typed["summary"]["silent_drop_total"] = "not-an-int"
    try:
        validate(typed, "evaluation-report.schema.json")
        raised = False
    except EvalSchemaError as e:
        raised = True
        assert "'not-an-int' is not of type 'integer', " \
            "'null'" in str(e)
        assert "path=['summary', 'silent_drop_total']" \
            in str(e)
    assert raised


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch312():
    src = _src()
    assert "成功数 + 成功率" in src
    assert "从 Manifest 对象提取" in src


# ---------- forbidden tokens 第五百八十五批（报告变体） ----------

def test_source_no_eval_batch312():
    assert "eval(" not in _src()


def test_source_no_exec_batch312():
    assert "exec(" not in _src()


def test_source_no_compile_batch312():
    assert "compile(" not in _src()


def test_source_no_globals_batch312():
    assert "globals(" not in _src()


def test_source_no_locals_batch312():
    assert "locals(" not in _src()


def test_source_no_os_system_batch312():
    assert "os.system" not in _src()


def test_source_no_popen_batch312():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch312():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch312():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch312():
    assert "socket" not in _src()


def test_source_no_requests_batch312():
    assert "requests" not in _src()


def test_source_no_urllib_batch312():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch312():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch312():
    assert "yield" not in _src()


def test_source_no_async_await_batch312():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch312():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch312():
    assert _src().count("subprocess.run") == 2
