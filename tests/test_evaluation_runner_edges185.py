"""evaluation/runner.py 第六百一十五轮 edges 测试（Round 1171）。

补强 edges184 未触及的角度（第五百四十三批，probe 实证）。

新角度（DOCX 列表样式流 / 重复 marker 首现）：
- **列表样式即段落**——List Bullet / List Number
  段落 → type=paragraph、style 原样入 metadata
  （'List Bullet'/'List Number'，level 0）——列
  表不成独立型（首锁）
- **列表全流**——intro+2 bullet+2 numbered 五段
  同 sequential 块 5 源；heading 软界切出第 2
  块 [heading+outro] 2 源
- **题界命中**——marker "two." after → 恰落 heading
  软界 → P/R/F1 全 1.0
- **重复 marker 首现**——"item." 出现两次（First/
  Second bullet），顺序搜索取首现 → GT 落块中
  距界 > tol → 全 0.0（首现语义首锁）
- forbidden tokens 第六百四十三批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None):
    (tmp_path / "s").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Intro before the list.")
    d.add_paragraph("First bullet item.", style="List Bullet")
    d.add_paragraph("Second bullet item.", style="List Bullet")
    d.add_paragraph("Numbered step one.", style="List Number")
    d.add_paragraph("Numbered step two.", style="List Number")
    d.add_heading("After Lists", level=1)
    d.add_paragraph("Outro after the lists.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    docs = [{"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}]
    if anchors is not None:
        (tmp_path / "a").mkdir(exist_ok=True)
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}), encoding="utf-8")
        docs[0]["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 列表样式即段落 ----------

def test_list_styles_element_meta_batch369(tmp_path):
    _board(tmp_path, "ls")
    doc, errors = process_single(
        tmp_path / "s" / "ls.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == ["paragraph"] * 5 + [
        "heading", "paragraph"]
    assert [e["metadata"]["style"] for e in els] == [
        "Normal", "List Bullet", "List Bullet",
        "List Number", "List Number", "Heading 1", "Normal"]
    assert all(e["metadata"]["level"] == 0
               for e in els if e["type"] == "paragraph")
    assert [e["source_locator"]["paragraph_index"]
            for e in els] == list(range(7))


# ---------- 列表全流 ----------

def test_list_flow_chunks_batch369(tmp_path):
    _board(tmp_path, "ls2")
    doc, errors = process_single(
        tmp_path / "s" / "ls2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "sequential"]
    assert len(chunks[0]["source_element_ids"]) == 5
    assert chunks[0]["text"] == (
        "Intro before the list. First bullet item. "
        "Second bullet item. Numbered step one. "
        "Numbered step two.")
    assert len(chunks[1]["source_element_ids"]) == 2
    assert chunks[1]["text"] == \
        "After Lists Outro after the lists."


# ---------- 题界命中 ----------

def test_list_junction_batch369(tmp_path):
    r = run_evaluation(_board(tmp_path, "ls3", [
        {"marker": "two.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


# ---------- 重复 marker 首现 ----------

def test_duplicate_marker_first_match_batch369(tmp_path):
    r = run_evaluation(_board(tmp_path, "ls4", [
        {"marker": "item.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.0, "reason": None}


# ---------- 指标 ----------

def test_list_by_type_batch369(tmp_path):
    r = run_evaluation(_board(tmp_path, "ls5"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 6, "heading": 1},
        "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch369():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("annotation") == 10
    assert src.count("chunk") == 9


# ---------- forbidden tokens 第六百四十三批 ----------

def test_source_no_eval_batch369():
    assert "eval(" not in _src()


def test_source_no_exec_batch369():
    assert "exec(" not in _src()


def test_source_no_compile_batch369():
    assert "compile(" not in _src()


def test_source_no_globals_batch369():
    assert "globals(" not in _src()


def test_source_no_locals_batch369():
    assert "locals(" not in _src()


def test_source_no_os_system_batch369():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch369():
    assert "subprocess" not in _src()


def test_source_no_popen_batch369():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch369():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch369():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch369():
    assert "socket" not in _src()


def test_source_no_requests_batch369():
    assert "requests" not in _src()


def test_source_no_urllib_batch369():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch369():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch369():
    assert "yield" not in _src()


def test_source_no_async_await_batch369():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch369():
    assert _src().count("open(") == 2
