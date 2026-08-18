"""evaluation/runner.py 第五百六十八轮 edges 测试（Round 1124）。

补强 edges142 未触及的角度（第五百批，probe 实证）。

新角度（负容差真跑）：
- **负容差全零**——tolerance_chars=-5 真跑：P/R/F1 三通道
  全 {0.0, reason None}——d <= -5 永不成立，精确命中也判
  失配（runner 层首锁；annotation_metrics 直调层旧锁
  edges104/112）
- **tol 0 对照保留部分匹配**——同一板 tol 0 → P 0.5 /
  R 1.0 / F1 0.6667——0 与负数的分界不是"更严"而是
  "全灭"，两档并跑显分界
- **负容差报告照过 Schema**——tol -5 报告过
  evaluation-report.schema.json——非法语义不产生非法报告
- forbidden tokens 第五百九十六批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA head.")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "AAA head.", "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1", "path": "samples/g.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _run(tmp_path, tol):
    out = tmp_path / f"r{tol}.json"
    r = run_evaluation(_board(tmp_path), out,
                       parser_name="fallback", max_chars=200,
                       tolerance_chars=tol)
    return r, out


# ---------- 负容差全零 ----------

def test_negative_tolerance_zeroes_all_batch323(tmp_path):
    r, _ = _run(tmp_path, -5)
    m = r["per_doc"][0]["metrics"]
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall",
              "chunk_boundary_f1"):
        assert m[k] == {"value": 0.0, "reason": None}


# ---------- tol 0 对照保留部分匹配 ----------

def test_zero_tolerance_partial_match_batch323(tmp_path):
    r, _ = _run(tmp_path, 0)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 负容差报告照过 Schema ----------

def test_negative_tolerance_report_schema_ok_batch323(tmp_path):
    r, out = _run(tmp_path, -5)
    validate(r, "evaluation-report.schema.json")
    validate(json.loads(out.read_text(encoding="utf-8")),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch323():
    src = _src()
    assert "给 output_path 让 pipeline 推 image_output_dir" in src
    assert "避免硬编码" in src


# ---------- forbidden tokens 第五百九十六批 ----------

def test_source_no_eval_batch323():
    assert "eval(" not in _src()


def test_source_no_exec_batch323():
    assert "exec(" not in _src()


def test_source_no_compile_batch323():
    assert "compile(" not in _src()


def test_source_no_globals_batch323():
    assert "globals(" not in _src()


def test_source_no_locals_batch323():
    assert "locals(" not in _src()


def test_source_no_os_system_batch323():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch323():
    assert "subprocess" not in _src()


def test_source_no_popen_batch323():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch323():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch323():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch323():
    assert "socket" not in _src()


def test_source_no_requests_batch323():
    assert "requests" not in _src()


def test_source_no_urllib_batch323():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch323():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch323():
    assert "yield" not in _src()


def test_source_no_async_await_batch323():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch323():
    assert _src().count("open(") == 2
