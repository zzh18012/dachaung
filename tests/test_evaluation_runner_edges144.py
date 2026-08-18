"""evaluation/runner.py 第五百六十九轮 edges 测试（Round 1125）。

补强 edges143 未触及的角度（第五百零一批，probe 实证）。

新角度（全通道全家桶真跑）：
- **单文档四通道并发**——一个 docx 同时挂：内嵌图片
  （add_picture）+ annotation（marker）+ expectations
  （paragraph 5）+ ef ghost——一次 run_evaluation 四通道
  全活：image ratio 1.0（真图落盘实存）、F1 0.6667（标注
  命中）、silent_drop 2、ef matches True（分通道旧锁各自
  单测，同板并发首锁）
- **images-* 目录真落盘**——输出旁出现 images-<sha> 目录
  （runner 复用 pipeline 命名约定）
- forbidden tokens 第五百九十七批（open 2）
"""

from __future__ import annotations

import base64
import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate

_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
    "nGNgYGBgAAAABQABh6FO1AAAAABJRU5ErkJggg==")


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    (tmp_path / "px.png").write_bytes(_PNG)
    d = Document()
    d.add_paragraph("AAA head start.")
    d.add_picture(str(tmp_path / "px.png"))
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "AAA head start.",
             "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/g.docx",
            "source_type": "docx",
            "annotation_file": "anns/a.json",
            "expectations": {"element_count_by_type":
                             {"paragraph": 5}}}],
        "expected_failures": [{
            "doc_id": "ef1", "path": "samples/ghost.docx",
            "expected_error_code": "file_not_found"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _run(tmp_path):
    out = tmp_path / "r.json"
    r = run_evaluation(_board(tmp_path), out,
                       parser_name="fallback", max_chars=200)
    return r, out


# ---------- 单文档四通道并发 ----------

def test_full_stack_image_channel_batch324(tmp_path):
    r, out = _run(tmp_path)
    m = r["per_doc"][0]["metrics"]
    assert m["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}
    img_dirs = list((out.parent / "_per_doc").glob("images-*"))
    assert len(img_dirs) == 1
    pngs = list(img_dirs[0].glob("image_*_para1_00.png"))
    assert len(pngs) == 1


def test_full_stack_boundary_and_drop_batch324(tmp_path):
    r, _ = _run(tmp_path)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}
    assert m["silent_drop_count"] == {"value": 2, "reason": None}
    assert r["summary"]["silent_drop_total"] == 2


def test_full_stack_ef_and_doc_count_batch324(tmp_path):
    r, _ = _run(tmp_path)
    assert r["expected_failures"][0]["matches"] is True
    assert len(r["per_doc"]) == 1


def test_full_stack_report_validates_batch324(tmp_path):
    r, out = _run(tmp_path)
    validate(r, "evaluation-report.schema.json")
    validate(json.loads(out.read_text(encoding="utf-8")),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch324():
    src = _src()
    assert "跑 process_single，返回" in src
    assert "清单 → 逐文档跑 pipeline" in src


# ---------- forbidden tokens 第五百九十七批 ----------

def test_source_no_eval_batch324():
    assert "eval(" not in _src()


def test_source_no_exec_batch324():
    assert "exec(" not in _src()


def test_source_no_compile_batch324():
    assert "compile(" not in _src()


def test_source_no_globals_batch324():
    assert "globals(" not in _src()


def test_source_no_locals_batch324():
    assert "locals(" not in _src()


def test_source_no_os_system_batch324():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch324():
    assert "subprocess" not in _src()


def test_source_no_popen_batch324():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch324():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch324():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch324():
    assert "socket" not in _src()


def test_source_no_requests_batch324():
    assert "requests" not in _src()


def test_source_no_urllib_batch324():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch324():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch324():
    assert "yield" not in _src()


def test_source_no_async_await_batch324():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch324():
    assert _src().count("open(") == 2
