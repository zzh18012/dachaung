r"""app/parsers/fallback_parser.py 边角测试 - 第三十八轮（Round 1429）。

新角度（probe 实证）现代 Word 结构包装层可见性（历史已锁
footnotes/headers/w:ins，未碰过 sdt 与文本框）：
- w:sdt（内容控件）包住的段落：完全不可见——只有前后正
  文段成元素、无告警
- run 内 w:pict/v:textbox/w:txbxContent 文本框：文本完全
  不可见（run 无 w:t，段落只剩 'before box'）、无告警
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _sdt_docx(tmp_path):
    d = Document()
    d.add_paragraph("before sdt")
    body = d.element.body
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "sdt wrapped paragraph"
    r.append(t)
    p.append(r)
    content.append(p)
    sdt.append(content)
    sectPr = body.find(
        "{http://schemas."
        "openxmlformats.org/"
        "wordprocessingml/2006/"
        "main}sectPr")
    sectPr.addprevious(sdt)
    p_ = tmp_path / "sdt.docx"
    d.save(str(p_))
    return p_


def _tb_docx(tmp_path):
    d = Document()
    para = d.add_paragraph("before box")
    run = para.add_run()
    xml = (
        '<w:pict %s '
        'xmlns:v="urn:schemas-'
        'microsoft-com:vml">'
        '<v:shape style='
        '"width:100pt;'
        'height:50pt">'
        '<v:textbox>'
        '<w:txbxContent>'
        '<w:p><w:r><w:t>'
        'text inside box'
        '</w:t></w:r></w:p>'
        '</w:txbxContent>'
        '</v:textbox>'
        '</v:shape>'
        '</w:pict>' % nsdecls('w'))
    run._r.append(parse_xml(xml))
    p = tmp_path / "tb.docx"
    d.save(str(p))
    return p


# ---------- w:sdt ----------

def test_sdt_invisible(tmp_path):
    p = _sdt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "before sdt"]


def test_sdt_text_nowhere(
        tmp_path):
    p = _sdt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    for e in doc.elements:
        assert "wrapped" \
            not in e.content
        assert "sdt wrapped" \
            not in e.content


def test_sdt_no_warnings(
        tmp_path):
    p = _sdt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_sdt_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _sdt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_sdt_pipeline_chunk(
        tmp_path):
    p = _sdt_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "before sdt"]


# ---------- 文本框 ----------

def test_tb_invisible(tmp_path):
    p = _tb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "before box"]


def test_tb_text_nowhere(
        tmp_path):
    p = _tb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    for e in doc.elements:
        assert "box" not in \
            e.content.replace(
                "before box", "")
        assert "inside" \
            not in e.content


def test_tb_no_warnings(
        tmp_path):
    p = _tb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_tb_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _tb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_tb_pipeline_chunk(
        tmp_path):
    p = _tb_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "before box"]
