"""evaluation/cli.py 第五百五十五轮 edges 测试（Round 1111）。

补强 edges138 未触及的角度（第四百八十七批，probe 实证）。

新角度（跨 schema 拒收 / inspect 零校验 / CLI 容差真值）：
- **manifest 喂 validate-report**：合法清单 JSON 喂报告
  校验子命令 → rc 1 + [FAIL] + "'report_version' is a
  required property @ path=[]" + 共 6 处——清单对报告
  schema 是六错齐发（跨 schema 拒收首锁）
- **inspect-doc 零校验**：评测报告 JSON 喂 inspect-doc →
  rc 0 照跑：type=unknown / elements=0 / document_id: ?——
  inspect-doc 不做任何 doc schema 校验（与 run 的两重
  schema 纪律分歧，本批锁定现状）
- **CLI 容差真值翻转**：--tolerance-chars 0 vs 30 同板
  同锚跑 run → 报告 f1 0.0 vs 0.6667——CLI 旗标不只是
  透传捕参（edges100 已锁 mock 捕参），本批锁真值入报告
- forbidden tokens 第五百八十三批（open 1）
"""

from __future__ import annotations

import inspect
import io
import contextlib
import json

from docx import Document

import evaluation.cli as cli_mod
from evaluation.cli import main


def _manifest(tmp_path, with_annotation):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("head TAIL")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    doc = {"doc_id": "d1", "path": "samples/g.docx",
           "source_type": "docx"}
    if with_annotation:
        (tmp_path / "anns").mkdir(exist_ok=True)
        (tmp_path / "anns" / "a.json").write_text(
            json.dumps({
                "annotation_version": "1.0", "doc_id": "d1",
                "chunk_boundary_anchors": [
                    {"marker": "head",
                     "position": "before"}]}),
            encoding="utf-8")
        doc["annotation_file"] = "anns/a.json"
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [doc]}), encoding="utf-8")
    return mf


# ---------- manifest 喂 validate-report ----------

def test_manifest_to_validate_report_rejected_batch310(
        tmp_path, capsys):
    mf = _manifest(tmp_path, with_annotation=False)
    rc = main(["validate-report", str(mf)])
    assert rc == 1
    err = capsys.readouterr().err
    assert "[FAIL]" in err
    assert "'report_version' is a required property" in err
    assert "path=[]" in err
    assert "6 处" in err


# ---------- inspect-doc 零校验 ----------

def test_inspect_doc_accepts_report_batch310(tmp_path):
    mf = _manifest(tmp_path, with_annotation=False)
    rc = main(["run", "--manifest", str(mf),
               "--output", str(tmp_path / "r.json")])
    assert rc == 0
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc2 = main(["inspect-doc", str(tmp_path / "r.json")])
    assert rc2 == 0
    out = buf.getvalue()
    assert "type=unknown" in out
    assert "elements=0" in out
    assert "document_id: ?" in out


# ---------- CLI 容差真值翻转 ----------

def test_cli_tolerance_flip_batch310(tmp_path):
    mf = _manifest(tmp_path, with_annotation=True)
    for tol, want in ((0, 0.0), (30, 0.6666666666666666)):
        out = tmp_path / f"r{tol}.json"
        rc = main(["run", "--manifest", str(mf),
                   "--output", str(out),
                   "--max-chars", "200",
                   "--tolerance-chars", str(tol)])
        assert rc == 0
        blob = json.loads(out.read_text(encoding="utf-8"))
        assert blob["per_doc"][0]["metrics"][
            "chunk_boundary_f1"] == {
            "value": want, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch310():
    src = _src()
    assert "清单加载失败" in src
    assert "报告自校验失败" in src


# ---------- forbidden tokens 第五百八十三批 ----------

def test_source_no_eval_batch310():
    assert "eval(" not in _src()


def test_source_no_exec_batch310():
    assert "exec(" not in _src()


def test_source_no_compile_batch310():
    assert "compile(" not in _src()


def test_source_no_globals_batch310():
    assert "globals(" not in _src()


def test_source_no_locals_batch310():
    assert "locals(" not in _src()


def test_source_no_os_system_batch310():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch310():
    assert "subprocess" not in _src()


def test_source_no_popen_batch310():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch310():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch310():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch310():
    assert "socket" not in _src()


def test_source_no_requests_batch310():
    assert "requests" not in _src()


def test_source_no_urllib_batch310():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch310():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch310():
    assert "yield" not in _src()


def test_source_no_async_await_batch310():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch310():
    assert _src().count("open(") == 1
