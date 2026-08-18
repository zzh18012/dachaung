"""evaluation/runner.py 第六百一十八轮 edges 测试（Round 1174）。

补强 edges187 未触及的角度（第五百四十六批，probe 实证）。

新角度（连排双题 / 尾题收块）：
- **连排双题**——H1 后紧跟 H2 再接正文 → chunks
  [seq("H One" 单源), seq("H Two"+正文 双源)]——
  第二题把第一题切出独立块（题对题也触发软界，
  首锁）；两题 level 1/2 各自保留
- **尾题收块**——文档以 heading 收尾（后无正文）→
  尾题自成末块 [seq(Intro), seq(Trailing Heading)]
  （文档末 flush 兜底首锁）
- **compliance 恒 1.0**——两种题排布下
  heading_boundary_compliance 均 1.0（题永居块首
  的结构不变量双证）
- forbidden tokens 第六百四十六批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, paras):
    (tmp_path / "s").mkdir(exist_ok=True)
    d = Document()
    for kind, text, lvl in paras:
        if kind == "h":
            d.add_heading(text, level=lvl)
        else:
            d.add_paragraph(text)
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"s/{doc_id}.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


_HH = [("h", "H One", 1), ("h", "H Two", 2),
       ("p", "Body after two headings.", None)]
_TH = [("p", "Intro paragraph first.", None),
       ("h", "Trailing Heading", 1)]


# ---------- 连排双题 ----------

def test_consecutive_headings_chunks_batch372(tmp_path):
    _board(tmp_path, "hh", _HH)
    doc, errors = process_single(
        tmp_path / "s" / "hh.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    dd = doc.to_dict()
    assert [(e["type"], e["metadata"]["level"])
            for e in dd["elements"]] == [
        ("heading", 1), ("heading", 2), ("paragraph", 0)]
    chunks = dd["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "sequential"]
    assert chunks[0]["text"] == "H One"
    assert len(chunks[0]["source_element_ids"]) == 1
    assert chunks[1]["text"] == \
        "H Two Body after two headings."
    assert len(chunks[1]["source_element_ids"]) == 2


def test_consecutive_headings_metrics_batch372(tmp_path):
    r = run_evaluation(_board(tmp_path, "hh2", _HH),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"heading": 2, "paragraph": 1},
        "reason": None}


# ---------- 尾题收块 ----------

def test_trailing_heading_chunks_batch372(tmp_path):
    _board(tmp_path, "th", _TH)
    doc, errors = process_single(
        tmp_path / "s" / "th.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "sequential"]
    assert chunks[0]["text"] == "Intro paragraph first."
    assert chunks[1]["text"] == "Trailing Heading"
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks)


def test_trailing_heading_metrics_batch372(tmp_path):
    r = run_evaluation(_board(tmp_path, "th2", _TH),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 1, "heading": 1},
        "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch372():
    src = _src()
    assert src.count("metrics") == 13
    assert src.count("chunk") == 9
    assert src.count("process_single") == 6


# ---------- forbidden tokens 第六百四十六批 ----------

def test_source_no_eval_batch372():
    assert "eval(" not in _src()


def test_source_no_exec_batch372():
    assert "exec(" not in _src()


def test_source_no_compile_batch372():
    assert "compile(" not in _src()


def test_source_no_globals_batch372():
    assert "globals(" not in _src()


def test_source_no_locals_batch372():
    assert "locals(" not in _src()


def test_source_no_os_system_batch372():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch372():
    assert "subprocess" not in _src()


def test_source_no_popen_batch372():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch372():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch372():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch372():
    assert "socket" not in _src()


def test_source_no_requests_batch372():
    assert "requests" not in _src()


def test_source_no_urllib_batch372():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch372():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch372():
    assert "yield" not in _src()


def test_source_no_async_await_batch372():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch372():
    assert _src().count("open(") == 2
