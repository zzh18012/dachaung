r"""evaluation/cli.py 边角第一百八十三轮（Round 1409）。

新角度（R1404-R1407 真实文件行为在评测层的组合）：
TJ 字距 PDF + 非零 MediaBox PDF + 嵌套表 docx + 分节
docx 四文件 manifest 穿 run → 报告 summary 精确值：
- success 4/4 rate 1.0；ect_sum 10 参与 4；sdt 0
- pdfloc/docxloc 各参与 2（类型互斥 not_evaluated 2）
- hbc 只参与有 heading 的 2 个 PDF（nest/sec 无 heading
  → not_evaluated 2）
- irer null no_image_elements；cbp/cbr/cbf null no_annotation
- validate-report rc 0
"""

from __future__ import annotations

import contextlib
import io
import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

from evaluation.cli import main


def _xref_pdf(objs):
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    maxid = max(objs)
    out += (f"xref\n0 {maxid + 1}\n"
            ).encode() \
        + b"0000000000 65535 f \n"
    for oid in range(1, maxid + 1):
        out += (
            f"{offsets[oid]:010d} "
            f"00000 n \n".encode()
            if oid in offsets
            else b"0000000000 65535 f \n")
    out += (f"trailer\n<< /Size "
            f"{maxid + 1} /Root 1 0 R "
            f">>\nstartxref\n{xref_pos}"
            f"\n%%EOF").encode()
    return bytes(out)


def _simple_page(content,
                 box="[0 0 612 792]"):
    return _xref_pdf({
        6: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] "
            b"/Count 1 >>"),
        3: (b"<< /Type /Page /Parent "
            b"2 0 R /MediaBox "
            + box.encode()
            + b" /Resources << /Font "
            b"<< /F1 6 0 R >> >> "
            b"/Contents 4 0 R >>"),
        4: (b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n"
            + content
            + b"\nendstream")})


def _stage(tmp_path):
    s = tmp_path / "samples"
    s.mkdir()
    (s / "tj.pdf").write_bytes(
        _simple_page(
            b"BT /F1 12 Tf 72 700 Td "
            b"[(K) -120 (er) 40 (ned "
            b"tex) 20 (t li) -30 (ne)] "
            b"TJ ET "
            b"BT /F1 12 Tf 72 640 Td "
            b"[(A) -250 (V second) 15 "
            b"( kerned words)] TJ ET"))
    (s / "box.pdf").write_bytes(
        _simple_page(
            b"BT /F1 12 Tf 100 600 Td "
            b"(Boxed heading) Tj ET "
            b"BT /F1 12 Tf 100 560 Td "
            b"(Boxed body text here.) "
            b"Tj ET",
            "[50 50 562 742]"))

    d = Document()
    d.add_paragraph("before outer")
    outer = d.add_table(rows=2,
                        cols=2)
    for i, txt in enumerate(
            ["oc00", "oc01",
             "oc10", "oc11"]):
        outer.cell(i // 2,
                   i % 2).text = txt
    inner = outer.cell(
        0, 0).add_table(rows=1,
                        cols=2)
    inner.cell(0, 0).text = "ic0"
    inner.cell(0, 1).text = "ic1"
    d.add_paragraph("after tables")
    d.save(str(s / "nest.docx"))

    d2 = Document()
    d2.add_paragraph("section one para")
    d2.add_section(
        WD_SECTION.NEW_PAGE)
    d2.add_paragraph(
        "section two para")
    d2.save(str(s / "sec.docx"))

    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "tj",
             "path": "samples/tj.pdf",
             "source_type": "pdf",
             "expectations": {
                 "element_count_by_type":
                     {"heading": 2}}},
            {"doc_id": "box",
             "path": "samples/box.pdf",
             "source_type": "pdf",
             "expectations": {
                 "element_count_by_type":
                     {"heading": 1,
                      "paragraph": 1}}},
            {"doc_id": "nest",
             "path": "samples/nest.docx",
             "source_type": "docx",
             "expectations": {
                 "element_count_by_type":
                     {"paragraph": 2,
                      "table": 1}}},
            {"doc_id": "sec",
             "path": "samples/sec.docx",
             "source_type": "docx",
             "expectations": {
                 "element_count_by_type":
                     {"paragraph": 3}}},
        ]}), encoding="utf-8")
    return mf


def _run(tmp_path):
    mf = _stage(tmp_path)
    rep = tmp_path / "r.json"
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["run", "--manifest",
                   str(mf),
                   "--output", str(rep),
                   "--parser", "fallback",
                   "--max-chars", "800"])
    data = json.loads(
        rep.read_text(encoding="utf-8"))
    return rc, data, rep


# ---------- run 总览 ----------

def test_run_rc0(tmp_path):
    rc, _, _ = _run(tmp_path)
    assert rc == 0


def test_success_all_four(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "success_rates"][
        "pipeline_success"] == {
        "success_count": 4,
        "total": 4, "rate": 1.0}


def test_ect_sum_ten(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "counts"][
        "element_count_total"] == {
        "sum": 10,
        "participating_docs": 4}


def test_silent_drop_zero(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "silent_drop_total"] == 0


def test_per_doc_sdcs(tmp_path):
    _, data, _ = _run(tmp_path)
    for d in data["per_doc"]:
        assert d["metrics"][
            "silent_drop_count"] == {
            "value": 0, "reason": None}


# ---------- ratio macro ----------

def test_schema_valid_macro(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "schema_valid"] == {
        "macro_average": 1.0,
        "participating_docs": 4,
        "not_evaluated": 0}


def test_pdfloc_macro(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "pdf_locator_valid_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 2}


def test_docxloc_macro(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "docx_locator_valid_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 2}


def test_hbc_only_heading_docs(
        tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "heading_boundary_compliance"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 2}


def test_irer_null(tmp_path):
    _, data, _ = _run(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "image_resource_exists_ratio"] == {
        "macro_average": None,
        "participating_docs": 0,
        "not_evaluated": 4}


def test_text_metrics_macro(tmp_path):
    _, data, _ = _run(tmp_path)
    rma = data["summary"][
        "ratio_macro_averages"]
    for k in ("chunk_reference_"
              "intact_ratio",
              "text_preservation_"
              "equal",
              "text_char_multiset_"
              "precision",
              "text_char_multiset_"
              "recall"):
        assert rma[k] == {
            "macro_average": 1.0,
            "participating_docs": 4,
            "not_evaluated": 0}


def test_chunk_boundary_nulls(
        tmp_path):
    _, data, _ = _run(tmp_path)
    rma = data["summary"][
        "ratio_macro_averages"]
    for k in ("chunk_boundary_"
              "precision",
              "chunk_boundary_"
              "recall",
              "chunk_boundary_f1"):
        assert rma[k] == {
            "macro_average": None,
            "participating_docs": 0,
            "not_evaluated": 4}


# ---------- per-doc 定位 ----------

def test_locators_by_type(tmp_path):
    _, data, _ = _run(tmp_path)
    by = {d["doc_id"]:
          d["metrics"]
          for d in data["per_doc"]}
    assert by["tj"][
        "pdf_locator_valid_ratio"
    ]["value"] == 1.0
    assert by["tj"][
        "docx_locator_valid_ratio"
    ]["reason"] == \
        "not_docx_document"
    assert by["nest"][
        "docx_locator_valid_ratio"
    ]["value"] == 1.0
    assert by["nest"][
        "pdf_locator_valid_ratio"
    ]["reason"] == \
        "not_pdf_document"


def test_ects_exact(tmp_path):
    _, data, _ = _run(tmp_path)
    by = {d["doc_id"]:
          d["metrics"][
              "element_count_by_type"]
          ["value"]
          for d in data["per_doc"]}
    assert by["tj"] == {
        "heading": 2}
    assert by["box"] == {
        "heading": 1,
        "paragraph": 1}
    assert by["nest"] == {
        "paragraph": 2,
        "table": 1}
    assert by["sec"] == {
        "paragraph": 3}


# ---------- validate-report ----------

def test_validate_report_rc0(
        tmp_path):
    _, _, rep = _run(tmp_path)
    buf = io.StringIO()
    with contextlib.redirect_stdout(
            buf):
        rc = main(["validate-report",
                   str(rep)])
    assert rc == 0
