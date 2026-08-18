"""evaluation/annotation_metrics.py 第五百七十七轮 edges 测试（Round 1238）。

补强 edges145 未触及的角度（第六百一十批，probe 实证）。

新角度（厨房水槽真板三块锚）：
- **占位 after 单锚**——marker
  "(空段落)" after → 首次出现恰
  收尾 chunk1 → d 0 中界 1，1 锚
  2 界 → P 0.5 / R 1.0 / F1 2/3
  （与 edges133 的 before 全
  1.0 成方位对照）
- **双占位 after 全中**——同锚
  ×2 → 前移 pos 找到两处真身，
  各中一界 → 全 1.0（文档真含
  两处占位符，同串双锚非标注
  错误首锁）
- **表格尾锚**——marker "--- |"
  after → 表格 markdown 收尾恰
  界 2 → d 0（表格内容作锚首锁）
- forbidden tokens 第七百零四批（open 0）
"""

from __future__ import annotations

import inspect
import tempfile
from pathlib import Path

import evaluation.annotation_metrics as am_mod
from evaluation.annotation_metrics import chunk_boundary_prf


def _base_doc(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from app.pipeline import process_single
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    p = tmp_path / "ks.docx"
    doc.save(str(p))
    d, errors = process_single(p, tmp_path / "o.json",
                               parser_name="fallback",
                               max_chars=120)
    assert errors == []
    return d.to_dict()


def _ann(*pairs):
    return {"chunk_boundary_anchors": [
        {"marker": m, "position": pos} for m, pos in pairs]}


# ---------- 占位 after 单锚 ----------

def test_one_placeholder_after_batch436(tmp_path):
    r = chunk_boundary_prf(
        _base_doc(tmp_path), _ann(("(空段落)", "after")))
    assert r["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert r["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert r["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 双占位 after 全中 ----------

def test_two_placeholder_after_all_hit_batch436(tmp_path):
    r = chunk_boundary_prf(
        _base_doc(tmp_path),
        _ann(("(空段落)", "after"), ("(空段落)", "after")))
    assert r["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert r["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert r["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_two_placeholder_no_missing_batch436(tmp_path):
    r = chunk_boundary_prf(
        _base_doc(tmp_path),
        _ann(("(空段落)", "after"), ("(空段落)", "after")))
    assert r.get("_missing_markers") is None


# ---------- 表格尾锚 ----------

def test_table_tail_marker_batch436(tmp_path):
    r = chunk_boundary_prf(
        _base_doc(tmp_path), _ann(("--- |", "after")))
    assert r["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert r["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert r["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


def test_board_three_chunks_batch436(tmp_path):
    dd = _base_doc(tmp_path)
    assert len(dd["chunks"]) == 3
    assert dd["chunks"][1]["text"] == \
        "| L | R |\n| --- | --- |"


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch436():
    src = _src()
    assert 'joined_raw = " ".join(norm_chunks)' in src
    assert "predicted.append(end)" in src
    assert "gt_positions.append(find_pos)" in src
    assert ("gt_positions.append(find_pos + len(marker))"
            in src)
    assert "if find_pos < 0:" in src


# ---------- forbidden tokens 第七百零四批 ----------

def test_source_no_eval_batch436():
    assert "eval(" not in _src()


def test_source_no_exec_batch436():
    assert "exec(" not in _src()


def test_source_no_compile_batch436():
    assert "compile(" not in _src()


def test_source_no_globals_batch436():
    assert "globals(" not in _src()


def test_source_no_locals_batch436():
    assert "locals(" not in _src()


def test_source_no_os_system_batch436():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch436():
    assert "subprocess" not in _src()


def test_source_no_popen_batch436():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch436():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch436():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch436():
    assert "socket" not in _src()


def test_source_no_requests_batch436():
    assert "requests" not in _src()


def test_source_no_urllib_batch436():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch436():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch436():
    assert "yield" not in _src()


def test_source_no_async_await_batch436():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch436():
    assert "open(" not in _src()
