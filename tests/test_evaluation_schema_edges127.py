"""evaluation/schema.py 第五百四十三轮 edges 测试（Round 1099）。

补强 edges124-126 未触及的角度（第四百七十五批，probe 实证）。

新角度（source_type 翻转的行为面：allOf if/then 单向执法）：
- **pdf 翻转拒绝**：真实 docx 文档翻 source_type 为
  "pdf"（locator 仍 docx 形）→ "'page' is a
  required property @ path=['elements', 0,
  'source_locator']"——allOf if/then 行为首锁
  （edges112/121 只做了 def 内省）
- **pdf 正形照过**：翻转后 locator 换 {page, bbox}
  → 通过——pdf_locator 只 required page
- **docx 宽容纳 pdf 形**：docx 文档塞 {page, bbox}
  locator → 照过——docx_locator additionalProperties
  True + minProperties 1，陌生键也满足 ≥1——
  执法不对称：pdf 侧有 required 咬合、docx 侧只数键
- **空 locator schema 拒绝带 path**：{} → "{} should
  be non-empty"，e.errors[0] path == ['elements',
  0, 'source_locator']——与 R1093 metrics 视角
  （docx_locator_valid_ratio 0.5）互为表里：同一
  变异，schema 铡、metrics 摊薄
- forbidden tokens 第五百七十批（open 2）
"""

from __future__ import annotations

import copy
import inspect
import pathlib
import tempfile

from docx import Document

import evaluation.schema as schema_mod
from app.pipeline import process_single
from evaluation.schema import (
    EvalSchemaError, load_schema, validate)


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("BBB second paragraph body.")
    d.save(str(p))
    doc, errors = process_single(
        p, tmp_path / "s.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    return doc.to_dict()


def _expect_reject(tmp_path, mut, frag, path):
    r = copy.deepcopy(_real_doc(tmp_path))
    mut(r)
    try:
        validate(r, "document.schema.json")
        raised = False
    except EvalSchemaError as e:
        raised = True
        assert frag in str(e)
        assert len(e.errors) == 1
        assert e.errors[0]["path"] == path
    assert raised


# ---------- pdf 翻转拒绝 ----------

def test_pdf_flip_rejected_batch298(tmp_path):
    r = copy.deepcopy(_real_doc(tmp_path))
    r["source_type"] = "pdf"
    try:
        validate(r, "document.schema.json")
        raised = False
    except EvalSchemaError as e:
        raised = True
        assert "'page' is a required property" in str(e)
        assert len(e.errors) == 2
        assert e.errors[0]["path"] == [
            "elements", 0, "source_locator"]
        assert e.errors[1]["path"] == [
            "elements", 1, "source_locator"]
    assert raised


# ---------- pdf 正形照过 ----------

def test_pdf_flip_proper_locator_passes_batch298(tmp_path):
    r = copy.deepcopy(_real_doc(tmp_path))
    r["source_type"] = "pdf"
    for el in r["elements"]:
        el["source_locator"] = {
            "page": 1, "bbox": [0.0, 0.0, 10.0, 10.0]}
    validate(r, "document.schema.json")


# ---------- docx 宽容纳 pdf 形 ----------

def test_docx_accepts_pdf_shaped_locator_batch298(tmp_path):
    r = copy.deepcopy(_real_doc(tmp_path))
    r["elements"][0]["source_locator"] = {
        "page": 1, "bbox": [0.0, 0.0, 1.0, 1.0]}
    validate(r, "document.schema.json")


# ---------- 空 locator schema 拒绝带 path ----------

def test_empty_locator_rejected_with_path_batch298(tmp_path):
    def mut(r):
        r["elements"][0]["source_locator"] = {}
    _expect_reject(
        tmp_path, mut, "{} should be non-empty",
        ["elements", 0, "source_locator"])


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch298():
    src = _src()
    assert "def validate(instance" in src
    assert "校验失败" in src


# ---------- forbidden tokens 第五百七十批 ----------

def test_source_no_eval_batch298():
    assert "eval(" not in _src()


def test_source_no_exec_batch298():
    assert "exec(" not in _src()


def test_source_no_compile_batch298():
    assert "compile(" not in _src()


def test_source_no_globals_batch298():
    assert "globals(" not in _src()


def test_source_no_locals_batch298():
    assert "locals(" not in _src()


def test_source_no_os_system_batch298():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch298():
    assert "subprocess" not in _src()


def test_source_no_popen_batch298():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch298():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch298():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch298():
    assert "socket" not in _src()


def test_source_no_requests_batch298():
    assert "requests" not in _src()


def test_source_no_urllib_batch298():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch298():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch298():
    assert "yield" not in _src()


def test_source_no_async_await_batch298():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch298():
    assert _src().count("open(") == 2
