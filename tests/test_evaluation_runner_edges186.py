"""evaluation/runner.py 第六百一十六轮 edges 测试（Round 1172）。

补强 edges185 未触及的角度（第五百四十四批，probe 实证）。

新角度（caption 正则变体 / 后置空格刀口）：
- **五变体全 caption**——表 1（中文关键词）、图２
  （全角数字+全角空格）、fig 3（IGNORECASE 小写）、
  Table7 后置空格（数字后空格即满足 [\.、:\s]
  分隔类——刀口反直觉首锁）、Fig.4（点分隔）
- **黏字脱轨**——"Table7x"（数字后直接字母）正则
  失配 → DOCX 无 short_line 兜底 → paragraph
  （PDF 侧 short_line 是 PDF 专属，DOCX 只认
  标题样式——两源分类差异首锁）
- **块布局**——5 caption 各 isolated_caption 单源
  + 尾部 [Table7x+Regular] sequential 双源
- **30 字容差内两锚**——marker "digit" after（GT
  距界恰 30 字）与 "glued" before（距 24）均命中
  → P 1/5 / R 1.0 / F1 0.33333333333333337
  （浮点尾差原样锁定）
- forbidden tokens 第六百四十四批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None):
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "a").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("表 1 中文表题示例")
    d.add_paragraph("图２　全角数字图题")
    d.add_paragraph("fig 3 lowercase caption")
    d.add_paragraph("Table7 no separator after digits")
    d.add_paragraph("Fig.4 dot separator caption")
    d.add_paragraph("Table7x letter glued to digit")
    d.add_paragraph("Regular paragraph sentence.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    docs = [{"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}]
    if anchors is not None:
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}), encoding="utf-8")
        docs[0]["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 五变体全 caption ----------

def test_caption_variants_batch370(tmp_path):
    _board(tmp_path, "cv")
    doc, errors = process_single(
        tmp_path / "s" / "cv.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == ["caption"] * 5 + [
        "paragraph", "paragraph"]
    assert els[0]["content"] == "表 1 中文表题示例"
    assert els[1]["content"] == "图２　全角数字图题"
    assert els[2]["content"] == "fig 3 lowercase caption"
    assert els[3]["content"] == \
        "Table7 no separator after digits"
    assert els[4]["content"] == "Fig.4 dot separator caption"
    assert els[5]["content"] == \
        "Table7x letter glued to digit"
    assert els[6]["content"] == "Regular paragraph sentence."


# ---------- 块布局 ----------

def test_caption_chunks_batch370(tmp_path):
    _board(tmp_path, "cv2")
    doc, errors = process_single(
        tmp_path / "s" / "cv2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "isolated_caption"] * 5 + ["sequential"]
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks[:5])
    assert len(chunks[5]["source_element_ids"]) == 2
    assert chunks[5]["text"] == (
        "Table7x letter glued to digit "
        "Regular paragraph sentence.")


# ---------- 30 字容差内两锚 ----------

def test_caption_prf_digit_after_batch370(tmp_path):
    r = run_evaluation(_board(tmp_path, "cv3", [
        {"marker": "digit", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.2, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.33333333333333337, "reason": None}


def test_caption_prf_glued_before_batch370(tmp_path):
    r = run_evaluation(_board(tmp_path, "cv4", [
        {"marker": "glued", "position": "before"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.2, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.33333333333333337, "reason": None}


# ---------- 指标 ----------

def test_caption_by_type_batch370(tmp_path):
    r = run_evaluation(_board(tmp_path, "cv5"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"caption": 5, "paragraph": 2},
        "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch370():
    src = _src()
    assert src.count("process_single") == 6
    assert src.count("manifest") == 5
    assert src.count("error_code") == 4


# ---------- forbidden tokens 第六百四十四批 ----------

def test_source_no_eval_batch370():
    assert "eval(" not in _src()


def test_source_no_exec_batch370():
    assert "exec(" not in _src()


def test_source_no_compile_batch370():
    assert "compile(" not in _src()


def test_source_no_globals_batch370():
    assert "globals(" not in _src()


def test_source_no_locals_batch370():
    assert "locals(" not in _src()


def test_source_no_os_system_batch370():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch370():
    assert "subprocess" not in _src()


def test_source_no_popen_batch370():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch370():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch370():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch370():
    assert "socket" not in _src()


def test_source_no_requests_batch370():
    assert "requests" not in _src()


def test_source_no_urllib_batch370():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch370():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch370():
    assert "yield" not in _src()


def test_source_no_async_await_batch370():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch370():
    assert _src().count("open(") == 2
