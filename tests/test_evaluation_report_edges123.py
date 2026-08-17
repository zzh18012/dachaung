"""evaluation/report.py 第五百一十八轮 edges 测试（Round 1074）。

补强 edges119-122 未触及的角度（第四百五十批，probe 实证）。

新角度（多文档 silent 求和的真实 run：3+2+null → 5）：
- 三份真实 docx（各 2 段）不同 expectations：d1 期望
  5 段 → silent 3、d2 期望 4 段 → silent 2、d3 无
  expectations → null——**逐文档 silent 三态同屏**
  [3, 2, None]，汇总 silent_drop_total **5**（null 不
  参与、非零值逐个相加——合成板未见过的真实多文档
  求和形态）
- 同板 counts {sum 6, participating 3}（2 段 × 3 文档
  全参与）与 success {3, 3, 1.0}——silent 参与度（2/3）
  与 counts 参与度（3/3）在同一 run 里分道
- forbidden tokens 第五百四十五批（report 变体：15 项
  去 subprocess + open 0 + subprocess.run == 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    for name in ("a", "b", "c"):
        d = Document()
        d.add_paragraph("AAA first paragraph body.")
        d.add_paragraph("BBB second paragraph body.")
        d.save(str(tmp_path / "samples" / f"{name}.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/a.docx",
             "source_type": "docx",
             "expectations": {"element_count_by_type":
                              {"paragraph": 5}}},
            {"doc_id": "d2", "path": "samples/b.docx",
             "source_type": "docx",
             "expectations": {"element_count_by_type":
                              {"paragraph": 4}}},
            {"doc_id": "d3", "path": "samples/c.docx",
             "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json",
                          max_chars=200)


# ---------- 逐文档 silent 三态 ----------

def test_multi_doc_silent_values_batch273(tmp_path):
    rep = _run(tmp_path)
    silents = [p["metrics"]["silent_drop_count"]
               for p in rep["per_doc"]]
    assert silents == [
        {"value": 3, "reason": None},
        {"value": 2, "reason": None},
        {"value": None,
         "reason": "no_expectations"}]


# ---------- 求和：null 剔除、余值相加 ----------

def test_silent_sum_excludes_null_batch273(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["silent_drop_total"] == 5


# ---------- counts / success 参与度分道 ----------

def test_counts_success_diverge_batch273(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["counts"] == {
        "element_count_total": {"sum": 6,
                                "participating_docs": 3}}
    assert rep["summary"]["success_rates"] == {
        "pipeline_success": {"success_count": 3,
                             "total": 3, "rate": 1.0}}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch273():
    src = _src()
    assert ('summary["silent_drop_total"] = '
            'sum(silent_vals) if silent_vals else None'
            in src)
    assert "silent_vals" in src


# ---------- forbidden tokens 第五百四十五批（report 变体） ----------

def test_source_no_eval_batch273():
    assert "eval(" not in _src()


def test_source_no_exec_batch273():
    assert "exec(" not in _src()


def test_source_no_compile_batch273():
    assert "compile(" not in _src()


def test_source_no_globals_batch273():
    assert "globals(" not in _src()


def test_source_no_locals_batch273():
    assert "locals(" not in _src()


def test_source_no_os_system_batch273():
    assert "os.system" not in _src()


def test_source_no_popen_batch273():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch273():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch273():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch273():
    assert "socket" not in _src()


def test_source_no_requests_batch273():
    assert "requests" not in _src()


def test_source_no_urllib_batch273():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch273():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch273():
    assert "yield" not in _src()


def test_source_no_async_await_batch273():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch273():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch273():
    assert _src().count("subprocess.run") == 2
