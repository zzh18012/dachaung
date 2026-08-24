r"""app/parsers/fallback_parser.py 边角测试 - 第三十三轮（Round 1423）。

新角度（probe 实证）docx 单元格内图片 + 纯空白段（历史图
片全在正文段落 / 空段只锁过 add_section 版）：
- cell 内 add_picture：**完全静默丢失**——无 image 元素、无
  落盘文件、无告警；该格渲染成空 '  '
- 纯空白段 '   ' → '(空段落)' 占位（metadata empty True）
- locator：正文段 paragraph_index 跳过表格（0/1/2），表格
  独立 table_index 0
"""

from __future__ import annotations

import io
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    return (sig
            + chunk(b"IHDR",
                    struct.pack(
                        ">IIBBBBB",
                        1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT",
                    zlib.compress(
                        b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b""))


def _doc(tmp_path):
    d = Document()
    d.add_paragraph("before table")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "cell text"
    c = t.cell(0, 1).paragraphs[0]
    c.add_run().add_picture(
        io.BytesIO(_png()),
        width=914400)
    d.add_paragraph("   ")
    d.add_paragraph("after ws")
    p = tmp_path / "tic.docx"
    d.save(str(p))
    return p


def _parse(tmp_path):
    p = _doc(tmp_path)
    return FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))


# ---------- 元素结构 ----------

def test_types(tmp_path):
    doc = _parse(tmp_path)
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "table",
        "paragraph", "paragraph"]


def test_image_in_cell_invisible(
        tmp_path):
    doc = _parse(tmp_path)
    assert not any(
        e.type == "image"
        for e in doc.elements)


def test_image_in_cell_no_file(
        tmp_path):
    _parse(tmp_path)
    assert not list(
        (tmp_path / "imgs").glob("*"))


def test_no_warnings(tmp_path):
    doc = _parse(tmp_path)
    assert doc.warnings == []


def test_table_content_exact(
        tmp_path):
    doc = _parse(tmp_path)
    assert doc.elements[
        1].content == \
        "| cell text |  |\n" \
        "| --- | --- |"


def test_table_metadata(
        tmp_path):
    doc = _parse(tmp_path)
    assert doc.elements[
        1].metadata == {
        "row_count": 1,
        "col_count": 2,
        "source": "python-docx"}


# ---------- 空白段占位 ----------

def test_ws_paragraph_placeholder(
        tmp_path):
    doc = _parse(tmp_path)
    assert doc.elements[
        2].content == "(空段落)"
    assert doc.elements[
        2].metadata["empty"] is True


def test_locators_skip_table(
        tmp_path):
    doc = _parse(tmp_path)
    assert [e.source_locator
            for e in doc.elements] == [
        {"paragraph_index": 0,
         "section": 0},
        {"table_index": 0,
         "section": 0},
        {"paragraph_index": 1,
         "section": 0},
        {"paragraph_index": 2,
         "section": 0}]


# ---------- 管线 ----------

def test_schema_valid(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path)
    assert is_valid(doc.to_dict())


def test_pipeline_chunks(tmp_path):
    p = _doc(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "before table",
        "| cell text |  |\n"
        "| --- | --- |",
        "(空段落) after ws"]
    assert [len(c.source_element_ids)
            for c in doc.chunks] == [
        1, 1, 2]


def test_ws_placeholder_chunk_text(
        tmp_path):
    p = _doc(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[2].text == (
        "(空段落) after ws")
