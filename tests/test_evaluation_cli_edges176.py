"""evaluation/cli.py 第六百七十七轮 edges 测试（Round 1346）。

补强 edges175 未触及的角度（第七百一十八批，probe 实证）。

新角度（docx-only ef 走 CLI run）：
- **docx 行**——
  devset 行
  'pdf=0 docx=1'
  （docx-only 板
  首锁）
- **ef 命中透传**
  ——nope.docx
  expected
  file_not_found →
  matches true
  （CLI 全链 ef 首锁）
- **rc 0 四行**——
  报告落盘 + 成功
  1/1
- forbidden tokens 第五百九十八批（open 1）
"""

from __future__ import annotations

import inspect
import io
import json
import sys
from contextlib import redirect_stderr, \
    redirect_stdout

import pytest

import evaluation.cli as cli_mod
from docx import Document
from evaluation.cli import main


@pytest.fixture(autouse=True)
def _restore_argv():
    saved = sys.argv
    yield
    sys.argv = saved


def _run(tmp_path):
    d = Document()
    d.add_heading("Doc", level=1)
    d.add_paragraph("hello world")
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx"}],
        "expected_failures": [
            {"doc_id": "efx", "path": "nope.docx",
             "expected_error_code":
                 "file_not_found",
             "source_type": "docx"}]}),
        encoding="utf-8")
    sys.argv = ["evaluation.cli", "run",
                "--manifest",
                str(tmp_path / "m.json"),
                "--output",
                str(tmp_path / "r.json"),
                "--parser", "fallback",
                "--max-chars", "800"]
    out, err = io.StringIO(), io.StringIO()
    with redirect_stdout(out), redirect_stderr(err):
        rc = main()
    return rc, out.getvalue(), err.getvalue()


# ---------- docx 行 ----------

def test_devset_docx_line_batch544(tmp_path):
    _, out, _ = _run(tmp_path)
    assert ("devset_status=incomplete "
            "file_count=1 groups=1 pdf=0 "
            "docx=1") in out


# ---------- ef 命中透传 ----------

def test_ef_matches_true_batch544(tmp_path):
    _run(tmp_path)
    rep = json.loads(
        (tmp_path / "r.json").read_text(
            encoding="utf-8"))
    assert rep["expected_failures"][0] == {
        "doc_id": "efx",
        "expected_error_code": "file_not_found",
        "actual_error_code": "file_not_found",
        "matches": True}


# ---------- rc 0 四行 ----------

def test_rc_zero_batch544(tmp_path):
    rc, _, err = _run(tmp_path)
    assert rc == 0
    assert err == ""


def test_documents_line_batch544(tmp_path):
    _, out, _ = _run(tmp_path)
    assert "documents=1（成功 1，失败 0）" in out


def test_ok_line_batch544(tmp_path):
    _, out, _ = _run(tmp_path)
    assert out.startswith("[OK] 评测完成：")
    assert str(tmp_path / "r.json") in out


def test_report_success_batch544(tmp_path):
    _run(tmp_path)
    rep = json.loads(
        (tmp_path / "r.json").read_text(
            encoding="utf-8"))
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 1, "total": 1,
        "rate": 1.0}


def test_report_devset_docx_batch544(tmp_path):
    _run(tmp_path)
    rep = json.loads(
        (tmp_path / "r.json").read_text(
            encoding="utf-8"))
    assert rep["devset"]["docx_count"] == 1
    assert rep["devset"]["pdf_count"] == 0


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_counts_batch544():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


# ---------- forbidden tokens 第五百九十八批 ----------

def test_source_no_eval_batch544():
    assert "eval(" not in _src()


def test_source_no_exec_batch544():
    assert "exec(" not in _src()


def test_source_no_compile_batch544():
    assert "compile(" not in _src()


def test_source_no_globals_batch544():
    assert "globals(" not in _src()


def test_source_no_locals_batch544():
    assert "locals(" not in _src()


def test_source_no_os_system_batch544():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch544():
    assert "subprocess" not in _src()


def test_source_no_popen_batch544():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch544():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch544():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch544():
    assert "socket" not in _src()


def test_source_no_requests_batch544():
    assert "requests" not in _src()


def test_source_no_urllib_batch544():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch544():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch544():
    assert "yield" not in _src()


def test_source_no_async_await_batch544():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch544():
    assert ".call(" not in _src()


def test_source_open_count_is_1_batch544():
    assert _src().count("open(") == 1
