"""evaluation/runner.py 第六百二十三轮 edges 测试（Round 1179）。

补强 edges191 未触及的角度（第五百五十一批，probe 实证）。

新角度（空 PDF / 混合成败 devset）：
- **空 PDF 失败通道**——零文本零表格零图片
  PDF → errors[no_extracted_elements]（details
  警告 pdf_no_text_extracted——与 DOCX 侧
  docx_no_content 成镜像首锁）
- **混合成败聚合**——坏 PDF + 好 DOCX 同
  devset：success {1, 2, 0.5}；counts 只算好
  文档 {sum: 2, participating_docs: 1}
- **失败文档不评**——schema_valid /
  text_preservation / docx_locator 均
  {1.0, 1 参与, 1 未评}（失败方计入 not_
  evaluated）；pdf_locator {null, 0, 2}（坏
  PDF 失败 + 好 DOCX 异源双排除）
- **error_code 分流**——bad 文档 error_code
  为 no_extracted_elements、good 文档为 null
- forbidden tokens 第六百五十一批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _empty_pdf() -> bytes:
    out = bytearray(b"%PDF-1.4\n")
    s = b"1 w 0 G\n"
    objects = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 400]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objects[num] + b"endobj\n")
    xref_pos = len(out)
    out += b"xref\n0 6\n"
    out += b"0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xref_pos).encode() + b"\n%%EOF\n")
    return bytes(out)


def _board(tmp_path):
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "s" / "empty.pdf").write_bytes(_empty_pdf())
    d = Document()
    d.add_paragraph("Good paragraph one here.")
    d.add_paragraph("Good paragraph two here.")
    d.save(str(tmp_path / "s" / "good.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "bad", "path": "s/empty.pdf",
             "source_type": "pdf"},
            {"doc_id": "good", "path": "s/good.docx",
             "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 空 PDF 失败通道 ----------

def test_empty_pdf_pipeline_error_batch377(tmp_path):
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "s" / "e.pdf").write_bytes(_empty_pdf())
    doc, errors = process_single(
        tmp_path / "s" / "e.pdf", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert doc is None
    assert len(errors) == 1
    e = errors[0]
    assert e.code == "no_extracted_elements"
    assert e.details["source_type"] == "pdf"
    assert e.details["warnings"][0]["code"] == \
        "pdf_no_text_extracted"


# ---------- 混合成败聚合 ----------

def test_mixed_success_rate_batch377(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {"success_count": 1,
                                "total": 2, "rate": 0.5}
    assert r["summary"]["counts"][
        "element_count_total"] == {"sum": 2,
                                   "participating_docs": 1}


# ---------- 失败文档不评 ----------

def test_mixed_not_evaluated_batch377(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    ra = r["summary"]["ratio_macro_averages"]
    for k in ("schema_valid", "text_preservation_equal",
              "docx_locator_valid_ratio"):
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 1,
                         "not_evaluated": 1}, k
    assert ra["pdf_locator_valid_ratio"] == {
        "macro_average": None, "participating_docs": 0,
        "not_evaluated": 2}


# ---------- error_code 分流 ----------

def test_mixed_error_code_split_batch377(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    by_id = {pd["doc_id"]: pd["metrics"]
             for pd in r["per_doc"]}
    assert by_id["bad"]["pipeline_success"] == {
        "value": False, "reason": None}
    assert by_id["bad"]["error_code"] == {
        "value": "no_extracted_elements", "reason": None}
    assert by_id["good"]["pipeline_success"] == {
        "value": True, "reason": None}
    assert by_id["good"]["error_code"] == {
        "value": None, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch377():
    src = _src()
    assert src.count("metrics") == 13
    assert src.count("manifest") == 5
    assert src.count("run_evaluation") == 2


# ---------- forbidden tokens 第六百五十一批 ----------

def test_source_no_eval_batch377():
    assert "eval(" not in _src()


def test_source_no_exec_batch377():
    assert "exec(" not in _src()


def test_source_no_compile_batch377():
    assert "compile(" not in _src()


def test_source_no_globals_batch377():
    assert "globals(" not in _src()


def test_source_no_locals_batch377():
    assert "locals(" not in _src()


def test_source_no_os_system_batch377():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch377():
    assert "subprocess" not in _src()


def test_source_no_popen_batch377():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch377():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch377():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch377():
    assert "socket" not in _src()


def test_source_no_requests_batch377():
    assert "requests" not in _src()


def test_source_no_urllib_batch377():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch377():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch377():
    assert "yield" not in _src()


def test_source_no_async_await_batch377():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch377():
    assert _src().count("open(") == 2
