r"""app/parsers/fallback_parser.py 边角测试 - 第四十五轮（Round 1437）。

新角度（probe 实证）分类启发式全景（caption > heading >
paragraph 优先级链，历史只被动见到 heading/paragraph 输出，
从未主动触发 caption 或否定条件）：
- PDF：'Table 1 ...' / 'Figure 2: ...' / 'Fig.3 ...' /
  小写 'table 9 ...'（IGNORECASE）→ caption，meta
  {'heuristic': 'caption_regex'}
- PDF 短行**句末标点否决**：'Short line.' / 'Short line!' /
  'Is this a heading?' → paragraph（<=80 但 endswith 句末符）
- PDF 长度阈值精确 80：80×a → heading {'level': 0,
  'heuristic': 'short_line'}，81×a → paragraph；字号无关
  （6pt/48pt 短词都是 heading）
- docx：'表3、...' / '图4 ...'（Unicode 路径 CJK 正常）→
  caption；**caption 压过 Heading 1 样式**（type=caption 但
  meta.style 保留 'Heading 1'）
- docx List Bullet/List Number 样式：纯文本段落无项目符号
- docx 表格单元格内 Heading 1 样式：被压平进表格 markdown
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _build(content):
    objs = {
        5: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] "
            b"/Count 1 >>"),
        3: (b"<< /Type /Page /Parent "
            b"2 0 R /MediaBox "
            b"[0 0 612 792] "
            b"/Resources << /Font "
            b"<< /F1 5 0 R >> >> "
            b"/Contents 4 0 R >>"),
        4: (b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n" + content
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    out += b"xref\n0 6\n" \
        b"0000000000 65535 f \n"
    for oid in range(1, 6):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 6 "
            b"/Root 1 0 R >>\n"
            b"startxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    return bytes(out)


def _pdf(tmp_path, name, text,
         size=12):
    c = (b"BT /F1 "
         + str(size).encode()
         + b" Tf 72 700 Td ("
         + text.encode("utf-8")
         + b") Tj ET")
    p = tmp_path / name
    p.write_bytes(_build(c))
    return p


# ---------- PDF caption ----------

def test_pdf_caption_table(
        tmp_path):
    p = _pdf(tmp_path, "cap1.pdf",
             "Table 1 Sales overview")
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "caption"
    assert doc.elements[
        0].metadata == {
        "heuristic": "caption_regex"}


def test_pdf_caption_figure(
        tmp_path):
    p = _pdf(tmp_path, "cap2.pdf",
             "Figure 2: A chart")
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "caption"


def test_pdf_caption_figdot(
        tmp_path):
    p = _pdf(tmp_path, "cap3.pdf",
             "Fig.3 Quarterly data")
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "caption"


def test_pdf_caption_lowercase(
        tmp_path):
    p = _pdf(tmp_path, "cap4.pdf",
             "table 9 lower case t")
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "caption"


# ---------- PDF 句末否决 ----------

def test_pdf_period_blocks(
        tmp_path):
    p = _pdf(tmp_path, "per.pdf",
             "Short line.")
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "paragraph"
    assert doc.elements[
        0].metadata == {}


def test_pdf_bang_qmark(
        tmp_path):
    for name, text in (
            ("b.pdf", "Short line!"),
            ("q.pdf",
             "Is this a heading?")):
        p = _pdf(tmp_path, name,
                 text)
        doc = FallbackParser().parse(
            p, compute_file_hash(p))
        assert doc.elements[
            0].type == "paragraph"


# ---------- PDF 长度阈值 ----------

def test_pdf_len80_heading(
        tmp_path):
    p = _pdf(tmp_path, "h80.pdf",
             "a" * 80)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "heading"
    assert doc.elements[
        0].metadata == {
        "level": 0,
        "heuristic": "short_line"}


def test_pdf_len81_paragraph(
        tmp_path):
    p = _pdf(tmp_path, "p81.pdf",
             "a" * 81)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "paragraph"


def test_pdf_size_irrelevant(
        tmp_path):
    for size in (6, 48):
        p = _pdf(tmp_path,
                 "s%d.pdf" % size,
                 "Sized", size)
        doc = FallbackParser().parse(
            p, compute_file_hash(p))
        assert doc.elements[
            0].type == "heading"


# ---------- docx caption ----------

def test_docx_caption_normal(
        tmp_path):
    d = Document()
    d.add_paragraph(
        "Table 1 docx caption text")
    p = tmp_path / "cap.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "caption"


def test_docx_caption_cjk(
        tmp_path):
    d = Document()
    d.add_paragraph("表3、中文表格标题")
    d.add_paragraph("图4 中文图题")
    p = tmp_path / "cjk.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "caption", "caption"]


def test_docx_caption_beats_heading(
        tmp_path):
    d = Document()
    d.add_paragraph(
        "Table 2 heading style caption",
        style="Heading 1")
    p = tmp_path / "ch.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].type == "caption"
    assert doc.elements[
        0].metadata["style"] == \
        "Heading 1"


# ---------- docx 列表 / 单元格 ----------

def test_docx_list_styles_plain(
        tmp_path):
    d = Document()
    d.add_paragraph(
        "Bullet item one",
        style="List Bullet")
    d.add_paragraph(
        "Numbered item",
        style="List Number")
    p = tmp_path / "list.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "paragraph"]
    assert [e.content
            for e in doc.elements] == [
        "Bullet item one",
        "Numbered item"]


def test_docx_heading_in_cell(
        tmp_path):
    d = Document()
    d.add_paragraph("Outside para")
    t = d.add_table(rows=2, cols=1)
    cell_p = t.cell(0, 0).paragraphs[0]
    cell_p.style = d.styles[
        "Heading 1"]
    cell_p.add_run("Heading in cell")
    t.cell(1, 0).text = "plain cell"
    p = tmp_path / "cellh.docx"
    d.save(str(p))
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "table"]
    assert doc.elements[
        1].content == \
        "| Heading in cell |\n" \
        "| --- |\n| plain cell |"


# ---------- 通用 ----------

def test_caption_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _pdf(tmp_path, "cs.pdf",
             "Table 1 Sales overview")
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_caption_chunk(
        tmp_path):
    p = _pdf(tmp_path, "cc.pdf",
             "Table 1 Sales overview")
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[
        0].text == \
        "Table 1 Sales overview"


def test_classify_no_warnings(
        tmp_path):
    for name, text in (
            ("w1.pdf",
             "Table 1 Sales overview"),
            ("w2.pdf", "Short line."),
            ("w3.pdf", "a" * 81)):
        p = _pdf(tmp_path, name, text)
        doc = FallbackParser().parse(
            p, compute_file_hash(p))
        assert doc.warnings == []
