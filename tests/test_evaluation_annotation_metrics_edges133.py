"""evaluation/annotation_metrics.py 第五百零七轮 edges 测试（Round 1063）。

补强 edges130-132 未触及的角度（第四百三十九批，probe 实证）。

新角度（parser 合成占位符 "(空段落)" 作标注锚，真实嵌图文档）：
- 真实图片 docx 的空承载段被 parser 渲染成 "(空段落)"
  纯文本——标注 marker **可以锚定占位符**：mc 80 双 chunk
  板上 chunk1 恰以占位符收尾，marker "(空段落)" before →
  P/R/F1 全 1.0（占位符成为合法 gold 边界）
- **占位符独立成 chunk**：mc 40 下 5 chunk 板中 chunks[2]
  == "(空段落)"（chunker 把合成文本当普通文本切）；单锚
  → P 0.25 / R 1.0 / F1 0.4（4 预测边界命中 1）
- **单 chunk 早退对 marker 存在性全盲**：< 2 chunk 分支
  不做任何 marker 匹配——锚里混入不存在的 PNGIMAGE 也
  不进 missing（_missing_markers 键缺席）、R 固定 0.0；
  同板零锚对照 → 三 null（R 也 null）——R 0.0 vs null
  的分界线在"有无锚"而非"锚是否可寻"
- forbidden tokens 第五百三十四批（open 0）
"""

from __future__ import annotations

import inspect
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.annotation_metrics as am_mod
from app.pipeline import process_single
from evaluation.annotation_metrics import chunk_boundary_prf


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _doc(tmp_path, mc):
    d = Document()
    d.add_paragraph("AAA first paragraph body with "
                    "enough text to split nicely here.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("BBB third paragraph body with "
                    "enough text to split nicely here.")
    src = tmp_path / "img.docx"
    d.save(str(src))
    doc, errors = process_single(
        src, tmp_path / "s.json", parser_name="fallback",
        max_chars=mc, write_json=False)
    assert errors == []
    return doc.to_dict()


def _ann(anchors):
    return {"annotation_version": "1.0", "doc_id": "x",
            "chunk_boundary_anchors": anchors}


def _prf(dd, anchors):
    out = chunk_boundary_prf(dd, _ann(anchors),
                             tolerance_chars=30)
    return {k: out[k]["value"] for k in (
        "chunk_boundary_precision",
        "chunk_boundary_recall",
        "chunk_boundary_f1")}, out


# ---------- 占位符作 gold 锚 ----------

def test_placeholder_anchor_gold_batch262(tmp_path):
    dd = _doc(tmp_path, 80)
    assert dd["chunks"][0]["text"].endswith("(空段落)")
    vals, _ = _prf(dd, [{"marker": "(空段落)",
                         "position": "before"}])
    assert vals == {"chunk_boundary_precision": 1.0,
                    "chunk_boundary_recall": 1.0,
                    "chunk_boundary_f1": 1.0}


# ---------- 干净对照组 ----------

def test_bbb_anchor_gold_batch262(tmp_path):
    vals, _ = _prf(_doc(tmp_path, 80),
                   [{"marker": "BBB third",
                     "position": "before"}])
    assert vals == {"chunk_boundary_precision": 1.0,
                    "chunk_boundary_recall": 1.0,
                    "chunk_boundary_f1": 1.0}


# ---------- 单 chunk 早退：对 marker 存在性全盲 ----------

def test_single_chunk_blind_to_markers_batch262(tmp_path):
    dd = _doc(tmp_path, 200)
    assert len(dd["chunks"]) == 1
    vals, out = _prf(dd, [
        {"marker": "(空段落)", "position": "before"},
        {"marker": "PNGIMAGE", "position": "before"}])
    assert vals == {"chunk_boundary_precision": None,
                    "chunk_boundary_recall": 0.0,
                    "chunk_boundary_f1": None}
    assert "_missing_markers" not in out


# ---------- 单 chunk 零锚：三 null 对照 ----------

def test_single_chunk_no_anchors_null_r_batch262(tmp_path):
    vals, _ = _prf(_doc(tmp_path, 200), [])
    assert vals == {"chunk_boundary_precision": None,
                    "chunk_boundary_recall": None,
                    "chunk_boundary_f1": None}


# ---------- 占位符独立成 chunk ----------

def test_placeholder_own_chunk_batch262(tmp_path):
    dd = _doc(tmp_path, 40)
    assert dd["chunks"][2]["text"] == "(空段落)"
    vals, _ = _prf(dd, [{"marker": "(空段落)",
                         "position": "before"}])
    assert vals == {"chunk_boundary_precision": 0.25,
                    "chunk_boundary_recall": 1.0,
                    "chunk_boundary_f1": 0.4}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch262():
    src = _src()
    assert "if not chunks or len(chunks) < 2:" in src
    assert "search_from = find_pos + len(marker)" in src


# ---------- forbidden tokens 第五百三十四批 ----------

def test_source_no_eval_batch262():
    assert "eval(" not in _src()


def test_source_no_exec_batch262():
    assert "exec(" not in _src()


def test_source_no_compile_batch262():
    assert "compile(" not in _src()


def test_source_no_globals_batch262():
    assert "globals(" not in _src()


def test_source_no_locals_batch262():
    assert "locals(" not in _src()


def test_source_no_os_system_batch262():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch262():
    assert "subprocess" not in _src()


def test_source_no_popen_batch262():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch262():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch262():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch262():
    assert "socket" not in _src()


def test_source_no_requests_batch262():
    assert "requests" not in _src()


def test_source_no_urllib_batch262():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch262():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch262():
    assert "yield" not in _src()


def test_source_no_async_await_batch262():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch262():
    assert "open(" not in _src()
