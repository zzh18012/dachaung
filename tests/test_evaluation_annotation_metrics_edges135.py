"""evaluation/annotation_metrics.py 第五百二十一轮 edges 测试（Round 1077）。

补强 edges130-134 未触及的角度（第四百五十三批，probe 实证）。

新角度（真实双图双题注板：占位符/题注各自成 chunk +
CJK marker 逐字命中）：
- 图+题注交替 doc（无正文段）真实切分：4 chunk 恰为
  ['(空段落)', 'Figure 1: ...', '(空段落)', '图 2 ...']
  ——图片占位符与题注**各自独立成 chunk**（image 无
  文本、caption 短文本互不相并），3 条预测边界
- 英文锚 "Figure 1:" before → gt 落题注 chunk 起点 →
  P 1/3 / R 1.0 / F1 0.5；**CJK 锚 "图 2" 数字完全
  相同**——中文 marker 在规范化流里逐字 find 命中
- 双锚 → P 2/3 / R 1.0 / F1 0.8——1/3→2/3 的命中
  阶梯（3 边界中 2 条被锚定），missing 恒空
- forbidden tokens 第五百四十八批（open 0）
"""

from __future__ import annotations

import inspect
import io
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.annotation_metrics as am_mod
from app.pipeline import process_single
from evaluation.annotation_metrics import chunk_boundary_prf


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _doc(tmp_path):
    d = Document()
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("Figure 1: a real caption line.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("图 2 中文题注内容")
    src = tmp_path / "cap.docx"
    d.save(str(src))
    doc, errors = process_single(
        src, tmp_path / "s.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    return doc.to_dict()


def _prf(dd, anchors):
    ann = {"annotation_version": "1.0", "doc_id": "x",
           "chunk_boundary_anchors": anchors}
    out = chunk_boundary_prf(dd, ann,
                             tolerance_chars=30)
    return {k: out[k]["value"] for k in (
        "chunk_boundary_precision",
        "chunk_boundary_recall",
        "chunk_boundary_f1")}, out


# ---------- 占位符/题注各自成 chunk ----------

def test_caption_doc_chunk_shape_batch276(tmp_path):
    dd = _doc(tmp_path)
    texts = [c["text"] for c in dd["chunks"]]
    assert texts == ["(空段落)",
                     "Figure 1: a real caption line.",
                     "(空段落)",
                     "图 2 中文题注内容"]


# ---------- 英文题注锚 ----------

def test_english_caption_anchor_batch276(tmp_path):
    vals, _ = _prf(_doc(tmp_path),
                   [{"marker": "Figure 1:",
                     "position": "before"}])
    assert vals == {
        "chunk_boundary_precision":
            0.3333333333333333,
        "chunk_boundary_recall": 1.0,
        "chunk_boundary_f1": 0.5}


# ---------- CJK 锚逐字命中 ----------

def test_cjk_marker_verbatim_batch276(tmp_path):
    vals, _ = _prf(_doc(tmp_path),
                   [{"marker": "图 2",
                     "position": "before"}])
    assert vals == {
        "chunk_boundary_precision":
            0.3333333333333333,
        "chunk_boundary_recall": 1.0,
        "chunk_boundary_f1": 0.5}


# ---------- 双锚命中阶梯 ----------

def test_dual_anchor_ladder_batch276(tmp_path):
    vals, out = _prf(_doc(tmp_path), [
        {"marker": "Figure 1:", "position": "before"},
        {"marker": "图 2", "position": "before"}])
    assert vals == {
        "chunk_boundary_precision":
            0.6666666666666666,
        "chunk_boundary_recall": 1.0,
        "chunk_boundary_f1": 0.8}
    assert "_missing_markers" not in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_key_lines_batch276():
    src = _src()
    assert "gt_positions.append(find_pos)" in src
    assert ('out["chunk_boundary_precision"] = '
            '_ratio(matched / num_pred)') in src


# ---------- forbidden tokens 第五百四十八批 ----------

def test_source_no_eval_batch276():
    assert "eval(" not in _src()


def test_source_no_exec_batch276():
    assert "exec(" not in _src()


def test_source_no_compile_batch276():
    assert "compile(" not in _src()


def test_source_no_globals_batch276():
    assert "globals(" not in _src()


def test_source_no_locals_batch276():
    assert "locals(" not in _src()


def test_source_no_os_system_batch276():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch276():
    assert "subprocess" not in _src()


def test_source_no_popen_batch276():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch276():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch276():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch276():
    assert "socket" not in _src()


def test_source_no_requests_batch276():
    assert "requests" not in _src()


def test_source_no_urllib_batch276():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch276():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch276():
    assert "yield" not in _src()


def test_source_no_async_await_batch276():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch276():
    assert "open(" not in _src()
