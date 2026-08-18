"""evaluation/runner.py 第六百五十五轮 edges 测试（Round 1233）。

补强 edges219 未触及的角度（第六百零五批，probe 实证）。

新角度（run 拼接 / 空白占位段 / 三段合一）：
- **run 拼接**——一段三 run（"Hel" +
  "lo wo" + "rld run split"）→ 单
  元素 content "Hello world run
  split"（parser 按 run 串接文本
  首锁，历史 add_run 只用于图片/
  链接）
- **空白占位段**——"   "（纯三空
  格）段 → "(空段落)"（占位判定
  基于 strip 而非 == ""，区别于
  分节符插入的占位来源）
- **三段合一**——mc100 下三段顺
  序合并单块 42 字符 3 源
- **summary counts 只收 ect**——
  by_type 不进 counts 聚合
- forbidden tokens 第七百批里程碑（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Before ws.")
    doc.add_paragraph("   ")
    para = doc.add_paragraph()
    para.add_run("Hel")
    para.add_run("lo wo")
    para.add_run("rld run split")
    (tmp_path / "s").mkdir(exist_ok=True)
    p = tmp_path / "s" / "rs.docx"
    doc.save(str(p))
    return p


def _board(tmp_path):
    _docx(tmp_path)
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "rs", "path": "s/rs.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- run 拼接 ----------

def test_run_split_joins_runs_batch431(tmp_path):
    from app.pipeline import process_single
    doc, errors = process_single(
        _docx(tmp_path), tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    assert errors == []
    assert doc.to_dict()["elements"][2]["content"] == \
        "Hello world run split"


def test_run_split_single_element_batch431(tmp_path):
    from app.pipeline import process_single
    doc, errors = process_single(
        _docx(tmp_path), tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert len(els) == 3
    assert [e["type"] for e in els] == ["paragraph"] * 3


# ---------- 空白占位段 ----------

def test_whitespace_paragraph_placeholder_batch431(tmp_path):
    from app.pipeline import process_single
    doc, errors = process_single(
        _docx(tmp_path), tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    assert errors == []
    assert doc.to_dict()["elements"][1]["content"] == "(空段落)"


def test_paragraph_index_continuous_batch431(tmp_path):
    from app.pipeline import process_single
    doc, errors = process_single(
        _docx(tmp_path), tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    assert errors == []
    dd = doc.to_dict()
    assert [e["source_locator"]["paragraph_index"]
            for e in dd["elements"]] == [0, 1, 2]
    assert all(e["source_locator"]["section"] == 0
               for e in dd["elements"])


# ---------- 三段合一 ----------

def test_merged_single_chunk_batch431(tmp_path):
    from app.pipeline import process_single
    doc, errors = process_single(
        _docx(tmp_path), tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    assert errors == []
    dd = doc.to_dict()
    chunks = dd["chunks"]
    assert len(chunks) == 1
    assert chunks[0]["text"] == \
        "Before ws. (空段落) Hello world run split"
    assert [s.split("::")[-1]
            for s in chunks[0]["source_element_ids"]] == \
        ["e0000", "e0001", "e0002"]


# ---------- runner 级指标 ----------

def test_metrics_ect_three_batch431(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_total"] == {"value": 3,
                                        "reason": None}


def test_metrics_by_type_para3_batch431(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 3}, "reason": None}


def test_metrics_tpe_placeholder_counts_batch431(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["schema_valid"] == {"value": True, "reason": None}


def test_metrics_docx_locator_batch431(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


def test_summary_counts_only_ect_batch431(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=100)
    assert sorted(r["summary"]["counts"].keys()) == \
        ["element_count_total"]
    assert r["summary"]["counts"]["element_count_total"] == {
        "sum": 3, "participating_docs": 1}


def test_success_bucket_batch431(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=100)
    assert r["summary"]["success_rates"]["pipeline_success"] == {
        "success_count": 1, "total": 1, "rate": 1.0}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch431():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第七百批里程碑 ----------

def test_source_no_eval_batch431():
    assert "eval(" not in _src()


def test_source_no_exec_batch431():
    assert "exec(" not in _src()


def test_source_no_compile_batch431():
    assert "compile(" not in _src()


def test_source_no_globals_batch431():
    assert "globals(" not in _src()


def test_source_no_locals_batch431():
    assert "locals(" not in _src()


def test_source_no_os_system_batch431():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch431():
    assert "subprocess" not in _src()


def test_source_no_popen_batch431():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch431():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch431():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch431():
    assert "socket" not in _src()


def test_source_no_requests_batch431():
    assert "requests" not in _src()


def test_source_no_urllib_batch431():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch431():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch431():
    assert "yield" not in _src()


def test_source_no_async_await_batch431():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch431():
    assert _src().count("open(") == 2
