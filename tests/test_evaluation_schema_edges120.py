"""evaluation/schema.py 第四百九十四轮 edges 测试（Round 1050）。

补强 edges119 未触及的角度（第四百二十六批，probe 实证）。

新角度（schema 判决与管线判决同 dict 分裂）：
- _load_annotation 只 json.load、从不调 validate——
  schema 对 annotation 的判决在 runner 边界是
  纯建议性的。同一份 position "bogus" 标注：
  annotation schema 拒绝（恰 1 错 @
  ['chunk_boundary_anchors', 0, 'position']，enum
  ['before', 'after']），而真实管线照收并以 after
  语义算出 P/R/F1 全 1.0——两个判决同屏
- 坏 JSON 标注与缺席标注文件走同一静默降级：run
  仍成功（pipeline True）、boundary 三键全 null
  no_annotation——真实文件基线（edges106 是 patch
  板）
- forbidden tokens 第五百二十一批（open 2）
"""

from __future__ import annotations

import inspect
import json

import pytest
from docx import Document

import evaluation.schema as schema_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import EvalSchemaError, validate


_BOGUS = {"annotation_version": "1.0", "doc_id": "d1",
          "chunk_boundary_anchors": [
              {"marker": "CCC", "position": "bogus"}]}


def _run(tmp_path, ann_content):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    (tmp_path / "anns").mkdir()
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("CCC third paragraph body.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    if ann_content is not None:
        (tmp_path / "anns" / "ann.json").write_text(
            ann_content, encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/a.docx",
                       "source_type": "docx",
                       "annotation_file":
                           "anns/ann.json"}],
        "expected_failures": []}), encoding="utf-8")
    rep = run_evaluation(load_manifest(mf, tmp_path),
                         tmp_path / "o.json", max_chars=40)
    return rep["per_doc"][0]["metrics"]


# ---------- schema 侧判决 ----------

def test_bogus_position_schema_rejects_batch248():
    with pytest.raises(EvalSchemaError) as ei:
        validate(_BOGUS, "annotation.schema.json")
    assert len(ei.value.errors) == 1
    assert ei.value.errors[0]["path"] == [
        "chunk_boundary_anchors", 0, "position"]
    assert ei.value.errors[0]["message"] == (
        "'bogus' is not one of ['before', 'after']")


# ---------- 管线侧判决：照收并计算 ----------

def test_bogus_position_pipeline_accepts_batch248(tmp_path):
    m = _run(tmp_path, json.dumps(_BOGUS))
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}
    assert m["chunk_boundary_precision"] == {"value": 1.0,
                                             "reason": None}
    assert m["chunk_boundary_recall"] == {"value": 1.0,
                                          "reason": None}
    assert m["chunk_boundary_f1"] == {"value": 1.0,
                                      "reason": None}


def test_verdict_pair_same_dict_batch248(tmp_path):
    rejected = False
    try:
        validate(_BOGUS, "annotation.schema.json")
    except EvalSchemaError:
        rejected = True
    m = _run(tmp_path, json.dumps(_BOGUS))
    assert rejected
    assert m["chunk_boundary_f1"] == {"value": 1.0,
                                      "reason": None}


# ---------- 静默降级两路 ----------

def test_badjson_silent_degradation_batch248(tmp_path):
    m = _run(tmp_path, "{not json")
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall",
              "chunk_boundary_f1"):
        assert m[k] == {"value": None,
                        "reason": "no_annotation"}


def test_missing_annotation_same_nulls_batch248(tmp_path):
    m = _run(tmp_path, None)
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": None, "reason": "no_annotation"}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch248():
    src = _src()
    assert "Draft202012Validator(schema)" in src
    assert "validator.iter_errors(instance)" in src
    assert "head = errors[0]" in src


# ---------- forbidden tokens 第五百二十一批 ----------

def test_source_no_eval_batch248():
    assert "eval(" not in _src()


def test_source_no_exec_batch248():
    assert "exec(" not in _src()


def test_source_no_compile_batch248():
    assert "compile(" not in _src()


def test_source_no_globals_batch248():
    assert "globals(" not in _src()


def test_source_no_locals_batch248():
    assert "locals(" not in _src()


def test_source_no_os_system_batch248():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch248():
    assert "subprocess" not in _src()


def test_source_no_popen_batch248():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch248():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch248():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch248():
    assert "socket" not in _src()


def test_source_no_requests_batch248():
    assert "requests" not in _src()


def test_source_no_urllib_batch248():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch248():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch248():
    assert "yield" not in _src()


def test_source_no_async_await_batch248():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch248():
    assert _src().count("open(") == 2
