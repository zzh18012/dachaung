"""evaluation/cli.py 第六百一十三轮 edges 测试（Round 1169）。

补强 cli edges144 未触及的角度（第五百四十一批，probe 实证）。

新角度（DOCX 五型板 CLI 全链）：
- **五型板 run**——heading/paragraph/image/
  caption/table 单 DOCX 经 CLI run → rc 0、报告
  by_type 五键齐（DOCX 通道五型首锁，与 edges144
  的 PDF 五型 CLI 成对照）
- **inspect-doc 五型 counts 行**——"counts:
  elements=5 chunks=3"——DOCX 检视通道计数行
- **自产自校**——同报告 validate-report rc 0
- forbidden tokens 第六百四十一批（open 1）
"""

from __future__ import annotations

import inspect
import io
import json

import evaluation.cli as cli_mod
from evaluation.cli import main


PNG = (b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01'
       b'\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89'
       b'\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01'
       b'\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82')


def _board(tmp_path):
    from docx import Document
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("Doc Heading L1", level=1)
    p = d.add_paragraph("Body text before the image. ")
    p.add_run().add_picture(io.BytesIO(PNG))
    d.add_paragraph("Figure 2: docx caption text below.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "X1"
    t.cell(0, 1).text = "Y1"
    d.save(str(tmp_path / "samples" / "t.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "df", "path": "samples/t.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _run_cli(capsys, argv):
    rc = main(argv)
    out = capsys.readouterr().out
    return rc, out


# ---------- 五型板 run ----------

def test_cli_run_docx_five_batch367(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "200"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"][
        "element_count_by_type"] == {
        "value": {"heading": 1, "paragraph": 1,
                  "image": 1, "caption": 1, "table": 1},
        "reason": None}


def test_cli_run_docx_five_stdout_batch367(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "200"])
    assert rc == 0
    assert "documents=1（成功 1，失败 0）" in out
    assert "pdf=0 docx=1" in out


# ---------- inspect-doc 五型 counts 行 ----------

def test_cli_inspect_docx_five_counts_batch367(tmp_path, capsys):
    from app.pipeline import process_single
    _board(tmp_path)
    doc, errors = process_single(
        tmp_path / "samples" / "t.docx", tmp_path / "doc.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    rc, out = _run_cli(capsys, [
        "inspect-doc", str(tmp_path / "doc.json")])
    assert rc == 0
    assert "counts:      elements=5 chunks=3" in out


# ---------- 自产自校 ----------

def test_cli_validate_docx_five_report_batch367(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "200"])
    assert rc == 0
    rc2, out2 = _run_cli(capsys, [
        "validate-report", str(tmp_path / "r.json")])
    assert rc2 == 0
    assert "[OK]" in out2


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_identifier_counts_batch367():
    src = _src()
    assert src.count("main") == 3
    assert src.count("tolerance_chars") == 4
    assert src.count("inspect") == 9


# ---------- forbidden tokens 第六百四十一批 ----------

def test_source_no_eval_batch367():
    assert "eval(" not in _src()


def test_source_no_exec_batch367():
    assert "exec(" not in _src()


def test_source_no_compile_batch367():
    assert "compile(" not in _src()


def test_source_no_globals_batch367():
    assert "globals(" not in _src()


def test_source_no_locals_batch367():
    assert "locals(" not in _src()


def test_source_no_os_system_batch367():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch367():
    assert "subprocess" not in _src()


def test_source_no_popen_batch367():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch367():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch367():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch367():
    assert "socket" not in _src()


def test_source_no_requests_batch367():
    assert "requests" not in _src()


def test_source_no_urllib_batch367():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch367():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch367():
    assert "yield" not in _src()


def test_source_no_async_await_batch367():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch367():
    assert _src().count("open(") == 1
