"""evaluation/manifest.py 第四百九十六轮 edges 测试（Round 1052）。

补强 edges129 未触及的角度（第四百二十八批，probe 实证）。

新角度（真实互配对 + 富标注全字段 + 加载器内容盲区）：
- 真实好 docx（d1）与真实损坏 docx（d2）互为
  paired_with：加载器只看存在性不看内容——损坏
  文件照常入册（失败留给 runner 揭示）
- content_group_count 1：互配对（d1↔d2）只计一组
- 双真值 sha256 同时逐字往返（好文件与坏字节
  各自的 hash 都只是数据）
- 富标注文件（annotator / date / heading_order /
  figure_caption_pairs / anchors 带 reason 字段——
  annotation schema 全部 6 个合法字段一次用齐）过
  RS；annotation_resolved 单侧接线（d1 有 d2 None）
- categories 三值并集排序 ['alpha','beta','gamma']
- forbidden tokens 第五百二十三批（open 1）
"""

from __future__ import annotations

import hashlib
import inspect
import json

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import load_manifest
from evaluation.schema import validate_file


def _load(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    (tmp_path / "anns").mkdir()
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA body.")
    d.save(str(tmp_path / "samples" / "good.docx"))
    bad = tmp_path / "samples" / "bad.docx"
    bad.write_bytes(b"corrupt")
    sha_g = hashlib.sha256(
        (tmp_path / "samples" / "good.docx")
        .read_bytes()).hexdigest()
    sha_b = hashlib.sha256(bad.read_bytes()).hexdigest()
    rich = tmp_path / "anns" / "rich.json"
    rich.write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "annotator": "reviewer_a", "date": "2026-08-17",
        "heading_order": [{"level": 1,
                           "text": "Real Title"}],
        "figure_caption_pairs": [{"figure_marker": "f1",
                                  "caption_text": "cap"}],
        "chunk_boundary_anchors": [
            {"marker": "AAA", "position": "after",
             "reason": "para boundary"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/good.docx",
             "source_type": "docx", "sha256": sha_g,
             "categories": ["beta", "alpha"],
             "paired_with": "d2",
             "annotation_file": "anns/rich.json"},
            {"doc_id": "d2", "path": "samples/bad.docx",
             "source_type": "docx", "sha256": sha_b,
             "categories": ["gamma"],
             "paired_with": "d1"}],
        "expected_failures": []}), encoding="utf-8")
    m = load_manifest(mf)
    return m, rich, sha_g, sha_b


# ---------- 加载器内容盲区 ----------

def test_corrupt_partner_loads_batch250(tmp_path):
    m, _, _, sha_b = _load(tmp_path)
    assert len(m.documents) == 2
    assert m.documents[1].resolved_path.name == "bad.docx"
    assert m.documents[1].sha256 == sha_b


# ---------- 互配对计组 ----------

def test_mutual_pair_one_group_batch250(tmp_path):
    m, _, _, _ = _load(tmp_path)
    assert m.documents[0].paired_with == "d2"
    assert m.documents[1].paired_with == "d1"
    assert m.content_group_count == 1
    assert m.file_count == 2
    assert m.docx_count == 2
    assert m.pdf_count == 0


# ---------- 富标注全字段 ----------

def test_rich_annotation_all_fields_batch250(tmp_path):
    _, rich, _, _ = _load(tmp_path)
    validate_file(rich, "annotation.schema.json")


def test_annotation_asymmetric_wiring_batch250(tmp_path):
    m, rich, _, _ = _load(tmp_path)
    assert m.documents[0].annotation_resolved == rich
    assert m.documents[1].annotation_resolved is None
    assert m.documents[1].annotation_file_str is None


# ---------- 双真值 sha + 类别并集 ----------

def test_dual_true_sha_and_categories_batch250(tmp_path):
    m, _, sha_g, sha_b = _load(tmp_path)
    assert m.documents[0].sha256 == sha_g
    assert m.documents[1].sha256 == sha_b
    assert len(sha_g) == len(sha_b) == 64
    assert m.categories_covered == ["alpha", "beta",
                                    "gamma"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch250():
    src = _src()
    assert "paired_with=d.get(\"paired_with\")" in src
    assert "categories=tuple(d.get(\"categories\", []))" in src


# ---------- forbidden tokens 第五百二十三批 ----------

def test_source_no_eval_batch250():
    assert "eval(" not in _src()


def test_source_no_exec_batch250():
    assert "exec(" not in _src()


def test_source_no_compile_batch250():
    assert "compile(" not in _src()


def test_source_no_globals_batch250():
    assert "globals(" not in _src()


def test_source_no_locals_batch250():
    assert "locals(" not in _src()


def test_source_no_os_system_batch250():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch250():
    assert "subprocess" not in _src()


def test_source_no_popen_batch250():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch250():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch250():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch250():
    assert "socket" not in _src()


def test_source_no_requests_batch250():
    assert "requests" not in _src()


def test_source_no_urllib_batch250():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch250():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch250():
    assert "yield" not in _src()


def test_source_no_async_await_batch250():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch250():
    assert _src().count("open(") == 1
