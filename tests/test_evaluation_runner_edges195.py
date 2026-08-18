"""evaluation/runner.py 第六百二十六轮 edges 测试（Round 1182）。

补强 edges194 未触及的角度（第五百五十四批，probe 实证）。

新角度（Title/深级标题样式判定）：
- **Title 即题**——style="Title" → heading
  level 1（"title" 特判首锁）；style="Heading 9"
  → level 9（深级透传）
- **Quote/Subtitle 非题**——style="Quote" /
  "Subtitle" → paragraph level 0（引文副题
  不入题类首锁）
- **两级题切流**——chunks [seq(Title 单源),
  seq(H9+Quote+Subtitle+Normal 四源)]——题
  对题软界与深级无关
- **流首容差命中**——marker "Title" before →
  GT 落流起点、距界 1 恰 22 字 ≤ 30 → 全 1.0
- forbidden tokens 第六百五十四批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None):
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "a").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Title style paragraph", style="Title")
    d.add_paragraph("Heading nine text here",
                    style="Heading 9")
    d.add_paragraph("Quote style paragraph", style="Quote")
    d.add_paragraph("Subtitle style paragraph",
                    style="Subtitle")
    d.add_paragraph("Normal tail paragraph.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    docs = [{"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}]
    if anchors is not None:
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        docs[0]["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- Title 即题 ----------

def test_title_style_heading_batch380(tmp_path):
    _board(tmp_path, "st")
    doc, errors = process_single(
        tmp_path / "s" / "st.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert els[0]["type"] == "heading"
    assert els[0]["metadata"] == {
        "level": 1, "style": "Title", "empty": False}


# ---------- 深级透传 / 非题样式 ----------

def test_heading9_and_nonheading_batch380(tmp_path):
    _board(tmp_path, "st2")
    doc, errors = process_single(
        tmp_path / "s" / "st2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "heading", "heading", "paragraph", "paragraph",
        "paragraph"]
    assert els[1]["metadata"] == {
        "level": 9, "style": "Heading 9", "empty": False}
    assert els[2]["metadata"]["style"] == "Quote"
    assert els[3]["metadata"]["style"] == "Subtitle"
    assert all(e["metadata"]["level"] == 0
               for e in els[2:])


# ---------- 两级题切流 ----------

def test_style_chunks_batch380(tmp_path):
    _board(tmp_path, "st3")
    doc, errors = process_single(
        tmp_path / "s" / "st3.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "sequential"]
    assert chunks[0]["text"] == "Title style paragraph"
    assert len(chunks[0]["source_element_ids"]) == 1
    assert chunks[1]["text"] == (
        "Heading nine text here Quote style paragraph "
        "Subtitle style paragraph Normal tail paragraph.")
    assert len(chunks[1]["source_element_ids"]) == 4


# ---------- 流首容差命中 ----------

def test_title_stream_start_batch380(tmp_path):
    r = run_evaluation(_board(tmp_path, "st4", [
        {"marker": "Title", "position": "before"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_h9_boundary_batch380(tmp_path):
    r = run_evaluation(_board(tmp_path, "st5", [
        {"marker": "here", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


# ---------- 指标 ----------

def test_style_by_type_batch380(tmp_path):
    r = run_evaluation(_board(tmp_path, "st6"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"heading": 2, "paragraph": 3},
        "reason": None}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch380():
    src = _src()
    assert src.count("process_single") == 6
    assert src.count("per_doc") == 12
    assert src.count("chunk") == 9


# ---------- forbidden tokens 第六百五十四批 ----------

def test_source_no_eval_batch380():
    assert "eval(" not in _src()


def test_source_no_exec_batch380():
    assert "exec(" not in _src()


def test_source_no_compile_batch380():
    assert "compile(" not in _src()


def test_source_no_globals_batch380():
    assert "globals(" not in _src()


def test_source_no_locals_batch380():
    assert "locals(" not in _src()


def test_source_no_os_system_batch380():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch380():
    assert "subprocess" not in _src()


def test_source_no_popen_batch380():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch380():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch380():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch380():
    assert "socket" not in _src()


def test_source_no_requests_batch380():
    assert "requests" not in _src()


def test_source_no_urllib_batch380():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch380():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch380():
    assert "yield" not in _src()


def test_source_no_async_await_batch380():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch380():
    assert _src().count("open(") == 2
