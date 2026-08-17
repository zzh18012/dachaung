"""evaluation/schema.py 第五百零八轮 edges 测试（Round 1064）。

补强 edges119-121 未触及的角度（第四百四十批，probe 实证）。

新角度（真实 image 元素 vs anyOf 双通道，真实嵌图文档）：
- 真实嵌图 docx 解析产物过 document.schema.json——
  image 元素 content None / resource_path 绝对路径，
  locator 四键 {paragraph_index, section,
  relationship_id, target_partname}——docx_locator 的
  additionalProperties:true 延伸到 rel 族键（R1057 的
  page/bbox 之外的第三组跨族键实证）
- **anyOf 违例的原始面目**：双剥（content 与
  resource_path 同置 None）→ 恰 1 错 @ ['elements', 2]，
  message 以**实例 repr 开头**（"{'element_id'..."）、以
  "is not valid under any of the given schemas" 收尾，
  schema_path 恰 ['properties','elements','items','anyOf']
  ——jsonschema anyOf 原始消息形态首次锁定
- **anyOf 的 OR 语义**：content 单独在场（resource_path
  仍 None）即恢复通过——任一通道可满足
- forbidden tokens 第五百三十五批（open 2）
"""

from __future__ import annotations

import copy
import inspect
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.schema as schema_mod
from app.pipeline import process_single
from evaluation.schema import EvalSchemaError, validate


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _img_doc(tmp_path):
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_picture(BytesIO(_png_bytes()))
    d.save(str(tmp_path / "img.docx"))
    doc, errors = process_single(
        tmp_path / "img.docx", tmp_path / "s.json",
        parser_name="fallback", max_chars=200,
        write_json=False)
    assert errors == []
    return doc.to_dict()


def _mut(dd, fn):
    r = copy.deepcopy(dd)
    fn(r)
    return r


# ---------- 真实嵌图产物全绿 ----------

def test_real_image_doc_passes_batch263(tmp_path):
    dd = _img_doc(tmp_path)
    validate(dd, "document.schema.json")
    img = dd["elements"][2]
    assert img["type"] == "image"
    assert img["content"] is None
    assert isinstance(img["resource_path"], str)


# ---------- locator rel 族键放行 ----------

def test_image_locator_rich_keys_pass_batch263(tmp_path):
    dd = _img_doc(tmp_path)
    loc = dd["elements"][2]["source_locator"]
    assert set(loc) == {"paragraph_index", "section",
                        "relationship_id",
                        "target_partname"}
    assert loc["target_partname"].startswith(
        "/word/media/image")
    validate(dd, "document.schema.json")


# ---------- anyOf 双剥违例原始形态 ----------

def test_both_content_and_resource_none_rejected_batch263(
        tmp_path):
    dd = _img_doc(tmp_path)
    r = _mut(dd, lambda x: x["elements"][2].__setitem__(
        "resource_path", None))
    try:
        validate(r, "document.schema.json")
        raised = False
    except EvalSchemaError as e:
        raised = True
        assert len(e.errors) == 1
        err = e.errors[0]
        assert err["path"] == ["elements", 2]
        assert err["message"].startswith("{'element_id'")
        assert err["message"].endswith(
            "is not valid under any of the given schemas")
        assert err["schema_path"] == [
            "properties", "elements", "items", "anyOf"]
    assert raised


# ---------- anyOf 的 OR 语义 ----------

def test_resource_only_restores_pass_batch263(tmp_path):
    dd = _img_doc(tmp_path)
    r = _mut(dd, lambda x: (
        x["elements"][2].__setitem__("resource_path", None),
        x["elements"][2].__setitem__("content", "alt"),
    ))
    validate(r, "document.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch263():
    src = _src()
    assert '"message": err.message' in src
    assert "key=lambda e: list(e.absolute_path)" in src


# ---------- forbidden tokens 第五百三十五批 ----------

def test_source_no_eval_batch263():
    assert "eval(" not in _src()


def test_source_no_exec_batch263():
    assert "exec(" not in _src()


def test_source_no_compile_batch263():
    assert "compile(" not in _src()


def test_source_no_globals_batch263():
    assert "globals(" not in _src()


def test_source_no_locals_batch263():
    assert "locals(" not in _src()


def test_source_no_os_system_batch263():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch263():
    assert "subprocess" not in _src()


def test_source_no_popen_batch263():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch263():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch263():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch263():
    assert "socket" not in _src()


def test_source_no_requests_batch263():
    assert "requests" not in _src()


def test_source_no_urllib_batch263():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch263():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch263():
    assert "yield" not in _src()


def test_source_no_async_await_batch263():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch263():
    assert _src().count("open(") == 2
