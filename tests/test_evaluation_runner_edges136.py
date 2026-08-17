"""evaluation/runner.py 第五百三十三轮 edges 测试（Round 1089）。

补强 edges133-135 未触及的角度（第四百六十五批，probe 实证）。

新角度（heading 文本锚点 + 标注身份端到端不设防）：
- **heading 文本做锚点**：锚 marker "Late Title"（before）
  挂在 heading 元素文本上——heading 恒 LEAD 其 chunk
  （R1079 结构保证）→ gt 恰落强制断点 → P/R/F1 全
  1.0——标题锚天然零距离命中
- **annotation doc_id 端到端不校验**：同一板标注
  doc_id 写 "WRONG-ID"（document 是 d1）照常生效全
  1.0——runner 129 行 _load_annotation(doc.
  annotation_resolved) 拿到就用，身份对齐全靠上游
  manifest
- 对照：同板去掉 annotation_file → boundary 三元组
  null no_annotation，而 heading_boundary_compliance
  仍 1.0——heading 度量不依赖标注通道
- forbidden tokens 第五百六十批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _run(tmp_path, ann_doc_id="d1"):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    (tmp_path / "anns").mkdir()
    d = Document()
    d.add_paragraph("AAA intro paragraph before the heading.")
    d.add_heading("Late Title", level=1)
    d.add_paragraph("BBB body after the heading one.")
    d.add_paragraph("CCC body after the heading two.")
    d.save(str(tmp_path / "samples" / "h.docx"))
    entry = {"doc_id": "d1", "path": "samples/h.docx",
             "source_type": "docx"}
    if ann_doc_id is not None:
        (tmp_path / "anns" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0",
            "doc_id": ann_doc_id,
            "chunk_boundary_anchors": [
                {"marker": "Late Title",
                 "position": "before"}]}), encoding="utf-8")
        entry["annotation_file"] = "anns/a.json"
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [entry],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(tmp_path / "m.json",
                                        tmp_path),
                          tmp_path / "o.json", max_chars=200)


# ---------- heading 文本锚点 ----------

def test_heading_text_anchor_gold_batch288(tmp_path):
    m = _run(tmp_path)["per_doc"][0]["metrics"]
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall",
              "chunk_boundary_f1"):
        assert m[k] == {"value": 1.0, "reason": None}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


# ---------- annotation doc_id 端到端不校验 ----------

def test_annotation_doc_id_mismatch_e2e_batch288(tmp_path):
    m = _run(tmp_path, ann_doc_id="WRONG-ID")[
        "per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


# ---------- 无标注对照：heading 度量独立 ----------

def test_heading_metric_annotation_free_batch288(tmp_path):
    m = _run(tmp_path, ann_doc_id=None)[
        "per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": None, "reason": "no_annotation"}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch288():
    src = _src()
    assert ("annotation = _load_annotation("
            "doc.annotation_resolved)") in src
    assert "def _load_annotation(" in src


# ---------- forbidden tokens 第五百六十批 ----------

def test_source_no_eval_batch288():
    assert "eval(" not in _src()


def test_source_no_exec_batch288():
    assert "exec(" not in _src()


def test_source_no_compile_batch288():
    assert "compile(" not in _src()


def test_source_no_globals_batch288():
    assert "globals(" not in _src()


def test_source_no_locals_batch288():
    assert "locals(" not in _src()


def test_source_no_os_system_batch288():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch288():
    assert "subprocess" not in _src()


def test_source_no_popen_batch288():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch288():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch288():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch288():
    assert "socket" not in _src()


def test_source_no_requests_batch288():
    assert "requests" not in _src()


def test_source_no_urllib_batch288():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch288():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch288():
    assert "yield" not in _src()


def test_source_no_async_await_batch288():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch288():
    assert _src().count("open(") == 2
