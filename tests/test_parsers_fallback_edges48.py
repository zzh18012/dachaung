r"""app/parsers/fallback_parser.py 边角测试 - 第四十八轮（Round 1440）。

新角度（probe 实证）docx 嵌套表 / 分节 / 合并单元格（历史
docx 表全是单层平表、单 section、无 merge）：
- cell.add_table 嵌套内表**完全不可见**：外表 markdown 只有
  4 格原文，'nested deep' 无任何元素承载、无告警（对照
  R1423 单元格内图片同款静默丢失）
- add_section(NEW_PAGE) 分节：**locator 无 section 区分**
  （两段都 section: 0）；分节符产生幽灵 '(空段落)'
  （paragraph_index 1，metadata empty True）；索引连续
  0/1/2
- cell(0,0).merge(cell(0,1)) 行向合并：文本进**两个**格位
  '| A\nB | A\nB |'（\n 连接、markdown 层重复），第二行
  不受影响 '| C | D |'
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import \
    WD_SECTION

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _nest_docx(tmp_path):
    d = Document()
    d.add_paragraph("outer intro")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "top left"
    t.cell(0, 1).text = "top right"
    t.cell(1, 0).text = "bottom left"
    t.cell(1, 1).text = "bottom right"
    inner = t.cell(1, 1).add_table(
        rows=1, cols=1)
    inner.cell(0, 0).text = \
        "nested deep"
    p = tmp_path / "nest.docx"
    d.save(str(p))
    return p


def _sec_docx(tmp_path):
    d = Document()
    d.add_paragraph("Section one para")
    d.add_section(WD_SECTION.NEW_PAGE)
    d.add_paragraph("Section two para")
    p = tmp_path / "sec.docx"
    d.save(str(p))
    return p


def _merge_docx(tmp_path):
    d = Document()
    d.add_paragraph("merge doc")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "C"
    t.cell(1, 1).text = "D"
    t.cell(0, 0).merge(t.cell(0, 1))
    p = tmp_path / "merge.docx"
    d.save(str(p))
    return p


# ---------- 嵌套表 ----------

def test_nested_table_invisible(
        tmp_path):
    p = _nest_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "table"]
    assert doc.elements[
        1].content == (
        "| top left | top right |"
        "\n| --- | --- |"
        "\n| bottom left |"
        " bottom right |")


def test_nested_text_nowhere(
        tmp_path):
    p = _nest_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    for e in doc.elements:
        assert "nested" not in (
            e.content or "")
    assert doc.warnings == []


def test_nested_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _nest_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


# ---------- 分节 ----------

def test_section_no_split(tmp_path):
    p = _sec_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "Section one para",
        "(空段落)",
        "Section two para"]
    assert [e.source_locator["section"]
            for e in doc.elements] == [
        0, 0, 0]


def test_section_ghost_empty(
        tmp_path):
    p = _sec_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    ghost = doc.elements[1]
    assert ghost.source_locator == {
        "paragraph_index": 1,
        "section": 0}
    assert ghost.metadata[
        "empty"] is True


def test_section_indices(tmp_path):
    p = _sec_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.source_locator[
        "paragraph_index"]
        for e in doc.elements] == [
        0, 1, 2]


# ---------- 合并单元格 ----------

def test_merge_row_duplicated(
        tmp_path):
    p = _merge_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        1].content == (
        "| A\nB | A\nB |"
        "\n| --- | --- |"
        "\n| C | D |")


def test_merge_not_warned(
        tmp_path):
    p = _merge_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_merge_chunks(tmp_path):
    p = _merge_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert ("| A\nB | A\nB |\n"
            "| --- | --- |\n"
            "| C | D |") in [
        c.text for c in doc.chunks]


# ---------- 通用 ----------

def test_section_chunks(tmp_path):
    p = _sec_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "Section one para "
        "(空段落) "
        "Section two para"]
    assert len(doc.chunks[0]
               .source_element_ids) == 3


def test_section_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _sec_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_merge_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _merge_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())
