"""evaluation/runner.py 第五百六十七轮 edges 测试（Round 1123）。

补强 edges141 未触及的角度（第四百九十九批，probe 实证）。

新角度（expectations 通道真跑精确值）：
- **缺口计数真跑**——真实 docx（2 段落）挂 expectations
  {paragraph: 5} → silent_drop_count 3 + summary
  silent_drop_total 3（旧锁 edges11 用 Fake 入口只断言
  键存在，真跑精确值首锁）
- **盈余不计负**——{paragraph: 1} → 0 + 0（surplus 归零，
  max(0, expected - actual)）
- **缺席类型计缺口**——{paragraph: 1, table: 2} → 2 + 2
  （table 无元素按 0 实际计）
- **无 expectations 双 null**——per-doc null
  no_expectations + total null
- forbidden tokens 第五百九十五批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _board(tmp_path, exp):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA head start.")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    doc = {"doc_id": "d1", "path": "samples/g.docx",
           "source_type": "docx"}
    if exp is not None:
        doc["expectations"] = {"element_count_by_type": exp}
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [doc]}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _run(tmp_path, exp):
    r = run_evaluation(_board(tmp_path, exp), tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    return (r["per_doc"][0]["metrics"]["silent_drop_count"],
            r["summary"]["silent_drop_total"])


# ---------- 缺口计数真跑 ----------

def test_deficit_drop_exact_batch322(tmp_path):
    per_doc, total = _run(tmp_path, {"paragraph": 5})
    assert per_doc == {"value": 3, "reason": None}
    assert total == 3


# ---------- 盈余不计负 ----------

def test_surplus_clamped_zero_batch322(tmp_path):
    per_doc, total = _run(tmp_path, {"paragraph": 1})
    assert per_doc == {"value": 0, "reason": None}
    assert total == 0


# ---------- 缺席类型计缺口 ----------

def test_absent_type_counts_batch322(tmp_path):
    per_doc, total = _run(tmp_path, {"paragraph": 1, "table": 2})
    assert per_doc == {"value": 2, "reason": None}
    assert total == 2


# ---------- 无 expectations 双 null ----------

def test_no_expectations_dual_null_batch322(tmp_path):
    per_doc, total = _run(tmp_path, None)
    assert per_doc == {"value": None,
                       "reason": "no_expectations"}
    assert total is None


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch322():
    src = _src()
    assert "与 pipeline 内部规则保持一致" in src
    assert "images-<sha>" in src


# ---------- forbidden tokens 第五百九十五批 ----------

def test_source_no_eval_batch322():
    assert "eval(" not in _src()


def test_source_no_exec_batch322():
    assert "exec(" not in _src()


def test_source_no_compile_batch322():
    assert "compile(" not in _src()


def test_source_no_globals_batch322():
    assert "globals(" not in _src()


def test_source_no_locals_batch322():
    assert "locals(" not in _src()


def test_source_no_os_system_batch322():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch322():
    assert "subprocess" not in _src()


def test_source_no_popen_batch322():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch322():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch322():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch322():
    assert "socket" not in _src()


def test_source_no_requests_batch322():
    assert "requests" not in _src()


def test_source_no_urllib_batch322():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch322():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch322():
    assert "yield" not in _src()


def test_source_no_async_await_batch322():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch322():
    assert _src().count("open(") == 2
