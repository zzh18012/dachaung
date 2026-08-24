r"""app/parsers/fallback_parser.py 边角测试 - 第二十九轮（Round 1418）。

新角度（probe 实证）docx 外围部件与修订可见性（历史只
锁过正文/表格/图片/题注/hyperlink run）：
- word/footnotes.xml（zip 手术注入 + footnoteReference）：
  完全不可见、无告警——正文含 reference 前后文本但脚注文
  本无处出现
- word/header1.xml（headerReference 挂 sectPr）：完全不可
  见、无告警
- w:ins 修订插入 run：静默丢弃（段落只剩 'normal'）
"""

from __future__ import annotations

import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _ins_docx(tmp_path):
    d = Document()
    p = d.add_paragraph()
    p.add_run("normal ")
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "a")
    ins.set(qn("w:date"),
            "2026-01-01T00:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "inserted revision"
    r.append(t)
    ins.append(r)
    p._p.append(ins)
    p_ = tmp_path / "ins.docx"
    d.save(str(p_))
    return p_


def _fn_docx(tmp_path):
    d = Document()
    d.add_paragraph("Main body text.")
    base = tmp_path / "fn_base.docx"
    d.save(str(base))
    fn_xml = (
        '<?xml version="1.0" '
        'encoding="UTF-8" '
        'standalone="yes"?>\n'
        '<w:footnotes xmlns:w='
        '"http://schemas.openxmlformats'
        '.org/wordprocessingml/2006'
        '/main">'
        '<w:footnote w:type='
        '"separator" w:id="-1">'
        '<w:p><w:r><w:t/></w:r>'
        '</w:p></w:footnote>'
        '<w:footnote w:type='
        '"continuationSeparator" '
        'w:id="0"><w:p><w:r><w:t/>'
        '</w:r></w:p></w:footnote>'
        '<w:footnote w:id="1">'
        '<w:p><w:r><w:t>Footnote '
        'hidden text.</w:t></w:r>'
        '</w:p></w:footnote>'
        '</w:footnotes>')
    out = tmp_path / "fn.docx"
    src = zipfile.ZipFile(base)
    with zipfile.ZipFile(
            out, "w",
            zipfile.ZIP_DEFLATED) as z:
        for item in src.infolist():
            data = src.read(
                item.filename)
            if item.filename == (
                    "[Content_"
                    "Types].xml"):
                data = data.replace(
                    b"</Types>",
                    b'<Override PartName='
                    b'"/word/footnotes.xml" '
                    b'ContentType='
                    b'"application/vnd.'
                    b'openxmlformats-'
                    b'officedocument.'
                    b'wordprocessingml.'
                    b'footnotes+xml"/>'
                    b'</Types>')
            elif item.filename == (
                    "word/_rels/"
                    "document.xml.rels"):
                data = data.replace(
                    b"</Relationships>",
                    b'<Relationship Id='
                    b'"rIdFn1" Type='
                    b'"http://schemas.'
                    b'openxmlformats.org/'
                    b'officeDocument/2006/'
                    b'relationships/'
                    b'footnotes" Target='
                    b'"footnotes.xml"/>'
                    b'</Relationships>')
            elif item.filename == (
                    "word/document.xml"):
                data = data.replace(
                    "<w:t>Main body "
                    "text.</w:t>"
                    .encode(),
                    '<w:t>Main body '
                    'text.</w:t></w:r>'
                    '<w:r><w:footnote'
                    'Reference w:id="1"/>'
                    '</w:r><w:r><w:t '
                    'xml:space='
                    '"preserve"> tail'
                    '</w:t>'.encode())
            z.writestr(item, data)
        z.writestr(
            "word/footnotes.xml",
            fn_xml)
    src.close()
    return out


def _hdr_docx(tmp_path):
    d = Document()
    d.add_paragraph("Body under header.")
    base = tmp_path / "hdr_base.docx"
    d.save(str(base))
    hdr_xml = (
        '<?xml version="1.0" '
        'encoding="UTF-8" '
        'standalone="yes"?>\n'
        '<w:hdr xmlns:w='
        '"http://schemas.openxmlformats'
        '.org/wordprocessingml/2006'
        '/main" xmlns:r='
        '"http://schemas.openxmlformats'
        '.org/officeDocument/2006'
        '/relationships">'
        '<w:p><w:r><w:t>Header '
        'hidden text.</w:t></w:r>'
        '</w:p></w:hdr>')
    out = tmp_path / "hdr.docx"
    src = zipfile.ZipFile(base)
    with zipfile.ZipFile(
            out, "w",
            zipfile.ZIP_DEFLATED) as z:
        for item in src.infolist():
            data = src.read(
                item.filename)
            if item.filename == (
                    "[Content_"
                    "Types].xml"):
                data = data.replace(
                    b"</Types>",
                    b'<Override PartName='
                    b'"/word/header1.xml" '
                    b'ContentType='
                    b'"application/vnd.'
                    b'openxmlformats-'
                    b'officedocument.'
                    b'wordprocessingml.'
                    b'header+xml"/>'
                    b'</Types>')
            elif item.filename == (
                    "word/_rels/"
                    "document.xml.rels"):
                data = data.replace(
                    b"</Relationships>",
                    b'<Relationship Id='
                    b'"rIdHdr1" Type='
                    b'"http://schemas.'
                    b'openxmlformats.org/'
                    b'officeDocument/2006/'
                    b'relationships/'
                    b'header" Target='
                    b'"header1.xml"/>'
                    b'</Relationships>')
            elif item.filename == (
                    "word/document.xml"):
                data = data.replace(
                    b"<w:sectPr",
                    b'<w:sectPr>'
                    b'<w:headerReference '
                    b'w:type="default" '
                    b'r:id="rIdHdr1"/>')
            z.writestr(item, data)
        z.writestr(
            "word/header1.xml",
            hdr_xml)
    src.close()
    return out


# ---------- footnotes ----------

def test_footnote_invisible(
        tmp_path):
    p = _fn_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    for e in doc.elements:
        assert "Footnote" \
            not in e.content
        assert "hidden" \
            not in e.content


def test_footnote_body_intact(
        tmp_path):
    p = _fn_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "Main body text. tail"]


def test_footnote_one_element(
        tmp_path):
    p = _fn_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert len(doc.elements) == 1
    assert doc.elements[
        0].type == "paragraph"


def test_footnote_locator(
        tmp_path):
    p = _fn_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].source_locator == {
        "paragraph_index": 0,
        "section": 0}


def test_footnote_no_warnings(
        tmp_path):
    p = _fn_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_footnote_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _fn_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_footnote_pipeline_chunk(
        tmp_path):
    p = _fn_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert len(doc.chunks) == 1
    assert doc.chunks[0].text == (
        "Main body text. tail")
    assert len(doc.chunks[0]
               .source_element_ids) == 1


# ---------- header ----------

def test_header_invisible(
        tmp_path):
    p = _hdr_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    for e in doc.elements:
        assert "Header" \
            not in e.content
        assert "hidden" \
            not in e.content


def test_header_body_intact(
        tmp_path):
    p = _hdr_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "Body under header."]


def test_header_no_warnings(
        tmp_path):
    p = _hdr_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_header_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _hdr_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


# ---------- w:ins 修订插入 ----------

def test_ins_revision_dropped(
        tmp_path):
    p = _ins_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "normal"]


def test_ins_single_element(
        tmp_path):
    p = _ins_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert len(doc.elements) == 1
    assert doc.warnings == []


def test_ins_pipeline_chunk(
        tmp_path):
    p = _ins_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[0].text == (
        "normal")
