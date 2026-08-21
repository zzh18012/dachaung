"""evaluation/cli.py 第六百七十一轮 edges 测试（Round 1310）。

补强 edges169 未触及的角度（第六百八十二批，probe 实证）。

新角度（真 DOCX inspect-doc 头部 / null 原因渲染）：
- **docx 头部块**——file/
  document_id(doc- 前缀)/
  source type=docx/
  parser vpdfplumber=/
  counts elements=2
  chunks=11 五行格式
  首锁
- **locator 分型行**——
  dlvr 1.0000 (ok) +
  plvr null (not_pdf_
  document) 同面板
- **error_code (None)**——
  null 分支渲染原始
  reason（None 不落
  'ok'——与其他分支
  or 'ok' 不对称首锁）
- **cbp 三连 null**——
  no_annotation 三行
- **irer null**——
  no_image_elements
- **run 行**——complete
  状态 + pdf=1 docx=1
  混合计数行
- **文档不存在**——rc 2
  + [ERROR] 文档不存在
- forbidden tokens 第五百九十二批（open 1）
"""

from __future__ import annotations

import inspect
import json
import sys

import pytest

import evaluation.cli as cli_mod
from docx import Document
from app.pipeline import process_single
from evaluation.cli import main


@pytest.fixture(autouse=True)
def _restore_argv():
    saved = sys.argv
    yield
    sys.argv = saved


def _doc_json(tmp_path):
    d = Document()
    d.add_heading("DocTitle", level=1)
    d.add_paragraph(" ".join("Sent%d." % i
                             for i in range(40)))
    d.save(str(tmp_path / "c.docx"))
    doc, errors = process_single(tmp_path / "c.docx",
                                 tmp_path / "o.json",
                                 parser_name="fallback",
                                 max_chars=32)
    assert errors == []
    return tmp_path / "o.json"


def _inspect_out(tmp_path, capsys):
    oj = _doc_json(tmp_path)
    sys.argv = ["evaluation.cli", "inspect-doc", str(oj)]
    rc = main()
    out = capsys.readouterr().out
    return rc, out, oj


# ---------- docx 头部块 ----------

def test_header_file_line_batch508(tmp_path, capsys):
    _, out, oj = _inspect_out(tmp_path, capsys)
    assert f"file:        {oj}" in out


def test_header_document_id_prefix_batch508(
        tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert "document_id: doc-" in out


def test_header_source_type_docx_batch508(
        tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert "  type=docx" in out


def test_header_parser_line_batch508(tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert "parser:      fallback vpdfplumber=" in out


def test_header_counts_line_batch508(tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert "counts:      elements=2 chunks=11" in out


# ---------- locator 分型行 ----------

def test_dlvr_line_batch508(tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert (f"  {'docx_locator_valid_ratio':36}"
            " 1.0000  (ok)") in out


def test_plvr_null_line_batch508(tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert (f"  {'pdf_locator_valid_ratio':36}"
            " null  (not_pdf_document)") in out


# ---------- error_code (None) ----------

def test_error_code_none_reason_batch508(
        tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert (f"  {'error_code':36}"
            " null  (None)") in out


def test_no_error_code_ok_line_batch508(
        tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert "error_code null  (ok)" not in out


# ---------- cbp 三连 null / irer null ----------

def test_cbp_null_trio_batch508(tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    for name in ("chunk_boundary_precision",
                 "chunk_boundary_recall",
                 "chunk_boundary_f1"):
        assert (f"  {name:36}"
                " null  (no_annotation)") in out


def test_irer_null_line_batch508(tmp_path, capsys):
    _, out, _ = _inspect_out(tmp_path, capsys)
    assert (f"  {'image_resource_exists_ratio':36}"
            " null  (no_image_elements)") in out


def test_inspect_rc_zero_batch508(tmp_path, capsys):
    rc, _, _ = _inspect_out(tmp_path, capsys)
    assert rc == 0


# ---------- run 行（complete + docx 计数） ----------

def _run_cli(tmp_path, capsys):
    d = Document()
    d.add_heading("DocTitle", level=1)
    d.add_paragraph(" ".join("Sent%d." % i
                             for i in range(40)))
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "complete",
        "documents": [
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx"}]}),
        encoding="utf-8")
    sys.argv = ["evaluation.cli", "run", "--manifest",
                str(tmp_path / "m.json"),
                "--output", str(tmp_path / "r.json"),
                "--parser", "fallback",
                "--max-chars", "32"]
    rc = main()
    return rc, capsys.readouterr().out


def test_run_complete_docx_line_batch508(
        tmp_path, capsys):
    _, out = _run_cli(tmp_path, capsys)
    assert ("devset_status=complete file_count=1 "
            "groups=1 pdf=0 docx=1") in out


def test_run_docx_success_line_batch508(
        tmp_path, capsys):
    _, out = _run_cli(tmp_path, capsys)
    assert "documents=1（成功 1，失败 0）" in out


# ---------- 文档不存在 ----------

def test_inspect_missing_doc_batch508(tmp_path,
                                      capsys):
    sys.argv = ["evaluation.cli", "inspect-doc",
                str(tmp_path / "nope.json")]
    rc = main()
    err = capsys.readouterr().err
    assert rc == 2
    assert "[ERROR] 文档不存在:" in err


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_counts_batch508():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


def test_source_null_branch_batch508():
    src = _src()
    assert ('return f"  {name:36} null  ({reason})"'
            in src)


# ---------- forbidden tokens 第五百九十二批 ----------

def test_source_no_eval_batch508():
    assert "eval(" not in _src()


def test_source_no_exec_batch508():
    assert "exec(" not in _src()


def test_source_no_compile_batch508():
    assert "compile(" not in _src()


def test_source_no_globals_batch508():
    assert "globals(" not in _src()


def test_source_no_locals_batch508():
    assert "locals(" not in _src()


def test_source_no_os_system_batch508():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch508():
    assert "subprocess" not in _src()


def test_source_no_popen_batch508():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch508():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch508():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch508():
    assert "socket" not in _src()


def test_source_no_requests_batch508():
    assert "requests" not in _src()


def test_source_no_urllib_batch508():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch508():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch508():
    assert "yield" not in _src()


def test_source_no_async_await_batch508():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch508():
    assert ".call(" not in _src()


def test_source_open_count_is_1_batch508():
    assert _src().count("open(") == 1
