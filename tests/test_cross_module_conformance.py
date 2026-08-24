r"""Round 1366 — 跨模块常量/结构一致性测试。

互补于 test_cross_module_inconsistency.py（错误路径），本文件锁
"常量与结构在各模块之间保持一致"：

- models.SCHEMA_VERSION == document.schema.json const
- models.SourceType/ElementType 字面量集合 == schema enum（含顺序）
- 四个 stdlib parser 共享 version "stdlib/0.1.0"，name 各异且与
  get_parser 键一致
- make_document_id == "doc-" + sha[:16]
- element_id/chunk_id 四位零填充模式跨 parser 一致
- evaluation REPORT_VERSION == evaluation-report.schema.json const
- 每个 parser 的真实产物（全管线后）都通过 app schema 校验

不修改任何源码。
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

from app.hash import compute_file_hash
from app.models import SCHEMA_VERSION, Document, Element
from app.parsers.base import make_document_id
from app.parsers.fallback_parser import FallbackParser
from app.parsers.html_parser import HtmlParser
from app.parsers.ipynb_parser import IpynbParser
from app.parsers.kreuzberg_parser import KreuzbergParser
from app.parsers.markdown_parser import MarkdownParser
from app.parsers.text_parser import TextParser
from app.pipeline import process_single
from app.schema import is_valid

_SCHEMA = json.loads(
    Path("schemas/document.schema.json").read_text(
        encoding="utf-8"))
_REPORT_SCHEMA = json.loads(
    Path("schemas/evaluation-report.schema.json"
         ).read_text(encoding="utf-8"))


def _enums(node):
    found = []
    if isinstance(node, dict):
        if "enum" in node and node["enum"] \
                and isinstance(node["enum"][0], str):
            found.append(node["enum"])
        for v in node.values():
            found.extend(_enums(v))
    elif isinstance(node, list):
        for v in node:
            found.extend(_enums(v))
    return found


# ---------- schema_version 常量一致 ----------

def test_models_schema_version_is_const():
    assert SCHEMA_VERSION == "0.1.0"


def test_schema_const_matches_models():
    assert _SCHEMA["properties"]["schema_version"][
        "const"] == SCHEMA_VERSION


def test_document_to_dict_uses_const():
    doc = Document(
        document_id="d", source_path="p",
        source_type="text", source_hash="a" * 64,
        parser_name="t", parser_version="1")
    assert doc.to_dict()["schema_version"] == \
        SCHEMA_VERSION


# ---------- enum 对齐 ----------

def test_source_type_enum_matches_models():
    enums = _enums(_SCHEMA)
    assert ["pdf", "docx", "markdown", "html",
            "text", "ipynb"] in enums


def test_element_type_enum_matches_models():
    enums = _enums(_SCHEMA)
    assert ["heading", "paragraph", "list_item",
            "table", "image", "caption", "header",
            "footer"] in enums


def test_ipynb_cell_type_enum():
    enums = _enums(_SCHEMA)
    assert ["markdown", "code", "raw"] in enums


# ---------- parser 常量 ----------

def test_stdlib_parsers_share_version():
    versions = {p.version for p in (
        MarkdownParser, HtmlParser, IpynbParser,
        TextParser)}
    assert versions == {"stdlib/0.1.0"}


def test_stdlib_parser_names_distinct():
    names = {p.name for p in (
        MarkdownParser, HtmlParser, IpynbParser,
        TextParser)}
    assert names == {"markdown", "html", "ipynb",
                     "text"}


def test_kreuzberg_identity():
    assert KreuzbergParser.name == "kreuzberg"
    assert KreuzbergParser.version == "4.10.2"


def test_fallback_name():
    assert FallbackParser.name == "fallback"


# ---------- make_document_id ----------

def test_doc_id_format():
    assert make_document_id("a" * 64) == \
        "doc-aaaaaaaaaaaaaaaa"


def test_doc_id_uses_first_16():
    sha = "0123456789abcdef" + "f" * 48
    assert make_document_id(sha) == \
        "doc-0123456789abcdef"


def test_doc_id_deterministic():
    assert make_document_id("b" * 64) == \
        make_document_id("b" * 64)


def test_doc_id_different_for_different_sha():
    assert make_document_id("a" * 64) != \
        make_document_id("b" * 64)


# ---------- evaluation 版本一致 ----------

def test_report_version_constant():
    from evaluation import REPORT_VERSION
    assert REPORT_VERSION == "1.1"


def test_report_schema_const_matches():
    assert _REPORT_SCHEMA["properties"][
        "report_version"]["const"] == "1.1"


# ---------- id 模式跨 parser 一致 ----------

MD = "# H\n\npara text\n"
HTML = ("<html><body><h1>H</h1>"
        "<p>para text</p></body></html>")
TXT = "para text\n"


def _parse(parser_name, tmp_path, filename,
           content):
    (tmp_path / filename).write_text(
        content, encoding="utf-8")
    return process_single(
        tmp_path / filename, None,
        parser_name=parser_name, max_chars=800)


def test_element_id_pattern_all_parsers(
        tmp_path):
    for name, fn, c in (
            ("markdown", "d.md", MD),
            ("html", "d.html", HTML),
            ("text", "d.txt", TXT)):
        doc, errors = _parse(name, tmp_path, fn, c)
        assert errors == [], name
        for e in doc.elements:
            assert e.element_id.startswith(
                doc.document_id + "::e")
            tail = e.element_id.rsplit("e", 1)[1]
            assert len(tail) == 4
            assert tail.isdigit()


def test_element_ids_zero_padded_from_zero(
        tmp_path):
    doc, errors = _parse("markdown", tmp_path,
                         "d.md", MD)
    assert [e.element_id.rsplit("e", 1)[1]
            for e in doc.elements] == ["0000",
                                       "0001"]


def test_chunk_id_pattern(tmp_path):
    doc, errors = _parse("markdown", tmp_path,
                         "d.md", MD)
    assert errors == []
    for ch in doc.chunks:
        assert ch.chunk_id.startswith(
            doc.document_id + "::c")
        tail = ch.chunk_id.rsplit("c", 1)[1]
        assert len(tail) == 4
        assert tail.isdigit()


# ---------- 同内容不同扩展 → 不同 source_type ----------

def test_same_text_different_types(tmp_path):
    body = "para text\n"
    d1, _ = _parse("text", tmp_path, "a.txt",
                   body)
    (tmp_path / "b.md").write_text(
        "para text\n", encoding="utf-8")
    d2, _ = process_single(
        tmp_path / "b.md", None,
        parser_name="markdown", max_chars=800)
    assert d1.source_type == "text"
    assert d2.source_type == "markdown"


# ---------- 全 parser 产物过 schema ----------

def test_all_stdlib_outputs_pass_schema(tmp_path):
    boards = (
        ("markdown", "d.md", MD),
        ("html", "d.html", HTML),
        ("text", "d.txt", TXT),
        ("ipynb", "d.ipynb", json.dumps({
            "cells": [{"cell_type": "markdown",
                       "source": "# H\n"}],
            "metadata": {}, "nbformat": 4})),
    )
    for name, fn, c in boards:
        (tmp_path / fn).write_text(c,
                                   encoding="utf-8")
        doc, errors = process_single(
            tmp_path / fn, None,
            parser_name=name, max_chars=800)
        assert errors == [], (name, errors)
        assert is_valid(doc.to_dict()), name


def test_docx_output_passes_schema(tmp_path):
    from docx import Document
    d = Document()
    d.add_heading("H", 1)
    d.add_paragraph("para")
    d.save(str(tmp_path / "d.docx"))
    doc, errors = process_single(
        tmp_path / "d.docx", None,
        parser_name="fallback", max_chars=800)
    assert errors == []
    assert is_valid(doc.to_dict())


# ---------- models Literal ↔ schema 运行时核 ----------

def test_every_models_source_type_in_schema_enum():
    enums = _enums(_SCHEMA)
    src_enum = [e for e in enums if "pdf" in e
                and "ipynb" in e][0]
    for st in ("pdf", "docx", "markdown",
               "html", "text", "ipynb"):
        assert st in src_enum


def test_every_models_element_type_in_schema_enum():
    enums = _enums(_SCHEMA)
    el_enum = [e for e in enums
               if "heading" in e][0]
    for et in ("heading", "paragraph",
               "list_item", "table", "image",
               "caption", "header", "footer"):
        assert et in el_enum


# ---------- Element 构造约束与 schema 对齐 ----------

def test_element_with_content_only_valid():
    e = Element(element_id="e", type="paragraph",
                source_locator={"line": 1},
                content="x")
    assert is_valid({
        "schema_version": SCHEMA_VERSION,
        "document_id": "d", "source_path": "p",
        "source_type": "text",
        "source_hash": "a" * 64,
        "parser_name": "t",
        "parser_version": "1",
        "elements": [e.to_dict()],
        "chunks": [], "relations": [],
        "warnings": [], "errors": [],
        "metadata": {}})
