"""evaluation/runner.py 第六百三十四轮 edges 测试（Round 1190）。

补强 edges200 未触及的角度（第五百六十二批，probe 实证）。

新角度（DOCX 嵌套表格的静默丢文）：
- **嵌套表不可见**——cell(0,0).add_table
  造嵌套表 → N1/N2 完全不进元素流
  （document.tables 只列顶层表 +
  cell.text 只取段落，双层滤除首锁）
- **外层表 markdown**——"| A1 | B1 |\\n
  | --- | --- |\\n| A2 | B2 |"，
  嵌套内容不入格
- **期望差即静默丢**——expectations
  {table: 2}（人计嵌套）→ silent_drop
  1；{table: 1} → 0
- **嵌套文锚不可寻**——marker "N1"
  不在流 → R no_ground_truth_
  anchors_in_stream 三态
- forbidden tokens 第六百六十二批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, expectations=None, anchors=None):
    from docx import Document
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "a").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Intro before the outer table.")
    outer = d.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "A1"
    outer.cell(0, 1).text = "B1"
    outer.cell(1, 0).text = "A2"
    outer.cell(1, 1).text = "B2"
    nested = outer.cell(0, 0).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "N1"
    nested.cell(0, 1).text = "N2"
    d.add_paragraph("Outro after the outer table.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    entry = {"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}
    if expectations is not None:
        entry["expectations"] = expectations
    if anchors is not None:
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        entry["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": [entry]}),
                  encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


_MD = "| A1 | B1 |\n| --- | --- |\n| A2 | B2 |"


# ---------- 嵌套表不可见 ----------

def test_nested_table_elements_batch388(tmp_path):
    _board(tmp_path, "ne")
    doc, errors = process_single(
        tmp_path / "s" / "ne.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "paragraph", "table", "paragraph"]
    assert els[1]["content"] == _MD
    assert not any("N1" in (e["content"] or "")
                   for e in els)
    assert not any("N2" in (e["content"] or "")
                   for e in els)


def test_nested_table_locators_batch388(tmp_path):
    _board(tmp_path, "ne2")
    doc, errors = process_single(
        tmp_path / "s" / "ne2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    els = doc.to_dict()["elements"]
    assert els[1]["source_locator"] == {
        "table_index": 0, "section": 0}
    assert els[0]["source_locator"]["paragraph_index"] == 0
    assert els[2]["source_locator"]["paragraph_index"] == 1


def test_nested_table_chunks_batch388(tmp_path):
    _board(tmp_path, "ne3")
    doc, errors = process_single(
        tmp_path / "s" / "ne3.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "isolated_table", "sequential"]
    assert [c["text"] for c in chunks] == [
        "Intro before the outer table.", _MD,
        "Outro after the outer table."]
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks)


# ---------- 期望差即静默丢 ----------

def test_silent_drop_match_batch388(tmp_path):
    r = run_evaluation(_board(tmp_path, "ne4", {
        "element_count_by_type": {"paragraph": 2, "table": 1}}),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["silent_drop_count"] == {"value": 0,
                                      "reason": None}


def test_silent_drop_nested_counted_batch388(tmp_path):
    r = run_evaluation(_board(tmp_path, "ne5", {
        "element_count_by_type": {"paragraph": 2, "table": 2}}),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["silent_drop_count"] == {"value": 1,
                                      "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 2, "table": 1},
        "reason": None}


# ---------- 锚 ----------

def test_nested_anchor_table_batch388(tmp_path):
    r = run_evaluation(_board(tmp_path, "ne6", None, [
        {"marker": "table.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


def test_nested_anchor_missing_batch388(tmp_path):
    r = run_evaluation(_board(tmp_path, "ne7", None, [
        {"marker": "N1", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": None,
        "reason": "no_ground_truth_anchors_in_stream"}
    assert m["chunk_boundary_f1"] == {
        "value": None,
        "reason": "precision_or_recall_not_evaluated"}


# ---------- 指标 ----------

def test_nested_metrics_batch388(tmp_path):
    r = run_evaluation(_board(tmp_path, "ne8"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["pdf_locator_valid_ratio"] == {
        "value": None, "reason": "not_pdf_document"}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch388():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百六十二批 ----------

def test_source_no_eval_batch388():
    assert "eval(" not in _src()


def test_source_no_exec_batch388():
    assert "exec(" not in _src()


def test_source_no_compile_batch388():
    assert "compile(" not in _src()


def test_source_no_globals_batch388():
    assert "globals(" not in _src()


def test_source_no_locals_batch388():
    assert "locals(" not in _src()


def test_source_no_os_system_batch388():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch388():
    assert "subprocess" not in _src()


def test_source_no_popen_batch388():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch388():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch388():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch388():
    assert "socket" not in _src()


def test_source_no_requests_batch388():
    assert "requests" not in _src()


def test_source_no_urllib_batch388():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch388():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch388():
    assert "yield" not in _src()


def test_source_no_async_await_batch388():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch388():
    assert _src().count("open(") == 2
