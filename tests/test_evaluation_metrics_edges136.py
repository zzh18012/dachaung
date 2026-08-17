"""evaluation/metrics.py 第五百五十一轮 edges 测试（Round 1107）。

补强 edges135 未触及的角度（第四百八十三批，probe 实证）。

新角度（schema-metrics 分歧 / heading 合规查 id / pdf bbox 豁免）：
- **schema-metrics 分歧**：docx 文档 elements[0].source_locator
  改 {page, bbox} → document.schema.json 照过（docx_locator
  def additionalProperties True，页键不违 schema）但
  docx_locator_valid_ratio == 0.5——_docx_locator_ratio 拒
  page/bbox 键：同一突变两本账分歧（schema 收、metrics 打折）
- **heading 合规查 id 不查文本**：intro+heading+body 三段
  mc 200 → heading 块 ids [e0001, e0002]，compliance 1.0；
  反转 ids → 0.0——合规只看 source_element_ids[0] 是否等于
  heading 的 element_id，块内文本顺序无关
- **pdf bbox 类型豁免**：paragraph {page, bbox} + image
  {page} 无 bbox → pdf_locator_valid_ratio 1.0——bbox 只对
  文本类型（heading/paragraph/caption/list_item）强制，
  image 豁免；换 caption page-only → 0.5（豁免名单外）
- forbidden tokens 第五百七十八批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from evaluation.metrics import compute_automatic_metrics
from evaluation.schema import validate


def _run(tmp_path, name, build, max_chars=200):
    p = tmp_path / name
    d = Document()
    build(d)
    d.save(str(p))
    doc, errors = process_single(
        p, tmp_path / "s.json", parser_name="fallback",
        max_chars=max_chars, write_json=False)
    assert errors == []
    return doc.to_dict()


# ---------- schema-metrics 分歧 ----------

def test_schema_metrics_divergence_batch306(tmp_path):
    r = _run(
        tmp_path, "a.docx",
        lambda d: (d.add_paragraph("AAA first body."),
                   d.add_paragraph("BBB second body.")))
    r["elements"][0]["source_locator"] = {
        "page": 1, "bbox": [0.0, 0.0, 1.0, 1.0]}
    validate(r, "document.schema.json")
    out = compute_automatic_metrics(r, None, "docx", None)
    assert out["docx_locator_valid_ratio"] == {
        "value": 0.5, "reason": None}


# ---------- heading 合规查 id ----------

def test_heading_compliance_id_swap_batch306(tmp_path):
    r = _run(
        tmp_path, "h.docx",
        lambda d: (d.add_paragraph("Intro text."),
                   d.add_heading("Late Title", level=1),
                   d.add_paragraph("Body after.")))
    out1 = compute_automatic_metrics(r, None, "docx", None)
    assert out1["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}
    head_id = [e["element_id"] for e in r["elements"]
               if e["type"] == "heading"][0]
    for c in r["chunks"]:
        if head_id in c["source_element_ids"]:
            c["source_element_ids"] = list(
                reversed(c["source_element_ids"]))
            break
    out2 = compute_automatic_metrics(r, None, "docx", None)
    assert out2["heading_boundary_compliance"] == {
        "value": 0.0, "reason": None}


# ---------- pdf bbox 类型豁免 ----------

def _pdf_doc(els):
    return {"document_id": "d", "elements": els, "chunks": []}


def test_pdf_bbox_image_exempt_batch306():
    els = [
        {"element_id": "e0", "type": "paragraph",
         "content": "pdf body",
         "source_locator": {
             "page": 1, "bbox": [0.0, 0.0, 1.0, 1.0]}},
        {"element_id": "e1", "type": "image",
         "resource_path": "img.png",
         "source_locator": {"page": 2}},
    ]
    out = compute_automatic_metrics(
        _pdf_doc(els), None, "pdf", None)
    assert out["pdf_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


def test_pdf_bbox_caption_required_batch306():
    els = [
        {"element_id": "e0", "type": "paragraph",
         "content": "pdf body",
         "source_locator": {
             "page": 1, "bbox": [0.0, 0.0, 1.0, 1.0]}},
        {"element_id": "e1", "type": "caption",
         "content": "cap",
         "source_locator": {"page": 3}},
    ]
    out = compute_automatic_metrics(
        _pdf_doc(els), None, "pdf", None)
    assert out["pdf_locator_valid_ratio"] == {
        "value": 0.5, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch306():
    src = _src()
    assert "至少一个结构键" in src
    assert "heading 是 chunk 的首元素" in src


# ---------- forbidden tokens 第五百七十八批 ----------

def test_source_no_eval_batch306():
    assert "eval(" not in _src()


def test_source_no_exec_batch306():
    assert "exec(" not in _src()


def test_source_no_compile_batch306():
    assert "compile(" not in _src()


def test_source_no_globals_batch306():
    assert "globals(" not in _src()


def test_source_no_locals_batch306():
    assert "locals(" not in _src()


def test_source_no_os_system_batch306():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch306():
    assert "subprocess" not in _src()


def test_source_no_popen_batch306():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch306():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch306():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch306():
    assert "socket" not in _src()


def test_source_no_requests_batch306():
    assert "requests" not in _src()


def test_source_no_urllib_batch306():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch306():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch306():
    assert "yield" not in _src()


def test_source_no_async_await_batch306():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch306():
    assert "open(" not in _src()
