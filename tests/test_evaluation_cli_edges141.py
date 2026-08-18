"""evaluation/cli.py 第五百六十六轮 edges 测试（Round 1122）。

补强 edges140 未触及的角度（第四百九十八批，probe 实证）。

新角度（双 main 调用链 / 重跑覆盖）：
- **run → validate-report 链**：同一进程两次 main()——先
  run 真实 docx 板产报告（rc 0），再 validate-report 校验
  同一文件 → rc 0 + [OK]——CLI 自产自校闭环（首锁；
  runner 层 API 版已锁 edges141）
- **重跑覆盖不累积**：同一 output 连跑两次 → 第二次 rc 0，
  per_doc 仍只有 1 条（新报告整体覆盖，不追加），且覆盖后
  validate-report 照过——重跑幂等（首锁）
- forbidden tokens 第五百九十四批（open 1）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document

import evaluation.cli as cli_mod
from evaluation.cli import main


def _manifest(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA head start.")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1", "path": "samples/g.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _run(tmp_path, out):
    buf_o, buf_e = io.StringIO(), io.StringIO()
    with contextlib.redirect_stdout(buf_o), \
            contextlib.redirect_stderr(buf_e):
        rc = main(["run", "--manifest",
                   str(_manifest(tmp_path)),
                   "--output", str(out),
                   "--parser", "fallback",
                   "--max-chars", "200"])
    return rc


def _validate(out):
    buf_o, buf_e = io.StringIO(), io.StringIO()
    with contextlib.redirect_stdout(buf_o), \
            contextlib.redirect_stderr(buf_e):
        rc = main(["validate-report", str(out)])
    return rc, buf_o.getvalue(), buf_e.getvalue()


# ---------- run → validate-report 链 ----------

def test_run_then_validate_report_chain_batch321(tmp_path):
    out = tmp_path / "r.json"
    assert _run(tmp_path, out) == 0
    rc, ok_out, _ = _validate(out)
    assert rc == 0
    assert "[OK]" in ok_out


# ---------- 重跑覆盖不累积 ----------

def test_run_twice_overwrites_cleanly_batch321(tmp_path):
    out = tmp_path / "r.json"
    assert _run(tmp_path, out) == 0
    assert _run(tmp_path, out) == 0
    r = json.loads(out.read_text(encoding="utf-8"))
    assert r["devset"]["file_count"] == 1
    assert len(r["per_doc"]) == 1
    rc, ok_out, _ = _validate(out)
    assert rc == 0
    assert "[OK]" in ok_out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch321():
    src = _src()
    assert "不写报告，不强制 manifest" in src
    assert "待校验的报告 JSON 路径" in src


# ---------- forbidden tokens 第五百九十四批 ----------

def test_source_no_eval_batch321():
    assert "eval(" not in _src()


def test_source_no_exec_batch321():
    assert "exec(" not in _src()


def test_source_no_compile_batch321():
    assert "compile(" not in _src()


def test_source_no_globals_batch321():
    assert "globals(" not in _src()


def test_source_no_locals_batch321():
    assert "locals(" not in _src()


def test_source_no_os_system_batch321():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch321():
    assert "subprocess" not in _src()


def test_source_no_popen_batch321():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch321():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch321():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch321():
    assert "socket" not in _src()


def test_source_no_requests_batch321():
    assert "requests" not in _src()


def test_source_no_urllib_batch321():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch321():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch321():
    assert "yield" not in _src()


def test_source_no_async_await_batch321():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch321():
    assert _src().count("open(") == 1
