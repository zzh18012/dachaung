"""evaluation/cli.py 第四百八十五轮 edges 测试（Round 1041）。

补强 edges128 未触及的角度（第四百一十七批，probe 实证）。

新角度（--max-chars 旗标翻转真实 docx 结局）：
- edges98 已锁真实 docx 默认 max_chars 全绿 run；本批
  同一真实文件仅改 --max-chars：50 → 全绿、30 → 真实
  结构分块器 chunker_failed，两方向同屏对比
- 真实管线失败时 CLI run 仍退出 0（结构化失败是一次
  成功的评测）：成功/失败计数行翻转
  （成功 1，失败 0 → 成功 0，失败 1），stderr 恒空
- 失败管线产出的报告仍过 evaluation-report RS，且
  validate-report 子命令对同一报告 rc 0 PASS——
  "报告合法性"与"管线成败"正交，真实文件双向锁
- forbidden tokens 第五百一十二批（open 1）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document

import evaluation.cli as cli_mod
from evaluation.cli import main


def _setup(tmp_path):
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_paragraph("Hello world paragraph one.")
    d.add_paragraph("Second paragraph here.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    mf = tmp_path / "manifest.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/a.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _run(tmp_path, mc):
    out, err = io.StringIO(), io.StringIO()
    rp = tmp_path / f"r{mc}.json"
    with contextlib.redirect_stdout(out), \
            contextlib.redirect_stderr(err):
        rc = main(["run", "--manifest",
                   str(tmp_path / "manifest.json"),
                   "--output", str(rp),
                   "--max-chars", str(mc)])
    return rc, out.getvalue(), err.getvalue(), rp


# ---------- 全绿门 ----------

def test_mc50_counts_green_batch239(tmp_path):
    _setup(tmp_path)
    rc, out, err, rp = _run(tmp_path, 50)
    assert rc == 0
    assert err == ""
    lines = out.splitlines()
    assert lines[1] == "      documents=1（成功 1，失败 0）"
    rep = json.loads(rp.read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"]["error_code"] == {
        "value": None, "reason": None}
    assert rep["summary"]["success_rates"][
        "pipeline_success"]["rate"] == 1.0


# ---------- 旗标翻转门 ----------

def test_mc30_exit_zero_fail_count_batch239(tmp_path):
    _setup(tmp_path)
    rc, out, err, _ = _run(tmp_path, 30)
    assert rc == 0
    assert err == ""
    assert "documents=1（成功 0，失败 1）" in out


def test_flag_flip_rate_contrast_batch239(tmp_path):
    _setup(tmp_path)
    rates = []
    for mc in (50, 30):
        _, _, _, rp = _run(tmp_path, mc)
        rep = json.loads(rp.read_text(encoding="utf-8"))
        rates.append(rep["summary"]["success_rates"][
            "pipeline_success"]["rate"])
    assert rates == [1.0, 0.0]


# ---------- 失败报告合法性正交 ----------

def test_mc30_failed_report_rs_and_vr_batch239(tmp_path):
    _setup(tmp_path)
    rc, _, _, rp = _run(tmp_path, 30)
    assert rc == 0
    rep = json.loads(rp.read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"]["error_code"] == {
        "value": "chunker_failed", "reason": None}
    out, err = io.StringIO(), io.StringIO()
    with contextlib.redirect_stdout(out), \
            contextlib.redirect_stderr(err):
        vrc = main(["validate-report", str(rp)])
    assert vrc == 0
    assert "通过 evaluation-report Schema 校验" in \
        out.getvalue()


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch239():
    src = _src()
    assert "n_fail = n_docs - n_ok" in src
    assert ('if r["metrics"].get("pipeline_success", {})'
            in src)


# ---------- forbidden tokens 第五百一十二批 ----------

def test_source_no_eval_batch239():
    assert "eval(" not in _src()


def test_source_no_exec_batch239():
    assert "exec(" not in _src()


def test_source_no_compile_batch239():
    assert "compile(" not in _src()


def test_source_no_globals_batch239():
    assert "globals(" not in _src()


def test_source_no_locals_batch239():
    assert "locals(" not in _src()


def test_source_no_os_system_batch239():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch239():
    assert "subprocess" not in _src()


def test_source_no_popen_batch239():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch239():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch239():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch239():
    assert "socket" not in _src()


def test_source_no_requests_batch239():
    assert "requests" not in _src()


def test_source_no_urllib_batch239():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch239():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch239():
    assert "yield" not in _src()


def test_source_no_async_await_batch239():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch239():
    assert _src().count("open(") == 1
