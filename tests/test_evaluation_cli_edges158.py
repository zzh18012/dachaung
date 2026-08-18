"""evaluation/cli.py 第六百五十四轮 edges 测试（Round 1242）。

补强 cli edges157 未触及的角度（第六百一十四批，probe 实证）。

新角度（标注板 + 容差归零敏感性）：
- **tol 0 劈开双锚对**——双
  "(空段落)" after 锚：第一处
  恰收尾 chunk1（d 0 恒中），
  第二处落 chunk3 内距界 2 约
  21 字符 → tol 0 漏 → P/R/F1
  全 0.5（容差敏感性真板首锁，
  与 edges143 的手造 tol 板相
  对）
- **默认 tol 30 全中**——同板
  无 flag → 全 1.0（21 ≤ 30）
- **标注报告过 validate-report**
- forbidden tokens 第五百八十批（open 1）
"""

from __future__ import annotations

import inspect
import json

import evaluation.cli as cli_mod
from evaluation.cli import main


def _docx(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    (tmp_path / "s").mkdir(exist_ok=True)
    p = tmp_path / "s" / "ks.docx"
    doc.save(str(p))
    return p


def _board(tmp_path):
    _docx(tmp_path)
    (tmp_path / "a").mkdir(exist_ok=True)
    (tmp_path / "a" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "ks",
        "chunk_boundary_anchors": [
            {"marker": "(空段落)", "position": "after"},
            {"marker": "(空段落)", "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "ks", "path": "s/ks.docx",
                       "source_type": "docx",
                       "annotation_file": "a/a.json"}]}),
        encoding="utf-8")
    return mf


def _run_cli(capsys, argv):
    rc = main(argv)
    out = capsys.readouterr().out
    return rc, out


# ---------- tol 0 劈开双锚对 ----------

def test_cli_run_tol0_half_batch440(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120",
        "--tolerance", "0"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    m = rep["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.5, "reason": None}


def test_cli_run_tol0_stdout_batch440(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120",
        "--tolerance", "0"])
    assert rc == 0
    assert "[OK]" in out
    assert "documents=1（成功 1，失败 0）" in out


# ---------- 默认 tol 30 全中 ----------

def test_cli_run_default_tol_all_hit_batch440(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    m = rep["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_cli_run_devset_line_batch440(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120",
        "--tolerance", "0"])
    assert rc == 0
    assert "devset_status=incomplete file_count=1 groups=1 pdf=0 docx=1" \
        in out


# ---------- 标注报告过 validate-report ----------

def test_cli_validate_annotated_report_batch440(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120",
        "--tolerance", "0"])
    assert rc == 0
    rc2, out2 = _run_cli(capsys, [
        "validate-report", str(tmp_path / "r.json")])
    assert rc2 == 0
    assert "[OK]" in out2


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_identifier_counts_batch440():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


# ---------- forbidden tokens 第五百八十批 ----------

def test_source_no_eval_batch440():
    assert "eval(" not in _src()


def test_source_no_exec_batch440():
    assert "exec(" not in _src()


def test_source_no_compile_batch440():
    assert "compile(" not in _src()


def test_source_no_globals_batch440():
    assert "globals(" not in _src()


def test_source_no_locals_batch440():
    assert "locals(" not in _src()


def test_source_no_os_system_batch440():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch440():
    assert "subprocess" not in _src()


def test_source_no_popen_batch440():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch440():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch440():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch440():
    assert "socket" not in _src()


def test_source_no_requests_batch440():
    assert "requests" not in _src()


def test_source_no_urllib_batch440():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch440():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch440():
    assert "yield" not in _src()


def test_source_no_async_await_batch440():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch440():
    assert _src().count("open(") == 1
