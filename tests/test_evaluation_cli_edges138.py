"""evaluation/cli.py 第五百四十八轮 edges 测试（Round 1104）。

补强 edges135-137 未触及的角度（第四百八十批，probe 实证）。

新角度（kreuzberg 真跑端到端 + 双 parser 指标等价）：
- **kreuzberg 真跑端到端**：--parser kreuzberg 对
  真实 docx 走完整管线 → rc 0 + [OK] + success
  {1, 1, 1.0} + element_count_total 2——旧 cli
  kreuzberg 测试全是 mock 捕参（edges112）或
  合成 dict（edges34），真跑首锁；provenance
  parser_name "kreuzberg" + parser_version
  "4.10.2"（真实版本号入档）
- **双 parser 指标等价**：同一文档 fallback 与
  kreuzberg 各跑一次 → 大四指标值完全一致
  （ect 2 / docx_loc 1.0 / chunk_ref 1.0 /
  text_eq True）——两 parser 在指标层可互换，
  差异只在 provenance
- **kreuzberg 报告过验**：真跑报告 validate-report
  rc 0——report schema 的 parser_name 接受
  "kreuzberg"（自由字符串）
- forbidden tokens 第五百七十五批（open 1）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.cli as cli_mod
from evaluation.cli import main


def _manifest(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA kreuzberg probe body.")
    d.add_paragraph("BBB second line.")
    d.save(str(tmp_path / "samples" / "g.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{
            "doc_id": "d1", "path": "samples/g.docx",
            "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _run(tmp_path, parser):
    out = tmp_path / f"r-{parser}.json"
    rc = main(["run", "--manifest", str(_manifest(tmp_path)),
               "--output", str(out), "--parser", parser])
    assert rc == 0
    return json.loads(out.read_text(encoding="utf-8"))


# ---------- kreuzberg 真跑端到端 ----------

def test_kreuzberg_real_run_batch303(tmp_path):
    blob = _run(tmp_path, "kreuzberg")
    m = blob["per_doc"][0]["metrics"]
    assert m["pipeline_success"] == {
        "value": True, "reason": None}
    assert m["element_count_total"] == {
        "value": 2, "reason": None}
    prov = blob["provenance"]
    assert prov["parser_name"] == "kreuzberg"
    assert prov["parser_version"] == "4.10.2"


# ---------- 双 parser 指标等价 ----------

def test_dual_parser_metric_equivalence_batch303(tmp_path):
    fb = _run(tmp_path, "fallback")
    kz = _run(tmp_path, "kreuzberg")
    for k in ("element_count_total",
              "docx_locator_valid_ratio",
              "chunk_reference_intact_ratio",
              "text_preservation_equal"):
        assert fb["per_doc"][0]["metrics"][k] == \
            kz["per_doc"][0]["metrics"][k]
    assert fb["provenance"]["parser_name"] == "fallback"
    assert kz["provenance"]["parser_name"] == \
        "kreuzberg"


# ---------- kreuzberg 报告过验 ----------

def test_kreuzberg_report_validates_batch303(tmp_path,
                                             capsys):
    blob = _run(tmp_path, "kreuzberg")
    out = tmp_path / "r-kreuzberg.json"
    rc = main(["validate-report", str(out)])
    assert rc == 0
    assert "通过 evaluation-report Schema 校验" in \
        capsys.readouterr().out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch303():
    src = _src()
    assert ('sub.add_parser("run", help="跑评测，'
            '生成报告 JSON")') in src
    assert "default=30," in src


# ---------- forbidden tokens 第五百七十五批 ----------

def test_source_no_eval_batch303():
    assert "eval(" not in _src()


def test_source_no_exec_batch303():
    assert "exec(" not in _src()


def test_source_no_compile_batch303():
    assert "compile(" not in _src()


def test_source_no_globals_batch303():
    assert "globals(" not in _src()


def test_source_no_locals_batch303():
    assert "locals(" not in _src()


def test_source_no_os_system_batch303():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch303():
    assert "subprocess" not in _src()


def test_source_no_popen_batch303():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch303():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch303():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch303():
    assert "socket" not in _src()


def test_source_no_requests_batch303():
    assert "requests" not in _src()


def test_source_no_urllib_batch303():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch303():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch303():
    assert "yield" not in _src()


def test_source_no_async_await_batch303():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch303():
    assert _src().count("open(") == 1
