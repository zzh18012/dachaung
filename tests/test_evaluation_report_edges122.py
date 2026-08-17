"""evaluation/report.py 第五百一十一轮 edges 测试（Round 1067）。

补强 edges119-121 未触及的角度（第四百四十三批，probe 实证）。

新角度（参与度三层塔：全/半/零，真实富板对素板）：
- 富板（heading+嵌图+标注+expectations）对素板（2 段、
  无标注无 expectations）同 run——12 个 ratio 指标分成
  三层：**全参与 2/2**（schema/docx_locator/chunk_ref/
  text_* 五项）、**半参与 1/1**（image——真图点亮 1.0、
  素板 null no_image_elements；heading_boundary——富板
  有标题、素板无；boundary 三项——富板有标注锚、素板
  无标注）、**零参与 0/2**（pdf_locator 门控）
- silent 混账：d1 {0, None}（精确相抵）+ d2 {None,
  no_expectations} → silent_drop_total **0 而非 None**
  ——0 参与求和、null 不参与，零值合法入账
- counts {sum 7, participating 2}：富板 5 元素（heading
  1 + para 3 + image 1）+ 素板 2
- forbidden tokens 第五百三十八批（report 变体：15 项
  去 subprocess + open 0 + subprocess.run == 2）
"""

from __future__ import annotations

import inspect
import json
import struct
import zlib
from io import BytesIO

from docx import Document

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA first paragraph body with "
                    "enough text to split nicely here.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("BBB third paragraph body with "
                    "enough text to split nicely here.")
    d.save(str(tmp_path / "samples" / "rich.docx"))
    d2 = Document()
    d2.add_paragraph("CCC first paragraph body.")
    d2.add_paragraph("DDD second paragraph body.")
    d2.save(str(tmp_path / "samples" / "plain.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "(空段落)",
             "position": "before"}]}), encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/rich.docx",
             "source_type": "docx",
             "expectations": {"element_count_by_type":
                              {"image": 1}},
             "annotation_file": "anns/a.json"},
            {"doc_id": "d2", "path": "samples/plain.docx",
             "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json", max_chars=80)


# ---------- 参与度三层塔 ----------

def test_participation_three_tiers_batch266(tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    full = {"schema_valid", "docx_locator_valid_ratio",
            "chunk_reference_intact_ratio",
            "text_preservation_equal",
            "text_char_multiset_precision",
            "text_char_multiset_recall"}
    half = {"image_resource_exists_ratio",
            "heading_boundary_compliance",
            "chunk_boundary_precision",
            "chunk_boundary_recall",
            "chunk_boundary_f1"}
    assert set(ra) == full | half | {"pdf_locator_valid_ratio"}
    for k in full:
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 2,
                         "not_evaluated": 0}, k
    for k in half:
        assert ra[k] == {"macro_average": 1.0,
                         "participating_docs": 1,
                         "not_evaluated": 1}, k
    assert ra["pdf_locator_valid_ratio"] == {
        "macro_average": None, "participating_docs": 0,
        "not_evaluated": 2}


# ---------- 半参与层的真实出处 ----------

def test_partial_image_tier_batch266(tmp_path):
    rep = _run(tmp_path)
    assert rep["per_doc"][1]["metrics"][
        "image_resource_exists_ratio"] == {
        "value": None, "reason": "no_image_elements"}


# ---------- silent 混账：0 与 null 共存 ----------

def test_silent_mixed_zero_and_null_batch266(tmp_path):
    rep = _run(tmp_path)
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 0,
                                 "reason": None}
    assert rep["per_doc"][1]["metrics"][
        "silent_drop_count"] == {
        "value": None, "reason": "no_expectations"}
    assert rep["summary"]["silent_drop_total"] == 0


# ---------- counts 求和 ----------

def test_counts_sum_seven_batch266(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["counts"] == {
        "element_count_total": {"sum": 7,
                                "participating_docs": 2}}


# ---------- 成功全绿 ----------

def test_success_full_batch266(tmp_path):
    rep = _run(tmp_path)
    assert rep["summary"]["success_rates"] == {
        "pipeline_success": {"success_count": 2,
                             "total": 2, "rate": 1.0}}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch266():
    src = _src()
    assert ("not_eval = len(per_doc_results) - len(values)"
            in src)
    assert '"not_evaluated": not_eval' in src


# ---------- forbidden tokens 第五百三十八批（report 变体） ----------

def test_source_no_eval_batch266():
    assert "eval(" not in _src()


def test_source_no_exec_batch266():
    assert "exec(" not in _src()


def test_source_no_compile_batch266():
    assert "compile(" not in _src()


def test_source_no_globals_batch266():
    assert "globals(" not in _src()


def test_source_no_locals_batch266():
    assert "locals(" not in _src()


def test_source_no_os_system_batch266():
    assert "os.system" not in _src()


def test_source_no_popen_batch266():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch266():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch266():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch266():
    assert "socket" not in _src()


def test_source_no_requests_batch266():
    assert "requests" not in _src()


def test_source_no_urllib_batch266():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch266():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch266():
    assert "yield" not in _src()


def test_source_no_async_await_batch266():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch266():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch266():
    assert _src().count("subprocess.run") == 2
