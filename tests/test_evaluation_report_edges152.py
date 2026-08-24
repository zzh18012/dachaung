"""evaluation/report.py 第五百八十三轮 edges 测试（Round 1353）。

补强 edges151 未触及的角度（第七百二十五批，probe 实证）。

新角度（unicode categories 全链排序 / 混合 expectations 板）：
- **unicode 排序**——
  ['测试','alpha']+
  ['中文'] → ['alpha',
  '中文','测试']
  （ASCII 先于 CJK、
  中文 U+4E2D <
  测试 U+6D4B 全链
  首锁；edges10 仅
  FakeManifest 预排）
- **混合 sdt**——
  doc1 无 expectations
  {None} + doc2 sdc 2
  → total 2（None
  不参与但成功不减）
- **双 docx 板**——
  files 2 / groups 2 /
  success 2/2
- forbidden tokens 第七百九十三批（open 0）
"""

from __future__ import annotations

import inspect
import json
import tempfile
from pathlib import Path

import evaluation.report as report_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _run(tmp_path):
    d = Document()
    d.add_paragraph("one")
    d.save(str(tmp_path / "a.docx"))
    d = Document()
    d.add_paragraph("two")
    d.save(str(tmp_path / "b.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "a.docx",
             "source_type": "docx",
             "categories": ["测试", "alpha"]},
            {"doc_id": "d2", "path": "b.docx",
             "source_type": "docx",
             "categories": ["中文"],
             "expectations": {
                 "element_count_by_type": {
                     "paragraph": 3}}},
        ]}), encoding="utf-8")
    mf = load_manifest(tmp_path / "m.json",
                       project_root=tmp_path)
    return run_evaluation(mf, tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=800)


# ---------- unicode 排序 ----------

def test_categories_sorted_batch551(tmp_path):
    assert _run(tmp_path)["devset"][
        "categories_covered"] == [
        "alpha", "中文", "测试"]


def test_ascii_before_cjk_batch551(tmp_path):
    cats = _run(tmp_path)["devset"][
        "categories_covered"]
    assert cats.index("alpha") < cats.index("中文")


def test_cjk_codepoint_order_batch551(tmp_path):
    cats = _run(tmp_path)["devset"][
        "categories_covered"]
    assert cats.index("中文") < cats.index("测试")


def test_categories_len_three_batch551(tmp_path):
    assert len(_run(tmp_path)["devset"][
        "categories_covered"]) == 3


# ---------- 混合 sdt ----------

def test_mixed_sdt_two_batch551(tmp_path):
    assert _run(tmp_path)["summary"][
        "silent_drop_total"] == 2


def test_doc1_sdc_no_expectations_batch551(
        tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "silent_drop_count"] == {
        "value": None,
        "reason": "no_expectations"}


def test_doc2_sdc_two_batch551(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "silent_drop_count"] == {
        "value": 2, "reason": None}


def test_mixed_sdt_per_doc_none_and_two_batch551(
        tmp_path):
    r = _run(tmp_path)
    sdcs = [p["metrics"]["silent_drop_count"][
        "value"] for p in r["per_doc"]]
    assert sdcs == [None, 2]


def test_sdt_equals_participating_sum_batch551(
        tmp_path):
    r = _run(tmp_path)
    vals = [p["metrics"]["silent_drop_count"][
        "value"] for p in r["per_doc"]]
    assert r["summary"]["silent_drop_total"] \
        == sum(v for v in vals if v is not None)


# ---------- 双 docx 板 ----------

def test_devset_full_dict_batch551(tmp_path):
    assert _run(tmp_path)["devset"] == {
        "status": "incomplete",
        "file_count": 2,
        "content_group_count": 2,
        "pdf_count": 0, "docx_count": 2,
        "categories_covered": [
            "alpha", "中文", "测试"]}


def test_success_two_of_two_batch551(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 2, "total": 2,
        "rate": 1.0}


def test_ect_sum_two_batch551(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"][
        "element_count_total"] == {
        "sum": 2, "participating_docs": 2}


def test_each_doc_one_paragraph_batch551(
        tmp_path):
    r = _run(tmp_path)
    for p in r["per_doc"]:
        assert p["metrics"][
            "element_count_by_type"] == {
            "value": {"paragraph": 1},
            "reason": None}


def test_hbc_none_two_batch551(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "heading_boundary_compliance"] == {
        "macro_average": None,
        "participating_docs": 0,
        "not_evaluated": 2}


def test_tpe_macro_one_batch551(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "text_preservation_equal"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 0}


# ---------- 报告合法性 ----------

def test_report_schema_batch551(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


def test_report_on_disk_round_trip_batch551(
        tmp_path):
    r = _run(tmp_path)
    on_disk = json.loads(
        (tmp_path / "r.json").read_text(
            encoding="utf-8"))
    assert on_disk == r


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_silent_sum_line_batch551():
    src = _src()
    assert "silent_drop_total" in src
    assert 'summary["silent_drop_total"]' in src


def test_source_manifest_categories_batch551():
    src = _src()
    assert "manifest.categories_covered" in src
    assert "manifest.content_group_count" in src


# ---------- forbidden tokens 第七百九十三批 ----------

def test_source_no_eval_batch551():
    assert "eval(" not in _src()


def test_source_no_exec_batch551():
    assert "exec(" not in _src()


def test_source_no_compile_batch551():
    assert "compile(" not in _src()


def test_source_no_globals_batch551():
    assert "globals(" not in _src()


def test_source_no_locals_batch551():
    assert "locals(" not in _src()


def test_source_no_os_system_batch551():
    assert "os.system" not in _src()


def test_source_subprocess_run_two_batch551():
    assert _src().count("subprocess.run") == 2


def test_source_no_popen_batch551():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch551():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch551():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch551():
    assert "socket" not in _src()


def test_source_no_requests_batch551():
    assert "requests" not in _src()


def test_source_no_urllib_batch551():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch551():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch551():
    assert "yield" not in _src()


def test_source_no_async_await_batch551():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch551():
    assert ".call(" not in _src()


def test_source_open_count_is_0_batch551():
    assert _src().count("open(") == 0
