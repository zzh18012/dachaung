"""evaluation/report.py 第五百六十五轮 edges 测试（Round 1239）。

补强 edges133 未触及的角度（第六百一十一批，probe 实证）。

新角度（混合真板单清单聚合全像）：
- **12 ratio 键全锁**——真实
  DOCX（水槽板）+ PDF（40 词板）
  单清单跑 run_evaluation，
  summary.ratio_macro_averages
  恰 12 键（真实数据驱动首锁，
  前史全手造条目）
- **locator 分源参评聚合**——
  pdf_locator {1.0, 1, 1} 与
  docx_locator {1.0, 1, 1} 镜像
  （聚合层的分源分裂）
- **hbc 单参评**——仅 DOCX 有
  heading → {1.0, 1, 1}（PDF 无
  heading → not_evaluated）
- **双参评布尔**——schema_valid /
  tpe {1.0, 2, 0}
- **counts 跨源求和**——ect 7+1
  = sum 8 / participating 2
- forbidden tokens 第七百零五批（open 0，subprocess.run 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _board(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    (tmp_path / "s").mkdir(exist_ok=True)
    doc.save(str(tmp_path / "s" / "ks.docx"))

    words = " ".join("w%02d" % i for i in range(40))
    s = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
         % words).encode()
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: (b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    (tmp_path / "s" / "u5.pdf").write_bytes(bytes(out))

    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [
            {"doc_id": "ks", "path": "s/ks.docx",
             "source_type": "docx"},
            {"doc_id": "u5", "path": "s/u5.pdf",
             "source_type": "pdf"}]}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _summary(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=120)
    return r["summary"]


# ---------- 12 ratio 键全锁 ----------

def test_ratio_keys_twelve_batch437(tmp_path):
    s = _summary(tmp_path)
    assert sorted(s["ratio_macro_averages"].keys()) == [
        "chunk_boundary_f1", "chunk_boundary_precision",
        "chunk_boundary_recall", "chunk_reference_intact_ratio",
        "docx_locator_valid_ratio",
        "heading_boundary_compliance",
        "image_resource_exists_ratio",
        "pdf_locator_valid_ratio", "schema_valid",
        "text_char_multiset_precision",
        "text_char_multiset_recall",
        "text_preservation_equal"]


# ---------- locator 分源参评聚合 ----------

def test_pdf_locator_split_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["ratio_macro_averages"]["pdf_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


def test_docx_locator_split_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["ratio_macro_averages"]["docx_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


# ---------- hbc 单参评 ----------

def test_hbc_single_participant_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["ratio_macro_averages"][
        "heading_boundary_compliance"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


# ---------- 双参评布尔 ----------

def test_both_docs_boolean_ratios_batch437(tmp_path):
    s = _summary(tmp_path)
    for name in ("schema_valid", "text_preservation_equal"):
        assert s["ratio_macro_averages"][name] == {
            "macro_average": 1.0, "participating_docs": 2,
            "not_evaluated": 0}


# ---------- 零参评桶 ----------

def test_image_zero_participating_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["ratio_macro_averages"][
        "image_resource_exists_ratio"] == {
        "macro_average": None, "participating_docs": 0,
        "not_evaluated": 2}


def test_boundary_zero_participating_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["ratio_macro_averages"]["chunk_boundary_f1"] == {
        "macro_average": None, "participating_docs": 0,
        "not_evaluated": 2}


# ---------- counts 跨源求和 ----------

def test_counts_sum_eight_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["counts"]["element_count_total"] == {
        "sum": 8, "participating_docs": 2}


def test_success_two_of_two_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["success_rates"]["pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}


def test_silent_drop_none_batch437(tmp_path):
    s = _summary(tmp_path)
    assert s["silent_drop_total"] is None


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch437():
    src = _src()
    assert "def aggregate_summary(" in src
    assert "_RATIO_METRICS" in src


# ---------- forbidden tokens 第七百零五批 ----------

def test_source_no_eval_batch437():
    assert "eval(" not in _src()


def test_source_no_exec_batch437():
    assert "exec(" not in _src()


def test_source_no_compile_batch437():
    assert "compile(" not in _src()


def test_source_no_globals_batch437():
    assert "globals(" not in _src()


def test_source_no_locals_batch437():
    assert "locals(" not in _src()


def test_source_no_os_system_batch437():
    assert "os.system" not in _src()


def test_source_no_subprocess_call_batch437():
    assert ".call(" not in _src()


def test_source_no_popen_batch437():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch437():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch437():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch437():
    assert "socket" not in _src()


def test_source_no_requests_batch437():
    assert "requests" not in _src()


def test_source_no_urllib_batch437():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch437():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch437():
    assert "yield" not in _src()


def test_source_no_async_await_batch437():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch437():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch437():
    assert _src().count("subprocess.run") == 2
