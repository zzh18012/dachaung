r"""app/parsers/fallback_parser.py 边角测试 - 第十四轮（Round 1390）。

新角度（probe 实证）：真 docx 内嵌 inline 图片穿真实提取
（python-docx add_picture + 手工 1x1 PNG，历史图片测试全是
假 XML 单元级）：
- image 元素 + 同段 '(空段落)' 占位符（图所在段落无文本）
- image_output_dir 时 resource_path 是落盘绝对路径，
  文件名 image_<hash16>_paraN_MM.png，文件真实存在
- 不带 dir → '(unsaved)' 占位符、不落盘、无告警
- locator 带 relationship_id/target_partname
- metadata {byte_size, ext, extracted_to_disk}
- 两图两段 → 计数后缀 _00/_01 各自落盘
- 管线层 image 不进任何 chunk；irer 0.0（unsaved 路径）
"""

from __future__ import annotations

import io
import re
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser


def _make_png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    ihdr = struct.pack(
        ">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return (sig + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat)
            + chunk(b"IEND", b""))


def _image_docx(tmp_path, name="img.docx",
                pictures=1):
    d = Document()
    d.add_heading("Image Doc", 1)
    d.add_paragraph("Before image.")
    for _ in range(pictures):
        d.add_picture(
            io.BytesIO(_make_png()),
            width=914400)
    d.add_paragraph("After image.")
    p = tmp_path / name
    d.save(str(p))
    return p


def _parse_images(tmp_path, pictures=1):
    p = _image_docx(tmp_path,
                    pictures=pictures)
    img_dir = tmp_path / "imgs"
    img_dir.mkdir()
    doc = FallbackParser(
        image_output_dir=str(img_dir)
    ).parse(p, compute_file_hash(p))
    return doc, img_dir


# ---------- 元素结构 ----------

def test_image_element_types(tmp_path):
    doc, _ = _parse_images(tmp_path)
    assert [e.type for e in doc.elements
            ] == ["heading", "paragraph",
                  "paragraph", "image",
                  "paragraph"]


def test_image_paragraph_placeholder(
        tmp_path):
    doc, _ = _parse_images(tmp_path)
    assert doc.elements[2].content == \
        "(空段落)"
    assert doc.elements[2].metadata[
        "empty"] is True


def test_image_same_paragraph_index(
        tmp_path):
    doc, _ = _parse_images(tmp_path)
    assert doc.elements[2].source_locator[
        "paragraph_index"] == 2
    assert doc.elements[3].source_locator[
        "paragraph_index"] == 2


def test_image_content_none(tmp_path):
    doc, _ = _parse_images(tmp_path)
    img = doc.elements[3]
    assert img.content is None
    assert img.resource_path is not None


# ---------- 落盘文件 ----------

def test_image_resource_path_in_dir(
        tmp_path):
    doc, img_dir = _parse_images(tmp_path)
    rp = doc.elements[3].resource_path
    assert rp.startswith(str(img_dir))


def test_image_filename_pattern(
        tmp_path):
    doc, _ = _parse_images(tmp_path)
    name = Path(
        doc.elements[3].resource_path
    ).name
    assert re.fullmatch(
        r"image_[0-9a-f]{16}"
        r"_para2_00\.png", name)


def test_image_file_exists(tmp_path):
    doc, _ = _parse_images(tmp_path)
    assert Path(
        doc.elements[3].resource_path
    ).is_file()


def test_image_file_bytes(tmp_path):
    doc, _ = _parse_images(tmp_path)
    assert Path(
        doc.elements[3].resource_path
    ).read_bytes() == _make_png()


def test_image_metadata(tmp_path):
    doc, _ = _parse_images(tmp_path)
    assert doc.elements[3].metadata == {
        "byte_size": 69, "ext": "png",
        "extracted_to_disk": True}


# ---------- locator ----------

def test_image_locator_rid(tmp_path):
    doc, _ = _parse_images(tmp_path)
    loc = doc.elements[3].source_locator
    assert loc["relationship_id"] == \
        "rId9"


def test_image_locator_target(tmp_path):
    doc, _ = _parse_images(tmp_path)
    loc = doc.elements[3].source_locator
    assert loc["target_partname"] == \
        "/word/media/image1.png"


# ---------- 两图计数 ----------

def test_two_images_two_files(tmp_path):
    doc, img_dir = _parse_images(
        tmp_path, pictures=2)
    names = sorted(
        f.name for f in img_dir.iterdir())
    assert len(names) == 2
    assert names[0].endswith(
        "_para2_00.png")
    assert names[1].endswith(
        "_para3_01.png")


def test_two_images_elements(tmp_path):
    doc, _ = _parse_images(
        tmp_path, pictures=2)
    imgs = [e for e in doc.elements
            if e.type == "image"]
    assert len(imgs) == 2
    assert imgs[0].resource_path != \
        imgs[1].resource_path


# ---------- 无 dir 占位 ----------

def test_no_dir_unsaved_placeholder(
        tmp_path):
    p = _image_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    imgs = [e for e in doc.elements
            if e.type == "image"]
    assert imgs[0].resource_path == \
        "(unsaved)"


def test_no_dir_no_file_written(
        tmp_path):
    p = _image_docx(tmp_path)
    FallbackParser().parse(
        p, compute_file_hash(p))
    assert not (tmp_path
                / "media").exists()
    assert list(tmp_path.glob(
        "**/*.png")) == []


def test_no_dir_no_warnings(tmp_path):
    p = _image_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


# ---------- schema + 管线 + 指标 ----------

def test_image_docx_passes_schema(
        tmp_path):
    from app.schema import is_valid
    doc, _ = _parse_images(tmp_path)
    assert is_valid(doc.to_dict())


def test_image_not_in_chunks(tmp_path):
    from app.pipeline import \
        process_single
    p = _image_docx(tmp_path)
    doc, errors = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    assert errors == []
    img_id = [e.element_id
              for e in doc.elements
              if e.type == "image"][0]
    assert all(
        img_id not in c.source_element_ids
        for c in doc.chunks)


def test_image_metrics_irer_zero(
        tmp_path):
    """process_single 不传 image_output_dir
    → resource_path '(unsaved)' → 存在率 0.0。"""
    from evaluation.metrics import \
        compute_automatic_metrics
    from app.pipeline import \
        process_single
    p = _image_docx(tmp_path)
    doc, _ = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "image_resource_exists_ratio"] \
        == {"value": 0.0, "reason": None}
    assert m[
        "chunk_reference_intact_ratio"] \
        == {"value": 1.0, "reason": None}
