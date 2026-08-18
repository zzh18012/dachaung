"""evaluation/runner.py 第五百六十五轮 edges 测试（Round 1121）。

补强 edges140 未触及的角度（第四百九十七批，probe 实证）。

新角度（真实运行产物回环三锁）：
- **产物过报告 Schema**：真实 docx 板（含 ef ghost）跑
  run_evaluation → 输出 JSON 直接过
  evaluation-report.schema.json——运行时拼装的报告永远
  与 Schema 同步（首锁；旧锁只测手拼报告过 Schema）
- **返回值与落盘文件全等**：run_evaluation 返回的 dict 与
  json.loads(输出文件) 完全相等——内存视图与磁盘视图无
  漂移（首锁）
- **产物过 validate-report CLI**：main(["validate-report",
  out]) → rc 0 + [OK]——评测自产的报告经自家 CLI 校验
  零告警（runner 层首锁；报告内容断言旧锁在 cli 侧）
- forbidden tokens 第五百九十三批（open 2）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA head start.")
    d.add_paragraph("B" * 250)
    d.save(str(tmp_path / "samples" / "g1.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/g1.docx",
                       "source_type": "docx"}],
        "expected_failures": [{
            "doc_id": "ef1", "path": "samples/ghost.docx",
            "expected_error_code": "file_not_found"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _run(tmp_path):
    out = tmp_path / "r.json"
    r = run_evaluation(_board(tmp_path), out,
                       parser_name="fallback", max_chars=200)
    return r, out


# ---------- 产物过报告 Schema ----------

def test_output_passes_report_schema_batch320(tmp_path):
    r, out = _run(tmp_path)
    validate(json.loads(out.read_text(encoding="utf-8")),
             "evaluation-report.schema.json")
    validate(r, "evaluation-report.schema.json")


# ---------- 返回值与落盘文件全等 ----------

def test_returned_equals_written_batch320(tmp_path):
    r, out = _run(tmp_path)
    on_disk = json.loads(out.read_text(encoding="utf-8"))
    assert r == on_disk


# ---------- 产物过 validate-report CLI ----------

def test_output_passes_validate_report_cli_batch320(tmp_path):
    _, out = _run(tmp_path)
    from evaluation.cli import main
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["validate-report", str(out)])
    assert rc == 0
    assert "[OK]" in buf.getvalue()


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch320():
    src = _src()
    assert "用 doc_id 作目录名" in src
    assert "不再从 document_id 反推" in src


# ---------- forbidden tokens 第五百九十三批 ----------

def test_source_no_eval_batch320():
    assert "eval(" not in _src()


def test_source_no_exec_batch320():
    assert "exec(" not in _src()


def test_source_no_compile_batch320():
    assert "compile(" not in _src()


def test_source_no_globals_batch320():
    assert "globals(" not in _src()


def test_source_no_locals_batch320():
    assert "locals(" not in _src()


def test_source_no_os_system_batch320():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch320():
    assert "subprocess" not in _src()


def test_source_no_popen_batch320():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch320():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch320():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch320():
    assert "socket" not in _src()


def test_source_no_requests_batch320():
    assert "requests" not in _src()


def test_source_no_urllib_batch320():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch320():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch320():
    assert "yield" not in _src()


def test_source_no_async_await_batch320():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch320():
    assert _src().count("open(") == 2
