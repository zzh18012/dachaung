"""evaluation/metrics.py 第五百六十二轮 edges 测试（Round 1237）。

补强 edges140 未触及的角度（第六百零九批，probe 实证）。

新角度（真实异构板直跑 compute_automatic_metrics）：
- **expectations 精确三类型**——
  {heading 1, paragraph 5, table 1}
  → silent_drop 0（真实混合板
  零漏首锁）
- **欠计钳 0**——paragraph 3 <
  实际 5 → 0（max(0,·) 真板复证）
- **从未发射类型**——caption 1
  → drop 1（actual 0 全额计）
- **过计 1**——heading 2 > 1 → 1
- **真板悬空引用**——chunk2 指
  向不存在 id → intact 2/3（历史
  悬空全在手造板）
- **error 透传不噬元素指标**——
  error dict 在场 ect 7 / docx_loc
  1.0 照出（DOCX 变体首锁）
- forbidden tokens 第七百零三批（open 0）
"""

from __future__ import annotations

import copy
import inspect
import tempfile
from pathlib import Path

import evaluation.metrics as metrics_mod
from evaluation.metrics import compute_automatic_metrics


def _base_doc(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from app.pipeline import process_single
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    p = tmp_path / "ks.docx"
    doc.save(str(p))
    d, errors = process_single(p, tmp_path / "o.json",
                               parser_name="fallback",
                               max_chars=120)
    assert errors == []
    return d.to_dict()


def _mut(base_doc, fn):
    d = copy.deepcopy(base_doc)
    fn(d)
    return d


# ---------- expectations 矩阵 ----------

def test_expectations_exact_zero_batch435(tmp_path):
    m = compute_automatic_metrics(
        _base_doc(tmp_path), None, "docx",
        {"element_count_by_type": {
            "heading": 1, "paragraph": 5, "table": 1}})
    assert m["silent_drop_count"] == {"value": 0, "reason": None}


def test_expectations_under_clamped_batch435(tmp_path):
    m = compute_automatic_metrics(
        _base_doc(tmp_path), None, "docx",
        {"element_count_by_type": {"paragraph": 3}})
    assert m["silent_drop_count"] == {"value": 0, "reason": None}


def test_expectations_caption_drop_batch435(tmp_path):
    m = compute_automatic_metrics(
        _base_doc(tmp_path), None, "docx",
        {"element_count_by_type": {"caption": 1}})
    assert m["silent_drop_count"] == {"value": 1, "reason": None}


def test_expectations_heading_over_batch435(tmp_path):
    m = compute_automatic_metrics(
        _base_doc(tmp_path), None, "docx",
        {"element_count_by_type": {"heading": 2}})
    assert m["silent_drop_count"] == {"value": 1, "reason": None}


def test_no_expectations_null_batch435(tmp_path):
    m = compute_automatic_metrics(
        _base_doc(tmp_path), None, "docx", None)
    assert m["silent_drop_count"] == {
        "value": None, "reason": "no_expectations"}


# ---------- 真板悬空引用 ----------

def test_one_dangling_two_thirds_batch435(tmp_path):
    d = _mut(_base_doc(tmp_path),
             lambda x: x["chunks"][2].update(
                 {"source_element_ids": ["nope"]}))
    m = compute_automatic_metrics(d, None, "docx", None)
    assert m["chunk_reference_intact_ratio"] == {
        "value": 0.6666666666666666, "reason": None}


def test_all_dangling_zero_batch435(tmp_path):
    def _all(d):
        for c in d["chunks"]:
            c["source_element_ids"] = ["nope"]
    m = compute_automatic_metrics(
        _mut(_base_doc(tmp_path), _all), None, "docx", None)
    assert m["chunk_reference_intact_ratio"] == {
        "value": 0.0, "reason": None}


# ---------- error 透传 ----------

def test_error_keeps_element_metrics_batch435(tmp_path):
    err = {"code": "parser_failed", "message": "boom"}
    m = compute_automatic_metrics(
        _base_doc(tmp_path), err, "docx", None)
    assert m["element_count_total"] == {"value": 7,
                                        "reason": None}
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"heading": 1, "paragraph": 5,
                  "table": 1}, "reason": None}


# ---------- 真板保持性 ----------

def test_real_board_preservation_batch435(tmp_path):
    m = compute_automatic_metrics(
        _base_doc(tmp_path), None, "docx", None)
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert m["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch435():
    src = _src()
    assert '_PDF_BBOX_REQUIRED_TYPES = ("heading", ' \
           '"paragraph", "caption", "list_item")' in src
    assert 'if e.get("type") == "heading"' in src


# ---------- forbidden tokens 第七百零三批 ----------

def test_source_no_eval_batch435():
    assert "eval(" not in _src()


def test_source_no_exec_batch435():
    assert "exec(" not in _src()


def test_source_no_compile_batch435():
    assert "compile(" not in _src()


def test_source_no_globals_batch435():
    assert "globals(" not in _src()


def test_source_no_locals_batch435():
    assert "locals(" not in _src()


def test_source_no_os_system_batch435():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch435():
    assert "subprocess" not in _src()


def test_source_no_popen_batch435():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch435():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch435():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch435():
    assert "socket" not in _src()


def test_source_no_requests_batch435():
    assert "requests" not in _src()


def test_source_no_urllib_batch435():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch435():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch435():
    assert "yield" not in _src()


def test_source_no_async_await_batch435():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch435():
    assert "open(" not in _src()
