"""evaluation/runner.py 第五百四十轮 edges 测试（Round 1096）。

补强 edges134-136 未触及的角度（第四百七十二批，probe 实证）。

新角度（同一文件的账本双重奏：重复路径 / 双角色同屏）：
- **同一路径双 doc_id**：good.docx 同时挂 d1/d2 →
  per_doc 两条、metrics 逐字相等（确定性）、counts
  {sum 4, participating 2}——**语料层求和双计**
  （manifest 不查路径重复，同一文档入两次账）、
  success {2, 2, 1.0}、devset file_count 照 2 记
- **双角色同屏**：bad.docx 既作 document 又作
  expected_failure 同一 run——per_doc 1 条
  docx_open_failed + ef matches True + success
  {0, 1, 0.0}（edges130 是 either/or 镜像，同屏
  首锁：失败被两条通道同时记账）
- **深嵌套输出目录**：output "deep/nested/r.json"
  → parents=True 自建、报告落盘存在
- forbidden tokens 第五百六十七批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _write_board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("BBB second paragraph body.")
    d.save(str(tmp_path / "samples" / "good.docx"))
    (tmp_path / "samples" / "bad.docx").write_bytes(b"garbage")


def _manifest(tmp_path, documents, ef=None):
    m = {"manifest_version": "1.0",
         "devset_status": "incomplete",
         "documents": documents}
    if ef is not None:
        m["expected_failures"] = ef
    mf = tmp_path / "manifest.json"
    mf.write_text(json.dumps(m), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 同一路径双 doc_id ----------

def test_dup_path_double_counts_batch295(tmp_path):
    _write_board(tmp_path)
    rep = run_evaluation(
        _manifest(tmp_path, [
            {"doc_id": "d1", "path": "samples/good.docx",
             "source_type": "docx"},
            {"doc_id": "d2", "path": "samples/good.docx",
             "source_type": "docx"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    ids = [p["doc_id"] for p in rep["per_doc"]]
    assert ids == ["d1", "d2"]
    assert rep["per_doc"][0]["metrics"] == \
        rep["per_doc"][1]["metrics"]
    assert rep["summary"]["counts"][
        "element_count_total"] == {
        "sum": 4, "participating_docs": 2}
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}
    assert rep["devset"]["file_count"] == 2
    for p in rep["per_doc"]:
        assert isinstance(
            p["wall_time_seconds"]["total"], float)


# ---------- 双角色同屏 ----------

def test_same_file_both_roles_batch295(tmp_path):
    _write_board(tmp_path)
    rep = run_evaluation(
        _manifest(tmp_path, [
            {"doc_id": "dBad", "path": "samples/bad.docx",
             "source_type": "docx"}],
            ef=[{"doc_id": "fBad",
                 "path": "samples/bad.docx",
                 "source_type": "docx",
                 "expected_error_code":
                     "docx_open_failed"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    assert rep["per_doc"][0]["metrics"][
        "error_code"] == {
        "value": "docx_open_failed", "reason": None}
    assert rep["expected_failures"] == [{
        "doc_id": "fBad",
        "expected_error_code": "docx_open_failed",
        "actual_error_code": "docx_open_failed",
        "matches": True}]
    assert rep["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 0, "total": 1, "rate": 0.0}


# ---------- 深嵌套输出目录 ----------

def test_nested_output_dir_created_batch295(tmp_path):
    _write_board(tmp_path)
    out = tmp_path / "deep" / "nested" / "r.json"
    run_evaluation(
        _manifest(tmp_path, [
            {"doc_id": "d1", "path": "samples/good.docx",
             "source_type": "docx"}]),
        out, parser_name="fallback", max_chars=200)
    assert out.exists()
    blob = json.loads(out.read_text(encoding="utf-8"))
    assert blob["summary"]["success_rates"][
        "pipeline_success"]["rate"] == 1.0


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch295():
    src = _src()
    assert ("if parser_version and not "
            "parser_version_for_prov:") in src
    assert ('"_annotation_present": annotation is not '
            'None') in src


# ---------- forbidden tokens 第五百六十七批 ----------

def test_source_no_eval_batch295():
    assert "eval(" not in _src()


def test_source_no_exec_batch295():
    assert "exec(" not in _src()


def test_source_no_compile_batch295():
    assert "compile(" not in _src()


def test_source_no_globals_batch295():
    assert "globals(" not in _src()


def test_source_no_locals_batch295():
    assert "locals(" not in _src()


def test_source_no_os_system_batch295():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch295():
    assert "subprocess" not in _src()


def test_source_no_popen_batch295():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch295():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch295():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch295():
    assert "socket" not in _src()


def test_source_no_requests_batch295():
    assert "requests" not in _src()


def test_source_no_urllib_batch295():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch295():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch295():
    assert "yield" not in _src()


def test_source_no_async_await_batch295():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch295():
    assert _src().count("open(") == 2
