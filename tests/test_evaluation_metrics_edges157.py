"""evaluation/metrics.py 第五百七十八轮 edges 测试（Round 1332）。

补强 edges156 未触及的角度（第七百零四批，probe 实证）。

新角度（locator 摘除阶梯 / docx 篡改保真 / schema_valid 翻 False）：
- **plvr 半降**——
  单元素 locator 摘
  page（留 bbox）或
  摘 bbox（留 page）
  → 均恰好
  {0.5, None}
- **plvr 归零**——
  双元素 locator
  全空 {} →
  {0.0, None}
- **schema_valid 翻
  False**——locator
  破坏后
  {value: False,
  reason: None}
  （首次真板翻锁）
- **docx 篡改保真**
  ——chunk0 text 换
  'ZZZ.' → tpe False
  / tcmp 0.8421 /
  tcmr 0.4103（docx
  几何首锁，异于
  pdf 0.9693/0.8367）
- **dlvr 三分之二**
  ——3 元素摘 1 →
  {0.6667, None}
- **locator 结构锁**
  ——pdf {page,bbox}、
  docx {paragraph_
  index,section}
- forbidden tokens 第七百七十六批（open 0）
"""

from __future__ import annotations

import copy
import inspect
import tempfile
from pathlib import Path

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from docx import Document
from evaluation.metrics import \
    compute_automatic_metrics


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
COMBO = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
         % ("A" * 80)
         + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
         % LONG).encode()


def _pdf():
    with tempfile.TemporaryDirectory() as td:
        tp = Path(td)
        (tp / "c.pdf").write_bytes(_wrap(COMBO))
        doc, errors = process_single(
            tp / "c.pdf", tp / "o.json",
            parser_name="fallback", max_chars=32)
        assert errors == []
        return doc.to_dict()


def _docx():
    with tempfile.TemporaryDirectory() as td:
        tp = Path(td)
        d = Document()
        d.add_heading("PicDoc", level=1)
        d.add_paragraph("Text before picture.")
        d.add_paragraph("Second para here.")
        d.save(str(tp / "c.docx"))
        doc, errors = process_single(
            tp / "c.docx", tp / "o.json",
            parser_name="fallback", max_chars=32)
        assert errors == []
        return doc.to_dict()


def _mp(dd):
    return compute_automatic_metrics(dd, None, "pdf",
                                     None)


def _md(dd):
    return compute_automatic_metrics(dd, None, "docx",
                                     None)


# ---------- locator 结构锁 ----------

def test_pdf_locator_shape_batch530():
    b = _pdf()
    for e in b["elements"]:
        assert set(e["source_locator"]) == {
            "page", "bbox"}


def test_docx_locator_shape_batch530():
    b = _docx()
    for e in b["elements"]:
        assert set(e["source_locator"]) == {
            "paragraph_index", "section"}


# ---------- plvr 半降 ----------

def test_plvr_no_page_half_batch530():
    d = copy.deepcopy(_pdf())
    loc = d["elements"][0]["source_locator"]
    d["elements"][0]["source_locator"] = {
        "bbox": loc["bbox"]}
    assert _mp(d)[
        "pdf_locator_valid_ratio"] == {
        "value": 0.5, "reason": None}


def test_plvr_no_bbox_half_batch530():
    d = copy.deepcopy(_pdf())
    d["elements"][1]["source_locator"] = {"page": 1}
    assert _mp(d)[
        "pdf_locator_valid_ratio"] == {
        "value": 0.5, "reason": None}


# ---------- plvr 归零 ----------

def test_plvr_empty_zero_batch530():
    d = copy.deepcopy(_pdf())
    d["elements"][0]["source_locator"] = {}
    d["elements"][1]["source_locator"] = {}
    assert _mp(d)[
        "pdf_locator_valid_ratio"] == {
        "value": 0.0, "reason": None}


# ---------- schema_valid 翻 False ----------

def test_schema_valid_false_after_strip_batch530():
    d = copy.deepcopy(_pdf())
    d["elements"][0]["source_locator"] = {}
    d["elements"][1]["source_locator"] = {}
    assert _mp(d)["schema_valid"] == {
        "value": False, "reason": None}


def test_schema_valid_true_baseline_batch530():
    assert _mp(copy.deepcopy(_pdf()))[
        "schema_valid"] == {"value": True,
                            "reason": None}


# ---------- docx 篡改保真 ----------

def test_docx_corrupt_tpe_false_batch530():
    d = copy.deepcopy(_docx())
    d["chunks"][0]["text"] = "ZZZ."
    assert _md(d)["text_preservation_equal"] == {
        "value": False, "reason": None}


def test_docx_corrupt_tcmp_batch530():
    d = copy.deepcopy(_docx())
    d["chunks"][0]["text"] = "ZZZ."
    assert _md(d)[
        "text_char_multiset_precision"] == {
        "value": 0.8421052631578947,
        "reason": None}


def test_docx_corrupt_tcmr_batch530():
    d = copy.deepcopy(_docx())
    d["chunks"][0]["text"] = "ZZZ."
    assert _md(d)[
        "text_char_multiset_recall"] == {
        "value": 0.41025641025641024,
        "reason": None}


# ---------- dlvr 三分之二 ----------

def test_dlvr_two_thirds_batch530():
    d = copy.deepcopy(_docx())
    d["elements"][0]["source_locator"] = {}
    assert _md(d)[
        "docx_locator_valid_ratio"] == {
        "value": 0.6666666666666666,
        "reason": None}


def test_dlvr_baseline_one_batch530():
    assert _md(copy.deepcopy(_docx()))[
        "docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch530():
    src = _src()
    assert "st_size" in src
    assert "no_image_elements" in src


# ---------- forbidden tokens 第七百七十六批 ----------

def test_source_no_eval_batch530():
    assert "eval(" not in _src()


def test_source_no_exec_batch530():
    assert "exec(" not in _src()


def test_source_no_compile_batch530():
    assert "compile(" not in _src()


def test_source_no_globals_batch530():
    assert "globals(" not in _src()


def test_source_no_locals_batch530():
    assert "locals(" not in _src()


def test_source_no_os_system_batch530():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch530():
    assert "subprocess" not in _src()


def test_source_no_popen_batch530():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch530():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch530():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch530():
    assert "socket" not in _src()


def test_source_no_requests_batch530():
    assert "requests" not in _src()


def test_source_no_urllib_batch530():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch530():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch530():
    assert "yield" not in _src()


def test_source_no_async_await_batch530():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_0_batch530():
    assert _src().count("open(") == 0
