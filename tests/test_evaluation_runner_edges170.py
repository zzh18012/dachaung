"""evaluation/runner.py 第五百九十七轮 edges 测试（Round 1153）。

补强 edges169 未触及的角度（第五百二十五批，probe 实证）。

新角度（DOCX 真标题样式 × 容差刀锋）：
- **Heading 1/2 真样式**——python-docx add_heading →
  heading 元素 metadata {'level': 1, 'style':
  'Heading 1', 'empty': False}（runner 级首锁，
  历史 DOCX 板全是 Normal 段落）
- **heading 软边界合并**——每 heading 与其后段落
  合块（2 chunks 各 2 源）；heading_boundary_
  compliance 仍 1.0——heading 只须起块不必收块
- **容差刀锋 30/31**——marker "Section Title" after
  落合并块内部距块界恰 31 字：tol 30 → F1 0.0、
  tol 31 → F1 1.0（闭区间容差边界首锁）
- forbidden tokens 第六百二十五批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, marker="Section Title"):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("Section Title", level=1)
    d.add_paragraph("Body follows the heading here.")
    d.add_heading("Sub Section", level=2)
    d.add_paragraph("More body text after sub.")
    d.save(str(tmp_path / "samples" / f"{doc_id}.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": doc_id,
        "chunk_boundary_anchors": [
            {"marker": marker, "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"samples/{doc_id}.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- Heading 1/2 真样式 ----------

def test_docx_heading_levels_batch351(tmp_path):
    _board(tmp_path, "hl")
    doc, errors = process_single(
        tmp_path / "samples" / "hl.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == ["heading", "paragraph",
                                        "heading", "paragraph"]
    assert els[0]["content"] == "Section Title"
    assert els[0]["metadata"] == {"level": 1, "style": "Heading 1",
                                  "empty": False}
    assert els[2]["metadata"] == {"level": 2, "style": "Heading 2",
                                  "empty": False}


def test_docx_heading_locators_batch351(tmp_path):
    _board(tmp_path, "hl2")
    doc, errors = process_single(
        tmp_path / "samples" / "hl2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    els = doc.to_dict()["elements"]
    assert [e["source_locator"]["paragraph_index"]
            for e in els] == [0, 1, 2, 3]
    assert all(e["source_locator"]["section"] == 0 for e in els)


# ---------- heading 软边界合并 ----------

def test_docx_heading_soft_merge_batch351(tmp_path):
    _board(tmp_path, "hl3")
    doc, errors = process_single(
        tmp_path / "samples" / "hl3.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert len(chunks) == 2
    assert chunks[0]["text"] == \
        "Section Title Body follows the heading here."
    assert chunks[1]["text"] == \
        "Sub Section More body text after sub."
    assert all(len(c["source_element_ids"]) == 2 for c in chunks)


def test_docx_heading_compliance_batch351(tmp_path):
    r = run_evaluation(_board(tmp_path, "hl4"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"heading": 2, "paragraph": 2}, "reason": None}


# ---------- 容差刀锋 30/31 ----------

def test_heading_marker_miss_at_30_batch351(tmp_path):
    r = run_evaluation(_board(tmp_path, "hl5"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200, tolerance_chars=30)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_f1"] == {"value": 0.0, "reason": None}


def test_heading_marker_hit_at_31_batch351(tmp_path):
    r = run_evaluation(_board(tmp_path, "hl6"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200, tolerance_chars=31)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_f1"] == {"value": 1.0, "reason": None}


def test_paragraph_end_marker_hit_batch351(tmp_path):
    r = run_evaluation(_board(tmp_path, "hl7", marker="here."),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_f1"] == {"value": 1.0, "reason": None}
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch351():
    src = _src()
    assert src.count("annotation") == 10
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9


# ---------- forbidden tokens 第六百二十五批 ----------

def test_source_no_eval_batch351():
    assert "eval(" not in _src()


def test_source_no_exec_batch351():
    assert "exec(" not in _src()


def test_source_no_compile_batch351():
    assert "compile(" not in _src()


def test_source_no_globals_batch351():
    assert "globals(" not in _src()


def test_source_no_locals_batch351():
    assert "locals(" not in _src()


def test_source_no_os_system_batch351():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch351():
    assert "subprocess" not in _src()


def test_source_no_popen_batch351():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch351():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch351():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch351():
    assert "socket" not in _src()


def test_source_no_requests_batch351():
    assert "requests" not in _src()


def test_source_no_urllib_batch351():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch351():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch351():
    assert "yield" not in _src()


def test_source_no_async_await_batch351():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch351():
    assert _src().count("open(") == 2
