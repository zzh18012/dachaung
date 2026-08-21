"""evaluation/report.py 第五百八十二轮 edges 测试（Round 1347）。

补强 edges150 未触及的角度（第七百一十九批，probe 实证）。

新角度（全链配对板 / categories 去重 / 分型参与互补）：
- **配对全链**——pdf+docx
  双向 paired_with +
  游离 docx →
  files 3 / groups 2
  （run_evaluation
  全链 group 收缩
  首锁；edges75/90
  仅 build_devset_section
  单元面）
- **categories 去重**
  ——alpha 两文档
  重复 → ['alpha',
  'beta','gamma']
  排序并集
- **分型参与互补**
  ——plvr {1.0,1,2}
  vs dlvr {1.0,2,1}
  （参与数和恒 3）
- **无标题全排除**
  ——三板皆无
  heading → hbc
  {None, 0, 3}
- forbidden tokens 第七百八十八批（open 0）
"""

from __future__ import annotations

import inspect
import json
import tempfile
from pathlib import Path

import evaluation.report as report_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
ONEP = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
        % LONG).encode()


def _run(tmp_path):
    (tmp_path / "c.pdf").write_bytes(_wrap(ONEP))
    d = Document()
    d.add_paragraph("hello world")
    d.save(str(tmp_path / "a.docx"))
    d = Document()
    d.add_paragraph("second doc")
    d.save(str(tmp_path / "b.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "g1", "path": "c.pdf",
             "source_type": "pdf",
             "categories": ["alpha"],
             "paired_with": "d1"},
            {"doc_id": "d1", "path": "a.docx",
             "source_type": "docx",
             "categories": ["alpha", "beta"],
             "paired_with": "g1"},
            {"doc_id": "d2", "path": "b.docx",
             "source_type": "docx",
             "categories": ["gamma"]},
        ]}), encoding="utf-8")
    mf = load_manifest(tmp_path / "m.json",
                       project_root=tmp_path)
    return run_evaluation(mf, tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=800)


# ---------- 配对全链 ----------

def test_devset_full_dict_batch545(tmp_path):
    assert _run(tmp_path)["devset"] == {
        "status": "incomplete", "file_count": 3,
        "content_group_count": 2, "pdf_count": 1,
        "docx_count": 2,
        "categories_covered": ["alpha", "beta", "gamma"]}


def test_pair_collapses_groups_batch545(tmp_path):
    d = _run(tmp_path)["devset"]
    assert d["file_count"] == 3
    assert d["content_group_count"] == 2


def test_pdf_docx_pair_counts_batch545(tmp_path):
    d = _run(tmp_path)["devset"]
    assert d["pdf_count"] == 1
    assert d["docx_count"] == 2


# ---------- categories 去重 ----------

def test_categories_dedup_sorted_batch545(tmp_path):
    assert _run(tmp_path)["devset"][
        "categories_covered"] == [
        "alpha", "beta", "gamma"]


def test_categories_len_three_batch545(tmp_path):
    cats = _run(tmp_path)["devset"][
        "categories_covered"]
    assert len(cats) == 3


# ---------- 分型参与互补 ----------

def test_plvr_pdf_only_batch545(tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    assert ra["pdf_locator_valid_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 1,
        "not_evaluated": 2}


def test_dlvr_docx_only_batch545(tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    assert ra["docx_locator_valid_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 1}


def test_locator_participation_sum_batch545(
        tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    p = (ra["pdf_locator_valid_ratio"]
         ["participating_docs"]
         + ra["docx_locator_valid_ratio"]
         ["participating_docs"])
    assert p == 3


def test_tpe_all_three_batch545(tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    assert ra["text_preservation_equal"] == {
        "macro_average": 1.0,
        "participating_docs": 3,
        "not_evaluated": 0}


# ---------- 无标题全排除 ----------

def test_hbc_none_zero_three_batch545(
        tmp_path):
    ra = _run(tmp_path)["summary"][
        "ratio_macro_averages"]
    assert ra["heading_boundary_compliance"] == {
        "macro_average": None,
        "participating_docs": 0,
        "not_evaluated": 3}


# ---------- 聚合 ----------

def test_success_three_of_three_batch545(
        tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 3, "total": 3,
        "rate": 1.0}


def test_per_doc_id_order_batch545(tmp_path):
    assert [p["doc_id"]
            for p in _run(tmp_path)["per_doc"]] == [
        "g1", "d1", "d2"]


def test_ect_sum_three_docs_batch545(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"][
        "element_count_total"] == {
        "sum": 3, "participating_docs": 3}


def test_counts_only_ect_no_expectations_batch545(
        tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"] == {
        "element_count_total": {
            "sum": 3, "participating_docs": 3}}


# ---------- 报告合法性 ----------

def test_report_schema_batch545(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


def test_report_on_disk_round_trip_batch545(
        tmp_path):
    r = _run(tmp_path)
    on_disk = json.loads(
        (tmp_path / "r.json").read_text(
            encoding="utf-8"))
    assert on_disk == r


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_manifest_props_batch545():
    src = _src()
    assert "manifest.content_group_count" in src
    assert "manifest.categories_covered" in src


def test_source_key_lines_batch545():
    src = _src()
    assert "def build_devset_section(" in src
    assert "_RATIO_METRICS" in src
    assert ("not_eval = len(per_doc_results) "
            "- len(values)") in src


# ---------- forbidden tokens 第七百八十八批 ----------

def test_source_no_eval_batch545():
    assert "eval(" not in _src()


def test_source_no_exec_batch545():
    assert "exec(" not in _src()


def test_source_no_compile_batch545():
    assert "compile(" not in _src()


def test_source_no_globals_batch545():
    assert "globals(" not in _src()


def test_source_no_locals_batch545():
    assert "locals(" not in _src()


def test_source_no_os_system_batch545():
    assert "os.system" not in _src()


def test_source_subprocess_run_two_batch545():
    assert _src().count("subprocess.run") == 2


def test_source_no_popen_batch545():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch545():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch545():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch545():
    assert "socket" not in _src()


def test_source_no_requests_batch545():
    assert "requests" not in _src()


def test_source_no_urllib_batch545():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch545():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch545():
    assert "yield" not in _src()


def test_source_no_async_await_batch545():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch545():
    assert ".call(" not in _src()


def test_source_open_count_is_0_batch545():
    assert _src().count("open(") == 0
