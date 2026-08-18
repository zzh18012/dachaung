"""evaluation/runner.py 第六百四十一轮 edges 测试（Round 1197）。

补强 edges207 未触及的角度（第五百六十九批，probe 实证）。

新角度（DOCX 巨词硬切的镜像）：
- **DOCX 无白界硬切**——131 个 W 无
  空白段 → chunks [100, 31] 恰在
  max_chars 处劈开（与 edges161 的
  PDF forced_char 成镜像；DOCX 通道
  strategy 标签为 long_paragraph_
  sentence_split 首锁）
- **四块结构**——[seq 45, lps 100,
  lps 31, seq 40] 各 1 源
- **流尾锚零界**——"word ends."
  after（流末无界）→ 全 0.0
- forbidden tokens 第六百六十九批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None):
    from docx import Document
    (tmp_path / "s").mkdir(exist_ok=True)
    (tmp_path / "a").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph(
        "Intro sentence before the giant word arrives.")
    d.add_paragraph("W" * 131)
    d.add_paragraph(
        "Tail sentence after the giant word ends.")
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    entry = {"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}
    if anchors is not None:
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}),
            encoding="utf-8")
        entry["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": [entry]}),
                  encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- DOCX 无白界硬切 ----------

def test_giant_word_elements_batch395(tmp_path):
    _board(tmp_path, "gw")
    doc, errors = process_single(
        tmp_path / "s" / "gw.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == ["paragraph"] * 3
    assert [len(e["content"]) for e in els] == [45, 131, 40]


def test_giant_word_chunks_batch395(tmp_path):
    _board(tmp_path, "gw2")
    doc, errors = process_single(
        tmp_path / "s" / "gw2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "sequential", "long_paragraph_sentence_split",
        "long_paragraph_sentence_split", "sequential"]
    assert [len(c["text"]) for c in chunks] == [45, 100, 31, 40]
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks)


def test_giant_word_lengths_batch395(tmp_path):
    _board(tmp_path, "gw3")
    doc, errors = process_single(
        tmp_path / "s" / "gw3.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=100)
    chunks = doc.to_dict()["chunks"]
    assert chunks[1]["text"] == "W" * 100
    assert chunks[2]["text"] == "W" * 31
    assert chunks[0]["text"] == \
        "Intro sentence before the giant word arrives."
    assert chunks[3]["text"] == \
        "Tail sentence after the giant word ends."


# ---------- 锚 ----------

def test_arrives_anchor_batch395(tmp_path):
    r = run_evaluation(_board(tmp_path, "a1", [
        {"marker": "arrives.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.3333333333333333, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.5, "reason": None}


def test_wordends_anchor_batch395(tmp_path):
    r = run_evaluation(_board(tmp_path, "a2", [
        {"marker": "word ends.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.0, "reason": None}


# ---------- 指标 ----------

def test_giant_word_metrics_batch395(tmp_path):
    r = run_evaluation(_board(tmp_path, "mx"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=100)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 3}, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert m["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_reference_intact_ratio"] == {
        "value": 1.0, "reason": None}
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch395():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百六十九批 ----------

def test_source_no_eval_batch395():
    assert "eval(" not in _src()


def test_source_no_exec_batch395():
    assert "exec(" not in _src()


def test_source_no_compile_batch395():
    assert "compile(" not in _src()


def test_source_no_globals_batch395():
    assert "globals(" not in _src()


def test_source_no_locals_batch395():
    assert "locals(" not in _src()


def test_source_no_os_system_batch395():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch395():
    assert "subprocess" not in _src()


def test_source_no_popen_batch395():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch395():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch395():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch395():
    assert "socket" not in _src()


def test_source_no_requests_batch395():
    assert "requests" not in _src()


def test_source_no_urllib_batch395():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch395():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch395():
    assert "yield" not in _src()


def test_source_no_async_await_batch395():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch395():
    assert _src().count("open(") == 2
