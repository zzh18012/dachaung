"""evaluation/metrics.py 第五百五十八轮 edges 测试（Round 1114）。

补强 edges136 未触及的角度（第四百九十批，probe 实证）。

新角度（文本保留三刀：重复 / 截半 / 清空）：
- **chunk 文本翻倍**：chunks[0].text * 2 → equal False +
  multiset precision 0.5 + recall 1.0——重复只罚 precision
  （实际流比期望多一倍字符），期望字符全在故 recall 满
- **chunk 文本截半**：text[:len//2] → equal False +
  precision 1.0 + recall 0.48148148148148145——丢失只罚
  recall（实际是期望子集故 precision 满），27/56 精确值
- **chunk 文本清空（chunks 仍在）**：text "" → equal
  False + precision null empty_actual + recall 0.0——
  空 chunk 而非缺 chunks 键的 empty_actual 路径（与
  edges105 缺键路径互补）
- forbidden tokens 第五百八十六批（open 0）
"""

from __future__ import annotations

import copy
import inspect
import tempfile
import pathlib

from docx import Document

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from evaluation.metrics import compute_automatic_metrics


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_paragraph("AAA first body.")
    d.add_paragraph("BBB second body.")
    d.save(str(p))
    doc, errors = process_single(
        p, tmp_path / "s.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    return doc.to_dict()


def _mutated(tmp_path, mut):
    r = copy.deepcopy(_real_doc(tmp_path))
    mut(r)
    return compute_automatic_metrics(r, None, "docx", None)


# ---------- chunk 文本翻倍 ----------

def test_text_doubled_batch313(tmp_path):
    out = _mutated(
        tmp_path,
        lambda r: r["chunks"][0].__setitem__(
            "text", r["chunks"][0]["text"] * 2))
    assert out["text_preservation_equal"] == {
        "value": False, "reason": None}
    assert out["text_char_multiset_precision"] == {
        "value": 0.5, "reason": None}
    assert out["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}


# ---------- chunk 文本截半 ----------

def test_text_half_truncated_batch313(tmp_path):
    def halve(r):
        t = r["chunks"][0]["text"]
        r["chunks"][0]["text"] = t[: len(t) // 2]
    out = _mutated(tmp_path, halve)
    assert out["text_preservation_equal"] == {
        "value": False, "reason": None}
    assert out["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert out["text_char_multiset_recall"] == {
        "value": 0.48148148148148145, "reason": None}


# ---------- chunk 文本清空（chunks 仍在） ----------

def test_text_emptied_batch313(tmp_path):
    out = _mutated(
        tmp_path,
        lambda r: r["chunks"][0].__setitem__("text", ""))
    assert out["text_preservation_equal"] == {
        "value": False, "reason": None}
    assert out["text_char_multiset_precision"] == {
        "value": None, "reason": "empty_actual"}
    assert out["text_char_multiset_recall"] == {
        "value": 0.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch313():
    src = _src()
    assert "source_spans（在 element content" in src
    assert "不能用于验证空白排版" in src


# ---------- forbidden tokens 第五百八十六批 ----------

def test_source_no_eval_batch313():
    assert "eval(" not in _src()


def test_source_no_exec_batch313():
    assert "exec(" not in _src()


def test_source_no_compile_batch313():
    assert "compile(" not in _src()


def test_source_no_globals_batch313():
    assert "globals(" not in _src()


def test_source_no_locals_batch313():
    assert "locals(" not in _src()


def test_source_no_os_system_batch313():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch313():
    assert "subprocess" not in _src()


def test_source_no_popen_batch313():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch313():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch313():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch313():
    assert "socket" not in _src()


def test_source_no_requests_batch313():
    assert "requests" not in _src()


def test_source_no_urllib_batch313():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch313():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch313():
    assert "yield" not in _src()


def test_source_no_async_await_batch313():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch313():
    assert "open(" not in _src()
