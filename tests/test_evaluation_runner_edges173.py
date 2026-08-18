"""evaluation/runner.py 第六百零一轮 edges 测试（Round 1157）。

补强 edges172 未触及的角度（第五百二十九批，probe 实证）。

新角度（多锚一对一阵营 / anchor 文档序）：
- **同界双锚只配一对**——{"here." after} 与
  {"Second Head" before} 指向同一块界（d=0 与
  d=1）→ 一对一贪心只许一配 → P/R/F1 全 0.5
- **乱序 anchor 判失**——"First Head" 列于
  "here." 之后：顺序搜索起点已推进 → 判
  missing 无 GT → P 0.5 / R 1.0 / F1 0.667
  （anchor 必须按文档序出现，真跑首锁）
- **三锚双中分数值**——3 GT 对 2 pred 贪心距离
  配对 → P 1.0 / R 0.667 / F1 0.8（分数 P/R/F1
  真跑首锁）
- forbidden tokens 第六百二十九批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, anchors):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("First Head", level=1)
    d.add_paragraph("Alpha body text ends here.")
    d.add_heading("Second Head", level=2)
    d.add_paragraph("Beta body text wraps up.")
    d.add_heading("Third Head", level=2)
    d.add_paragraph("Gamma body text finale.")
    d.save(str(tmp_path / "samples" / "h.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "hd",
        "chunk_boundary_anchors": anchors}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "hd",
                       "path": "samples/h.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _prf(r):
    m = r["per_doc"][0]["metrics"]
    return (m["chunk_boundary_precision"],
            m["chunk_boundary_recall"],
            m["chunk_boundary_f1"])


# ---------- 板型基线 ----------

def test_board_chunk_shape_batch355(tmp_path):
    _board(tmp_path, [{"marker": "here.",
                       "position": "after"}])
    doc, errors = process_single(
        tmp_path / "samples" / "h.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    chunks = doc.to_dict()["chunks"]
    assert len(chunks) == 3
    assert chunks[0]["text"] == \
        "First Head Alpha body text ends here."
    assert chunks[1]["text"] == \
        "Second Head Beta body text wraps up."
    assert chunks[2]["text"] == \
        "Third Head Gamma body text finale."


# ---------- 同界双锚只配一对 ----------

def test_same_boundary_two_anchors_batch355(tmp_path):
    r = run_evaluation(_board(tmp_path, [
        {"marker": "here.", "position": "after"},
        {"marker": "Second Head", "position": "before"},
    ]), tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    p, rec, f1 = _prf(r)
    assert p == {"value": 0.5, "reason": None}
    assert rec == {"value": 0.5, "reason": None}
    assert f1 == {"value": 0.5, "reason": None}


# ---------- 乱序 anchor 判失 ----------

def test_out_of_order_anchor_missing_batch355(tmp_path):
    r = run_evaluation(_board(tmp_path, [
        {"marker": "here.", "position": "after"},
        {"marker": "First Head", "position": "before"},
    ]), tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    p, rec, f1 = _prf(r)
    assert p == {"value": 0.5, "reason": None}
    assert rec == {"value": 1.0, "reason": None}
    assert f1 == {"value": 0.6666666666666666,
                  "reason": None}


# ---------- 三锚双中分数值 ----------

def test_triple_anchors_partial_batch355(tmp_path):
    r = run_evaluation(_board(tmp_path, [
        {"marker": "here.", "position": "after"},
        {"marker": "Second Head", "position": "before"},
        {"marker": "wraps up.", "position": "after"},
    ]), tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    p, rec, f1 = _prf(r)
    assert p == {"value": 1.0, "reason": None}
    assert rec == {"value": 0.6666666666666666,
                   "reason": None}
    assert f1 == {"value": 0.8, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch355():
    src = _src()
    assert src.count("process_single") == 6
    assert src.count("run_evaluation") == 2
    assert src.count("chunk") == 9


# ---------- forbidden tokens 第六百二十九批 ----------

def test_source_no_eval_batch355():
    assert "eval(" not in _src()


def test_source_no_exec_batch355():
    assert "exec(" not in _src()


def test_source_no_compile_batch355():
    assert "compile(" not in _src()


def test_source_no_globals_batch355():
    assert "globals(" not in _src()


def test_source_no_locals_batch355():
    assert "locals(" not in _src()


def test_source_no_os_system_batch355():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch355():
    assert "subprocess" not in _src()


def test_source_no_popen_batch355():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch355():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch355():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch355():
    assert "socket" not in _src()


def test_source_no_requests_batch355():
    assert "requests" not in _src()


def test_source_no_urllib_batch355():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch355():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch355():
    assert "yield" not in _src()


def test_source_no_async_await_batch355():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch355():
    assert _src().count("open(") == 2
