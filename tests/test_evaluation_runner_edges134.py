"""evaluation/runner.py 第五百一十九轮 edges 测试（Round 1075）。

补强 edges129-133 未触及的角度（第四百五十一批，probe 实证）。

新角度（孪生文档的内容寻址图片目录坍缩）：
- 三份嵌图 docx：twin1 与 twin2 **内容逐字节相同**
  （同文同图）→ 同 source_hash → 落进**同一个**
  images-<sha> 目录、同一文件名——第二次写入覆盖
  第一次（同内容覆盖无害）；other 内容不同 → 独立
  目录。**3 docs → 2 files / 2 dirs**——图片落盘是
  内容寻址的，与 doc_id 无关
- 坍缩对指标透明：三文档 image ratio 全 1.0、
  success {3, 3, 1.0}——共享不损任何度量
- 文件名内嵌 sha 前缀：images-<sha16> 目录下躺
  image_<sha16>_para1_00.png——内容寻址双重可见
- forbidden tokens 第五百四十六批（open 2）
"""

from __future__ import annotations

import inspect
import io
import json
import struct
import zlib

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _build(path, marker):
    d = Document()
    d.add_paragraph(f"AAA {marker} paragraph body.")
    d.add_picture(io.BytesIO(_png_bytes()))
    d.save(str(path))


def _run(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    _build(tmp_path / "samples" / "twin1.docx", "twin")
    _build(tmp_path / "samples" / "twin2.docx", "twin")
    _build(tmp_path / "samples" / "other.docx", "other")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/twin1.docx",
             "source_type": "docx"},
            {"doc_id": "d2", "path": "samples/twin2.docx",
             "source_type": "docx"},
            {"doc_id": "d3", "path": "samples/other.docx",
             "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json",
                          max_chars=200)


def _inventory(tmp_path):
    per = tmp_path / "_per_doc"
    files = [p for p in per.rglob("*") if p.is_file()]
    dirs = [p for p in per.iterdir() if p.is_dir()]
    return files, dirs


# ---------- 孪生坍缩：3 docs → 2 files ----------

def test_twin_docs_share_image_dir_batch274(tmp_path):
    _run(tmp_path)
    files, dirs = _inventory(tmp_path)
    assert len(files) == 2
    assert len(dirs) == 2
    counts = {d.name: len(list(d.iterdir()))
              for d in dirs}
    assert sorted(counts.values()) == [1, 1]


# ---------- 坍缩对指标透明 ----------

def test_ratios_all_lit_despite_sharing_batch274(
        tmp_path):
    rep = _run(tmp_path)
    assert [p["metrics"]["image_resource_exists_ratio"][
        "value"] for p in rep["per_doc"]] == [1.0, 1.0,
                                              1.0]
    assert rep["summary"]["success_rates"] == {
        "pipeline_success": {"success_count": 3,
                             "total": 3, "rate": 1.0}}


# ---------- 文件名内嵌 sha 前缀 ----------

def test_image_filename_sha_prefix_batch274(tmp_path):
    _run(tmp_path)
    files, _ = _inventory(tmp_path)
    for f in files:
        sha = f.parent.name[len("images-"):]
        assert f.name == f"image_{sha}_para1_00.png"


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch274():
    src = _src()
    assert ("image_output_dir_for(out_stub, "
            "document.source_hash)") in src
    assert "out_stub.unlink()" in src


# ---------- forbidden tokens 第五百四十六批 ----------

def test_source_no_eval_batch274():
    assert "eval(" not in _src()


def test_source_no_exec_batch274():
    assert "exec(" not in _src()


def test_source_no_compile_batch274():
    assert "compile(" not in _src()


def test_source_no_globals_batch274():
    assert "globals(" not in _src()


def test_source_no_locals_batch274():
    assert "locals(" not in _src()


def test_source_no_os_system_batch274():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch274():
    assert "subprocess" not in _src()


def test_source_no_popen_batch274():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch274():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch274():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch274():
    assert "socket" not in _src()


def test_source_no_requests_batch274():
    assert "requests" not in _src()


def test_source_no_urllib_batch274():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch274():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch274():
    assert "yield" not in _src()


def test_source_no_async_await_batch274():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch274():
    assert _src().count("open(") == 2
