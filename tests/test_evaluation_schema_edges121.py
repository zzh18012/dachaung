"""evaluation/schema.py 第五百零一轮 edges 测试（Round 1057）。

补强 edges120 未触及的角度（第四百三十三批，probe 实证）。

新角度（locator 定义的严格性倒挂，真实文档实证）：
- document.schema.json 两 locator def 均
  additionalProperties: true——docx_locator 只要求
  minProperties 1（结构键 ≥ 1）、pdf_locator 只要求
  page——**跨族键不禁**：docx 元素带 page/bbox 照过
  schema，而 metrics 的 _docx_locator_ratio 会把
  带 page 的元素计无效——schema 宽、指标严的倒挂
  在 def 层与真实文档双重锁死
- 真实三类型 docx（heading+段+table）解析产物过
  document.schema.json——真实富板对业务 schema 的
  首次全绿
- pdf 分支正要求：source_type 改 pdf 后全部元素
  缺 page → 多错（每元素一条 required）
- 真实基底突变：chunk source_element_ids 置空 →
  恰 1 错 "[] should be non-empty"（"每 chunk 非空
  引用"不变量的真实文件版）
- forbidden tokens 第五百二十八批（open 2）
"""

from __future__ import annotations

import copy
import inspect

from docx import Document

import evaluation.schema as schema_mod
from app.pipeline import process_single
from evaluation.schema import (
    EvalSchemaError, load_schema, validate)


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA first paragraph body.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "cell one"
    d.save(str(p))
    doc, errors = process_single(p, tmp_path / "s.json",
                                 parser_name="fallback",
                                 max_chars=200,
                                 write_json=False)
    assert errors == []
    return doc.to_dict()


def _mut(dd, fn):
    r = copy.deepcopy(dd)
    fn(r)
    return r


# ---------- 真实富板全绿 ----------

def test_real_rich_doc_passes_batch255(tmp_path):
    validate(_real_doc(tmp_path), "document.schema.json")


# ---------- 跨族键不禁（schema 宽） ----------

def test_docx_page_extra_allowed_batch255(tmp_path):
    dd = _real_doc(tmp_path)
    r = _mut(dd, lambda x: x["elements"][0]
             ["source_locator"].update(
                 {"page": 1, "bbox": [0, 0, 1, 1]}))
    validate(r, "document.schema.json")


# ---------- pdf 分支正要求 ----------

def test_pdf_branch_requires_page_batch255(tmp_path):
    dd = _real_doc(tmp_path)
    r = _mut(dd, lambda x: x.__setitem__("source_type",
                                         "pdf"))
    try:
        validate(r, "document.schema.json")
        raised = False
    except EvalSchemaError as e:
        raised = True
        assert len(e.errors) == len(r["elements"])
        assert all("'page' is a required property"
                   in err["message"] for err in e.errors)
    assert raised


# ---------- 真实基底：空引用拒绝 ----------

def test_empty_chunk_refs_rejected_batch255(tmp_path):
    dd = _real_doc(tmp_path)
    r = _mut(dd, lambda x: x["chunks"][0].__setitem__(
        "source_element_ids", []))
    try:
        validate(r, "document.schema.json")
        raised = False
    except EvalSchemaError as e:
        raised = True
        assert len(e.errors) == 1
        assert e.errors[0]["path"] == [
            "chunks", 0, "source_element_ids"]
        assert e.errors[0]["message"] == \
            "[] should be non-empty"
    assert raised


# ---------- def 层开放性直锁 ----------

def test_locator_defs_open_batch255():
    s = load_schema("document.schema.json")
    docx_loc = s["$defs"]["docx_locator"]
    pdf_loc = s["$defs"]["pdf_locator"]
    assert docx_loc["additionalProperties"] is True
    assert docx_loc["minProperties"] == 1
    assert pdf_loc["additionalProperties"] is True
    assert pdf_loc["required"] == ["page"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch255():
    src = _src()
    assert "def load_schema(name" in src
    assert "Draft202012Validator(schema)" in src


# ---------- forbidden tokens 第五百二十八批 ----------

def test_source_no_eval_batch255():
    assert "eval(" not in _src()


def test_source_no_exec_batch255():
    assert "exec(" not in _src()


def test_source_no_compile_batch255():
    assert "compile(" not in _src()


def test_source_no_globals_batch255():
    assert "globals(" not in _src()


def test_source_no_locals_batch255():
    assert "locals(" not in _src()


def test_source_no_os_system_batch255():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch255():
    assert "subprocess" not in _src()


def test_source_no_popen_batch255():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch255():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch255():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch255():
    assert "socket" not in _src()


def test_source_no_requests_batch255():
    assert "requests" not in _src()


def test_source_no_urllib_batch255():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch255():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch255():
    assert "yield" not in _src()


def test_source_no_async_await_batch255():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch255():
    assert _src().count("open(") == 2
