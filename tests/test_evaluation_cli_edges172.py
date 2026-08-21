"""evaluation/cli.py 第六百七十三轮 edges 测试（Round 1322）。

补强 edges171 未触及的角度（第六百九十四批，probe 实证）。

新角度（图片 DOCX inspect / JSON 错误路径）：
- **图片 DOCX inspect**——
  counts elements=5
  chunks=2；ecbt 行
  'heading=1, image=1,
  paragraph=3'（键排
  序）；irer 1.0000
  (ok) 行首锁
- **双 locator 行**——
  dlvr 1.0000 +
  plvr null (not_pdf_
  document) 同面板
- **JSON 解析失败**——
  非法 JSON → rc 1 +
  [ERROR] JSON 解析
  失败
- **顶层非对象**——
  数组 JSON → rc 1 +
  [ERROR] JSON 顶层
  不是对象
- **error_code
  (None)**——null 原
  因渲染复核
- forbidden tokens 第五百九十四批（open 1）
"""

from __future__ import annotations

import inspect
import struct
import sys
import zlib

import pytest

import evaluation.cli as cli_mod
from app.pipeline import process_single
from docx import Document
from evaluation.cli import main


@pytest.fixture(autouse=True)
def _restore_argv():
    saved = sys.argv
    yield
    sys.argv = saved


def _chunk(t: bytes, data: bytes) -> bytes:
    c = struct.pack(">I", len(data)) + t + data
    return c + struct.pack(
        ">I", zlib.crc32(t + data) & 0xffffffff)


PNG = (b"\x89PNG\r\n\x1a\n"
       + _chunk(b"IHDR", struct.pack(">IIBBBBB",
                                     1, 1, 8, 2, 0, 0, 0))
       + _chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
       + _chunk(b"IEND", b""))


def _doc(tmp_path):
    (tmp_path / "img.png").write_bytes(PNG)
    d = Document()
    d.add_heading("PicDoc", level=1)
    d.add_paragraph("Text before picture.")
    d.add_picture(str(tmp_path / "img.png"))
    d.add_paragraph("Text after picture.")
    d.save(str(tmp_path / "c.docx"))
    doc, errors = process_single(tmp_path / "c.docx",
                                 tmp_path / "o.json",
                                 parser_name="fallback",
                                 max_chars=32)
    assert errors == []
    return tmp_path / "o.json"


def _inspect(tmp_path, capsys):
    oj = _doc(tmp_path)
    sys.argv = ["evaluation.cli", "inspect-doc",
                str(oj)]
    rc = main()
    return rc, capsys.readouterr().out


# ---------- 图片 DOCX inspect ----------

def test_counts_line_batch520(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert "counts:      elements=5 chunks=2" in out


def test_ecbt_line_batch520(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'element_count_by_type':36}"
            " heading=1, image=1, paragraph=3"
            "  (ok)") in out


def test_irer_line_batch520(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'image_resource_exists_ratio':36}"
            " 1.0000  (ok)") in out


def test_ect_line_batch520(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'element_count_total':36}"
            " 5  (ok)") in out


# ---------- 双 locator 行 ----------

def test_dlvr_line_batch520(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'docx_locator_valid_ratio':36}"
            " 1.0000  (ok)") in out


def test_plvr_null_line_batch520(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'pdf_locator_valid_ratio':36}"
            " null  (not_pdf_document)") in out


def test_rc_zero_batch520(tmp_path, capsys):
    rc, _ = _inspect(tmp_path, capsys)
    assert rc == 0


# ---------- JSON 解析失败 ----------

def test_invalid_json_batch520(tmp_path, capsys):
    (tmp_path / "bad.json").write_text("{not json",
                                       encoding="utf-8")
    sys.argv = ["evaluation.cli", "inspect-doc",
                str(tmp_path / "bad.json")]
    rc = main()
    err = capsys.readouterr().err
    assert rc == 1
    assert "[ERROR] JSON 解析失败:" in err


# ---------- 顶层非对象 ----------

def test_array_json_batch520(tmp_path, capsys):
    (tmp_path / "arr.json").write_text("[1,2]",
                                       encoding="utf-8")
    sys.argv = ["evaluation.cli", "inspect-doc",
                str(tmp_path / "arr.json")]
    rc = main()
    err = capsys.readouterr().err
    assert rc == 1
    assert "[ERROR] JSON 顶层不是对象" in err


# ---------- error_code (None) 复核 ----------

def test_error_code_none_line_batch520(tmp_path,
                                       capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'error_code':36}"
            " null  (None)") in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_counts_batch520():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


def test_source_error_paths_batch520():
    src = _src()
    assert "[ERROR] JSON 解析失败" in src
    assert "[ERROR] JSON 顶层不是对象" in src


# ---------- forbidden tokens 第五百九十四批 ----------

def test_source_no_eval_batch520():
    assert "eval(" not in _src()


def test_source_no_exec_batch520():
    assert "exec(" not in _src()


def test_source_no_compile_batch520():
    assert "compile(" not in _src()


def test_source_no_globals_batch520():
    assert "globals(" not in _src()


def test_source_no_locals_batch520():
    assert "locals(" not in _src()


def test_source_no_os_system_batch520():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch520():
    assert "subprocess" not in _src()


def test_source_no_popen_batch520():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch520():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch520():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch520():
    assert "socket" not in _src()


def test_source_no_requests_batch520():
    assert "requests" not in _src()


def test_source_no_urllib_batch520():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch520():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch520():
    assert "yield" not in _src()


def test_source_no_async_await_batch520():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch520():
    assert ".call(" not in _src()


def test_source_open_count_is_1_batch520():
    assert _src().count("open(") == 1
