r"""app/parsers/fallback_parser.py 边角测试 - 第二十七轮（Round 1415）。

新角度（probe 实证）单容器多图命名（历史只锁过跨段/跨页
单图）：
- docx 同一段落两图 + 夹文本 → 文本段落 ('between' strip)
  + 两个 image 元素殿后；命名 image_<sha>_para1_00.png /
  _para1_01.png（同段顺序后缀 00/01）
- pdf 同页两 XObject → 各自 bbox（40x40 @ x=100/300）；
  命名 image_<sha>_p1_00.png / _p1_01.png；204 字节/枚
"""

from __future__ import annotations

import io
import re
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser


def _make_png():
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(typ, data):
        c = (struct.pack(">I", len(data))
             + typ + data)
        return c + struct.pack(
            ">I",
            zlib.crc32(typ + data)
            & 0xFFFFFFFF)

    return (sig
            + chunk(b"IHDR",
                    struct.pack(
                        ">IIBBBBB",
                        1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT",
                    zlib.compress(
                        b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b""))


def _two_docx(tmp_path):
    d = Document()
    d.add_paragraph("intro")
    p = d.add_paragraph()
    p.add_run().add_picture(
        io.BytesIO(_make_png()),
        width=914400)
    p.add_run(" between ")
    p.add_run().add_picture(
        io.BytesIO(_make_png()),
        width=914400)
    p = tmp_path / "two.docx"
    d.save(str(p))
    return p


def _two_pdf(tmp_path):
    c1 = (
        b"BT /F1 12 Tf 72 700 Td "
        b"(Two images page) Tj ET "
        b"q 40 0 0 40 100 500 cm "
        b"/Im1 Do Q "
        b"q 40 0 0 40 300 500 cm "
        b"/Im1 Do Q")
    img = b"\x00\xff\x00"
    objs = {
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] "
            b"/Count 1 >>"),
        3: (b"<< /Type /Page /Parent "
            b"2 0 R /MediaBox "
            b"[0 0 612 792] "
            b"/Resources << /Font "
            b"<< /F1 6 0 R >> "
            b"/XObject << /Im1 "
            b"5 0 R >> >> /Contents "
            b"4 0 R >>"),
        5: (b"<< /Type /XObject "
            b"/Subtype /Image "
            b"/Width 1 /Height 1 "
            b"/ColorSpace /DeviceRGB "
            b"/BitsPerComponent 8 "
            b"/Length 3 >>\nstream\n"
            + img + b"\nendstream"),
        6: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        4: (b"<< /Length "
            + str(len(c1)).encode()
            + b" >>\nstream\n" + c1
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    out += (b"xref\n0 7\n"
            b"0000000000 65535 f \n")
    for oid in range(1, 7):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 7 "
            b"/Root 1 0 R >>\n"
            b"startxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    p = tmp_path / "two.pdf"
    p.write_bytes(bytes(out))
    return p


# ---------- docx 同段两图 ----------

def test_docx_elements(tmp_path):
    p = _two_docx(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "paragraph",
        "image", "image"]


def test_docx_inter_text_paragraph(
        tmp_path):
    p = _two_docx(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))
    assert doc.elements[
        1].content == "between"


def test_docx_two_image_files(
        tmp_path):
    p = _two_docx(tmp_path)
    FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))
    names = sorted(
        f.name for f in
        (tmp_path / "imgs")
        .glob("*.png"))
    assert len(names) == 2
    for n in names:
        assert re.fullmatch(
            r"image_[0-9a-f]{16}"
            r"_para1_0[01].png",
            n)
    assert names[0][:-6] == \
        names[1][:-6]
    assert names[0].endswith(
        "_para1_00.png")
    assert names[1].endswith(
        "_para1_01.png")


def test_docx_image_sizes(tmp_path):
    p = _two_docx(tmp_path)
    FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))
    for f in (tmp_path / "imgs").glob("*.png"):
        assert f.stat().st_size == 69
        assert f.read_bytes()[:8] == \
            b"\x89PNG\r\n\x1a\n"


def test_docx_resource_paths(tmp_path):
    p = _two_docx(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))
    imgs = [e for e in doc.elements
            if e.type == "image"]
    assert imgs[0].resource_path\
        .endswith("_para1_00.png")
    assert imgs[1].resource_path\
        .endswith("_para1_01.png")


def test_docx_schema_valid(tmp_path):
    from app.schema import is_valid
    p = _two_docx(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgs")
    ).parse(p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


# ---------- pdf 同页两图 ----------

def test_pdf_elements(tmp_path):
    p = _two_pdf(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgsp")
    ).parse(p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "heading", "image",
        "image"]


def test_pdf_image_bboxes(tmp_path):
    p = _two_pdf(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgsp")
    ).parse(p, compute_file_hash(p))
    imgs = [e for e in doc.elements
            if e.type == "image"]
    assert imgs[0].source_locator[
        "bbox"] == [
        100.0, 252.0, 140.0,
        292.0]
    assert imgs[1].source_locator[
        "bbox"] == [
        300.0, 252.0, 340.0,
        292.0]


def test_pdf_two_image_files(
        tmp_path):
    p = _two_pdf(tmp_path)
    FallbackParser(
        image_output_dir=(
            tmp_path / "imgsp")
    ).parse(p, compute_file_hash(p))
    names = sorted(
        f.name for f in
        (tmp_path / "imgsp")
        .glob("*.png"))
    assert names == [
        "image_5a2457900a48d05e"
        "_p1_00.png",
        "image_5a2457900a48d05e"
        "_p1_01.png"]


def test_pdf_image_sizes(tmp_path):
    p = _two_pdf(tmp_path)
    FallbackParser(
        image_output_dir=(
            tmp_path / "imgsp")
    ).parse(p, compute_file_hash(p))
    for f in (tmp_path / "imgsp").glob("*.png"):
        assert f.stat().st_size == 204
        assert f.read_bytes()[:8] == \
            b"\x89PNG\r\n\x1a\n"


def test_pdf_heading_bbox(tmp_path):
    p = _two_pdf(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgsp")
    ).parse(p, compute_file_hash(p))
    assert doc.elements[
        0].source_locator["bbox"] == [
        72.0, 82.48400000000004,
        166.70399999999998,
        94.48400000000004]


def test_pdf_schema_valid(tmp_path):
    from app.schema import is_valid
    p = _two_pdf(tmp_path)
    doc = FallbackParser(
        image_output_dir=(
            tmp_path / "imgsp")
    ).parse(p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_no_warnings_both(tmp_path):
    d1 = FallbackParser(
        image_output_dir=(
            tmp_path / "i1")
    ).parse(_two_docx(tmp_path),
            compute_file_hash(
                _two_docx(tmp_path)))
    d2 = FallbackParser(
        image_output_dir=(
            tmp_path / "i2")
    ).parse(_two_pdf(tmp_path),
            compute_file_hash(
                _two_pdf(tmp_path)))
    assert d1.warnings == []
    assert d2.warnings == []
