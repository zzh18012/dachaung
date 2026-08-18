"""evaluation/runner.py 第六百零六轮 edges 测试（Round 1162）。

补强 edges177 未触及的角度（第五百三十四批，probe 实证）。

新角度（空段落占位符 / 五型标注）：
- **空段落占位符**——python-docx 空段 → 元素
  content "(空段落)" + empty: True——占位文本
  进元素流（首锁）
- **占位符入块**——三段（实/空/实）→ 单 chunk
  3 源，文本含字面 "(空段落)"——占位符参与拼接
- **占位符可作 marker**——marker "(空段落)"
  after 恰落 heading 前块界 → P/R/F1 全 1.0
  （占位文本进匹配流，奇观首锁）
- **五型双 marker 全中**——五型 DOCX 挂双锚
  （块界 0|1 与 1|2）→ 全 1.0；五型
  expectations 精确一致 → drop 0
- forbidden tokens 第六百三十四批（open 2）
"""

from __future__ import annotations

import inspect
import json
from io import BytesIO

import evaluation.runner as runner_mod
from docx import Document
from docx.shared import Inches
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single

_PNG = (b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01'
        b'\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89'
        b'\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01'
        b'\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82')


def _empty_board(tmp_path, doc_id, marker):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Alpha paragraph content.")
    d.add_paragraph("")
    d.add_heading("Beta heading after gap", level=1)
    d.add_paragraph("Beta body text follows.")
    d.save(str(tmp_path / "samples" / f"{doc_id}.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": doc_id,
        "chunk_boundary_anchors": [
            {"marker": marker, "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"samples/{doc_id}.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _five_type_board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    d = Document()
    d.add_heading("Doc Heading", level=1)
    p = d.add_paragraph("Body text with inline picture next.")
    p.add_run().add_picture(BytesIO(_PNG), width=Inches(1))
    d.add_paragraph("Figure 2: docx caption text below.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "X1"
    t.cell(0, 1).text = "Y1"
    d.save(str(tmp_path / "samples" / "mi.docx"))
    (tmp_path / "anns" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "mi",
        "chunk_boundary_anchors": [
            {"marker": "picture next.", "position": "after"},
            {"marker": "Figure 2: docx caption text below.",
             "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "mi",
                       "path": "samples/mi.docx",
                       "source_type": "docx",
                       "annotation_file": "anns/a.json",
                       "expectations": {
                           "element_count_by_type": {
                               "heading": 1, "paragraph": 1,
                               "image": 1, "caption": 1,
                               "table": 1}}}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 空段落占位符 ----------

def test_empty_paragraph_placeholder_batch360(tmp_path):
    _empty_board(tmp_path, "ep", "(空段落)")
    doc, errors = process_single(
        tmp_path / "samples" / "ep.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "paragraph", "paragraph", "heading", "paragraph"]
    assert els[1]["content"] == "(空段落)"
    assert els[1]["metadata"]["empty"] is True
    assert els[0]["metadata"]["empty"] is False


def test_placeholder_joins_chunk_batch360(tmp_path):
    _empty_board(tmp_path, "ep2", "(空段落)")
    doc, errors = process_single(
        tmp_path / "samples" / "ep2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert chunks[0]["text"] == "Alpha paragraph content. (空段落)"
    assert len(chunks[0]["source_element_ids"]) == 2


# ---------- 占位符可作 marker ----------

def test_placeholder_marker_boundary_batch360(tmp_path):
    r = run_evaluation(
        _empty_board(tmp_path, "ep3", "(空段落)"),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {"value": 1.0,
                                      "reason": None}


# ---------- 五型双 marker 全中 ----------

def test_five_type_annotation_hit_batch360(tmp_path):
    r = run_evaluation(_five_type_board(tmp_path),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {"value": 1.0,
                                      "reason": None}


def test_five_type_expectations_zero_drop_batch360(tmp_path):
    r = run_evaluation(_five_type_board(tmp_path),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["silent_drop_count"] == {"value": 0,
                                      "reason": None}
    assert r["summary"]["silent_drop_total"] == 0


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch360():
    src = _src()
    assert src.count("manifest") == 5
    assert src.count("error_code") == 4
    assert src.count("metrics") == 13


# ---------- forbidden tokens 第六百三十四批 ----------

def test_source_no_eval_batch360():
    assert "eval(" not in _src()


def test_source_no_exec_batch360():
    assert "exec(" not in _src()


def test_source_no_compile_batch360():
    assert "compile(" not in _src()


def test_source_no_globals_batch360():
    assert "globals(" not in _src()


def test_source_no_locals_batch360():
    assert "locals(" not in _src()


def test_source_no_os_system_batch360():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch360():
    assert "subprocess" not in _src()


def test_source_no_popen_batch360():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch360():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch360():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch360():
    assert "socket" not in _src()


def test_source_no_requests_batch360():
    assert "requests" not in _src()


def test_source_no_urllib_batch360():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch360():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch360():
    assert "yield" not in _src()


def test_source_no_async_await_batch360():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch360():
    assert _src().count("open(") == 2
