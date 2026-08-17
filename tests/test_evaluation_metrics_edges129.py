"""evaluation/metrics.py 第五百零二轮 edges 测试（Round 1058）。

补强 edges128 未触及的角度（第四百三十四批，probe 实证）。

新角度（markdown 装饰的上游起源 + 合并 chunk 保持性）：
- R1051 锁过 table chunk 是 markdown；本批补因果
  起源：**table 元素的 element.content 在 parser 阶段
  就已是 markdown**（"| cell one |\\n| --- |"，原始
  cell 文本 "cell one" 不复存在）——装饰在分块上游，
  chunker 只是透传
- 正因装饰上游，保持性指标全绿：text_char_multiset
  P/R 1.0 + text_preservation True + intact 1.0
  ——markdown 表格不破坏任何保持性不变量
- heading+paragraph 两元素在 mc 200 下合并单 chunk
  （"Real Title AAA first paragraph body."，单空格
  join），source_element_ids 双双在场——多元素一
  chunk 的真实形态
- forbidden tokens 第五百二十九批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

from app.pipeline import process_single
import evaluation.metrics as metrics_mod
from evaluation.metrics import compute_automatic_metrics


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA first paragraph body.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "cell one"
    d.save(str(p))
    doc, errors = process_single(p, tmp_path / "s.json",
                                 parser_name="fallback",
                                 max_chars=200,
                                 write_json=False)
    assert errors == []
    return doc.to_dict()


# ---------- markdown 装饰上游起源 ----------

def test_table_markdown_origin_batch256(tmp_path):
    dd = _real_doc(tmp_path)
    assert dd["elements"][2]["type"] == "table"
    assert dd["elements"][2]["content"] == \
        "| cell one |\n| --- |"
    assert dd["chunks"][1]["text"] == \
        dd["elements"][2]["content"]


# ---------- 装饰不破坏保持性 ----------

def test_markdown_preservation_clean_batch256(tmp_path):
    m = compute_automatic_metrics(_real_doc(tmp_path), None,
                                  "docx", None)
    assert m["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert m["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["chunk_reference_intact_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 多元素单 chunk ----------

def test_merged_heading_chunk_batch256(tmp_path):
    dd = _real_doc(tmp_path)
    first = dd["chunks"][0]
    assert first["text"] == \
        "Real Title AAA first paragraph body."
    assert first["source_element_ids"] == [
        dd["elements"][0]["element_id"],
        dd["elements"][1]["element_id"]]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch256():
    src = _src()
    assert "text_char_multiset_precision" in src
    assert "text_char_multiset_recall" in src


# ---------- forbidden tokens 第五百二十九批 ----------

def test_source_no_eval_batch256():
    assert "eval(" not in _src()


def test_source_no_exec_batch256():
    assert "exec(" not in _src()


def test_source_no_compile_batch256():
    assert "compile(" not in _src()


def test_source_no_globals_batch256():
    assert "globals(" not in _src()


def test_source_no_locals_batch256():
    assert "locals(" not in _src()


def test_source_no_os_system_batch256():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch256():
    assert "subprocess" not in _src()


def test_source_no_popen_batch256():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch256():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch256():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch256():
    assert "socket" not in _src()


def test_source_no_requests_batch256():
    assert "requests" not in _src()


def test_source_no_urllib_batch256():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch256():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch256():
    assert "yield" not in _src()


def test_source_no_async_await_batch256():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch256():
    assert "open(" not in _src()
