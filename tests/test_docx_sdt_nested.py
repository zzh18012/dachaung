"""w:sdt 嵌套内容欠提取修复测试（Stage 7 批次 14）。

背景：DC-REAL-002-DOCX 封面是 body 顶层 w:sdt（index 0），旧主循环
`for child in body.iterchildren()` 只认 w:p/w:tbl → sdt 整块静默跳过
（heading 10、image 0、marker 0/3 丢失 → silent_drop 7）。
修复：_iter_flow_elements 递归生成器（Option A，批次 14 技术裁决批准）。

裁决边界声明映射：
1. 无 w:sdtContent → 跳过（不产出元素）
2. sdt 嵌套 sdt → 递归处理
3. w:tc 内 sdt → 本批不处理（走 _rows_to_markdown 管线，已知边界）
4. heading 实测如实报告 → 本文件只验证机制，真实样本指标在评测侧验证
"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.schema import validate as validate_udm

docx = pytest.importorskip("docx", reason="python-docx 未安装")
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls, qn  # noqa: E402


def _sdt_shell() -> "docx.oxml.CT_Sdt":
    """空的 w:sdt 容器（sdtPr + 空 sdtContent），子元素由调用方移入。"""
    return parse_xml(
        f'<w:sdt {nsdecls("w")}><w:sdtPr><w:id w:val="1"/></w:sdtPr>'
        f'<w:sdtContent/></w:sdt>'
    )


def _append_sdt(d, sdt_el) -> None:
    """把 sdt 插到 sectPr 之前（保持 body 流式顺序合法）。"""
    body = d.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(sdt_el)
    else:
        body.append(sdt_el)


def _wrap_into_sdt(d, elements):
    """把若干 python-docx 对象（Paragraph/Table）移入一个新 sdt 并挂到 body。"""
    sdt = _sdt_shell()
    content = sdt.find(qn("w:sdtContent"))
    for el in elements:
        content.append(el._p if hasattr(el, "_p") else el._tbl)
    _append_sdt(d, sdt)
    return sdt


def _parse(tmp_path: Path, d):
    from app.parsers.fallback_parser import FallbackParser

    p = tmp_path / "synthetic.docx"
    d.save(str(p))
    return FallbackParser().parse(p, source_hash="a" * 64)


# ---------- 1. sdt 内 heading 提取（封面标题场景） ----------

def test_sdt_heading_extracted(tmp_path: Path):
    d = docx.Document()
    h = d.add_heading("Cover Title", level=1)
    _wrap_into_sdt(d, [h])
    d.add_paragraph("Body paragraph.")

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    assert els[0]["type"] == "heading"
    assert els[0]["content"] == "Cover Title"
    assert els[0]["metadata"]["level"] == 1
    # sdt 内段落照常占用 paragraph_index，后续段落连续
    assert els[0]["source_locator"]["paragraph_index"] == 0
    assert els[1]["type"] == "paragraph"
    assert els[1]["source_locator"]["paragraph_index"] == 1


# ---------- 2. 文档顺序：sdt 内容按 DFS 序插入 ----------

def test_sdt_paragraph_document_order(tmp_path: Path):
    d = docx.Document()
    pa = d.add_paragraph("A")
    pb = d.add_paragraph("B")
    pc = d.add_paragraph("C")
    pd = d.add_paragraph("D")
    # B、C 移入 sdt 并放回 A 之后 → 期望文档顺序 A, B, C, D
    sdt = _wrap_into_sdt(d, [pb, pc])
    pa._p.addnext(sdt)

    doc = _parse(tmp_path, d)
    contents = [e["content"] for e in doc.to_dict()["elements"]]
    assert contents == ["A", "B", "C", "D"]
    ids = [e["element_id"] for e in doc.to_dict()["elements"]]
    assert ids == [doc.document_id + f"::e{i:04d}" for i in range(4)]


# ---------- 3. sdt 内表格提取 ----------

def test_sdt_table_extracted(tmp_path: Path):
    d = docx.Document()
    tbl = d.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "AlphaCell"
    tbl.cell(0, 1).text = "BetaCell"
    _wrap_into_sdt(d, [tbl])
    d.add_table(rows=1, cols=1).cell(0, 0).text = "GammaCell"

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    tables = [e for e in els if e["type"] == "table"]
    assert len(tables) == 2
    # sdt 内表格参与 table_index 计数且顺序正确
    assert tables[0]["source_locator"]["table_index"] == 0
    assert tables[1]["source_locator"]["table_index"] == 1
    assert "AlphaCell" in tables[0]["content"]


# ---------- 4. sdt 嵌套 sdt → 递归 ----------

def test_nested_sdt_recursion(tmp_path: Path):
    d = docx.Document()
    deep = d.add_heading("Deep Heading", level=2)
    inner = _sdt_shell()
    inner.find(qn("w:sdtContent")).append(deep._p)
    outer = _sdt_shell()
    outer.find(qn("w:sdtContent")).append(inner)
    _append_sdt(d, outer)

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    assert [(e["type"], e["content"]) for e in els] == [
        ("heading", "Deep Heading")
    ]
    assert els[0]["metadata"]["level"] == 2


# ---------- 5. 无 sdtContent → 跳过 ----------

def test_sdt_without_sdtcontent_skipped(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("After")
    bare = parse_xml(
        f'<w:sdt {nsdecls("w")}><w:sdtPr><w:id w:val="2"/></w:sdtPr></w:sdt>'
    )
    _append_sdt(d, bare)

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    # 裸 sdt 不产出元素；后续段落 paragraph_index 从 0 开始不受影响
    assert [(e["type"], e["source_locator"]["paragraph_index"]) for e in els] == [
        ("paragraph", 0)
    ]
    # 不产生"无内容"警告：裸 sdt 是声明边界（跳过），不是提取失败
    assert doc.to_dict()["warnings"] == []


# ---------- 6. 空 sdtContent → 跳过 ----------

def test_sdt_empty_sdtcontent_no_elements(tmp_path: Path):
    d = docx.Document()
    d.add_paragraph("Solo")
    _append_sdt(d, _sdt_shell())  # sdtContent 存在但为空

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    assert [e["content"] for e in els] == ["Solo"]
    validate_udm(doc.to_dict())


# ---------- 7. 混合结构：sdt(heading+p+tbl) + body(p+tbl) 交错 ----------

def test_mixed_sdt_and_body_order_and_counters(tmp_path: Path):
    d = docx.Document()
    h = d.add_heading("Cover", level=1)
    p1 = d.add_paragraph("Inside sdt")
    t1 = d.add_table(rows=1, cols=1)
    t1.cell(0, 0).text = "TblIn"
    p2 = d.add_paragraph("Outside sdt")
    t2 = d.add_table(rows=1, cols=1)
    t2.cell(0, 0).text = "TblOut"
    sdt = _wrap_into_sdt(d, [h, p1, t1])  # 前三个进 sdt，后两个留在 body
    d.element.body.insert(0, sdt)          # sdt 放回最前（封面位置）

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    assert [(e["type"], e["content"]) for e in els[:2]] == [
        ("heading", "Cover"),
        ("paragraph", "Inside sdt"),
    ]
    assert els[2]["type"] == "table" and "TblIn" in els[2]["content"]
    assert (els[3]["type"], els[3]["content"]) == ("paragraph", "Outside sdt")
    assert els[4]["type"] == "table" and "TblOut" in els[4]["content"]
    para_idx = [
        e["source_locator"]["paragraph_index"]
        for e in els if e["type"] in ("paragraph", "heading")
    ]
    assert para_idx == [0, 1, 2]  # sdt 内外连续
    tbl_idx = [
        e["source_locator"]["table_index"]
        for e in els if e["type"] == "table"
    ]
    assert tbl_idx == [0, 1]
    validate_udm(doc.to_dict())


# ---------- 8. 无 sdt 文档零回归 ----------

def test_no_sdt_document_baseline_unchanged(tmp_path: Path):
    d = docx.Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("Plain body.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "Cell"

    doc = _parse(tmp_path, d)
    els = doc.to_dict()["elements"]
    assert [(e["type"]) for e in els] == ["heading", "paragraph", "table"]
    assert els[0]["source_locator"]["paragraph_index"] == 0
    assert els[1]["source_locator"]["paragraph_index"] == 1
    assert els[2]["source_locator"]["table_index"] == 0
    validate_udm(doc.to_dict())
