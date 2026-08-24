r"""evaluation/cli.py 边角测试 - 第一百八十轮（Round 1387）。

新角度（probe 实证）：混合真文件 manifest——手工合法 xref 的
真 PDF（heading/paragraph/caption 三类）+ python-docx 真文件，
run → report → validate-report 全链：
- stdout 摘要 documents=2（成功 2，失败 0）、pdf=1 docx=1
- pdf 文档 docx_locator null not_docx_document（反之亦然）
- 双 locator ratio 聚合各 1 参与 + 1 not_evaluated
- expectations 齐全 → silent_drop_total = 0（int）
- metric 键全集 21 键；report_version/evaluator_version 1.1
- validate-report rc 0 过自家 Schema
"""

from __future__ import annotations

import contextlib
import io
import json
import tempfile
from pathlib import Path

from docx import Document

from evaluation.cli import main


def _build_pdf(pages_lines):
    n_pages = len(pages_lines)
    page_ids = [3 + i * 2
                for i in range(n_pages)]
    content_ids = [4 + i * 2
                   for i in range(n_pages)]
    font_id = 3 + n_pages * 2
    objs = {font_id: b"<< /Type /Font "
                    b"/Subtype /Type1 "
                    b"/BaseFont /Helvetica >>"}
    objs[1] = (b"<< /Type /Catalog "
               b"/Pages 2 0 R >>")
    kids = " ".join(f"{pid} 0 R"
                    for pid in page_ids)
    objs[2] = (f"<< /Type /Pages /Kids "
               f"[{kids}] /Count "
               f"{n_pages} >>").encode()
    for i, lines in enumerate(pages_lines):
        pid = page_ids[i]
        cid = content_ids[i]
        objs[pid] = (
            f"<< /Type /Page /Parent 2 0 R "
            f"/MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 "
            f"{font_id} 0 R >> >> /Contents "
            f"{cid} 0 R >>").encode()
        blocks = []
        for (y, line) in lines:
            esc = line.replace(
                "\\", r"\\").replace(
                "(", r"\(").replace(
                ")", r"\)")
            blocks.append(
                f"BT /F1 12 Tf 72 {y} Td "
                f"({esc}) Tj ET")
        stream = " ".join(blocks).encode()
        objs[cid] = (
            b"<< /Length "
            + str(len(stream)).encode()
            + b" >>\nstream\n" + stream
            + b"\nendstream")
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n".encode()
                + objs[oid] + b"\nendobj\n")
    xref_pos = len(out)
    maxid = max(objs)
    out += (f"xref\n0 {maxid + 1}\n"
            .encode()
            + b"0000000000 65535 f \n")
    for oid in range(1, maxid + 1):
        out += (
            f"{offsets[oid]:010d} 00000 n \n"
            .encode() if oid in offsets
            else b"0000000000 65535 f \n")
    out += (f"trailer\n<< /Size "
            f"{maxid + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF"
            ).encode()
    return bytes(out)


def _run_mixed(tmp_path):
    (tmp_path / "samples").mkdir()
    (tmp_path / "samples" / "a.pdf").write_bytes(
        _build_pdf([[
            (700, "Mixed Real Heading"),
            (640, "PDF body line one ends "
                  "here."),
            (580, "Figure 1: real pdf "
                  "caption")]]))
    d = Document()
    d.add_heading("Docx Side Title", 1)
    d.add_paragraph("Docx body paragraph.")
    d.save(str(tmp_path / "samples"
               / "b.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "pdf1",
             "path": "samples/a.pdf",
             "source_type": "pdf",
             "expectations": {
                 "element_count_by_type": {
                     "heading": 1,
                     "paragraph": 1,
                     "caption": 1}}},
            {"doc_id": "docx1",
             "path": "samples/b.docx",
             "source_type": "docx",
             "expectations": {
                 "element_count_by_type": {
                     "heading": 1,
                     "paragraph": 1}}}]}),
        encoding="utf-8")
    rep = tmp_path / "r.json"
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["run", "--manifest", str(mf),
                   "--output", str(rep),
                   "--parser", "fallback",
                   "--max-chars", "800"])
    data = json.loads(
        rep.read_text(encoding="utf-8"))
    return rc, buf.getvalue(), data, rep


# ---------- run 摘要 ----------

def test_run_rc0(tmp_path):
    rc, _, _, _ = _run_mixed(tmp_path)
    assert rc == 0


def test_run_stdout_summary(tmp_path):
    _, so, _, _ = _run_mixed(tmp_path)
    assert "documents=2（成功 2，失败 0）" in so


def test_run_stdout_devset_counts(tmp_path):
    _, so, _, _ = _run_mixed(tmp_path)
    assert "pdf=1 docx=1" in so


# ---------- 报告结构 ----------

def test_report_top_keys(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert sorted(data.keys()) == [
        "devset", "expected_failures",
        "per_doc", "provenance",
        "report_version", "summary"]


def test_report_version(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["report_version"] == "1.1"


def test_evaluator_version(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["provenance"][
        "evaluator_version"] == "1.1"


def test_expected_failures_empty(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["expected_failures"] == []


def test_devset_block(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["devset"] == {
        "status": "incomplete",
        "file_count": 2,
        "content_group_count": 2,
        "pdf_count": 1,
        "docx_count": 1,
        "categories_covered": []}


# ---------- per_doc ----------

def test_per_doc_keys(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert all(sorted(pd.keys()) == [
        "doc_id", "metrics",
        "source_type",
        "wall_time_seconds"]
        for pd in data["per_doc"])


def test_per_doc_source_types(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert [pd["source_type"]
            for pd in data["per_doc"]] == [
        "pdf", "docx"]


def test_pdf_doc_ect(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["per_doc"][0]["metrics"][
        "element_count_by_type"] == {
        "value": {"heading": 1,
                  "paragraph": 1,
                  "caption": 1},
        "reason": None}


def test_docx_doc_ect(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["per_doc"][1]["metrics"][
        "element_count_by_type"] == {
        "value": {"heading": 1,
                  "paragraph": 1},
        "reason": None}


def test_pdf_doc_docx_locator_null(
        tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["per_doc"][0]["metrics"][
        "docx_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_docx_document"}


def test_docx_doc_pdf_locator_null(
        tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["per_doc"][1]["metrics"][
        "pdf_locator_valid_ratio"] == {
        "value": None,
        "reason": "not_pdf_document"}


def test_metric_keys_full_set(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert list(data["per_doc"][0]
                ["metrics"].keys()) == [
        "pipeline_success", "error_code",
        "schema_valid",
        "element_count_total",
        "element_count_by_type",
        "pdf_locator_valid_ratio",
        "docx_locator_valid_ratio",
        "image_resource_exists_ratio",
        "chunk_reference_intact_ratio",
        "text_preservation_equal",
        "text_char_multiset_precision",
        "text_char_multiset_recall",
        "heading_boundary_compliance",
        "silent_drop_count",
        "figure_caption_precision",
        "figure_caption_recall",
        "figure_caption_f1",
        "chunk_boundary_precision",
        "chunk_boundary_recall",
        "chunk_boundary_f1"]


def test_wall_time_not_instrumented(
        tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    w = data["per_doc"][0][
        "wall_time_seconds"]
    assert w["parse"] is None
    assert w["chunk"] is None
    assert w["parse_reason"] == \
        "not_instrumented"
    assert w["chunk_reason"] == \
        "not_instrumented"
    assert w["total"] > 0


# ---------- summary 聚合 ----------

def test_summary_counts(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"]["counts"][
        "element_count_total"] == {
        "sum": 5, "participating_docs": 2}


def test_summary_success(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"][
        "success_rates"][
        "pipeline_success"] == {
        "success_count": 2, "total": 2,
        "rate": 1.0}


def test_summary_pdf_locator(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "pdf_locator_valid_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 1,
        "not_evaluated": 1}


def test_summary_docx_locator(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "docx_locator_valid_ratio"] == {
        "macro_average": 1.0,
        "participating_docs": 1,
        "not_evaluated": 1}


def test_summary_hbc_both(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "heading_boundary_compliance"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 0}


def test_summary_tpe_both(tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"][
        "ratio_macro_averages"][
        "text_preservation_equal"] == {
        "macro_average": 1.0,
        "participating_docs": 2,
        "not_evaluated": 0}


def test_summary_silent_drop_zero(
        tmp_path):
    _, _, data, _ = _run_mixed(tmp_path)
    assert data["summary"][
        "silent_drop_total"] == 0


# ---------- validate-report ----------

def test_validate_report_ok(tmp_path):
    _, _, _, rep = _run_mixed(tmp_path)
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["validate-report", str(rep)])
    assert rc == 0
    assert "通过 evaluation-report Schema 校验" \
        in buf.getvalue()
