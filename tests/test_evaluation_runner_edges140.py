"""evaluation/runner.py 第五百六十一轮 edges 测试（Round 1117）。

补强 edges139 未触及的角度（第四百九十三批，probe 实证）。

新角度（共享标注分歧行 / doc_id 不校验）：
- **共享标注分歧行**：两文档挂同一 annotation_file，marker
  "AAA" 只在 d1 流中——d1 {P 0.5 / R 1.0 / F1 0.6667}、
  d2 {P 0.0 / R null no_ground_truth / F1 null}——同一
  标注文件在不同文档流上各自求值，同源不同命（manifest
  侧共享已锁 edges139，本批锁运行时分歧行）
- **annotation doc_id 不校验**：标注 doc_id "zzz-matches-
  nothing" 匹配零文档，marker 在两文档流中都在 → 两文档
  F1 同为 0.6667——runner 从不比对 annotation.doc_id 与
  文档 id，标注按文件挂载不按身份（首锁）
- forbidden tokens 第五百八十九批（open 2）
"""

from __future__ import annotations

import inspect
import json

from docx import Document

import evaluation.runner as runner_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _board(tmp_path, anchor_doc_id, marker_docs):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    for name, p1 in (("g1.docx", "AAA head start."),
                     ("g2.docx", "CCC start.")):
        d = Document()
        d.add_paragraph(p1)
        d.add_paragraph("B" * 250)
        d.save(str(tmp_path / "samples" / name))
    (tmp_path / "anns" / "a.json").write_text(
        json.dumps({
            "annotation_version": "1.0",
            "doc_id": anchor_doc_id,
            "chunk_boundary_anchors": [
                {"marker": "AAA", "position": "before"}]}),
        encoding="utf-8")
    docs = []
    for i, doc in enumerate(marker_docs, start=1):
        entry = {"doc_id": f"d{i}",
                 "path": f"samples/g{doc}.docx",
                 "source_type": "docx",
                 "annotation_file": "anns/a.json"}
        docs.append(entry)
    mf = tmp_path / "manifest.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 共享标注分歧行 ----------

def test_shared_annotation_divergence_batch316(tmp_path):
    man = _board(tmp_path, "d1", marker_docs=(1, 2))
    r = run_evaluation(man, tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    m1 = r["per_doc"][0]["metrics"]
    m2 = r["per_doc"][1]["metrics"]
    assert m1["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m1["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m1["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}
    assert m2["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m2["chunk_boundary_recall"] == {
        "value": None,
        "reason": "no_ground_truth_anchors_in_stream"}
    assert m2["chunk_boundary_f1"] == {
        "value": None,
        "reason": "precision_or_recall_not_evaluated"}


# ---------- annotation doc_id 不校验 ----------

def test_annotation_doc_id_not_checked_batch316(tmp_path):
    man = _board(tmp_path, "zzz-matches-nothing",
                 marker_docs=(1, 1))
    r = run_evaluation(man, tmp_path / "r.json",
                       parser_name="fallback", max_chars=200)
    for pd in r["per_doc"]:
        assert pd["metrics"]["chunk_boundary_f1"] == {
            "value": 0.6666666666666666, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_lines_batch316():
    src = _src()
    assert "评测预期失败用例" in src
    assert "让 pipeline 把图片写入" in src


# ---------- forbidden tokens 第五百八十九批 ----------

def test_source_no_eval_batch316():
    assert "eval(" not in _src()


def test_source_no_exec_batch316():
    assert "exec(" not in _src()


def test_source_no_compile_batch316():
    assert "compile(" not in _src()


def test_source_no_globals_batch316():
    assert "globals(" not in _src()


def test_source_no_locals_batch316():
    assert "locals(" not in _src()


def test_source_no_os_system_batch316():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch316():
    assert "subprocess" not in _src()


def test_source_no_popen_batch316():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch316():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch316():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch316():
    assert "socket" not in _src()


def test_source_no_requests_batch316():
    assert "requests" not in _src()


def test_source_no_urllib_batch316():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch316():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch316():
    assert "yield" not in _src()


def test_source_no_async_await_batch316():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch316():
    assert _src().count("open(") == 2
