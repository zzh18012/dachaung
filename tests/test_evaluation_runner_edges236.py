"""evaluation/runner.py 第六百七十一轮 edges 测试（Round 1321）。

补强 edges235 未触及的角度（第六百九十三批，probe 实证）。

新角度（富 devset 聚合面 / sha256 不核验）：
- **sha256 不核验**——
  g.pdf 声明 sha256
  'b'×64（必然错）→
  照常成功 error None
  （评测不校验文档哈
  希首锁）
- **irer 翻参与**——
  图片 docx 入 devset
  → image_resource_
  exists_ratio {1.0,
  1 参与, 1 未评}（从
  全未评到半参与首锁）
- **marker 命中正文**
  ——docx 锚 marker
  'before'（恰段落文
  本词）→ d1 三元组
  (1.0, 1.0, 1.0)；宏
  cbp (1/15+1)/2、f1
  (0.125+1)/2=0.5625
- **schema_valid 比
  例键**——{1.0, 2,
  0}（12 键中的
  schema_valid 显式
  首锁）
- **counts sum 7**——
  pdf 2 + docx 5 跨型
- **sdt 1**——docx
  image:2 期望
- forbidden tokens 第七百六十七批（open 2）
"""

from __future__ import annotations

import inspect
import json
import struct
import zlib

import evaluation.runner as runner_mod
from app.pipeline import process_single  # noqa: F401
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from evaluation.schema import validate


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


def _chunk(t: bytes, data: bytes) -> bytes:
    c = struct.pack(">I", len(data)) + t + data
    return c + struct.pack(
        ">I", zlib.crc32(t + data) & 0xffffffff)


PNG = (b"\x89PNG\r\n\x1a\n"
       + _chunk(b"IHDR", struct.pack(">IIBBBBB",
                                     1, 1, 8, 2, 0, 0, 0))
       + _chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
       + _chunk(b"IEND", b""))

LONG = " ".join("Word%d." % i for i in range(60))
STREAM = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
          % ("A" * 80)
          + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
          % LONG).encode()


def _board(tmp_path):
    (tmp_path / "g.pdf").write_bytes(_wrap(STREAM))
    (tmp_path / "img.png").write_bytes(PNG)
    d = Document()
    d.add_heading("PicDoc", level=1)
    d.add_paragraph("Text before picture.")
    d.add_picture(str(tmp_path / "img.png"))
    d.add_paragraph("Text after picture.")
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "ann").mkdir(exist_ok=True)
    (tmp_path / "ann" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "g1",
        "chunk_boundary_anchors": [
            {"marker": "Word3.", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "ann" / "b.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "d1",
        "chunk_boundary_anchors": [
            {"marker": "before", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "g1", "path": "g.pdf",
             "source_type": "pdf",
             "annotation_file": "ann/a.json",
             "sha256": "b" * 64},
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx",
             "annotation_file": "ann/b.json",
             "expectations": {
                 "element_count_by_type": {
                     "image": 2}}}]}),
        encoding="utf-8")
    return load_manifest((tmp_path / "m.json"),
                         project_root=tmp_path)


def _run(tmp_path):
    return run_evaluation(_board(tmp_path),
                          tmp_path / "r.json",
                          parser_name="fallback",
                          max_chars=32)


# ---------- sha256 不核验 ----------

def test_wrong_sha256_still_runs_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "error_code"] == {"value": None, "reason": None}


def test_wrong_sha256_success_true_batch519(
        tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "pipeline_success"] == {"value": True,
                                "reason": None}


# ---------- irer 翻参与 ----------

def test_irer_participating_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "image_resource_exists_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


def test_irer_d1_one_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


def test_irer_g1_not_evaluated_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][0]["metrics"][
        "image_resource_exists_ratio"] == {
        "value": None,
        "reason": "no_image_elements"}


# ---------- marker 命中正文 ----------

def test_d1_perfect_trio_batch519(tmp_path):
    r = _run(tmp_path)
    m = r["per_doc"][1]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 1.0, "reason": None}


def test_cbp_macro_half_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_precision"] == {
        "macro_average": (1 / 15 + 1.0) / 2,
        "participating_docs": 2,
        "not_evaluated": 0}


def test_cbf_macro_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"] == {
        "macro_average": 0.5625,
        "participating_docs": 2,
        "not_evaluated": 0}


# ---------- schema_valid 比例键 ----------

def test_schema_valid_ratio_key_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["ratio_macro_averages"][
        "schema_valid"] == {
        "macro_average": 1.0, "participating_docs": 2,
        "not_evaluated": 0}


# ---------- counts / sdt / success ----------

def test_counts_sum_seven_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["counts"][
        "element_count_total"] == {
        "sum": 7, "participating_docs": 2}


def test_sdt_one_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["silent_drop_total"] == 1


def test_d1_sdc_image_two_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["per_doc"][1]["metrics"][
        "silent_drop_count"] == {"value": 1,
                                 "reason": None}


def test_success_full_batch519(tmp_path):
    r = _run(tmp_path)
    assert r["summary"]["success_rates"][
        "pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}


# ---------- 报告合法性 ----------

def test_report_schema_batch519(tmp_path):
    validate(_run(tmp_path),
             "evaluation-report.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_key_counts_batch519():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("per_doc") == 12
    assert src.count("open(") == 2


# ---------- forbidden tokens 第七百六十七批 ----------

def test_source_no_eval_batch519():
    assert "eval(" not in _src()


def test_source_no_exec_batch519():
    assert "exec(" not in _src()


def test_source_no_compile_batch519():
    assert "compile(" not in _src()


def test_source_no_globals_batch519():
    assert "globals(" not in _src()


def test_source_no_locals_batch519():
    assert "locals(" not in _src()


def test_source_no_os_system_batch519():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch519():
    assert "subprocess" not in _src()


def test_source_no_popen_batch519():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch519():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch519():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch519():
    assert "socket" not in _src()


def test_source_no_requests_batch519():
    assert "requests" not in _src()


def test_source_no_urllib_batch519():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch519():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch519():
    assert "yield" not in _src()


def test_source_no_async_await_batch519():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch519():
    assert ".call(" not in _src()
