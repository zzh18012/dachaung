r"""app/parsers/fallback_parser.py 边角测试 - 第十五轮（Round 1391）。

新角度（probe 实证）：真 docx 题注分类与 hyperlink 段：
- 英文 'Figure 1:' / 'TABLE 2:' 段 → caption（style Normal）
- 中文 '图 3：'（全角冒号）→ paragraph——_CAPTION_RE 分隔符
  类 [.、:\s] 只含 ASCII 冒号/句点/顿号/空白，不含全角
  '：' 与 '，'
- '图 1:x'（ASCII 冒号）/'图 1 x'（空格）→ caption
- hyperlink run 文本并入段落（'link text follows:
  clicked part'）
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser


def _parse(tmp_path, builder,
           name="c.docx"):
    p = tmp_path / name
    builder(p)
    return FallbackParser().parse(
        p, compute_file_hash(p))


def _caps(p):
    d = Document()
    d.add_heading("Cap Doc", 1)
    d.add_paragraph(
        "Figure 1: real docx caption line")
    d.add_paragraph(
        "TABLE 2: second caption here")
    d.add_paragraph("图 3：中文题注行")
    d.add_paragraph("表 4：中文表题")
    d.save(str(p))


def _seps(p):
    d = Document()
    d.add_paragraph("图 1:x")
    d.add_paragraph("图 2，x")
    d.add_paragraph("图 3 x")
    d.save(str(p))


def _hyperlink(p):
    d = Document()
    d.add_heading("Link Doc", 1)
    para = d.add_paragraph()
    para.add_run("link text follows: ")
    hl = para._p.makeelement(
        qn('w:hyperlink'),
        {qn('r:id'): 'rId9'})
    r = hl.makeelement(qn('w:r'), {})
    t = r.makeelement(qn('w:t'), {})
    t.text = "clicked part"
    r.append(t)
    hl.append(r)
    para._p.append(hl)
    d.save(str(p))


# ---------- 英文题注 ----------

def test_figure_caption_type(tmp_path):
    doc = _parse(tmp_path, _caps)
    assert doc.elements[1].type == \
        "caption"
    assert doc.elements[1].content == (
        "Figure 1: real docx caption "
        "line")


def test_table_caption_type(tmp_path):
    doc = _parse(tmp_path, _caps)
    assert doc.elements[2].type == \
        "caption"
    assert doc.elements[2].content == (
        "TABLE 2: second caption here")


def test_caption_metadata_style_normal(
        tmp_path):
    doc = _parse(tmp_path, _caps)
    assert doc.elements[1].metadata == {
        "level": 0, "style": "Normal",
        "empty": False}


# ---------- 中文全角分隔符 ----------

def test_chinese_fullwidth_colon_not_caption(
        tmp_path):
    """全角 '：' 不在分隔符类 [.、:\s] 里。"""
    doc = _parse(tmp_path, _caps)
    assert doc.elements[3].type == \
        "paragraph"
    assert doc.elements[3].content == \
        "图 3：中文题注行"


def test_biao_fullwidth_colon_not_caption(
        tmp_path):
    doc = _parse(tmp_path, _caps)
    assert doc.elements[4].type == \
        "paragraph"
    assert doc.elements[4].content == \
        "表 4：中文表题"


# ---------- ASCII 分隔符命中 ----------

def test_ascii_colon_caption(tmp_path):
    doc = _parse(tmp_path, _seps)
    assert doc.elements[0].type == \
        "caption"
    assert doc.elements[0].content == \
        "图 1:x"


def test_fullwidth_comma_not_caption(
        tmp_path):
    doc = _parse(tmp_path, _seps)
    assert doc.elements[1].type == \
        "paragraph"


def test_space_separator_caption(
        tmp_path):
    doc = _parse(tmp_path, _seps)
    assert doc.elements[2].type == \
        "caption"
    assert doc.elements[2].content == \
        "图 3 x"


# ---------- hyperlink ----------

def test_hyperlink_text_merged(tmp_path):
    doc = _parse(tmp_path, _hyperlink)
    para = [e for e in doc.elements
            if e.type == "paragraph"][0]
    assert para.content == (
        "link text follows: clicked part")


def test_hyperlink_paragraph_count(
        tmp_path):
    doc = _parse(tmp_path, _hyperlink)
    assert [e.type
            for e in doc.elements] == [
        "heading", "paragraph"]


def test_hyperlink_no_warnings(tmp_path):
    doc = _parse(tmp_path, _hyperlink)
    assert doc.warnings == []


# ---------- 结构 ----------

def test_caps_paragraph_indexes(tmp_path):
    doc = _parse(tmp_path, _caps)
    assert [e.source_locator[
        "paragraph_index"]
        for e in doc.elements] == [
        0, 1, 2, 3, 4]


def test_caps_types_sequence(tmp_path):
    doc = _parse(tmp_path, _caps)
    assert [e.type for e in doc.elements
            ] == ["heading", "caption",
                  "caption", "paragraph",
                  "paragraph"]


def test_caps_passes_schema(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path, _caps)
    assert is_valid(doc.to_dict())


# ---------- 管线 + 指标 ----------

def test_caps_ect(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    from app.pipeline import \
        process_single
    p = tmp_path / "p.docx"
    _caps(p)
    doc, _ = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "element_count_by_type"][
        "value"] == {
        "heading": 1, "caption": 2,
        "paragraph": 2}


def test_caps_hbc_one(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    from app.pipeline import \
        process_single
    p = tmp_path / "h.docx"
    _caps(p)
    doc, _ = process_single(
        p, None, parser_name="fallback",
        max_chars=800)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "heading_boundary_compliance"] \
        == {"value": 1.0, "reason": None}
