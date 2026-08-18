"""evaluation/cli.py 第六百五十三轮 edges 测试（Round 1236）。

补强 cli edges156 未触及的角度（第六百零八批，probe 实证）。

新角度（厨房水槽板 CLI 全链）：
- **三类型 by_type 行**——
  "element_count_by_type
  heading=1, paragraph=5, table=1"
  （逗号拼接多类型格式首锁）
- **counts 行**——"counts:
  elements=7 chunks=3"
- **hbc 1.0000 (ok) 行**
- **docx 文档的 pdf_locator null
  行**——"null  (not_pdf_
  document)"（源反转对照）
- **silent/image null 行**
- **stdout devset 行**——pdf=0
  docx=1
- forbidden tokens 第五百七十九批（open 1）
"""

from __future__ import annotations

import inspect
import json

import evaluation.cli as cli_mod
from evaluation.cli import main


def _docx(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    (tmp_path / "s").mkdir(exist_ok=True)
    p = tmp_path / "s" / "ks.docx"
    doc.save(str(p))
    return p


def _board(tmp_path):
    _docx(tmp_path)
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "ks", "path": "s/ks.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _doc(tmp_path):
    from app.pipeline import process_single
    _board(tmp_path)
    doc, errors = process_single(
        tmp_path / "s" / "ks.docx", tmp_path / "doc.json",
        parser_name="fallback", max_chars=120)
    assert errors == []
    return tmp_path / "doc.json"


def _run_cli(capsys, argv):
    rc = main(argv)
    out = capsys.readouterr().out
    return rc, out


# ---------- run ----------

def test_cli_run_report_batch434(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    m = rep["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"heading": 1, "paragraph": 5,
                  "table": 1}, "reason": None}
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


def test_cli_run_stdout_batch434(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    assert "[OK]" in out
    assert "documents=1（成功 1，失败 0）" in out


def test_cli_run_stdout_devset_line_batch434(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    assert "devset_status=incomplete file_count=1 groups=1 pdf=0 docx=1" \
        in out


def test_cli_validate_report_batch434(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "120"])
    assert rc == 0
    rc2, out2 = _run_cli(capsys, [
        "validate-report", str(tmp_path / "r.json")])
    assert rc2 == 0
    assert "[OK]" in out2


# ---------- inspect-doc ----------

def test_cli_inspect_counts_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert "counts:      elements=7 chunks=3" in out


def test_cli_inspect_three_type_by_type_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("element_count_by_type                "
            "heading=1, paragraph=5, table=1  (ok)") in out


def test_cli_inspect_hbc_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("heading_boundary_compliance          "
            "1.0000  (ok)") in out


def test_cli_inspect_docx_locator_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("docx_locator_valid_ratio             "
            "1.0000  (ok)") in out


def test_cli_inspect_pdf_null_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("pdf_locator_valid_ratio              "
            "null  (not_pdf_document)") in out


def test_cli_inspect_silent_null_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("silent_drop_count                    "
            "null  (no_expectations)") in out


def test_cli_inspect_image_null_batch434(tmp_path, capsys):
    docp = _doc(tmp_path)
    rc, out = _run_cli(capsys, ["inspect-doc", str(docp)])
    assert rc == 0
    assert ("image_resource_exists_ratio          "
            "null  (no_image_elements)") in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_identifier_counts_batch434():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


# ---------- forbidden tokens 第五百七十九批 ----------

def test_source_no_eval_batch434():
    assert "eval(" not in _src()


def test_source_no_exec_batch434():
    assert "exec(" not in _src()


def test_source_no_compile_batch434():
    assert "compile(" not in _src()


def test_source_no_globals_batch434():
    assert "globals(" not in _src()


def test_source_no_locals_batch434():
    assert "locals(" not in _src()


def test_source_no_os_system_batch434():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch434():
    assert "subprocess" not in _src()


def test_source_no_popen_batch434():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch434():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch434():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch434():
    assert "socket" not in _src()


def test_source_no_requests_batch434():
    assert "requests" not in _src()


def test_source_no_urllib_batch434():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch434():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch434():
    assert "yield" not in _src()


def test_source_no_async_await_batch434():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch434():
    assert _src().count("open(") == 1
