"""evaluation/cli.py 第六百七十二轮 edges 测试（Round 1316）。

补强 edges170 未触及的角度（第六百八十八批，probe 实证）。

新角度（PDF 侧 inspect-doc / tolerance 透传 / 混合计数行）：
- **type=pdf 行**——
  PDF 文档 JSON
  inspect 头部（补
  edges170 的 docx
  侧）
- **counts 行**——
  elements=2
  chunks=16（组合板
  mc32）
- **plvr 1.0000**——
  PDF 侧 locator 行
  + dlvr null (not_
  docx_document) 同
  面板
- **_tolerance_chars
  透传**——默认 30；
  --tolerance-chars
  7 → 行值 7（旗标
  出现在指标面板首锁）
- **error_code
  (None)**——null
  原因渲染复核
- **混合计数行**——
  pdf=1 docx=1 同
  行（CLI 级混合首锁）
- forbidden tokens 第五百九十三批（open 1）
"""

from __future__ import annotations

import inspect
import json
import sys

import pytest

import evaluation.cli as cli_mod
from app.pipeline import process_single
from docx import Document
from evaluation.cli import main


@pytest.fixture(autouse=True)
def _restore_argv():
    saved = sys.argv
    yield
    sys.argv = saved


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
STREAM = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
          % ("A" * 80)
          + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
          % LONG).encode()


def _pdf_doc(tmp_path):
    (tmp_path / "c.pdf").write_bytes(_wrap(STREAM))
    doc, errors = process_single(tmp_path / "c.pdf",
                                 tmp_path / "o.json",
                                 parser_name="fallback",
                                 max_chars=32)
    assert errors == []
    return tmp_path / "o.json"


def _inspect(tmp_path, capsys, *extra):
    oj = _pdf_doc(tmp_path)
    sys.argv = ["evaluation.cli", "inspect-doc",
                str(oj), *extra]
    rc = main()
    return rc, capsys.readouterr().out


# ---------- type=pdf 行 ----------

def test_type_pdf_line_batch514(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert "  type=pdf" in out


def test_counts_line_batch514(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert "counts:      elements=2 chunks=16" in out


# ---------- locator 分型行 ----------

def test_plvr_line_batch514(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'pdf_locator_valid_ratio':36}"
            " 1.0000  (ok)") in out


def test_dlvr_null_line_batch514(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'docx_locator_valid_ratio':36}"
            " null  (not_docx_document)") in out


# ---------- tolerance 透传 ----------

def test_tolerance_default_30_batch514(tmp_path,
                                       capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'_tolerance_chars':36}"
            " 30  (ok)") in out


def test_tolerance_flag_7_batch514(tmp_path, capsys):
    _, out = _inspect(tmp_path, capsys,
                      "--tolerance-chars", "7")
    assert (f"  {'_tolerance_chars':36}"
            " 7  (ok)") in out


def test_tolerance_flag_rc_zero_batch514(tmp_path,
                                         capsys):
    rc, _ = _inspect(tmp_path, capsys,
                     "--tolerance-chars", "7")
    assert rc == 0


# ---------- error_code (None) 复核 ----------

def test_error_code_none_line_batch514(tmp_path,
                                       capsys):
    _, out = _inspect(tmp_path, capsys)
    assert (f"  {'error_code':36}"
            " null  (None)") in out


# ---------- 混合计数行 ----------

def _mixed_run(tmp_path, capsys):
    (tmp_path / "g.pdf").write_bytes(_wrap(STREAM))
    d = Document()
    d.add_heading("T", level=1)
    d.add_paragraph("Body text.")
    d.save(str(tmp_path / "c.docx"))
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "g1", "path": "g.pdf",
             "source_type": "pdf"},
            {"doc_id": "d1", "path": "c.docx",
             "source_type": "docx"}]}),
        encoding="utf-8")
    sys.argv = ["evaluation.cli", "run", "--manifest",
                str(tmp_path / "m.json"),
                "--output", str(tmp_path / "r.json"),
                "--parser", "fallback",
                "--max-chars", "32"]
    rc = main()
    return rc, capsys.readouterr().out


def test_mixed_pdf_docx_line_batch514(tmp_path,
                                      capsys):
    _, out = _mixed_run(tmp_path, capsys)
    assert ("devset_status=incomplete file_count=2 "
            "groups=2 pdf=1 docx=1") in out


def test_mixed_success_line_batch514(tmp_path,
                                     capsys):
    rc, out = _mixed_run(tmp_path, capsys)
    assert rc == 0
    assert "documents=2（成功 2，失败 0）" in out


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_counts_batch514():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


def test_source_tolerance_default_batch514():
    src = _src()
    assert '"--tolerance-chars"' in src
    assert "default=30" in src


# ---------- forbidden tokens 第五百九十三批 ----------

def test_source_no_eval_batch514():
    assert "eval(" not in _src()


def test_source_no_exec_batch514():
    assert "exec(" not in _src()


def test_source_no_compile_batch514():
    assert "compile(" not in _src()


def test_source_no_globals_batch514():
    assert "globals(" not in _src()


def test_source_no_locals_batch514():
    assert "locals(" not in _src()


def test_source_no_os_system_batch514():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch514():
    assert "subprocess" not in _src()


def test_source_no_popen_batch514():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch514():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch514():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch514():
    assert "socket" not in _src()


def test_source_no_requests_batch514():
    assert "requests" not in _src()


def test_source_no_urllib_batch514():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch514():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch514():
    assert "yield" not in _src()


def test_source_no_async_await_batch514():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_call_batch514():
    assert ".call(" not in _src()


def test_source_open_count_is_1_batch514():
    assert _src().count("open(") == 1
