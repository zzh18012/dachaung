"""evaluation/metrics.py 第五百零九轮 edges 测试（Round 1065）。

补强 edges127-129 未触及的角度（第四百四十一批，probe 实证）。

新角度（真实 image 元素的指标板 + 磁盘状态因果对）：
- 真实嵌图 docx 直跑 compute_automatic_metrics：
  ecbt {paragraph 3, image 1} / ect 4 / docx_locator 1.0
  （image locator 的 relationship_id/target_partname
  rel 族键不 invalidate——R1057 page/bbox 之外的
  第三组键实证）/ image ratio **image_base_dir=None
  也 1.0**（绝对 resource_path 使 base dir 形同虚设，
  与 R1062 inspect-doc 口径一致）
- **磁盘状态因果对**：同一份 doc dict，PNG 在盘 1.0、
  unlink 后 0.0、**截断成 0 字节也 0.0**——存在性检查
  是 `is_file() and st_size > 0`，空文件等同缺失
- **占位符计入保持性账本**：parser 合成的 "(空段落)"
  元素 type=paragraph，与真实文本同权参与
  text_preservation_equal / multiset P/R（全绿——
  合成文本也是"不能丢"的内容）
- forbidden tokens 第五百三十六批（open 0）
"""

from __future__ import annotations

import inspect
import struct
import zlib
from io import BytesIO
from pathlib import Path

from docx import Document

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from evaluation.metrics import compute_automatic_metrics


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _img_doc(tmp_path):
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.add_picture(BytesIO(_png_bytes()))
    d.add_paragraph("BBB third paragraph body.")
    d.save(str(tmp_path / "img.docx"))
    doc, errors = process_single(
        tmp_path / "img.docx", tmp_path / "s.json",
        parser_name="fallback", max_chars=200,
        write_json=False)
    assert errors == []
    return doc.to_dict()


def _m(dd):
    return compute_automatic_metrics(dd, None, "docx",
                                     None,
                                     image_base_dir=None)


# ---------- 真实 image 板 ----------

def test_image_doc_full_board_batch264(tmp_path):
    m = _m(_img_doc(tmp_path))
    assert m["element_count_total"] == {"value": 4,
                                        "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 3, "image": 1},
        "reason": None}
    assert m["docx_locator_valid_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 绝对路径使 base dir 失效 ----------

def test_image_ratio_base_none_still_lit_batch264(tmp_path):
    dd = _img_doc(tmp_path)
    assert Path(dd["elements"][2]["resource_path"]).is_absolute()
    assert _m(dd)["image_resource_exists_ratio"] == {
        "value": 1.0, "reason": None}


# ---------- 磁盘状态因果对：unlink ----------

def test_image_ratio_zero_after_unlink_batch264(tmp_path):
    dd = _img_doc(tmp_path)
    Path(dd["elements"][2]["resource_path"]).unlink()
    assert _m(dd)["image_resource_exists_ratio"] == {
        "value": 0.0, "reason": None}


# ---------- 磁盘状态因果对：0 字节截断 ----------

def test_image_ratio_zero_on_empty_file_batch264(tmp_path):
    dd = _img_doc(tmp_path)
    p = Path(dd["elements"][2]["resource_path"])
    p.write_bytes(b"")
    assert p.is_file()
    assert _m(dd)["image_resource_exists_ratio"] == {
        "value": 0.0, "reason": None}


# ---------- 占位符计入保持性账本 ----------

def test_placeholder_in_preservation_batch264(tmp_path):
    dd = _img_doc(tmp_path)
    assert dd["elements"][1]["type"] == "paragraph"
    assert dd["elements"][1]["content"] == "(空段落)"
    m = _m(dd)
    assert m["text_preservation_equal"] == {
        "value": True, "reason": None}
    assert m["text_char_multiset_precision"] == {
        "value": 1.0, "reason": None}
    assert m["text_char_multiset_recall"] == {
        "value": 1.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch264():
    src = _src()
    assert ("if p.is_file() and p.stat().st_size > 0:"
            in src)
    assert "image_base_dir" in src


# ---------- forbidden tokens 第五百三十六批 ----------

def test_source_no_eval_batch264():
    assert "eval(" not in _src()


def test_source_no_exec_batch264():
    assert "exec(" not in _src()


def test_source_no_compile_batch264():
    assert "compile(" not in _src()


def test_source_no_globals_batch264():
    assert "globals(" not in _src()


def test_source_no_locals_batch264():
    assert "locals(" not in _src()


def test_source_no_os_system_batch264():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch264():
    assert "subprocess" not in _src()


def test_source_no_popen_batch264():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch264():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch264():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch264():
    assert "socket" not in _src()


def test_source_no_requests_batch264():
    assert "requests" not in _src()


def test_source_no_urllib_batch264():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch264():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch264():
    assert "yield" not in _src()


def test_source_no_async_await_batch264():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch264():
    assert "open(" not in _src()
