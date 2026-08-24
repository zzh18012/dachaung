r"""app/parsers/fallback_parser.py 边角测试 - 第二十四轮（Round 1411）。

新角度（probe 实证）docx run 级内容（历史只锁过 hyperlink
run 与图片 run，普通混合格式 run 从未锁）：
- bold/underline/italic 混排 run 直接拼接、格式零痕迹
- 两 run 之间无任何分隔符（'after tab'+'second' →
  'after tabsecond'）
- run 内 \t 制表符原样保留
- add_break（w:br）→ '\n' 进 content 与 chunk
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _build(tmp_path):
    d = Document()
    p = d.add_paragraph()
    p.add_run("plain start then ")
    p.add_run("bold middle").bold = True
    p.add_run(" and ").underline = True
    p.add_run("italic tail").italic = True
    p2 = d.add_paragraph()
    p2.add_run("before tab\tafter tab")
    p2.add_run("second\ttab")
    p3 = d.add_paragraph()
    p3.add_run("line one")
    p3.add_run().add_break()
    p3.add_run("line two after break")
    p = tmp_path / "runs.docx"
    d.save(str(p))
    return p


def _parse(tmp_path):
    p = _build(tmp_path)
    return FallbackParser().parse(
        p, compute_file_hash(p))


# ---------- 混排 run ----------

def test_mixed_runs_concatenated(
        tmp_path):
    doc = _parse(tmp_path)
    assert doc.elements[
        0].content == (
        "plain start then "
        "bold middle and "
        "italic tail")


def test_no_formatting_marks(
        tmp_path):
    doc = _parse(tmp_path)
    for mark in ("**", "__", "*",
                 "[", "]"):
        assert mark not in \
            doc.elements[0].content


def test_all_paragraph_types(
        tmp_path):
    doc = _parse(tmp_path)
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "paragraph",
        "paragraph"]


def test_paragraph_indexes(tmp_path):
    doc = _parse(tmp_path)
    assert [e.source_locator[
        "paragraph_index"]
        for e in doc.elements
        ] == [0, 1, 2]


# ---------- tab 与无分隔拼接 ----------

def test_tabs_preserved(tmp_path):
    doc = _parse(tmp_path)
    assert doc.elements[
        1].content == (
        "before tab\tafter tab"
        "second\ttab")


def test_runs_no_separator(tmp_path):
    """'after tab' + 'second' 之间
    无空格无 tab——run 文本纯拼接。"""
    doc = _parse(tmp_path)
    assert "after tabsecond" in \
        doc.elements[1].content


# ---------- w:br 换行 ----------

def test_break_becomes_newline(
        tmp_path):
    doc = _parse(tmp_path)
    assert doc.elements[
        2].content == (
        "line one\n"
        "line two after break")


# ---------- 管线与校验 ----------

def test_schema_valid(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path)
    assert is_valid(doc.to_dict())


def test_no_warnings(tmp_path):
    doc = _parse(tmp_path)
    assert doc.warnings == []


def test_pipeline_single_chunk(
        tmp_path):
    p = _build(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert len(doc.chunks) == 1
    assert doc.chunks[0].text == (
        "plain start then "
        "bold middle and "
        "italic tail "
        "before tab\tafter tab"
        "second\ttab "
        "line one\n"
        "line two after break")
    assert len(doc.chunks[0]
               .source_element_ids) == 3


def test_metrics_green(tmp_path):
    from evaluation.metrics import \
        compute_automatic_metrics
    p = _build(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    m = compute_automatic_metrics(
        doc.to_dict(), None, "docx",
        None)
    assert m[
        "text_preservation_equal"] \
        == {"value": True,
            "reason": None}
    assert m[
        "docx_locator_valid_ratio"] \
        == {"value": 1.0,
            "reason": None}
