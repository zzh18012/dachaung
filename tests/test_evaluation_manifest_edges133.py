"""evaluation/manifest.py 第五百一十七轮 edges 测试（Round 1073）。

补强 edges129-132 未触及的角度（第四百四十九批，probe 实证）。

新角度（caption 类型进 expectations + 非 ASCII categories）：
- **caption 首次入 expectations**：真实双图双题注 doc
  （ecbt {paragraph 2, image 2, caption 2}）——
  {caption 2, image 2, paragraph 2} 精确相抵 → silent
  {0}、total 0；过索 caption 3 → silent 1；**欠索
  caption 1 → silent 0**（过量供给不记 drop——三向
  不对称账本在 caption 类型上复现）
- **非 ASCII categories**：["图类", "Alpha"] →
  categories_covered ["Alpha", "图类"]（sorted 码点
  序，ASCII 在前中文在后）——中英混类目排序首锁
- forbidden tokens 第五百四十四批（open 1）
"""

from __future__ import annotations

import inspect
import io
import json
import struct
import zlib

from docx import Document

import evaluation.manifest as manifest_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _png_bytes() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return (struct.pack(">I", len(data)) + c
                + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\x00\x00")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", idat) + chunk(b"IEND", b""))


def _setup(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_picture(io.BytesIO(_png_bytes()))
    d.add_paragraph("Figure 1: a real caption line.")
    d.add_picture(io.BytesIO(_png_bytes()))
    d.add_paragraph("Figure 2: another caption line.")
    d.save(str(tmp_path / "samples" / "cap.docx"))


def _run(tmp_path, expectations, categories=None):
    doc_entry = {
        "doc_id": "d1", "path": "samples/cap.docx",
        "source_type": "docx",
        "expectations": {"element_count_by_type":
                         expectations}}
    if categories is not None:
        doc_entry["categories"] = categories
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [doc_entry],
        "expected_failures": []}), encoding="utf-8")
    return run_evaluation(load_manifest(mf, tmp_path),
                          tmp_path / "o.json",
                          max_chars=200)


# ---------- caption 精确相抵 ----------

def test_caption_expectations_exact_batch272(tmp_path):
    _setup(tmp_path)
    rep = _run(tmp_path, {"caption": 2, "image": 2,
                          "paragraph": 2})
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 0,
                                 "reason": None}
    assert rep["summary"]["silent_drop_total"] == 0


# ---------- caption 过索 ----------

def test_caption_over_demand_batch272(tmp_path):
    _setup(tmp_path)
    rep = _run(tmp_path, {"caption": 3})
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 1,
                                 "reason": None}


# ---------- caption 欠索：过量不记 ----------

def test_caption_under_demand_zero_batch272(tmp_path):
    _setup(tmp_path)
    rep = _run(tmp_path, {"caption": 1})
    assert rep["per_doc"][0]["metrics"][
        "silent_drop_count"] == {"value": 0,
                                 "reason": None}


# ---------- 非 ASCII categories 码点排序 ----------

def test_cjk_categories_sorted_batch272(tmp_path):
    _setup(tmp_path)
    rep = _run(tmp_path, {"caption": 2},
               categories=["图类", "Alpha"])
    assert rep["devset"]["categories_covered"] == [
        "Alpha", "图类"]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(manifest_mod)


def test_source_key_lines_batch272():
    src = _src()
    assert ("categories=tuple(d.get(\"categories\", []))"
            in src)
    assert "categories_covered" in src


# ---------- forbidden tokens 第五百四十四批 ----------

def test_source_no_eval_batch272():
    assert "eval(" not in _src()


def test_source_no_exec_batch272():
    assert "exec(" not in _src()


def test_source_no_compile_batch272():
    assert "compile(" not in _src()


def test_source_no_globals_batch272():
    assert "globals(" not in _src()


def test_source_no_locals_batch272():
    assert "locals(" not in _src()


def test_source_no_os_system_batch272():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch272():
    assert "subprocess" not in _src()


def test_source_no_popen_batch272():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch272():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch272():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch272():
    assert "socket" not in _src()


def test_source_no_requests_batch272():
    assert "requests" not in _src()


def test_source_no_urllib_batch272():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch272():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch272():
    assert "yield" not in _src()


def test_source_no_async_await_batch272():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch272():
    assert _src().count("open(") == 1
