"""evaluation/runner.py 第五百九十五轮 edges 测试（Round 1151）。

补强 edges167 未触及的角度（第五百二十三批，probe 实证）。

新角度（跨源标注 devset 汇总）：
- **双源双标注全命中**——真文本 PDF + 真 DOCX 各挂
  marker（劈点恰 d=0）→ 两 doc F1 皆 1.0，summary
  chunk_boundary_f1 macro {1.0, participating 2,
  not_evaluated 0}——跨源标注宏平均首锁
- **locator 分源参评**——pdf_locator 参 1 免 1、
  docx_locator 参 1 免 1——混合 devset 里两指标各只
  计本源文档（首锁）
- forbidden tokens 第六百二十三批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _build_pdf(objects, n_obj) -> bytes:
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objects[num] + b"endobj\n")
    xref_pos = len(out)
    out += b"xref\n0 " + str(n_obj).encode() + b"\n"
    out += b"0000000000 65535 f \n"
    for num in range(1, n_obj):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size " + str(n_obj).encode()
            + b"/Root 1 0 R>>\nstartxref\n"
            + str(xref_pos).encode() + b"\n%%EOF\n")
    return bytes(out)


def _board(tmp_path):
    (tmp_path / "samples").mkdir(exist_ok=True)
    (tmp_path / "anns").mkdir(exist_ok=True)
    s = (b"BT /F1 12 Tf 10 80 Td "
         b"(Alpha beta gamma delta epsilon zeta tail.) Tj ET")
    (tmp_path / "samples" / "p.pdf").write_bytes(_build_pdf({
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 500 100]"
            b"/Resources<</Font<</F1 5 0 R>>>>"
            b"/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }, 6))
    d = Document()
    d.add_paragraph("First para with words here.")
    d.add_paragraph("Second para closes it out.")
    d.save(str(tmp_path / "samples" / "w.docx"))
    (tmp_path / "anns" / "p.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "pd",
        "chunk_boundary_anchors": [
            {"marker": "gamma delta", "position": "after"}]}),
        encoding="utf-8")
    (tmp_path / "anns" / "w.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "wd",
        "chunk_boundary_anchors": [
            {"marker": "words here", "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [
            {"doc_id": "pd", "path": "samples/p.pdf",
             "source_type": "pdf",
             "annotation_file": "anns/p.json"},
            {"doc_id": "wd", "path": "samples/w.docx",
             "source_type": "docx",
             "annotation_file": "anns/w.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 双源双标注全命中 ----------

def test_cross_source_annotation_hits_batch349(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=32)
    for pd_ in r["per_doc"]:
        assert pd_["metrics"]["chunk_boundary_f1"] == {
            "value": 1.0, "reason": None}
    f1 = r["summary"]["ratio_macro_averages"][
        "chunk_boundary_f1"]
    assert f1 == {"macro_average": 1.0,
                  "participating_docs": 2,
                  "not_evaluated": 0}


# ---------- locator 分源参评 ----------

def test_locator_split_participation_batch349(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=32)
    ratios = r["summary"]["ratio_macro_averages"]
    assert ratios["pdf_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}
    assert ratios["docx_locator_valid_ratio"] == {
        "macro_average": 1.0, "participating_docs": 1,
        "not_evaluated": 1}


def test_cross_source_success_batch349(tmp_path):
    r = run_evaluation(_board(tmp_path), tmp_path / "r.json",
                       parser_name="fallback", max_chars=32)
    assert r["summary"]["success_rates"]["pipeline_success"] == {
        "success_count": 2, "total": 2, "rate": 1.0}
    assert r["devset"]["pdf_count"] == 1
    assert r["devset"]["docx_count"] == 1


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch349():
    src = _src()
    assert src.count("process_single") == 6
    assert src.count("chunk") == 9
    assert src.count("metrics") == 13


# ---------- forbidden tokens 第六百二十三批 ----------

def test_source_no_eval_batch349():
    assert "eval(" not in _src()


def test_source_no_exec_batch349():
    assert "exec(" not in _src()


def test_source_no_compile_batch349():
    assert "compile(" not in _src()


def test_source_no_globals_batch349():
    assert "globals(" not in _src()


def test_source_no_locals_batch349():
    assert "locals(" not in _src()


def test_source_no_os_system_batch349():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch349():
    assert "subprocess" not in _src()


def test_source_no_popen_batch349():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch349():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch349():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch349():
    assert "socket" not in _src()


def test_source_no_requests_batch349():
    assert "requests" not in _src()


def test_source_no_urllib_batch349():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch349():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch349():
    assert "yield" not in _src()


def test_source_no_async_await_batch349():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch349():
    assert _src().count("open(") == 2
