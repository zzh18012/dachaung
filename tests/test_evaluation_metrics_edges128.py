"""evaluation/metrics.py 第四百九十五轮 edges 测试（Round 1051）。

补强 edges127 未触及的角度（第四百二十七批，probe 实证）。

新角度（真实三类型异构板：heading + paragraph + table）：
- 真实 python-docx 文档加 add_heading / 两段 /
  add_table 穿真实 parser：heading 被真实分类为
  "heading"（locator paragraph_index 0）、table 被
  分类为 "table"（locator **table_index** 0——结构
  键家族里 table 键首次以真实产物出现）
- ecbt {heading 1, paragraph 2, table 1} + total 4：
  真实三类型计数板
- heading_boundary_compliance 1.0：真实 heading 是
  首 chunk 的首元素 id——first-id 规则以真实文档
  满足（此前 0.0 板全是手造）
- 真实 table 渲染成 markdown 表格 chunk
  （"| cell one |"）——chunk 文本形态首次锁定
- docx_locator 1.0：paragraph_index 与 table_index
  两族结构键混合真实板全过
- forbidden tokens 第五百二十二批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

from app.pipeline import process_single
import evaluation.metrics as metrics_mod
from evaluation.metrics import compute_automatic_metrics


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_heading("Real Title", level=1)
    d.add_paragraph("AAA first paragraph body.")
    d.add_paragraph("CCC third paragraph body.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "cell one"
    d.save(str(p))
    doc, errors = process_single(p, tmp_path / "s.json",
                                 parser_name="fallback",
                                 max_chars=200,
                                 write_json=False)
    assert errors == []
    return doc.to_dict()


# ---------- 真实类型分类与 locator ----------

def test_real_type_classification_batch249(tmp_path):
    dd = _real_doc(tmp_path)
    els = dd["elements"]
    assert [e["type"] for e in els] == ["heading",
                                        "paragraph",
                                        "paragraph", "table"]
    assert els[0]["source_locator"] == {
        "paragraph_index": 0, "section": 0}
    assert els[3]["source_locator"] == {
        "table_index": 0, "section": 0}


# ---------- 三类型计数板 ----------

def test_real_three_type_counts_batch249(tmp_path):
    m = compute_automatic_metrics(_real_doc(tmp_path), None,
                                  "docx", None)
    assert m["element_count_by_type"] == {
        "value": {"heading": 1, "paragraph": 2,
                  "table": 1}, "reason": None}
    assert m["element_count_total"] == {"value": 4,
                                        "reason": None}


# ---------- first-id 规则真实满足 ----------

def test_real_heading_compliance_batch249(tmp_path):
    m = compute_automatic_metrics(_real_doc(tmp_path), None,
                                  "docx", None)
    assert m["heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


# ---------- markdown 表格 chunk ----------

def test_real_table_markdown_chunk_batch249(tmp_path):
    dd = _real_doc(tmp_path)
    table_chunks = [c for c in dd["chunks"]
                    if "cell one" in c["text"]]
    assert len(table_chunks) == 1
    assert table_chunks[0]["text"].startswith("| cell one |")


# ---------- 混合结构键 locator ----------

def test_real_mixed_structural_keys_batch249(tmp_path):
    m = compute_automatic_metrics(_real_doc(tmp_path), None,
                                  "docx", None)
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["chunk_reference_intact_ratio"] == {
        "value": 1.0, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch249():
    src = _src()
    assert "def compute_automatic_metrics(" in src
    assert "heading_boundary_compliance" in src
    assert "table_index" in src


# ---------- forbidden tokens 第五百二十二批 ----------

def test_source_no_eval_batch249():
    assert "eval(" not in _src()


def test_source_no_exec_batch249():
    assert "exec(" not in _src()


def test_source_no_compile_batch249():
    assert "compile(" not in _src()


def test_source_no_globals_batch249():
    assert "globals(" not in _src()


def test_source_no_locals_batch249():
    assert "locals(" not in _src()


def test_source_no_os_system_batch249():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch249():
    assert "subprocess" not in _src()


def test_source_no_popen_batch249():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch249():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch249():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch249():
    assert "socket" not in _src()


def test_source_no_requests_batch249():
    assert "requests" not in _src()


def test_source_no_urllib_batch249():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch249():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch249():
    assert "yield" not in _src()


def test_source_no_async_await_batch249():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch249():
    assert "open(" not in _src()
