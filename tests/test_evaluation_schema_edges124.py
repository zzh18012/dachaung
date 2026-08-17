"""evaluation/schema.py 第五百二十二轮 edges 测试（Round 1078）。

补强 edges119-123 未触及的角度（第四百五十四批，probe 实证）。

新角度（manifest 闭仓 + document 元素类型枚举名册）：
- manifest.schema.json 两个 def（document / expected_
  failure）均 additionalProperties: false——真实
  load_manifest 对 document 条目多塞一个 "note" 键
  即拒（'Additional properties are not allowed'）
  ——清单条目与报告 ef 条目同纪律，唯报告 metrics
  暗仓例外（R1071 的倒挂图谱补上 manifest 一角）
- def 层 required 名册：document [doc_id, path,
  source_type]、expected_failure [doc_id, path,
  expected_error_code]
- **document 元素类型枚举 8 值全名册**：[heading,
  paragraph, list_item, table, image, caption, header,
  footer]——caption 合法（真实产物旁证）；header /
  footer / list_item 合法但 fallback parser 从不产
  出——未达枚举值首锁；真实文档元素 type 改成
  "header" 照过 schema
- forbidden tokens 第五百四十九批（open 2）
"""

from __future__ import annotations

import copy
import inspect
import json

from docx import Document

import evaluation.schema as schema_mod
from app.pipeline import process_single
from evaluation.manifest import (
    ManifestError, load_manifest)
from evaluation.schema import (
    EvalSchemaError, load_schema, validate)


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.save(str(p))
    doc, errors = process_single(
        p, tmp_path / "s.json", parser_name="fallback",
        max_chars=200, write_json=False)
    assert errors == []
    return doc.to_dict()


def _manifest(tmp_path, extra):
    (tmp_path / "samples").mkdir(exist_ok=True)
    entry = {"doc_id": "d1", "path": "samples/a.docx",
             "source_type": "docx"}
    if extra:
        entry["note"] = "extra key"
    (tmp_path / "m.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [entry],
        "expected_failures": []}), encoding="utf-8")


# ---------- 真实加载：多键即拒 ----------

def test_manifest_doc_entry_closed_batch277(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("AAA first paragraph body.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    try:
        _manifest(tmp_path, extra=True)
        load_manifest(tmp_path / "m.json", tmp_path)
        raised = False
    except (ManifestError, EvalSchemaError) as e:
        raised = True
        assert "Additional properties are not allowed" \
            in str(e)
    assert raised


# ---------- def 层闭仓名册 ----------

def test_manifest_defs_closed_batch277():
    s = load_schema("manifest.schema.json")
    doc_def = s["$defs"]["document"]
    ef_def = s["$defs"]["expected_failure"]
    assert doc_def["additionalProperties"] is False
    assert doc_def["required"] == ["doc_id", "path",
                                   "source_type"]
    assert ef_def["additionalProperties"] is False
    assert ef_def["required"] == ["doc_id", "path",
                                  "expected_error_code"]


# ---------- 元素类型枚举 8 值名册 ----------

def test_element_type_enum_roster_batch277():
    s = load_schema("document.schema.json")
    enum = s["$defs"]["element"]["properties"]["type"][
        "enum"]
    assert enum == ["heading", "paragraph", "list_item",
                    "table", "image", "caption",
                    "header", "footer"]
    assert set(enum) >= {"caption", "header", "footer",
                         "list_item"}


# ---------- 未达枚举值照过 schema ----------

def test_unproduced_type_accepted_batch277(tmp_path):
    dd = _real_doc(tmp_path)
    r = copy.deepcopy(dd)
    r["elements"][0]["type"] = "header"
    validate(r, "document.schema.json")
    r2 = copy.deepcopy(dd)
    r2["elements"][0]["type"] = "footer"
    validate(r2, "document.schema.json")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch277():
    src = _src()
    assert "raise EvalSchemaError(" in src
    assert '"validate_file",' in src


# ---------- forbidden tokens 第五百四十九批 ----------

def test_source_no_eval_batch277():
    assert "eval(" not in _src()


def test_source_no_exec_batch277():
    assert "exec(" not in _src()


def test_source_no_compile_batch277():
    assert "compile(" not in _src()


def test_source_no_globals_batch277():
    assert "globals(" not in _src()


def test_source_no_locals_batch277():
    assert "locals(" not in _src()


def test_source_no_os_system_batch277():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch277():
    assert "subprocess" not in _src()


def test_source_no_popen_batch277():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch277():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch277():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch277():
    assert "socket" not in _src()


def test_source_no_requests_batch277():
    assert "requests" not in _src()


def test_source_no_urllib_batch277():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch277():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch277():
    assert "yield" not in _src()


def test_source_no_async_await_batch277():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch277():
    assert _src().count("open(") == 2
