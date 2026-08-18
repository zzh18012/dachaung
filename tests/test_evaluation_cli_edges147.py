"""evaluation/cli.py 第六百二十八轮 edges 测试（Round 1184）。

补强 cli edges146 未触及的角度（第五百五十六批，probe 实证）。

新角度（样式板 CLI 全链）：
- **样式板 run**——Title/Heading 9/Quote/
  Subtitle/Normal 五段 DOCX 经 CLI run →
  rc 0、by_type {heading: 2, paragraph: 3}
  （深级题与 Title 特判经 CLI 通道首锁）
- **inspect-doc 样式板 counts 行**——
  "counts:      elements=5 chunks=2"
- **自产自校**——同报告 validate-report rc 0
- forbidden tokens 第五百五十六批（open 1）
"""

from __future__ import annotations

import inspect
import json

import evaluation.cli as cli_mod
from evaluation.cli import main


def _board(tmp_path):
    from docx import Document
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("Title style paragraph", style="Title")
    d.add_paragraph("Heading nine text here",
                    style="Heading 9")
    d.add_paragraph("Quote style paragraph", style="Quote")
    d.add_paragraph("Subtitle style paragraph",
                    style="Subtitle")
    d.add_paragraph("Normal tail paragraph.")
    d.save(str(tmp_path / "samples" / "t.docx"))
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": "st", "path": "samples/t.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return mf


def _run_cli(capsys, argv):
    rc = main(argv)
    out = capsys.readouterr().out
    return rc, out


# ---------- 样式板 run ----------

def test_cli_run_style_board_batch382(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "200"])
    assert rc == 0
    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    assert rep["per_doc"][0]["metrics"][
        "element_count_by_type"] == {
        "value": {"heading": 2, "paragraph": 3},
        "reason": None}
    assert rep["per_doc"][0]["metrics"][
        "heading_boundary_compliance"] == {
        "value": 1.0, "reason": None}


def test_cli_run_style_stdout_batch382(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, out = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "200"])
    assert rc == 0
    assert "documents=1（成功 1，失败 0）" in out
    assert "pdf=0 docx=1" in out


# ---------- inspect-doc 样式板 counts 行 ----------

def test_cli_inspect_style_counts_batch382(tmp_path, capsys):
    from app.pipeline import process_single
    _board(tmp_path)
    doc, errors = process_single(
        tmp_path / "samples" / "t.docx", tmp_path / "doc.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    rc, out = _run_cli(capsys, [
        "inspect-doc", str(tmp_path / "doc.json")])
    assert rc == 0
    assert "counts:      elements=5 chunks=2" in out


# ---------- 自产自校 ----------

def test_cli_validate_style_report_batch382(tmp_path, capsys):
    mf = _board(tmp_path)
    rc, _ = _run_cli(capsys, [
        "run", "--manifest", str(mf),
        "--output", str(tmp_path / "r.json"),
        "--parser", "fallback", "--max-chars", "200"])
    assert rc == 0
    rc2, out2 = _run_cli(capsys, [
        "validate-report", str(tmp_path / "r.json")])
    assert rc2 == 0
    assert "[OK]" in out2


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_identifier_counts_batch382():
    src = _src()
    assert src.count("validate") == 7
    assert src.count("inspect") == 9
    assert src.count("report") == 14


# ---------- forbidden tokens 第五百五十六批 ----------

def test_source_no_eval_batch382():
    assert "eval(" not in _src()


def test_source_no_exec_batch382():
    assert "exec(" not in _src()


def test_source_no_compile_batch382():
    assert "compile(" not in _src()


def test_source_no_globals_batch382():
    assert "globals(" not in _src()


def test_source_no_locals_batch382():
    assert "locals(" not in _src()


def test_source_no_os_system_batch382():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch382():
    assert "subprocess" not in _src()


def test_source_no_popen_batch382():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch382():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch382():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch382():
    assert "socket" not in _src()


def test_source_no_requests_batch382():
    assert "requests" not in _src()


def test_source_no_urllib_batch382():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch382():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch382():
    assert "yield" not in _src()


def test_source_no_async_await_batch382():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch382():
    assert _src().count("open(") == 1
