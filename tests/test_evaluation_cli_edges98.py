"""evaluation/cli.py 第二百六十九轮 edges 测试（Round 825）。

补强 edges97 未触及的角度（第一百九十八批）。

新角度（**无 mock 真跑端到端**）：
- python-docx 在 tmp 现造双段 DOCX → manifest →
  main run 全链路：rc 0、真实 fallback parser、
  pipeline_success True、element_count_total 2、
  text_preservation_equal True、parser_name fallback
- 摘要行真数据：documents=1（成功 1，失败 0）/
  devset_status=incomplete file_count=1 groups=1
  pdf=0 docx=1 / git_commit=unknown git_dirty=False
  （非 git tmp：returncode 路径 dirty False）
- 生成的报告直接通过 validate-report rc 0
- forbidden tokens 第二百九十五批
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document as DocxDocument

import evaluation.cli as cli_mod
from evaluation.cli import main


def _cap():
    out, err = io.StringIO(), io.StringIO()
    return out, err, contextlib.redirect_stdout(out), \
        contextlib.redirect_stderr(err)


def _setup(tmp_path):
    (tmp_path / "samples").mkdir()
    doc = DocxDocument()
    doc.add_paragraph("Hello world paragraph one.")
    doc.add_paragraph("Second paragraph here.")
    doc.save(tmp_path / "samples" / "a.docx")
    mf = tmp_path / "manifest.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [
            {"doc_id": "d1", "path": "samples/a.docx",
             "source_type": "docx"}]}), encoding="utf-8")
    return mf


# ---------- 端到端 run ----------

def test_real_end_to_end_run_batch55(tmp_path):
    mf = _setup(tmp_path)
    out, err, co, ce = _cap()
    with co, ce:
        rc = main(["run", "--manifest", str(mf),
                   "--output", str(tmp_path / "r.json")])
    assert rc == 0
    assert err.getvalue() == ""

    rep = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    m = rep["per_doc"][0]["metrics"]
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}
    assert m["element_count_total"] == {"value": 2,
                                        "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert rep["provenance"]["parser_name"] == "fallback"
    assert rep["provenance"]["parser_version"] is not None

    lines = out.getvalue().splitlines()
    assert lines[1] == "      documents=1（成功 1，失败 0）"
    assert lines[2] == ("      devset_status=incomplete "
                        "file_count=1 groups=1 pdf=0 docx=1")
    assert lines[3] == \
        "      git_commit=unknown git_dirty=False"


# ---------- 端到端 validate-report ----------

def test_real_report_validates_batch55(tmp_path):
    mf = _setup(tmp_path)
    out, err, co, ce = _cap()
    with co, ce:
        main(["run", "--manifest", str(mf),
              "--output", str(tmp_path / "r.json")])
    out, err, co, ce = _cap()
    with co, ce:
        rc = main(["validate-report", str(tmp_path / "r.json")])
    assert rc == 0
    assert out.getvalue().strip().endswith(
        "通过 evaluation-report Schema 校验")


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch55():
    src = _src()
    assert "run_evaluation(" in src
    assert "validate_file(output_path, \"evaluation-report.schema.json\")" in src


# ---------- forbidden tokens 第二百九十五批 ----------

def test_source_no_eval_batch55():
    assert "eval(" not in _src()


def test_source_no_exec_batch55():
    assert "exec(" not in _src()


def test_source_no_compile_batch55():
    assert "compile(" not in _src()


def test_source_no_globals_batch55():
    assert "globals(" not in _src()


def test_source_no_locals_batch55():
    assert "locals(" not in _src()


def test_source_no_os_system_batch55():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch55():
    assert "subprocess" not in _src()


def test_source_no_popen_batch55():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch55():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch55():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch55():
    assert "socket" not in _src()


def test_source_no_requests_batch55():
    assert "requests" not in _src()


def test_source_no_urllib_batch55():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch55():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch55():
    assert "yield" not in _src()


def test_source_no_async_await_batch55():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch55():
    assert _src().count("open(") == 1
