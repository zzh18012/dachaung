"""evaluation/cli.py 第五百三十四轮 edges 测试（Round 1090）。

补强 edges133-135 未触及的角度（第四百六十六批，probe 实证）。

新角度（heading 度量在 inspect-doc 点亮 + complete 真跑过验）：
- 真实 heading 文档（intro + add_heading + 两段）经
  process_single 后 inspect-doc：
  "  heading_boundary_compliance          1.0000  (ok)"
  ——heading 度量首次在渲染层点亮（edges134 只见过
  null no_heading_elements）
- dict 行 "  element_count_by_type                "
  "heading=1, paragraph=2  (ok)"——heading 码点排首
  （edges118 的 _format_metric 单元层结论在真实
  文档渲染整装复现）
- counts 行 "elements=3 chunks=2"——R1079 的 heading
  强制断点形态（3 元素 2 chunk）浮出 CLI 渲染
- **complete 真跑过验**：devset_status complete 的
  真实 run 报告 → validate-report rc 0 [OK]——
  complete 不只是能跑（edges30 只验 run rc 0），
  产物报告也过 schema
- forbidden tokens 第五百六十一批（open 1）
"""

from __future__ import annotations

import contextlib
import inspect
import io
import json

from docx import Document

import evaluation.cli as cli_mod
from app.pipeline import process_single
from evaluation.cli import main
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _inspect(tmp_path):
    d = Document()
    d.add_paragraph("AAA intro paragraph before the heading.")
    d.add_heading("Late Title", level=1)
    d.add_paragraph("BBB body after the heading one.")
    d.save(str(tmp_path / "h.docx"))
    doc, errors = process_single(
        tmp_path / "h.docx", tmp_path / "doc.json",
        parser_name="fallback", max_chars=200,
        write_json=True)
    assert errors == []
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["inspect-doc", str(tmp_path / "doc.json")])
    return rc, buf.getvalue()


# ---------- heading 度量点亮 ----------

def test_heading_lit_render_batch289(tmp_path):
    rc, out = _inspect(tmp_path)
    assert rc == 0
    assert ("  heading_boundary_compliance"
            "          1.0000  (ok)") in out


# ---------- dict 行 heading 排首 ----------

def test_ecbt_heading_first_batch289(tmp_path):
    _, out = _inspect(tmp_path)
    assert ("  element_count_by_type"
            "                heading=1, paragraph=2"
            "  (ok)") in out


# ---------- counts 行：heading 强制断点形态 ----------

def test_counts_heading_shape_batch289(tmp_path):
    _, out = _inspect(tmp_path)
    assert "counts:      elements=3 chunks=2" in out


# ---------- complete 真跑过验 ----------

def test_complete_real_run_validates_batch289(tmp_path):
    (tmp_path / "pyproject.toml").write_text("",
                                             encoding="utf-8")
    (tmp_path / "samples").mkdir()
    d = Document()
    d.add_paragraph("AAA body paragraph.")
    d.save(str(tmp_path / "samples" / "a.docx"))
    (tmp_path / "mc.json").write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "complete",
        "documents": [{"doc_id": "d1",
                       "path": "samples/a.docx",
                       "source_type": "docx"}],
        "expected_failures": []}), encoding="utf-8")
    run_evaluation(load_manifest(tmp_path / "mc.json",
                                 tmp_path),
                   tmp_path / "rc.json", max_chars=200)
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        rc = main(["validate-report",
                   str(tmp_path / "rc.json")])
    assert rc == 0
    assert "[OK]" in buf.getvalue()


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(cli_mod)


def test_source_key_lines_batch289():
    src = _src()
    assert "def _run_inspect_doc(" in src
    assert "image_base_dir=None" in src


# ---------- forbidden tokens 第五百六十一批 ----------

def test_source_no_eval_batch289():
    assert "eval(" not in _src()


def test_source_no_exec_batch289():
    assert "exec(" not in _src()


def test_source_no_compile_batch289():
    assert "compile(" not in _src()


def test_source_no_globals_batch289():
    assert "globals(" not in _src()


def test_source_no_locals_batch289():
    assert "locals(" not in _src()


def test_source_no_os_system_batch289():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch289():
    assert "subprocess" not in _src()


def test_source_no_popen_batch289():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch289():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch289():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch289():
    assert "socket" not in _src()


def test_source_no_requests_batch289():
    assert "requests" not in _src()


def test_source_no_urllib_batch289():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch289():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch289():
    assert "yield" not in _src()


def test_source_no_async_await_batch289():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_1_batch289():
    assert _src().count("open(") == 1
