"""evaluation/runner.py 第六百零七轮 edges 测试（Round 1163）。

补强 edges178 未触及的角度（第五百三十五批，probe 实证）。

新角度（合并单元格 / 分节符幻影段）：
- **合并单元格文本重复列**——cell(1,0).merge(
  cell(1,1)) 后行遍历回显两格 → markdown
  "| Left | Right |…| MergedWide | MergedWide |"
  ——col_count 仍 2，合并文本双写（首锁）
- **分节符幻影空段**——add_section() 引入一个
  "(空段落)" 元素（paragraph_index 1）；后续段
  落 section 仍 0——分节不升 section 号（可观测
  语义首锁）
- **跨节文本同流**——三段（含幻影空段）单 chunk
  3 源；by_type {paragraph: 3} 含占位符
- forbidden tokens 第六百三十五批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _merged_board(tmp_path, doc_id):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Left"
    t.cell(0, 1).text = "Right"
    t.cell(1, 0).merge(t.cell(1, 1))
    t.cell(1, 0).text = "MergedWide"
    d.save(str(tmp_path / "samples" / f"{doc_id}.docx"))
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"samples/{doc_id}.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _section_board(tmp_path, doc_id):
    (tmp_path / "samples").mkdir(exist_ok=True)
    d = Document()
    d.add_paragraph("First section paragraph.")
    d.add_section()
    d.add_paragraph("Second section paragraph.")
    d.save(str(tmp_path / "samples" / f"{doc_id}.docx"))
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0",
        "devset_status": "incomplete",
        "documents": [{"doc_id": doc_id,
                       "path": f"samples/{doc_id}.docx",
                       "source_type": "docx"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 合并单元格文本重复列 ----------

def test_merged_cell_duplicated_batch361(tmp_path):
    _merged_board(tmp_path, "mc")
    doc, errors = process_single(
        tmp_path / "samples" / "mc.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == ["table"]
    assert els[0]["content"] == (
        "| Left | Right |\n| --- | --- |\n"
        "| MergedWide | MergedWide |")
    assert els[0]["metadata"]["col_count"] == 2
    assert els[0]["metadata"]["row_count"] == 2


def test_merged_table_metrics_batch361(tmp_path):
    r = run_evaluation(_merged_board(tmp_path, "mc2"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"table": 1}, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["pipeline_success"] == {"value": True,
                                     "reason": None}


# ---------- 分节符幻影空段 ----------

def test_section_break_phantom_empty_batch361(tmp_path):
    _section_board(tmp_path, "sec")
    doc, errors = process_single(
        tmp_path / "samples" / "sec.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["content"] for e in els] == [
        "First section paragraph.", "(空段落)",
        "Second section paragraph."]
    assert [e["source_locator"]["paragraph_index"]
            for e in els] == [0, 1, 2]
    assert all(e["source_locator"]["section"] == 0
               for e in els)


# ---------- 跨节文本同流 ----------

def test_section_break_flows_batch361(tmp_path):
    _section_board(tmp_path, "sec2")
    doc, errors = process_single(
        tmp_path / "samples" / "sec2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert len(chunks) == 1
    assert chunks[0]["text"] == (
        "First section paragraph. (空段落) "
        "Second section paragraph.")
    assert len(chunks[0]["source_element_ids"]) == 3


def test_section_board_metrics_batch361(tmp_path):
    r = run_evaluation(_section_board(tmp_path, "sec3"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 3}, "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch361():
    src = _src()
    assert src.count("process_single") == 6
    assert src.count("annotation") == 10
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百三十五批 ----------

def test_source_no_eval_batch361():
    assert "eval(" not in _src()


def test_source_no_exec_batch361():
    assert "exec(" not in _src()


def test_source_no_compile_batch361():
    assert "compile(" not in _src()


def test_source_no_globals_batch361():
    assert "globals(" not in _src()


def test_source_no_locals_batch361():
    assert "locals(" not in _src()


def test_source_no_os_system_batch361():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch361():
    assert "subprocess" not in _src()


def test_source_no_popen_batch361():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch361():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch361():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch361():
    assert "socket" not in _src()


def test_source_no_requests_batch361():
    assert "requests" not in _src()


def test_source_no_urllib_batch361():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch361():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch361():
    assert "yield" not in _src()


def test_source_no_async_await_batch361():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch361():
    assert _src().count("open(") == 2
