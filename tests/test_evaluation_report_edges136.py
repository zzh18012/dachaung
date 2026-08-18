"""evaluation/report.py 第五百六十七轮 edges 测试（Round 1258）。

补强 edges135 未触及的角度（第六百三十批，probe 实证）。

新角度（报告序列化往返 / 键序 / 容差缺席）：
- **往返相等**——盘上 r.json 读回
  == run_evaluation 返回值（默认
  json.dumps 序列化无损首锁）
- **文件键序**——['report_version',
  'provenance', 'devset', 'summary',
  'per_doc', 'expected_failures']（插
  入序非字母序首锁）
- **per_doc 保 manifest 序**——
  wg31 在 wg30 前列 → per_doc 同
  序（非字母序反转首锁）
- **provenance 九键**——dependencies/
  evaluator_version/git_commit/
  git_dirty/max_chars/parser_name/
  parser_version/report_version/
  run_timestamp_iso
- **容差公开缺席**——标注在场时
  chunk_boundary_* 出值，但
  _tolerance_chars 连 metrics 内也
  被剥，"tolerance_chars" 全报告
  无串
- forbidden tokens 第七百二十批（open 0）
"""

from __future__ import annotations

import inspect
import json

import evaluation.report as report_mod
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation


def _ks_board(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    doc.add_paragraph("Tail run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    doc.save(str(tmp_path / "ks.docx"))
    (tmp_path / "a").mkdir()
    (tmp_path / "a" / "a.json").write_text(json.dumps({
        "annotation_version": "1.0", "doc_id": "ks",
        "chunk_boundary_anchors": [
            {"marker": "(空段落)", "position": "after"}]}),
        encoding="utf-8")
    mf = tmp_path / "m.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [{"doc_id": "ks", "path": "ks.docx",
                       "source_type": "docx",
                       "annotation_file": "a/a.json"}]}),
        encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _two_page(y2: int) -> bytes:
    s1 = (("BT /F1 12 Tf 10 700 Td (Top line text here.) Tj ET\n"
           "BT /F1 12 Tf 10 %d Td (Lower line text here.) Tj ET\n"
           % y2).encode())
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R 6 0 R]/Count 2>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s1)).encode()
            + b">>stream\n" + s1 + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        6: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 7 0 R>>"),
        7: (b"<</Length " + str(len(s1)).encode()
            + b">>stream\n" + s1 + b"\nendstream "),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 8\n0000000000 65535 f \n"
    for num in range(1, 8):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 8/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


def _pair_board(tmp_path):
    for did, y2 in (("wg30", 670), ("wg31", 669)):
        (tmp_path / (did + ".pdf")).write_bytes(_two_page(y2))
    mf = tmp_path / "mp.json"
    mf.write_text(json.dumps({
        "manifest_version": "1.0", "devset_status": "incomplete",
        "documents": [
            {"doc_id": "wg31", "path": "wg31.pdf",
             "source_type": "pdf"},
            {"doc_id": "wg30", "path": "wg30.pdf",
             "source_type": "pdf"}]}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


def _pair_run(tmp_path):
    return run_evaluation(_pair_board(tmp_path), tmp_path / "r.json",
                          parser_name="fallback", max_chars=200)


# ---------- 往返相等 ----------

def test_round_trip_equal_batch456(tmp_path):
    r = _pair_run(tmp_path)
    on_disk = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"))
    assert on_disk == r


# ---------- 文件键序 ----------

def test_file_key_order_batch456(tmp_path):
    _pair_run(tmp_path)
    order = json.loads(
        (tmp_path / "r.json").read_text(encoding="utf-8"),
        object_pairs_hook=list)
    assert [k for k, _ in order] == [
        "report_version", "provenance", "devset", "summary",
        "per_doc", "expected_failures"]


def test_report_version_string_batch456(tmp_path):
    r = _pair_run(tmp_path)
    assert r["report_version"] == "1.1"


# ---------- per_doc 保 manifest 序 ----------

def test_per_doc_manifest_order_batch456(tmp_path):
    r = _pair_run(tmp_path)
    assert [p["doc_id"] for p in r["per_doc"]] == ["wg31", "wg30"]


# ---------- provenance 九键 ----------

def test_provenance_keys_batch456(tmp_path):
    r = _pair_run(tmp_path)
    assert sorted(r["provenance"].keys()) == [
        "dependencies", "evaluator_version", "git_commit",
        "git_dirty", "max_chars", "parser_name", "parser_version",
        "report_version", "run_timestamp_iso"]


def test_provenance_values_batch456(tmp_path):
    r = _pair_run(tmp_path)
    pv = r["provenance"]
    assert pv["evaluator_version"] == "1.1"
    assert pv["parser_name"] == "fallback"
    assert pv["max_chars"] == 200
    assert pv["report_version"] == "1.1"


# ---------- 容差公开缺席 ----------

def test_annotated_boundary_values_present_batch456(tmp_path):
    r = run_evaluation(_ks_board(tmp_path), tmp_path / "rk.json",
                       parser_name="fallback", max_chars=120)
    md = r["per_doc"][0]["metrics"]
    assert md["chunk_boundary_precision"] == {"value": 0.5,
                                              "reason": None}
    assert md["chunk_boundary_recall"] == {"value": 1.0,
                                           "reason": None}
    assert md["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


def test_tolerance_absent_public_batch456(tmp_path):
    r = run_evaluation(_ks_board(tmp_path), tmp_path / "rk.json",
                       parser_name="fallback", max_chars=120)
    assert "_tolerance_chars" not in r["per_doc"][0]["metrics"]
    assert "tolerance_chars" not in json.dumps(r)


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(report_mod)


def test_source_key_lines_batch456():
    src = _src()
    assert "def aggregate_summary(" in src
    assert "_RATIO_METRICS" in src


# ---------- forbidden tokens 第七百二十批 ----------

def test_source_no_eval_batch456():
    assert "eval(" not in _src()


def test_source_no_exec_batch456():
    assert "exec(" not in _src()


def test_source_no_compile_batch456():
    assert "compile(" not in _src()


def test_source_no_globals_batch456():
    assert "globals(" not in _src()


def test_source_no_locals_batch456():
    assert "locals(" not in _src()


def test_source_no_os_system_batch456():
    assert "os.system" not in _src()


def test_source_no_subprocess_call_batch456():
    assert ".call(" not in _src()


def test_source_no_popen_batch456():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch456():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch456():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch456():
    assert "socket" not in _src()


def test_source_no_requests_batch456():
    assert "requests" not in _src()


def test_source_no_urllib_batch456():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch456():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch456():
    assert "yield" not in _src()


def test_source_no_async_await_batch456():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_no_open_batch456():
    assert "open(" not in _src()


def test_source_subprocess_run_count_is_2_batch456():
    assert _src().count("subprocess.run") == 2
