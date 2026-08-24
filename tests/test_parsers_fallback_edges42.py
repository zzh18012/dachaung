r"""app/parsers/fallback_parser.py 边角测试 - 第四十二轮（Round 1434）。

新角度（probe 实证）嵌套 BT + 相邻表（R1424 锁过缺 ET，
未碰过嵌套 BT；docx 表格历史全被段落隔开）：
- BT 内再 BT：pdfminer 容忍——两段文本各自成元素、bbox
  正常、无告警
- docx 两个 add_table 紧邻（无段落间隔）：两个独立 table
  元素 table_index 0/1，无幽灵空段
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _build(content):
    objs = {
        5: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] "
            b"/Count 1 >>"),
        3: (b"<< /Type /Page /Parent "
            b"2 0 R /MediaBox "
            b"[0 0 612 792] "
            b"/Resources << /Font "
            b"<< /F1 5 0 R >> >> "
            b"/Contents 4 0 R >>"),
        4: (b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n" + content
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    out += b"xref\n0 6\n" \
        b"0000000000 65535 f \n"
    for oid in range(1, 6):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 6 "
            b"/Root 1 0 R >>\n"
            b"startxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    return bytes(out)


def _nbt_pdf(tmp_path):
    p = tmp_path / "nbt.pdf"
    p.write_bytes(_build(
        b"BT /F1 12 Tf 72 700 Td "
        b"(Outer nested text) Tj "
        b"BT /F1 12 Tf 72 600 Td "
        b"(Inner nested text) Tj "
        b"ET ET"))
    return p


def _tt_docx(tmp_path):
    d = Document()
    d.add_paragraph("intro")
    t1 = d.add_table(rows=1, cols=1)
    t1.cell(0, 0).text = "first"
    t2 = d.add_table(rows=1, cols=1)
    t2.cell(0, 0).text = "second"
    p = tmp_path / "tt.docx"
    d.save(str(p))
    return p


# ---------- 嵌套 BT ----------

def test_nbt_both_extracted(
        tmp_path):
    p = _nbt_pdf(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "Outer nested text",
        "Inner nested text"]


def test_nbt_bboxes(tmp_path):
    p = _nbt_pdf(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].source_locator["bbox"] == [
        72.0, 82.48400000000004,
        164.052, 94.48400000000004]
    assert doc.elements[
        1].source_locator["bbox"] == [
        72.0, 182.48400000000004,
        161.38799999999998,
        194.48400000000004]


def test_nbt_no_warnings(tmp_path):
    p = _nbt_pdf(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_nbt_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _nbt_pdf(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_nbt_chunks(tmp_path):
    p = _nbt_pdf(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "Outer nested text",
        "Inner nested text"]


# ---------- 相邻表 ----------

def test_tt_two_tables(tmp_path):
    p = _tt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.type
            for e in doc.elements] == [
        "paragraph", "table",
        "table"]


def test_tt_contents(tmp_path):
    p = _tt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        1].content == \
        "| first |\n| --- |"
    assert doc.elements[
        2].content == \
        "| second |\n| --- |"


def test_tt_table_indices(
        tmp_path):
    p = _tt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        1].source_locator == {
        "table_index": 0,
        "section": 0}
    assert doc.elements[
        2].source_locator == {
        "table_index": 1,
        "section": 0}


def test_tt_no_ghost_paragraph(
        tmp_path):
    p = _tt_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert len(doc.elements) == 3
    assert doc.warnings == []


def test_tt_chunks(tmp_path):
    p = _tt_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert [c.text
            for c in doc.chunks] == [
        "intro", "| first |\n"
        "| --- |",
        "| second |\n| --- |"]
