"""evaluation/metrics.py 第四百八十八轮 edges 测试（Round 1044）。

补强 edges126 未触及的角度（第四百二十批，probe 实证）。

新角度（真实 parser 产物直喂 compute + 因果更正）：
- metrics 测试 126 轮全部手工构造 document；本批用
  process_single 真实解析 python-docx 文档后把
  to_dict() 直喂 compute_automatic_metrics（不经
  runner，metrics 模块层首次见真实产物）
- 真实 locator 形状锁定：{"paragraph_index": 0,
  "section": 0}——fallback parser 是 0 基索引且双键
  同现；edges126 板 0.75 的真因是 table 带 page 键
  （docx 侧禁 page/bbox），非 idx 0——本批以真实
  idx 0 → ratio 1.0 更正因果记录（键在场即有效，
  无 >= 1 下限）
- 真实 silent_drop 三变体：expectations paragraph 5 →
  drop 3（欠供）、paragraph 1 → drop 0（过供触底）、
  heading 2 → drop 2（类型缺席）
- 同屏因果对照：真实双 idx 0 板 1.0 vs 手造 page 键
  板 0.0——docx locator 真正的失效规则
- forbidden tokens 第五百一十五批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

from app.pipeline import process_single
import evaluation.metrics as metrics_mod
from evaluation.metrics import compute_automatic_metrics


def _real_doc(tmp_path):
    p = tmp_path / "a.docx"
    d = Document()
    d.add_paragraph("Hello world paragraph one.")
    d.add_paragraph("Second paragraph here.")
    d.save(str(p))
    doc, errors = process_single(p, tmp_path / "stub.json",
                                 parser_name="fallback",
                                 max_chars=50,
                                 write_json=False)
    assert errors == []
    return doc.to_dict()


def _e(t, eid, loc):
    return {"type": t, "element_id": eid, "content": "x",
            "parent_id": None, "confidence": 0.9,
            "metadata": {}, "source_locator": loc}


def _doc(elements):
    return {"elements": elements,
            "chunks": [{"chunk_id": "c", "text": "x",
                        "source_element_ids": ["e1"],
                        "metadata": {}}],
            "source_type": "docx", "document_id": "x",
            "schema_version": "0.1.0",
            "source_path": "a", "source_hash": "a" * 64,
            "parser_name": "fb", "parser_version": "1",
            "relations": [], "warnings": [], "errors": [],
            "metadata": {}}


# ---------- 真实 parser 产物 ----------

def test_real_locator_shape_batch242(tmp_path):
    dd = _real_doc(tmp_path)
    assert dd["elements"][0]["source_locator"] == {
        "paragraph_index": 0, "section": 0}
    assert dd["elements"][1]["source_locator"] == {
        "paragraph_index": 1, "section": 0}


def test_real_docx_metrics_board_batch242(tmp_path):
    dd = _real_doc(tmp_path)
    m = compute_automatic_metrics(dd, None, "docx", None)
    assert m["element_count_total"] == {"value": 2,
                                        "reason": None}
    assert m["element_count_by_type"] == {
        "value": {"paragraph": 2}, "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["pdf_locator_valid_ratio"] == {
        "value": None, "reason": "not_pdf_document"}
    assert m["silent_drop_count"] == {
        "value": None, "reason": "no_expectations"}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}
    assert m["schema_valid"] == {"value": True, "reason": None}


# ---------- 真实 silent_drop 三变体 ----------

def test_real_silent_drop_variants_batch242(tmp_path):
    dd = _real_doc(tmp_path)
    under = compute_automatic_metrics(
        dd, None, "docx",
        {"element_count_by_type": {"paragraph": 5}})
    over = compute_automatic_metrics(
        dd, None, "docx",
        {"element_count_by_type": {"paragraph": 1}})
    absent = compute_automatic_metrics(
        dd, None, "docx",
        {"element_count_by_type": {"heading": 2}})
    assert under["silent_drop_count"] == {"value": 3,
                                          "reason": None}
    assert over["silent_drop_count"] == {"value": 0,
                                         "reason": None}
    assert absent["silent_drop_count"] == {"value": 2,
                                           "reason": None}


# ---------- 因果对照：键在场 vs page 禁键 ----------

def test_idx0_presence_semantics_batch242():
    doc = _doc([_e("paragraph", "p1",
                   {"paragraph_index": 0, "section": 0}),
                _e("paragraph", "p2",
                   {"paragraph_index": 0, "section": 0})])
    m = compute_automatic_metrics(doc, None, "docx", None)
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}


def test_page_key_invalidates_docx_batch242():
    doc = _doc([_e("table", "t1", {"page": 1})])
    m = compute_automatic_metrics(doc, None, "docx", None)
    assert m["docx_locator_valid_ratio"] == {"value": 0.0,
                                             "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch242():
    src = _src()
    assert 'if "page" in loc or "bbox" in loc:' in src
    assert ("if not any(k in loc for k in "
            "structural_keys):" in src)
    assert '"paragraph_index",' in src


# ---------- forbidden tokens 第五百一十五批 ----------

def test_source_no_eval_batch242():
    assert "eval(" not in _src()


def test_source_no_exec_batch242():
    assert "exec(" not in _src()


def test_source_no_compile_batch242():
    assert "compile(" not in _src()


def test_source_no_globals_batch242():
    assert "globals(" not in _src()


def test_source_no_locals_batch242():
    assert "locals(" not in _src()


def test_source_no_os_system_batch242():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch242():
    assert "subprocess" not in _src()


def test_source_no_popen_batch242():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch242():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch242():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch242():
    assert "socket" not in _src()


def test_source_no_requests_batch242():
    assert "requests" not in _src()


def test_source_no_urllib_batch242():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch242():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch242():
    assert "yield" not in _src()


def test_source_no_async_await_batch242():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch242():
    assert "open(" not in _src()
