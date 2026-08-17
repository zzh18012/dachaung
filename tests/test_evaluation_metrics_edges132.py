"""evaluation/metrics.py 第五百二十三轮 edges 测试（Round 1079）。

补强 edges130-131 未触及的角度（第四百五十五批，probe 实证）。

新角度（文档中段 heading 的结构保证：heading 恒居 chunk 首）：
- intro 段 + add_heading("Late Title", level=1) + 两段的
  真实文档：elements types ['paragraph', 'heading',
  'paragraph', 'paragraph']
- **heading 强制 chunk 断开**：mc 200 下 chunks[0] 的
  source_element_ids 恰为 [intro e0000]（intro 独占一
  chunk，尽管 200 字符预算远未用尽）；chunks[1] 以
  heading id 打头——heading 恒 LEAD 其 chunk
- 由此真实管线里 heading_boundary_compliance 恒可达
  1.0：mc 200 与 mc 40 双预算均 {value 1.0, reason
  None}——0.0 分支只能靠手工合成板触达
- 无 heading 对照：纯段落文档 → {value None, reason
  no_heading_elements}
- forbidden tokens 第五百五十批（open 0）
"""

from __future__ import annotations

import inspect

from docx import Document

import evaluation.metrics as metrics_mod
from app.pipeline import process_single
from evaluation.metrics import compute_automatic_metrics


def _build(tmp_path, max_chars=200, with_heading=True):
    d = Document()
    d.add_paragraph("AAA intro paragraph before the heading.")
    if with_heading:
        d.add_heading("Late Title", level=1)
    d.add_paragraph("BBB body after the heading one.")
    d.add_paragraph("CCC body after the heading two.")
    p = tmp_path / "h.docx"
    d.save(str(p))
    doc, errors = process_single(
        p, tmp_path / "s.json", parser_name="fallback",
        max_chars=max_chars, write_json=False)
    assert errors == []
    return doc


# ---------- 中段 heading：双预算恒 1.0 ----------

def test_mid_heading_compliance_one_batch278(tmp_path):
    for mc in (200, 40):
        doc = _build(tmp_path, max_chars=mc)
        m = compute_automatic_metrics(doc.to_dict(), None,
                                      "docx", None)
        assert m["heading_boundary_compliance"] == {
            "value": 1.0, "reason": None}


# ---------- heading 强制 chunk 断开 ----------

def test_heading_forces_chunk_break_batch278(tmp_path):
    doc = _build(tmp_path, max_chars=200)
    dd = doc.to_dict()
    assert [e["type"] for e in dd["elements"]] == [
        "paragraph", "heading", "paragraph", "paragraph"]
    ids = [e["element_id"] for e in dd["elements"]]
    assert dd["chunks"][0]["source_element_ids"] == [ids[0]]
    assert dd["chunks"][1]["source_element_ids"] == [
        ids[1], ids[2], ids[3]]
    assert dd["chunks"][1]["source_element_ids"][0] \
        == ids[1]


# ---------- 无 heading 对照 ----------

def test_no_heading_null_batch278(tmp_path):
    doc = _build(tmp_path, with_heading=False)
    m = compute_automatic_metrics(doc.to_dict(), None,
                                  "docx", None)
    assert m["heading_boundary_compliance"] == {
        "value": None, "reason": "no_heading_elements"}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(metrics_mod)


def test_source_key_lines_batch278():
    src = _src()
    assert ('metrics["heading_boundary_compliance"] = '
            '_heading_boundary_ratio(elements, chunks)'
            in src)


# ---------- forbidden tokens 第五百五十批 ----------

def test_source_no_eval_batch278():
    assert "eval(" not in _src()


def test_source_no_exec_batch278():
    assert "exec(" not in _src()


def test_source_no_compile_batch278():
    assert "compile(" not in _src()


def test_source_no_globals_batch278():
    assert "globals(" not in _src()


def test_source_no_locals_batch278():
    assert "locals(" not in _src()


def test_source_no_os_system_batch278():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch278():
    assert "subprocess" not in _src()


def test_source_no_popen_batch278():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch278():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch278():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch278():
    assert "socket" not in _src()


def test_source_no_requests_batch278():
    assert "requests" not in _src()


def test_source_no_urllib_batch278():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch278():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch278():
    assert "yield" not in _src()


def test_source_no_async_await_batch278():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_zero_batch278():
    assert "open(" not in _src()
