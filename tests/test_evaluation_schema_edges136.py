"""evaluation/schema.py 第五百七十轮 edges 测试（Round 1240）。

补强 edges135 未触及的角度（第六百一十二批，probe 实证）。

新角度（水槽 DOCX 真板变异 / 定位负空间）：
- **elements 缺键**——pop →
  "'elements' is a required
  property" @ []（elements 根空
  路首锁，与 chunks 缺键成对）
- **表格 locator 免 table_index**
  ——pop table_index → VALID
  （docx 定位分支无逐键必填，
  if/then 只特判 pdf 与非
  pdf/docx）
- **section 收字符串**——"0"
  照过（section 无类型约束）
- **chunk 纯空白文本过**——
  "   " VALID（非空判定字面，
  不 strip）
- **paragraph_index 负数**——
  -1 → "less than the minimum
  of 0" @ 深路径（文档层该路径
  首锁，前史 minimum 全在清单层）
- **table_index 浮点**——0.5 →
  "is not of type 'integer'"
- **表格元素双空**——content
  None + resource None → anyOf
  回拒 @ ['elements', 3]（表格
  位 anyOf 首锁）
- forbidden tokens 第七百零六批（open 2）
"""

from __future__ import annotations

import copy
import inspect
import tempfile
from pathlib import Path

import pytest

import evaluation.schema as schema_mod
from evaluation.schema import EvalSchemaError, validate


@pytest.fixture()
def base_doc(tmp_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from app.pipeline import process_single
    doc = Document()
    doc.add_heading("Chapter One Title", 1)
    doc.add_paragraph("First para under chapter one.")
    doc.add_paragraph("   ")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "L"
    t.cell(0, 1).text = "R"
    para = doc.add_paragraph()
    para.add_run("Tail ")
    para.add_run("run split")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section body text.")
    p = tmp_path / "ks.docx"
    doc.save(str(p))
    d, errors = process_single(
        p, tmp_path / "o.json", parser_name="fallback",
        max_chars=120)
    assert errors == []
    return d.to_dict()


def _mut(base_doc, fn):
    d = copy.deepcopy(base_doc)
    fn(d)
    return d


def _head(ei):
    return ei.value.errors[0]


# ---------- 根必填 ----------

def test_elements_missing_root_path_batch438(base_doc):
    with pytest.raises(EvalSchemaError) as ei:
        validate(_mut(base_doc, lambda x: x.pop("elements")),
                 "document.schema.json")
    assert _head(ei)["message"] == \
        "'elements' is a required property"
    assert _head(ei)["path"] == []


# ---------- 定位负空间 ----------

def test_table_locator_without_table_index_valid_batch438(base_doc):
    d = _mut(base_doc, lambda x: x["elements"][3][
        "source_locator"].pop("table_index"))
    validate(d, "document.schema.json")


def test_section_string_valid_batch438(base_doc):
    d = _mut(base_doc, lambda x: x["elements"][0][
        "source_locator"].update({"section": "0"}))
    validate(d, "document.schema.json")


def test_chunk_whitespace_text_valid_batch438(base_doc):
    d = _mut(base_doc, lambda x: x["chunks"][2].update(
        {"text": "   "}))
    validate(d, "document.schema.json")


# ---------- 定位键类型/边界 ----------

def test_paragraph_index_negative_batch438(base_doc):
    with pytest.raises(EvalSchemaError) as ei:
        validate(_mut(base_doc, lambda x: x["elements"][0][
            "source_locator"].update(
                {"paragraph_index": -1})),
            "document.schema.json")
    assert _head(ei)["message"] == \
        "-1 is less than the minimum of 0"
    assert _head(ei)["path"] == [
        "elements", 0, "source_locator", "paragraph_index"]


def test_table_index_float_batch438(base_doc):
    with pytest.raises(EvalSchemaError) as ei:
        validate(_mut(base_doc, lambda x: x["elements"][3][
            "source_locator"].update(
                {"table_index": 0.5})),
            "document.schema.json")
    assert _head(ei)["message"] == \
        "0.5 is not of type 'integer'"
    assert _head(ei)["path"] == [
        "elements", 3, "source_locator", "table_index"]


# ---------- 表格元素双空 ----------

def test_table_element_both_none_batch438(base_doc):
    with pytest.raises(EvalSchemaError) as ei:
        validate(_mut(base_doc, lambda x: x["elements"][3].update(
            {"content": None, "resource_path": None})),
            "document.schema.json")
    assert "is not valid under any of the given schemas" \
        in _head(ei)["message"]
    assert _head(ei)["path"] == ["elements", 3]


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(schema_mod)


def test_source_key_lines_batch438():
    src = _src()
    assert "class EvalSchemaError(Exception):" in src
    assert "def validate(instance" in src


# ---------- forbidden tokens 第七百零六批 ----------

def test_source_no_eval_batch438():
    assert "eval(" not in _src()


def test_source_no_exec_batch438():
    assert "exec(" not in _src()


def test_source_no_compile_batch438():
    assert "compile(" not in _src()


def test_source_no_globals_batch438():
    assert "globals(" not in _src()


def test_source_no_locals_batch438():
    assert "locals(" not in _src()


def test_source_no_os_system_batch438():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch438():
    assert "subprocess" not in _src()


def test_source_no_popen_batch438():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch438():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch438():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch438():
    assert "socket" not in _src()


def test_source_no_requests_batch438():
    assert "requests" not in _src()


def test_source_no_urllib_batch438():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch438():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch438():
    assert "yield" not in _src()


def test_source_no_async_await_batch438():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch438():
    assert _src().count("open(") == 2
