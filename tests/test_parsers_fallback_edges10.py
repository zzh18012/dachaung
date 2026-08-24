r"""app/parsers/fallback_parser.py 边角测试 - 第十轮（Round 1383）。

补强 docx 样式元数据面（probe 实证，'style'/'empty' 键在
fallback 测试零覆盖）：
- 每个段落/标题 metadata 三键 {level, style, empty}
- add_heading(x, 0) 是 Title 样式 → level 1
- H1-H9 全透传（level 9 不截断）
- 空段落不跳过——content 是占位符 '(空段落)' 且 empty=True
- Normal 段落 level 0；非内置样式（Intense Quote）原样保留
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser


def _parse(tmp_path, builder, name="d.docx"):
    p = tmp_path / name
    builder(p)
    return FallbackParser().parse(
        p, compute_file_hash(p))


def _levels(p):
    d = Document()
    d.add_heading("H1", 1)
    d.add_heading("H2", 2)
    d.add_heading("H3", 3)
    d.add_heading("H4", 4)
    d.add_heading("H9", 9)
    d.save(str(p))


def _styled(p):
    d = Document()
    d.add_heading("Title style", 0)
    d.add_paragraph("")
    d.add_paragraph("quote text",
                    style="Intense Quote")
    d.add_paragraph("normal one")
    d.save(str(p))


# ---------- heading 级别 ----------

def test_heading_levels_passthrough(tmp_path):
    doc = _parse(tmp_path, _levels)
    levels = [e.metadata["level"]
              for e in doc.elements]
    assert levels == [1, 2, 3, 4, 9]


def test_heading_level9_not_clamped(tmp_path):
    doc = _parse(tmp_path, _levels)
    assert doc.elements[4].metadata["level"] == 9


def test_heading_style_names(tmp_path):
    doc = _parse(tmp_path, _levels)
    styles = [e.metadata["style"]
              for e in doc.elements]
    assert styles == [
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 9"]


def test_headings_empty_false(tmp_path):
    doc = _parse(tmp_path, _levels)
    assert all(e.metadata["empty"] is False
               for e in doc.elements)


# ---------- Title 样式 ----------

def test_title_style_is_level1(tmp_path):
    doc = _parse(tmp_path, _styled)
    t = doc.elements[0]
    assert t.type == "heading"
    assert t.metadata == {
        "level": 1, "style": "Title",
        "empty": False}


# ---------- 空段落占位符 ----------

def test_empty_paragraph_placeholder(tmp_path):
    doc = _parse(tmp_path, _styled)
    empty = doc.elements[1]
    assert empty.type == "paragraph"
    assert empty.content == "(空段落)"


def test_empty_paragraph_flag(tmp_path):
    doc = _parse(tmp_path, _styled)
    assert doc.elements[1].metadata["empty"] \
        is True


def test_empty_paragraph_kept_not_skipped(
        tmp_path):
    doc = _parse(tmp_path, _styled)
    assert len(doc.elements) == 4


# ---------- Normal / 自定义样式 ----------

def test_normal_level_zero(tmp_path):
    doc = _parse(tmp_path, _styled)
    normal = doc.elements[3]
    assert normal.metadata["level"] == 0
    assert normal.metadata["style"] == \
        "Normal"


def test_custom_style_kept(tmp_path):
    doc = _parse(tmp_path, _styled)
    quote = doc.elements[2]
    assert quote.metadata["style"] == \
        "Intense Quote"


def test_non_heading_level_zero(tmp_path):
    doc = _parse(tmp_path, _styled)
    assert all(
        e.metadata["level"] == 0
        for e in doc.elements
        if e.type == "paragraph")


# ---------- locator ----------

def test_styled_paragraph_indexes(tmp_path):
    doc = _parse(tmp_path, _styled)
    idxs = [e.source_locator[
        "paragraph_index"]
        for e in doc.elements]
    assert idxs == [0, 1, 2, 3]


# ---------- schema ----------

def test_styled_passes_schema(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path, _styled)
    assert is_valid(doc.to_dict())


def test_levels_passes_schema(tmp_path):
    from app.schema import is_valid
    doc = _parse(tmp_path, _levels)
    assert is_valid(doc.to_dict())
