r"""app/parsers/fallback_parser.py 边角测试 - 第四十一轮（Round 1432）。

新角度（probe 实证）纵线聚类阈值 + docx 分页符（历史考察
全在横向）：
- 同 x 两行 dy=2pt：字符级交错 'ULopwpeerr lliinnee
  wwoorrddss'（与 R1431 横向交错同机制）
- dy=5/8/13pt：不分元素——两行词空格连接成单元素
  'Upper line words Lower line words'（跨行并词）
- docx run 级分页符（w:br type=page）：**无任何字符**——
  'before page breakafter page break' 直接拼接（对照 R1424
  文本 w:br → '\n'）
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from app.hash import compute_file_hash
from app.parsers.fallback_parser import \
    FallbackParser
from app.pipeline import process_single


def _build(content):
    objs = {
        5: (b"<< /Type /Font /Subtype "
            b"/Type1 /BaseFont "
            b"/Helvetica >>"),
        1: (b"<< /Type /Catalog "
            b"/Pages 2 0 R >>"),
        2: (b"<< /Type /Pages "
            b"/Kids [3 0 R] "
            b"/Count 1 >>"),
        3: (b"<< /Type /Page /Parent "
            b"2 0 R /MediaBox "
            b"[0 0 612 792] "
            b"/Resources << /Font "
            b"<< /F1 5 0 R >> >> "
            b"/Contents 4 0 R >>"),
        4: (b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n" + content
            + b"\nendstream"),
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for oid in sorted(objs):
        offsets[oid] = len(out)
        out += (f"{oid} 0 obj\n"
                .encode()
                + objs[oid]
                + b"\nendobj\n")
    xref_pos = len(out)
    out += b"xref\n0 6\n" \
        b"0000000000 65535 f \n"
    for oid in range(1, 6):
        out += ("%010d 00000 n \n"
                % offsets[oid]).encode()
    out += (b"trailer\n<< /Size 6 "
            b"/Root 1 0 R >>\n"
            b"startxref\n"
            + str(xref_pos).encode()
            + b"\n%%EOF")
    return bytes(out)


def _v_pdf(tmp_path, dy):
    c = (b"BT /F1 12 Tf 72 700 Td "
         b"(Upper line words) Tj ET "
         b"BT /F1 12 Tf 72 "
         + str(700 - dy).encode()
         + b" Td (Lower line words) "
         b"Tj ET")
    p = tmp_path / f"v{dy}.pdf"
    p.write_bytes(_build(c))
    return p


def _pb_docx(tmp_path):
    d = Document()
    para = d.add_paragraph()
    para.add_run("before page break")
    para.add_run().add_break(
        WD_BREAK.PAGE)
    para.add_run("after page break")
    p = tmp_path / "pb.docx"
    d.save(str(p))
    return p


# ---------- 纵线聚类 ----------

def test_dy2_interleaved(tmp_path):
    p = _v_pdf(tmp_path, 2)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.elements[
        0].content == \
        "ULopwpeerr lliinnee " \
        "wwoorrddss"


def test_dy5_single_element(
        tmp_path):
    p = _v_pdf(tmp_path, 5)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "Upper line words "
        "Lower line words"]


def test_dy8_single_element(
        tmp_path):
    p = _v_pdf(tmp_path, 8)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert len(doc.elements) == 1


def test_dy13_single_element(
        tmp_path):
    p = _v_pdf(tmp_path, 13)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert len(doc.elements) == 1
    assert doc.elements[
        0].content == \
        "Upper line words " \
        "Lower line words"


def test_dy13_chunk(tmp_path):
    p = _v_pdf(tmp_path, 13)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[0].text == (
        "Upper line words "
        "Lower line words")


# ---------- docx 分页符 ----------

def test_pb_no_separator(
        tmp_path):
    p = _pb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert [e.content
            for e in doc.elements] == [
        "before page break"
        "after page break"]


def test_pb_single_paragraph(
        tmp_path):
    p = _pb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert len(doc.elements) == 1
    assert doc.elements[
        0].type == "paragraph"
    assert "\n" not in \
        doc.elements[0].content


def test_pb_no_warnings(tmp_path):
    p = _pb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert doc.warnings == []


def test_pb_schema_valid(
        tmp_path):
    from app.schema import is_valid
    p = _pb_docx(tmp_path)
    doc = FallbackParser().parse(
        p, compute_file_hash(p))
    assert is_valid(doc.to_dict())


def test_pb_chunk(tmp_path):
    p = _pb_docx(tmp_path)
    doc, errors = process_single(
        p, None,
        parser_name="fallback",
        max_chars=800)
    assert errors == []
    assert doc.chunks[
        0].text == (
        "before page break"
        "after page break")


def test_v_no_warnings(tmp_path):
    for dy in (2, 13):
        p = _v_pdf(tmp_path, dy)
        doc = FallbackParser().parse(
            p, compute_file_hash(p))
        assert doc.warnings == []
