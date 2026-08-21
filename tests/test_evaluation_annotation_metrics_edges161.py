"""evaluation/annotation_metrics.py 第五百九十二轮 edges 测试（Round 1325）。

补强 edges160 未触及的角度（第六百九十七批，probe 实证）。

新角度（锚形态容忍 / 重复幂等 / docx 双锚）：
- **anchor 附加 reason 键**——
  {marker, position,
  reason} 三键 → 照常
  HIT (1/15, 1.0, 0.125)
  （schema 允许的 optional
  reason 运行时容忍首锁）
- **position "AFTER"**
  ——非 "before" 一律走
  after 分支（精确
  `== "before"` 比较
  锁）→ 同 HIT
- **重复 marker 幂等**——
  [Word3, Word5, Word5]
  → gt 2 不胀 3：与
  [Word3, Word5] 完全
  同值 {2/15, 1.0,
  4/17}（第二现
  search_from 顺移后
  missing 首锁）
- **标注附加键忽略**——
  annotator / date /
  heading_order /
  figure_caption_pairs
  齐上 → 只读
  chunk_boundary_anchors，
  照常 HIT
  （heading_order 零消费
  锁）
- **docx 双锚**——
  PicDoc + before 双
  marker → {1.0, 0.5,
  2/3}（2-chunk 板双锚
  几何首锁）
- forbidden tokens 第七百七十批（open 0）
"""

from __future__ import annotations

import inspect
import tempfile
from pathlib import Path

import evaluation.annotation_metrics as am_mod
from app.pipeline import process_single
from docx import Document
from evaluation.annotation_metrics import \
    chunk_boundary_prf


def _wrap(s: bytes) -> bytes:
    objs = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        2: b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        3: (b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 400 800]"
            b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>"),
        4: (b"<</Length " + str(len(s)).encode()
            + b">>stream\n" + s + b"\nendstream "),
        5: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += (str(num).encode() + b" 0 obj"
                + objs[num] + b"endobj\n")
    xp = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for num in range(1, 6):
        out += ("%010d 00000 n \n" % offsets[num]).encode()
    out += (b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n"
            + str(xp).encode() + b"\n%%EOF\n")
    return bytes(out)


LONG = " ".join("Word%d." % i for i in range(60))
COMBO = ("BT /F1 12 Tf 10 700 Td (%s) Tj ET\n"
         % ("A" * 80)
         + "BT /F1 12 Tf 10 660 Td (%s) Tj ET\n"
         % LONG).encode()

HIT = ({"value": 1 / 15, "reason": None},
       {"value": 1.0, "reason": None},
       {"value": 0.125, "reason": None})
PAIR = ({"value": 2 / 15, "reason": None},
        {"value": 1.0, "reason": None},
        {"value": 4 / 17, "reason": None})


def _pdf_doc():
    with tempfile.TemporaryDirectory() as td:
        tp = Path(td)
        (tp / "c.pdf").write_bytes(_wrap(COMBO))
        doc, errors = process_single(
            tp / "c.pdf", tp / "o.json",
            parser_name="fallback", max_chars=32)
        assert errors == []
        return doc.to_dict()


def _docx_doc():
    with tempfile.TemporaryDirectory() as td:
        tp = Path(td)
        d = Document()
        d.add_heading("PicDoc", level=1)
        d.add_paragraph("Text before picture.")
        d.add_paragraph("Second para here.")
        d.save(str(tp / "c.docx"))
        doc, errors = process_single(
            tp / "c.docx", tp / "o.json",
            parser_name="fallback", max_chars=32)
        assert errors == []
        return doc.to_dict()


def _trio(doc, anchors):
    ann = {"annotation_version": "1.0", "doc_id": "x",
           "chunk_boundary_anchors": anchors}
    r = chunk_boundary_prf(doc, ann, 30)
    return (r["chunk_boundary_precision"],
            r["chunk_boundary_recall"],
            r["chunk_boundary_f1"])


# ---------- anchor 附加 reason 键 ----------

def test_anchor_reason_extra_hit_batch523():
    assert _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "after",
         "reason": "checked"}]) == HIT


def test_anchor_reason_extra_recall_batch523():
    t = _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "after",
         "reason": "anything"}])
    assert t[1] == {"value": 1.0, "reason": None}


# ---------- position 非字面 before ----------

def test_position_upper_after_batch523():
    assert _trio(_pdf_doc(), [
        {"marker": "Word3.",
         "position": "AFTER"}]) == HIT


def test_position_after_default_equivalence_batch523():
    explicit = _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "after"}])
    upper = _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "AFTER"}])
    assert explicit == upper


# ---------- 重复 marker 幂等 ----------

def test_duplicate_marker_same_as_pair_batch523():
    dup = _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "after"},
        {"marker": "Word5.", "position": "after"},
        {"marker": "Word5.", "position": "after"}])
    assert dup == PAIR


def test_plain_pair_baseline_batch523():
    assert _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "after"},
        {"marker": "Word5.",
         "position": "after"}]) == PAIR


def test_duplicate_f1_four_seventeenths_batch523():
    t = _trio(_pdf_doc(), [
        {"marker": "Word3.", "position": "after"},
        {"marker": "Word5.", "position": "after"},
        {"marker": "Word5.", "position": "after"}])
    assert t[2]["value"] == 4 / 17


# ---------- 标注附加键忽略 ----------

def test_full_annotation_extra_keys_hit_batch523():
    ann = {
        "annotation_version": "1.0", "doc_id": "x",
        "annotator": "a", "date": "2026-01-01",
        "heading_order": [{"level": 1, "text": "H"}],
        "figure_caption_pairs": [
            {"figure_marker": "f",
             "caption_text": "c"}],
        "chunk_boundary_anchors": [
            {"marker": "Word3.",
             "position": "after"}]}
    r = chunk_boundary_prf(_pdf_doc(), ann, 30)
    assert r["chunk_boundary_precision"] == {
        "value": 1 / 15, "reason": None}
    assert r["chunk_boundary_f1"] == {
        "value": 0.125, "reason": None}


def test_heading_order_alone_missing_batch523():
    ann = {
        "annotation_version": "1.0", "doc_id": "x",
        "heading_order": [{"level": 1, "text": "H"}]}
    r = chunk_boundary_prf(_pdf_doc(), ann, 30)
    for k in ("chunk_boundary_precision",
              "chunk_boundary_recall",
              "chunk_boundary_f1"):
        assert r[k] == {
            "value": None,
            "reason": "no_ground_truth_anchors"}


# ---------- docx 双锚 ----------

def test_docx_dual_anchor_trio_batch523():
    assert _trio(_docx_doc(), [
        {"marker": "PicDoc", "position": "after"},
        {"marker": "before",
         "position": "after"}]) == (
        {"value": 1.0, "reason": None},
        {"value": 0.5, "reason": None},
        {"value": 2 / 3, "reason": None})


def test_docx_dual_anchor_precision_one_batch523():
    t = _trio(_docx_doc(), [
        {"marker": "PicDoc", "position": "after"},
        {"marker": "before", "position": "after"}])
    assert t[0]["value"] == 1.0
    assert t[1]["value"] == 0.5


def test_docx_chunks_two_batch523():
    assert len(_docx_doc()["chunks"]) == 2


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(am_mod)


def test_source_exact_before_compare_batch523():
    assert 'position == "before"' in _src()


def test_source_anchors_read_batch523():
    assert "chunk_boundary_anchors" in _src()


# ---------- forbidden tokens 第七百七十批 ----------

def test_source_no_eval_batch523():
    assert "eval(" not in _src()


def test_source_no_exec_batch523():
    assert "exec(" not in _src()


def test_source_no_compile_batch523():
    assert "compile(" not in _src()


def test_source_no_globals_batch523():
    assert "globals(" not in _src()


def test_source_no_locals_batch523():
    assert "locals(" not in _src()


def test_source_no_os_system_batch523():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch523():
    assert "subprocess" not in _src()


def test_source_no_popen_batch523():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch523():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch523():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch523():
    assert "socket" not in _src()


def test_source_no_requests_batch523():
    assert "requests" not in _src()


def test_source_no_urllib_batch523():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch523():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch523():
    assert "yield" not in _src()


def test_source_no_async_await_batch523():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_0_batch523():
    assert _src().count("open(") == 0
