"""evaluation/runner.py 第六百一十一轮 edges 测试（Round 1167）。

补强 edges181 未触及的角度（第五百三十九批，probe 实证）。

新角度（DOCX 双表 table_index 计数器）：
- **双表元素序**——表 A / 段落 / 表 B 同文档 →
  [table, paragraph, table] 文档序；table_index
  0/1 计数器跨表递增；段落仍 paragraph_index 0
  （段落计数独立于表，首锁）
- **双表三块**——3 chunks [isolated_table,
  sequential, isolated_table] 各 1 源；两表
  markdown 各自独立成块
- **中间段落界**——marker "tables." after → GT
  落段尾 = 第 2 界 → P 1/2 / R 1.0 / F1 2/3
- **流首无界**——marker "A1" before → GT 落流
  绝对起点，首界距 38 字 > tol 30 → 全 0.0
- forbidden tokens 第六百三十九批（open 2）
"""

from __future__ import annotations

import inspect
import json

import evaluation.runner as runner_mod
from docx import Document
from evaluation.manifest import load_manifest
from evaluation.runner import run_evaluation
from app.pipeline import process_single


def _board(tmp_path, doc_id, anchors=None):
    (tmp_path / "s").mkdir(exist_ok=True)
    d = Document()
    ta = d.add_table(rows=2, cols=2)
    ta.cell(0, 0).text = "A1"
    ta.cell(0, 1).text = "A2"
    ta.cell(1, 0).text = "B1"
    ta.cell(1, 1).text = "B2"
    d.add_paragraph("Text separating two tables.")
    tb = d.add_table(rows=1, cols=2)
    tb.cell(0, 0).text = "X1"
    tb.cell(0, 1).text = "Y1"
    d.save(str(tmp_path / "s" / f"{doc_id}.docx"))
    docs = [{"doc_id": doc_id, "path": f"s/{doc_id}.docx",
             "source_type": "docx"}]
    if anchors is not None:
        (tmp_path / "a").mkdir(exist_ok=True)
        (tmp_path / "a" / "a.json").write_text(json.dumps({
            "annotation_version": "1.0", "doc_id": doc_id,
            "chunk_boundary_anchors": anchors}), encoding="utf-8")
        docs[0]["annotation_file"] = "a/a.json"
    mf = tmp_path / f"m{doc_id}.json"
    mf.write_text(json.dumps({"manifest_version": "1.0",
                              "devset_status": "incomplete",
                              "documents": docs}), encoding="utf-8")
    return load_manifest(mf, project_root=tmp_path)


# ---------- 双表元素序 ----------

def test_two_tables_element_order_batch365(tmp_path):
    _board(tmp_path, "tt")
    doc, errors = process_single(
        tmp_path / "s" / "tt.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    assert errors == []
    els = doc.to_dict()["elements"]
    assert [e["type"] for e in els] == [
        "table", "paragraph", "table"]
    assert els[0]["content"] == (
        "| A1 | A2 |\n| --- | --- |\n| B1 | B2 |")
    assert els[1]["content"] == "Text separating two tables."
    assert els[2]["content"] == (
        "| X1 | Y1 |\n| --- | --- |")
    assert els[0]["source_locator"] == {"table_index": 0,
                                        "section": 0}
    assert els[1]["source_locator"] == {"paragraph_index": 0,
                                        "section": 0}
    assert els[2]["source_locator"] == {"table_index": 1,
                                        "section": 0}


# ---------- 双表三块 ----------

def test_two_tables_chunks_batch365(tmp_path):
    _board(tmp_path, "tt2")
    doc, errors = process_single(
        tmp_path / "s" / "tt2.docx", tmp_path / "o.json",
        parser_name="fallback", max_chars=200)
    chunks = doc.to_dict()["chunks"]
    assert [c["metadata"]["strategy"] for c in chunks] == [
        "isolated_table", "sequential", "isolated_table"]
    assert all(len(c["source_element_ids"]) == 1
               for c in chunks)
    assert chunks[1]["text"] == "Text separating two tables."


# ---------- 指标 ----------

def test_two_tables_metrics_batch365(tmp_path):
    r = run_evaluation(_board(tmp_path, "tt3"),
                       tmp_path / "r.json", parser_name="fallback",
                       max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["element_count_by_type"] == {
        "value": {"table": 2, "paragraph": 1},
        "reason": None}
    assert m["docx_locator_valid_ratio"] == {"value": 1.0,
                                             "reason": None}
    assert m["text_preservation_equal"] == {"value": True,
                                            "reason": None}


# ---------- 中间段落界 ----------

def test_mid_paragraph_junction_batch365(tmp_path):
    r = run_evaluation(_board(tmp_path, "tt4", [
        {"marker": "tables.", "position": "after"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.5, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 1.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.6666666666666666, "reason": None}


# ---------- 流首无界 ----------

def test_stream_start_miss_batch365(tmp_path):
    r = run_evaluation(_board(tmp_path, "tt5", [
        {"marker": "A1", "position": "before"}]),
        tmp_path / "r.json", parser_name="fallback",
        max_chars=200)
    m = r["per_doc"][0]["metrics"]
    assert m["chunk_boundary_precision"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_recall"] == {
        "value": 0.0, "reason": None}
    assert m["chunk_boundary_f1"] == {
        "value": 0.0, "reason": None}


# ---------- 源码补强 ----------

def _src():
    return inspect.getsource(runner_mod)


def test_source_identifier_counts_batch365():
    src = _src()
    assert src.count("run_evaluation") == 2
    assert src.count("metrics") == 13
    assert src.count("per_doc") == 12


# ---------- forbidden tokens 第六百三十九批 ----------

def test_source_no_eval_batch365():
    assert "eval(" not in _src()


def test_source_no_exec_batch365():
    assert "exec(" not in _src()


def test_source_no_compile_batch365():
    assert "compile(" not in _src()


def test_source_no_globals_batch365():
    assert "globals(" not in _src()


def test_source_no_locals_batch365():
    assert "locals(" not in _src()


def test_source_no_os_system_batch365():
    assert "os.system" not in _src()


def test_source_no_subprocess_batch365():
    assert "subprocess" not in _src()


def test_source_no_popen_batch365():
    assert "popen" not in _src()


def test_source_no_yaml_load_batch365():
    assert "yaml.load" not in _src()


def test_source_no_pickle_load_batch365():
    assert "pickle.load" not in _src()


def test_source_no_socket_batch365():
    assert "socket" not in _src()


def test_source_no_requests_batch365():
    assert "requests" not in _src()


def test_source_no_urllib_batch365():
    assert "urllib" not in _src()


def test_source_no_shutil_rmtree_batch365():
    assert "shutil.rmtree" not in _src()


def test_source_no_yield_batch365():
    assert "yield" not in _src()


def test_source_no_async_await_batch365():
    assert "async " not in _src()
    assert "await " not in _src()


def test_source_open_count_is_2_batch365():
    assert _src().count("open(") == 2
