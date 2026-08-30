# -*- coding: utf-8 -*-
"""Stage 6 批次 7 holdout 合成 docx 夹具一次性生成器（溯源脚本）。

纪律（docs/table-caption-relation-contract.md §6 + 2026-08-30 裁决⑤）：
- 夹具生成一次后字节固定，sha256 登记 ADOPTION.md；
- 运行时不得重新生成——本脚本带防重入守卫（目标已存在即拒绝），
  测试与 holdout 运行永远不调用它。

结构（供期望手工推导，生成后打印实际 body 大纲核对）：四 case 覆盖
裁决三类 + 前缀互斥负例——
- T1 表题注在上：caption "Table 1. ..." 紧邻 table 之前 → 期望 1 条
  table_has_caption（from=table, to=caption）
- T2 无题注：普通段落紧邻 table 之前 → 期望零
- T3 孤立表题注段落：caption 后无邻接表 → 期望零
- T4 图题注紧邻表上：caption "Figure 3. ..." 在 table 之前 → 前缀
  互斥，期望零
本夹具无 image（图/表前缀互斥用 caption 分类即可检验，无需图片字节）。
"""

from __future__ import annotations

import hashlib
import sys
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "samples" / "synthetic" / "holdout-table-caption"
OUT = OUT_DIR / "holdout-table-caption.docx"


def _fill(table, cells: list[list[str]]) -> None:
    for r, row in enumerate(cells):
        for c, text in enumerate(row):
            table.cell(r, c).text = text


def main() -> None:
    if OUT.exists():
        sys.exit(f"FATAL: 夹具已存在，禁止重新生成（裁决⑤ 字节固定）: {OUT}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    doc.add_heading("Synthetic table caption holdout", level=1)
    doc.add_paragraph("Intro paragraph.")
    # T1 表题注在上（命中）
    doc.add_paragraph("Table 1. Module status matrix")
    _fill(doc.add_table(rows=2, cols=2), [["S1A", "S1B"], ["S2A", "S2B"]])
    # T2 无题注（零）
    doc.add_paragraph("T2 table has no caption before it.")
    _fill(doc.add_table(rows=2, cols=2), [["U1A", "U1B"], ["U2A", "U2B"]])
    # T3 孤立表题注段落（无邻接表，零）。注意：分类用 _CAPTION_RE 的
    # 数字须紧跟前缀 token（仅空白），"表格 2、" 不会被分类为 caption，
    # 故用 "表 2、"（推导阶段发现并修正，见 expectations 推导注记）。
    doc.add_paragraph("表 2、孤立说明")
    doc.add_paragraph("Plain paragraph after orphan caption.")
    # T4 图题注紧邻表上（前缀互斥，零）
    doc.add_paragraph("Figure 3. demo")
    _fill(doc.add_table(rows=2, cols=2), [["V1A", "V1B"], ["V2A", "V2B"]])
    doc.add_paragraph("Closing paragraph.")
    doc.save(str(OUT))

    print(f"fixture: {OUT}")
    print(f"sha256:  {hashlib.sha256(OUT.read_bytes()).hexdigest()}")
    print("--- body outline (authored structure, for hand derivation) ---")
    from docx.oxml.ns import qn

    para_i = 0
    tbl_i = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            p = Paragraph(child, doc)
            print(f"w:p para_index={para_i} style={p.style.name!r} text={p.text!r}")
            para_i += 1
        elif child.tag == qn("w:tbl"):
            print(f"w:tbl table_index={tbl_i}")
            tbl_i += 1


if __name__ == "__main__":
    main()
