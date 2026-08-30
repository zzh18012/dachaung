# -*- coding: utf-8 -*-
"""Stage 6 批次 5 holdout 合成夹具一次性生成器（溯源脚本）。

纪律（docs/table-linearization-contract.md §5 + 2026-08-30 裁决⑤ 沿用）：
- 四个 fixture（docx/md/html/ipynb）生成一次后字节固定，sha256 登记
  ADOPTION.md；运行时不得重新生成——本脚本带防重入守卫（任一目标已
  存在即拒绝），测试与 holdout 运行永远不调用它。

覆盖面（契约 §5）：
- docx：多段落单元格（\\n→<br>）、合并单元格（重复语义）、含 | 单元格、
  空单元格、普通对照表；
- md：\\| 转义、<br>、参差行；
- html：含 | 单元格、th/thead（验证不特殊化）；
- ipynb：markdown cell 内 pipe 表格 + code/raw cell 无表格（裁决⑦）。
"""

from __future__ import annotations

import hashlib
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "samples" / "synthetic" / "holdout-table"

MD_TEXT = (
    "| h1 | h2 | h3 |\n"
    "| --- | --- | --- |\n"
    "| a\\|b | c | d |\n"
    "| x | y |\n"
    "| p<br>q |  | z |\n"
)

HTML_TEXT = (
    "<table>\n"
    "<tr><th>A</th><th>B</th></tr>\n"
    "<tr><td>x|y</td><td></td></tr>\n"
    "</table>\n"
)

IPYNB_DATA = {
    "nbformat": 4,
    "nbformat_minor": 5,
    "metadata": {},
    "cells": [
        {"cell_type": "markdown",
         "source": ["| h | k |\n", "| --- | --- |\n", "| v\\|w |  |\n"]},
        {"cell_type": "code", "source": ["x = 1\n"]},
        {"cell_type": "raw", "source": ["not a table\n"]},
    ],
}


def _build_docx(path: Path) -> None:
    import docx as pydocx

    d = pydocx.Document()
    d.add_paragraph("Intro.")
    t1 = d.add_table(rows=2, cols=2)
    c = t1.cell(0, 0)
    c.text = "p1"
    c.add_paragraph("p2")
    t1.cell(0, 1).text = "h2"
    t1.cell(1, 0).text = "a|b"
    t1.cell(1, 1).text = ""
    t2 = d.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "m"
    t2.cell(0, 0).merge(t2.cell(0, 1))
    t2.cell(1, 0).text = "x"
    t2.cell(1, 1).text = "y"
    d.add_paragraph("End.")
    d.save(str(path))


def main() -> None:
    targets = {
        "holdout-table.md": None,
        "holdout-table.html": None,
        "holdout-table.docx": _build_docx,
        "holdout-table.ipynb": None,
    }
    existing = [n for n in targets if (OUT_DIR / n).exists()]
    if existing:
        sys.exit(f"FATAL: 夹具已存在，禁止重新生成（裁决⑤ 字节固定）: {existing}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    (OUT_DIR / "holdout-table.md").write_text(MD_TEXT, encoding="utf-8", newline="")
    (OUT_DIR / "holdout-table.html").write_text(HTML_TEXT, encoding="utf-8", newline="")
    (OUT_DIR / "holdout-table.ipynb").write_text(
        json.dumps(IPYNB_DATA, ensure_ascii=False, indent=1) + "\n",
        encoding="utf-8", newline="",
    )
    _build_docx(OUT_DIR / "holdout-table.docx")

    for name in targets:
        p = OUT_DIR / name
        print(f"{name}  sha256={hashlib.sha256(p.read_bytes()).hexdigest()}  size={p.stat().st_size}")
    print("--- authored structure (for hand derivation) ---")
    print("[md] 5 行 pipe 表：header 3 列；行2 含 \\| 转义；行3 参差 2 列；行4 含 <br> 与空 cell")
    print("[html] 1 表：th 行 + td 行（含 | 与空 cell）")
    print("[ipynb] cell0 markdown pipe 表（含 \\| 与空 cell）；cell1 code；cell2 raw")
    print("[docx] p(Intro.) + t1(多段落 cell/h2/a|b/空) + t2(首行合并 m) + p(End.)")


if __name__ == "__main__":
    main()
