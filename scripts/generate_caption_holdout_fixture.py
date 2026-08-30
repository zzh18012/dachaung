# -*- coding: utf-8 -*-
"""Stage 6 批次 4 holdout 合成 docx 夹具一次性生成器（溯源脚本）。

纪律（docs/caption-relation-contract.md §6 + 2026-08-30 裁决⑤）：
- 夹具生成一次后字节固定，sha256 登记 ADOPTION.md；
- 运行时不得重新生成——本脚本带防重入守卫（目标已存在即拒绝），
  测试与 holdout 运行永远不调用它。

结构（供期望手工推导，生成后打印实际 body 大纲核对）：
- 两图各自紧邻下一段是图题注（Figure 1. / 图 2、）→ 期望 2 条 relation
- 一图紧邻下一段是表题注（Table 1.）→ 前缀集排除，期望零 relation
- 一图后无题注 → 期望零 relation
"""

from __future__ import annotations

import hashlib
import io
import sys
from pathlib import Path

import docx
from docx.shared import Inches
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "samples" / "synthetic" / "holdout-caption"
OUT = OUT_DIR / "holdout-caption.docx"


def _png(color: tuple[int, int, int]) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (8, 8), color).save(buf, format="PNG")
    return buf.getvalue()


def main() -> None:
    if OUT.exists():
        sys.exit(f"FATAL: 夹具已存在，禁止重新生成（裁决⑤ 字节固定）: {OUT}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    doc.add_heading("Synthetic caption holdout", level=1)
    doc.add_paragraph("Intro paragraph before any figure.")
    doc.add_paragraph().add_run().add_picture(
        io.BytesIO(_png((255, 0, 0))), width=Inches(1.5)
    )
    doc.add_paragraph("Figure 1. Flow overview")
    doc.add_paragraph("Body paragraph between figures.")
    doc.add_paragraph().add_run().add_picture(
        io.BytesIO(_png((0, 160, 0))), width=Inches(1.5)
    )
    doc.add_paragraph("图 2、架构示意")
    doc.add_paragraph().add_run().add_picture(
        io.BytesIO(_png((0, 0, 255))), width=Inches(1.5)
    )
    doc.add_paragraph("Table 1. Metrics")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A1"
    t.cell(0, 1).text = "B1"
    t.cell(1, 0).text = "A2"
    t.cell(1, 1).text = "B2"
    doc.add_paragraph().add_run().add_picture(
        io.BytesIO(_png((128, 128, 128))), width=Inches(1.5)
    )
    doc.add_paragraph("Closing paragraph with no caption.")
    doc.save(str(OUT))

    print(f"fixture: {OUT}")
    print(f"sha256:  {hashlib.sha256(OUT.read_bytes()).hexdigest()}")
    print("--- body outline (authored structure, for hand derivation) ---")
    from docx.oxml.ns import qn

    para_i = 0
    tbl_i = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            p = Paragraph(child, doc)
            n_drawings = len(list(child.iter(qn("w:drawing"))))
            tag = "w:p"
            extra = f" drawings={n_drawings}" if n_drawings else ""
            print(
                f"{tag} para_index={para_i} style={p.style.name!r} "
                f"text={p.text!r}{extra}"
            )
            para_i += 1
        elif child.tag == qn("w:tbl"):
            print(f"w:tbl table_index={tbl_i}")
            tbl_i += 1


if __name__ == "__main__":
    main()
